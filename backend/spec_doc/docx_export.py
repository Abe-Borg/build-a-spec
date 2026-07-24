"""``.docx`` export of the SectionFormat tree via python-docx.

Office-style SectionFormat layout: centered section header, PART headings,
``1.1  TITLE`` articles, hanging-indent paragraph levels (A. / 1. / a. /
1)) backed by genuine Word multilevel numbering in clean exports, END OF
SECTION — followed on a new page by the **assumptions
schedule**: every ``assumed`` block listed with its numbering so a senior
reviewer can audit each model default in one pass, plus the open-item
schedule ([TBD: ...] markers and ``needs_input`` blocks).
"""
from __future__ import annotations

import io
import itertools
import re
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .. import settings
from .diffing import ElementDiff, SectionDiff
from .model import (
    SpecSection,
    iter_paragraphs,
    open_questions,
)
from .word_numbering import SectionFormatNumbering

_LEVEL_INDENT = Inches(0.45)


def _style_base(document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    for docx_section in document.sections:
        docx_section.top_margin = Inches(1)
        docx_section.bottom_margin = Inches(1)
        docx_section.left_margin = Inches(1)
        docx_section.right_margin = Inches(1)


def _centered(document, text: str, *, bold: bool = True):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    return p


def _schedule_table(document, rows: list[tuple[str, str]], headers: tuple[str, str]):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, headers):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    for ref, text in rows:
        ref_cell, text_cell = table.add_row().cells
        ref_cell.text = ref
        text_cell.text = text
    table.columns[0].width = Inches(1.1)
    table.columns[1].width = Inches(5.4)
    return table


def build_docx(
    section: SpecSection,
    audit_result: dict | None = None,
    qc_result: dict | None = None,
    redline: SectionDiff | None = None,
    redline_date: str | None = None,
) -> bytes:
    """Render the section; a QC or audit closing carries the review trail.

    ``qc_result`` is the Batch 4 Final-QC dict (:meth:`QCResult.to_dict`);
    ``audit_result`` is the Phase 5 audit dict. When a QC result is present
    it supersedes the audit closing (the QC lenses cover the audit's ground
    and more); otherwise the audit closing is rendered as before. The
    rendering states which document version was reviewed.

    When ``redline`` (a :class:`SectionDiff`) is supplied (Batch 5), the body
    is rendered as genuine Word tracked changes (``w:ins``/``w:del`` with
    ``w:delText``, deleted/inserted paragraph marks) instead of the plain
    tree. **Accept All** reproduces ``section`` (the cur tree) exactly,
    numbering included; **Reject All** reproduces the diff's base tree's
    provision *text*. Unlike the clean renderer's genuine Word numbering,
    redline display numbering (A. / 1. / a.) remains a plain positional
    literal so a survivor whose position shifted keeps the current label,
    never a tracked mark (the frozen "moves are not marked" decision). The
    schedules below are always rendered plainly from ``section`` (the current
    document), never redlined. ``redline_date`` overrides the ISO-8601
    ``w:date`` stamp.
    """
    document = Document()
    _style_base(document)

    if redline is not None:
        _render_redline_body(
            document,
            section,
            redline,
            settings.APP_NAME,
            redline_date or _redline_now(),
        )
    else:
        _render_clean_body(document, section)

    # -- assumptions schedule ----------------------------------------------
    document.add_page_break()
    _centered(document, "ASSUMPTIONS SCHEDULE")
    document.add_paragraph(
        "The following provisions were drafted from defaults (the applicable "
        "codes, standards, and disciplinary norms in effect) without explicit "
        "confirmation. Each requires review by the design professional of "
        "record."
    )
    assumed = [
        (ref, p.text)
        for _part, _article, p, _depth, ref in iter_paragraphs(section)
        if p.status == "assumed"
    ]
    if assumed:
        _schedule_table(document, assumed, ("Ref", "Assumed provision"))
    else:
        document.add_paragraph("None — every provision is confirmed.")

    # -- imported provisions not yet reviewed (Phase 5 master import) ------
    imported = [
        (ref, p.text)
        for _part, _article, p, _depth, ref in iter_paragraphs(section)
        if p.status == "imported"
    ]
    if imported:
        document.add_paragraph()
        _centered(document, "IMPORTED PROVISIONS NOT YET REVIEWED")
        document.add_paragraph(
            "The following provisions were imported from a master "
            "specification and have not been confirmed or adapted for this "
            "project. Each requires review before issue."
        )
        _schedule_table(document, imported, ("Ref", "Imported provision"))

    items = open_questions(section)
    if items:
        document.add_paragraph()
        _centered(document, "OPEN ITEMS")
        rows = [
            (
                item["ref"],
                ("[TBD] " if item["kind"] == "tbd" else "[NEEDS INPUT] ")
                + item["label"],
            )
            for item in items
        ]
        _schedule_table(document, rows, ("Ref", "Open item"))

    # -- Final QC closing (Batch 4) supersedes the audit closing -----------
    if qc_result and (qc_result.get("findings") or qc_result.get("lens_statuses")):
        _render_qc_closing(document, qc_result, compact=True)
    # -- compliance audit closing section (Phase 5) ------------------------
    elif audit_result and audit_result.get("coverage"):
        document.add_page_break()
        _centered(document, "COMPLIANCE AUDIT SUMMARY")
        audited_at = audit_result.get("audited_at", "")
        version_index = audit_result.get("version_index")
        version_note = (
            f" of document version v{int(version_index) + 1}"
            if isinstance(version_index, int)
            else ""
        )
        document.add_paragraph(
            f"Advisory audit{version_note} against the researched project "
            f"requirements profile ({audited_at}). Grounded requirements "
            "only; this summary is not a substitute for review by a "
            "licensed design professional."
        )
        summary = str(audit_result.get("summary") or "").strip()
        if summary:
            document.add_paragraph(summary)
        coverage_rows = [
            (
                str(entry.get("status", "")).upper(),
                f"[{entry.get('requirement_id', '')}] "
                + str(entry.get("note") or entry.get("evidence_quote") or ""),
            )
            for entry in audit_result.get("coverage", [])
        ]
        if coverage_rows:
            _schedule_table(document, coverage_rows, ("Status", "Requirement"))
        findings = audit_result.get("findings") or []
        if findings:
            document.add_paragraph()
            _centered(document, "AUDIT FINDINGS")
            finding_rows = [
                (
                    str(finding.get("severity", "")).upper(),
                    str(finding.get("issue", ""))
                    + (
                        f" Suggestion: {finding['suggestion']}"
                        if finding.get("suggestion")
                        else ""
                    ),
                )
                for finding in findings
            ]
            _schedule_table(document, finding_rows, ("Severity", "Finding"))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_clean_body(document, section: SpecSection) -> None:
    """SectionFormat body with genuine Word-numbered provision levels."""
    numbering = SectionFormatNumbering(document)
    _centered(document, f"SECTION {section.number or '[TBD]'}")
    _centered(document, section.title or "[TBD: SECTION TITLE]")
    document.add_paragraph()

    for part in section.parts:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.add_run(part.title).bold = True
        if not part.articles:
            document.add_paragraph("(Not used.)")
        for a_idx, article in enumerate(part.articles):
            # A separate w:num per article is the restart boundary for A.;
            # even an empty article receives its deterministic instance.
            article_num_id = numbering.new_article()
            ap = document.add_paragraph()
            apf = ap.paragraph_format
            apf.space_before = Pt(10)
            apf.tab_stops.add_tab_stop(_LEVEL_INDENT, WD_TAB_ALIGNMENT.LEFT)
            ap.add_run(
                f"{part.number}.{a_idx + 1}\t{article.title.upper()}"
            ).bold = True

            def walk(paragraphs, depth: int) -> None:
                for para in paragraphs:
                    provision = document.add_paragraph(para.text)
                    provision.paragraph_format.space_after = Pt(6)
                    numbering.apply(
                        provision,
                        num_id=article_num_id,
                        level=depth,
                    )
                    walk(para.children, depth + 1)

            walk(article.paragraphs, 0)

    document.add_paragraph()
    _centered(document, f"END OF SECTION {section.number or ''}".rstrip())


def export_filename(section: SpecSection) -> str:
    stem = f"SECTION {section.number}" if section.number else "DRAFT SECTION"
    if section.title:
        stem += f" - {section.title}"
    stem = re.sub(r'[\\/:*?"<>|]+', "", stem).strip() or "DRAFT SECTION"
    return f"{stem}.docx"


def redline_filename(section: SpecSection) -> str:
    """The clean export name with a `` - REDLINE`` suffix before ``.docx``."""
    name = export_filename(section)
    return name[: -len(".docx")] + " - REDLINE.docx"


# ---------------------------------------------------------------------------
# Tracked-changes (redline) body — Batch 5
#
# python-docx has no tracked-changes API, so the w:ins / w:del / w:delText
# elements are built directly with docx.oxml, mirroring the shapes the
# importer's tests already manufacture (tests/test_importer.py). The killer
# invariant: re-importing this export (Accept-All resolution) reproduces the
# cur tree exactly; a Reject-All reading reproduces the base tree's provision
# TEXT. Display numbering (A. / 1.1 / a.) is positional and recomputes to the
# rendered view — it is a literal label, not tracked content, so a survivor
# whose position shifted shows the current number under both resolutions
# (the frozen "moves are not marked" decision). Reject-All is therefore
# text-faithful, not label-faithful, for position-shifted survivors.
# ---------------------------------------------------------------------------


def _redline_now() -> str:
    """ISO-8601 UTC stamp for w:date (e.g. ``2026-07-21T14:30:00Z``)."""
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _set_revision_attrs(element, ids, author: str, date: str) -> None:
    element.set(qn("w:id"), str(next(ids)))
    element.set(qn("w:author"), author)
    element.set(qn("w:date"), date)


def _content_run(text: str, *, del_text: bool = False, bold: bool = False):
    """A ``w:r`` carrying ``text``; tabs become ``w:tab`` for real indents.

    Uses ``w:delText`` inside deletions (required) and ``w:t`` otherwise;
    ``xml:space=preserve`` keeps token whitespace byte-exact.
    """
    run = OxmlElement("w:r")
    if bold:
        rpr = OxmlElement("w:rPr")
        rpr.append(OxmlElement("w:b"))
        run.append(rpr)
    tag = "w:delText" if del_text else "w:t"
    for i, segment in enumerate(text.split("\t")):
        if i > 0:
            run.append(OxmlElement("w:tab"))
        if segment:
            text_el = OxmlElement(tag)
            text_el.set(qn("xml:space"), "preserve")
            text_el.text = segment
            run.append(text_el)
    return run


def _append_equal(paragraph, text: str, *, bold: bool = False) -> None:
    paragraph._p.append(_content_run(text, bold=bold))


def _append_ins(paragraph, text: str, ids, author, date, *, bold=False) -> None:
    ins = OxmlElement("w:ins")
    _set_revision_attrs(ins, ids, author, date)
    ins.append(_content_run(text, bold=bold))
    paragraph._p.append(ins)


def _append_del(paragraph, text: str, ids, author, date, *, bold=False) -> None:
    dele = OxmlElement("w:del")
    _set_revision_attrs(dele, ids, author, date)
    dele.append(_content_run(text, del_text=True, bold=bold))
    paragraph._p.append(dele)


def _append_runs(paragraph, runs, ids, author, date, *, bold=False) -> None:
    for run in runs:
        if run.op == "equal":
            _append_equal(paragraph, run.text, bold=bold)
        elif run.op == "ins":
            _append_ins(paragraph, run.text, ids, author, date, bold=bold)
        else:  # del
            _append_del(paragraph, run.text, ids, author, date, bold=bold)


def _mark_paragraph(paragraph, tag: str, ids, author, date) -> None:
    """Flag the paragraph MARK as inserted/deleted (``w:pPr/w:rPr/<tag>``).

    A deleted paragraph mark is what makes Word collapse the whole paragraph
    on Accept; an inserted one makes Reject remove it. ``w:rPr`` is the last
    child of ``w:pPr`` per the schema, so appending it is valid.
    """
    ppr = paragraph._p.get_or_add_pPr()
    rpr = ppr.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        ppr.append(rpr)
    marker = OxmlElement(tag)
    _set_revision_attrs(marker, ids, author, date)
    rpr.append(marker)


def _redline_paragraph_format(paragraph, level: int) -> None:
    pf = paragraph.paragraph_format
    pf.left_indent = _LEVEL_INDENT * (level + 1)
    pf.first_line_indent = -_LEVEL_INDENT
    pf.tab_stops.add_tab_stop(_LEVEL_INDENT * (level + 1), WD_TAB_ALIGNMENT.LEFT)
    pf.space_after = Pt(6)


def _render_redline_section(document, element: ElementDiff, ids, author, date):
    # SECTION <number> line (centered). The clean body substitutes "[TBD]" for
    # an empty number, so an empty side must carry that placeholder too, or the
    # round-trip diverges on a from-scratch (vs-empty) redline.
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_equal(p, "SECTION ", bold=True)
    if element.number_base == element.number_cur:
        _append_equal(p, element.number_cur or "[TBD]", bold=True)
    else:
        _append_del(p, element.number_base or "[TBD]", ids, author, date, bold=True)
        _append_ins(p, element.number_cur or "[TBD]", ids, author, date, bold=True)

    # Section title line (centered): word-level diff when both sides have a
    # title, whole del/ins with the placeholder when a side is empty.
    q = document.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder = "[TBD: SECTION TITLE]"
    if element.base_text == element.cur_text:
        _append_equal(q, element.cur_text or placeholder, bold=True)
    elif not element.base_text or not element.cur_text:
        _append_del(q, element.base_text or placeholder, ids, author, date, bold=True)
        _append_ins(q, element.cur_text or placeholder, ids, author, date, bold=True)
    else:
        _append_runs(q, element.runs or [], ids, author, date, bold=True)
    document.add_paragraph()


def _render_redline_article(document, element: ElementDiff, ids, author, date):
    p = document.add_paragraph()
    apf = p.paragraph_format
    apf.space_before = Pt(10)
    apf.tab_stops.add_tab_stop(_LEVEL_INDENT, WD_TAB_ALIGNMENT.LEFT)
    number = element.ref_cur or element.ref_base
    if element.kind == "inserted":
        _append_ins(
            p, f"{number}\t{element.cur_text.upper()}", ids, author, date, bold=True
        )
        _mark_paragraph(p, "w:ins", ids, author, date)
    elif element.kind == "deleted":
        _append_del(
            p, f"{number}\t{element.base_text.upper()}", ids, author, date, bold=True
        )
        _mark_paragraph(p, "w:del", ids, author, date)
    elif element.kind == "changed":
        _append_equal(p, f"{number}\t", bold=True)
        upper = [
            type(run)(run.op, run.text.upper()) for run in (element.runs or [])
        ]
        _append_runs(p, upper, ids, author, date, bold=True)
    else:  # unchanged
        _append_equal(p, f"{number}\t{element.cur_text.upper()}", bold=True)


def _render_redline_paragraph(document, element: ElementDiff, ids, author, date):
    p = document.add_paragraph()
    _redline_paragraph_format(p, element.depth)
    label = element.label
    if element.kind == "inserted":
        _append_ins(p, f"{label}\t{element.cur_text}", ids, author, date)
        _mark_paragraph(p, "w:ins", ids, author, date)
    elif element.kind == "deleted":
        _append_del(p, f"{label}\t{element.base_text}", ids, author, date)
        _mark_paragraph(p, "w:del", ids, author, date)
    elif element.kind == "changed":
        _append_equal(p, f"{label}\t")
        _append_runs(p, element.runs or [], ids, author, date)
    else:  # unchanged
        _append_equal(p, f"{label}\t{element.cur_text}")


def _render_redline_body(
    document, section: SpecSection, redline: SectionDiff, author: str, date: str
) -> None:
    ids = itertools.count(1)
    # Per part, does it hold any article in cur / in base? The clean body
    # prints "(Not used.)" for an empty part, so a part that empties (or fills)
    # between the two versions must track that placeholder — else Accept-All
    # (or Reject-All) would not reproduce the clean export. An article row's
    # kind says which side it exists on: not-deleted => in cur; not-inserted
    # => in base.
    part_cur: dict[str, bool] = {}
    part_base: dict[str, bool] = {}
    current_part: str | None = None
    for element in redline.elements:
        if element.node_type == "part":
            current_part = element.uid
            part_cur.setdefault(current_part, False)
            part_base.setdefault(current_part, False)
        elif element.node_type == "article" and current_part is not None:
            if element.kind != "deleted":
                part_cur[current_part] = True
            if element.kind != "inserted":
                part_base[current_part] = True

    for element in redline.elements:
        if element.node_type == "section":
            _render_redline_section(document, element, ids, author, date)
        elif element.node_type == "part":
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.add_run(element.cur_text).bold = True
            has_cur = part_cur.get(element.uid, False)
            has_base = part_base.get(element.uid, False)
            if not has_cur and not has_base:
                document.add_paragraph("(Not used.)")  # empty both ways
            elif not has_cur:
                # Emptied in cur (articles all deleted): "(Not used.)" appears
                # on Accept, is gone on Reject (where the articles return).
                np = document.add_paragraph()
                _append_ins(np, "(Not used.)", ids, author, date)
                _mark_paragraph(np, "w:ins", ids, author, date)
            elif not has_base:
                # Filled in cur (empty in base): "(Not used.)" appears on
                # Reject, is gone on Accept (where the articles are present).
                np = document.add_paragraph()
                _append_del(np, "(Not used.)", ids, author, date)
                _mark_paragraph(np, "w:del", ids, author, date)
        elif element.node_type == "article":
            _render_redline_article(document, element, ids, author, date)
        else:  # paragraph
            _render_redline_paragraph(document, element, ids, author, date)

    document.add_paragraph()
    _centered(document, f"END OF SECTION {section.number or ''}".rstrip())


# ---------------------------------------------------------------------------
# Final QC memo (Batch 4)
# ---------------------------------------------------------------------------

import json
from urllib.parse import urlsplit

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import RGBColor

_SEVERITY_ORDER = ("critical", "high", "medium", "low")

_QC_BLUE = "2E74B5"
_QC_DARK_BLUE = "1F4D78"
_QC_INK = "0B2545"
_QC_MUTED = "59636E"
_QC_LIGHT_GRAY = "F2F4F7"
_QC_CALLOUT = "F4F6F9"
_QC_BORDER = "C9D3DF"
_QC_POSITIVE = "1F3A5F"
_QC_CAUTION = "7A5A00"
_QC_RISK = "9B1C1C"
_QC_LINK = "0563C1"
_QC_CONTENT_WIDTH_DXA = 9360
_QC_TABLE_INDENT_DXA = 120
_QC_EXPECTED_LENS_IDS = (
    "code_compliance",
    "coordination_consistency",
    "completeness",
    "enforceability_language",
    "provenance_hygiene",
)


def _qc_set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    """Apply renderer-stable Latin font settings to a run."""
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _qc_set_style_font(style, *, size: float, color: str, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), "Calibri")


def _qc_style(document, name: str, base: str = "Normal"):
    styles = document.styles
    if name in styles:
        return styles[name]
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles[base]
    return style


def _qc_configure_styles(document) -> None:
    """Resolve the standard_business_brief preset into concrete styles."""
    normal = document.styles["Normal"]
    _qc_set_style_font(normal, size=11, color="000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, _QC_BLUE, 16, 8),
        "Heading 2": (13, _QC_BLUE, 12, 6),
        "Heading 3": (12, _QC_DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        _qc_set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    title = document.styles["Title"]
    _qc_set_style_font(title, size=23, color="000000", bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.0
    title.paragraph_format.keep_with_next = True

    subtitle = document.styles["Subtitle"]
    _qc_set_style_font(subtitle, size=14, color="373737")
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.paragraph_format.line_spacing = 1.0
    subtitle.paragraph_format.keep_with_next = True

    metadata = _qc_style(document, "QC Metadata")
    _qc_set_style_font(metadata, size=11, color="000000")
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(2)
    metadata.paragraph_format.line_spacing = 1.0

    lead = _qc_style(document, "QC Lead")
    _qc_set_style_font(lead, size=11, color=_QC_INK)
    lead.paragraph_format.space_before = Pt(4)
    lead.paragraph_format.space_after = Pt(8)
    lead.paragraph_format.line_spacing = 1.10

    small = _qc_style(document, "QC Small")
    _qc_set_style_font(small, size=9, color=_QC_MUTED)
    small.paragraph_format.space_before = Pt(0)
    small.paragraph_format.space_after = Pt(3)
    small.paragraph_format.line_spacing = 1.0

    code = _qc_style(document, "QC Operation")
    _qc_set_style_font(code, size=10, color="202020")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.05)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.05

    citation = _qc_style(document, "QC Table Citation")
    _qc_set_style_font(citation, size=9, color=_QC_MUTED)
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)
    citation.paragraph_format.line_spacing = 1.0


def _qc_set_paragraph_border(
    paragraph,
    *,
    edge: str,
    color: str,
    size: int = 8,
    space: int = 1,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    old = p_bdr.find(qn(f"w:{edge}"))
    if old is not None:
        p_bdr.remove(old)
    border = OxmlElement(f"w:{edge}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def _qc_shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:shd"))
    if old is not None:
        p_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _qc_add_field(paragraph, instruction: str, display: str = "1"):
    run = paragraph.add_run()
    _qc_set_run_font(run, size=9, color=_QC_MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    return run


def _qc_setup_page(document, section: SpecSection) -> None:
    """Letter, one-inch margins, quiet running furniture, and page fields."""
    for docx_section in document.sections:
        docx_section.page_width = Inches(8.5)
        docx_section.page_height = Inches(11)
        docx_section.top_margin = Inches(1)
        docx_section.right_margin = Inches(1)
        docx_section.bottom_margin = Inches(1)
        docx_section.left_margin = Inches(1)
        docx_section.header_distance = Inches(0.492)
        docx_section.footer_distance = Inches(0.492)
        docx_section.different_first_page_header_footer = False

        header = docx_section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = ""
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)
        hp.paragraph_format.line_spacing = 1.0
        hp.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.5), WD_TAB_ALIGNMENT.RIGHT
        )
        left = hp.add_run("FINAL QC AUDIT TRAIL")
        _qc_set_run_font(left, size=8.5, color=_QC_MUTED, bold=True)
        right = hp.add_run(f"\tSECTION {section.number or '[TBD]'}")
        _qc_set_run_font(right, size=8.5, color=_QC_MUTED)

        footer = docx_section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.paragraph_format.space_before = Pt(3)
        fp.paragraph_format.space_after = Pt(0)
        fp.paragraph_format.line_spacing = 1.0
        fp.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.5), WD_TAB_ALIGNMENT.RIGHT
        )
        _qc_set_paragraph_border(
            fp, edge="top", color=_QC_BORDER, size=4, space=3
        )
        fr = fp.add_run(f"Build-a-Spec | Section {section.number or '[TBD]'}")
        _qc_set_run_font(fr, size=9, color=_QC_MUTED)
        page_label = fp.add_run("\tPage ")
        _qc_set_run_font(page_label, size=9, color=_QC_MUTED)
        _qc_add_field(fp, "PAGE")
        of_label = fp.add_run(" of ")
        _qc_set_run_font(of_label, size=9, color=_QC_MUTED)
        _qc_add_field(fp, "NUMPAGES")


def _qc_is_safe_http_url(value: object) -> bool:
    if not isinstance(value, str):
        return False
    url = value.strip()
    if not url or any(ord(char) < 32 for char in url) or any(
        char.isspace() for char in url
    ):
        return False
    try:
        parsed = urlsplit(url)
        _ = parsed.port
    except (ValueError, TypeError):
        return False
    return (
        parsed.scheme.lower() in {"http", "https"}
        and bool(parsed.hostname)
        and parsed.username is None
        and parsed.password is None
    )


def _qc_add_hyperlink(paragraph, url: str, label: str | None = None):
    """Add a clickable relationship only for a validated HTTP(S) target."""
    text = label or url
    if not _qc_is_safe_http_url(url):
        run = paragraph.add_run(text)
        _qc_set_run_font(run, color="000000")
        return run
    try:
        relationship_id = paragraph.part.relate_to(
            url.strip(), RT.HYPERLINK, is_external=True
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)
        run_el = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{attr}"), "Calibri")
        r_pr.append(r_fonts)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), _QC_LINK)
        r_pr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)
        run_el.append(r_pr)
        text_el = OxmlElement("w:t")
        if text[:1].isspace() or text[-1:].isspace():
            text_el.set(qn("xml:space"), "preserve")
        text_el.text = text
        run_el.append(text_el)
        hyperlink.append(run_el)
        paragraph._p.append(hyperlink)
        return hyperlink
    except (KeyError, TypeError, ValueError):
        run = paragraph.add_run(text)
        _qc_set_run_font(run, color="000000")
        return run


def _qc_heading(document, text: str, level: int = 1):
    p = document.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def _qc_text(value: object, fallback: str = "Not recorded") -> str:
    if value is None:
        return fallback
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, float):
        return f"{value:,.6f}".rstrip("0").rstrip(".")
    if isinstance(value, int):
        return f"{value:,}"
    text = str(value).strip()
    return text or fallback


def _qc_json(value: object) -> str:
    try:
        return json.dumps(
            value,
            ensure_ascii=True,
            sort_keys=True,
            separators=(", ", ": "),
            default=str,
        )
    except (TypeError, ValueError, OverflowError):
        return str(value)


def _qc_list(value: object) -> list:
    return value if isinstance(value, list) else []


def _qc_dict(value: object) -> dict:
    return value if isinstance(value, dict) else {}


def _qc_legacy_schema(document) -> bool:
    try:
        return int(getattr(document, "_qc_schema_version", 1) or 1) < 2
    except (TypeError, ValueError):
        return True


def _qc_add_label(
    document,
    label: str,
    value: object,
    *,
    style: str | None = None,
    color: str | None = None,
):
    p = document.add_paragraph(style=style)
    label_run = p.add_run(f"{label}: ")
    _qc_set_run_font(label_run, color=_QC_DARK_BLUE, bold=True)
    if isinstance(value, str) and _qc_is_safe_http_url(value):
        _qc_add_hyperlink(p, value, value)
    else:
        value_run = p.add_run(_qc_text(value))
        _qc_set_run_font(value_run, color=color or "000000")
    return p


def _qc_add_callout(
    document,
    label: str,
    text: str,
    *,
    accent: str = _QC_INK,
    fill: str = _QC_CALLOUT,
):
    p = document.add_paragraph(style="QC Lead")
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.04)
    _qc_shade_paragraph(p, fill)
    _qc_set_paragraph_border(p, edge="left", color=accent, size=18, space=6)
    label_run = p.add_run(f"{label}: ")
    _qc_set_run_font(label_run, color=accent, bold=True)
    body_run = p.add_run(text)
    _qc_set_run_font(body_run, color=_QC_INK)
    return p


def _qc_numbering_part(document):
    return document.part.numbering_part.element


def _qc_next_numbering_id(numbering, tag: str, attr: str) -> int:
    values: list[int] = []
    for child in numbering.findall(qn(f"w:{tag}")):
        raw = child.get(qn(f"w:{attr}"))
        try:
            values.append(int(raw))
        except (TypeError, ValueError):
            continue
    return max(values, default=0) + 1


def _qc_numbering_abstracts(document) -> dict[str, int]:
    cached = getattr(document, "_qc_list_abstracts", None)
    if isinstance(cached, dict):
        return cached
    numbering = _qc_numbering_part(document)
    abstract_ids: dict[str, int] = {}
    for kind in ("bullet", "decimal"):
        abstract_id = _qc_next_numbering_id(
            numbering, "abstractNum", "abstractNumId"
        )
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "\u2022" if kind == "bullet" else "%1.")
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Calibri")
        r_fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(r_fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        abstract_ids[kind] = abstract_id
    setattr(document, "_qc_list_abstracts", abstract_ids)
    return abstract_ids


def _qc_new_list_numbering(document, kind: str) -> int:
    numbering = _qc_numbering_part(document)
    abstract_id = _qc_numbering_abstracts(document)[kind]
    num_id = _qc_next_numbering_id(numbering, "num", "numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _qc_apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:numPr"))
    if old is not None:
        p_pr.remove(old)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    p_pr.insert(0, num_pr)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def _qc_add_bullets(document, items: list[str]) -> None:
    if not items:
        return
    num_id = _qc_new_list_numbering(document, "bullet")
    for item in items:
        p = document.add_paragraph()
        _qc_apply_numbering(p, num_id)
        p.add_run(str(item))


def _qc_add_numbered_steps(document, items: list[tuple[str, str]]) -> None:
    num_id = _qc_new_list_numbering(document, "decimal")
    for label, detail in items:
        p = document.add_paragraph()
        _qc_apply_numbering(p, num_id)
        label_run = p.add_run(f"{label}. ")
        _qc_set_run_font(label_run, color=_QC_DARK_BLUE, bold=True)
        p.add_run(detail)


def _qc_add_source_list(
    document,
    label: str,
    sources: object,
    *,
    empty: str = "No source record was persisted.",
) -> None:
    values = _qc_list(sources)
    if not values:
        _qc_add_label(document, label, empty)
        return
    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run(f"{label}:")
    _qc_set_run_font(run, color=_QC_DARK_BLUE, bold=True)
    num_id = _qc_new_list_numbering(document, "bullet")
    for value in values:
        p = document.add_paragraph()
        _qc_apply_numbering(p, num_id)
        if isinstance(value, str):
            source = value.strip()
            _qc_add_hyperlink(p, source, source)
            if source and not _qc_is_safe_http_url(source):
                note = p.add_run(" (not linked: target failed HTTP(S) safety checks)")
                _qc_set_run_font(note, color=_QC_MUTED, italic=True)
        elif isinstance(value, dict):
            url = str(
                value.get("url")
                or value.get("source_url")
                or value.get("uri")
                or ""
            ).strip()
            title = str(value.get("title") or "").strip()
            if url:
                if title:
                    title_run = p.add_run(f"{title} - ")
                    _qc_set_run_font(title_run, bold=True)
                _qc_add_hyperlink(p, url, url)
                remainder = {
                    key: item
                    for key, item in value.items()
                    if key not in {"url", "source_url", "uri", "title"}
                }
                if remainder:
                    detail = p.add_run(f" | {_qc_json(remainder)}")
                    _qc_set_run_font(detail, color=_QC_MUTED)
            else:
                p.add_run(_qc_json(value))
        else:
            p.add_run(_qc_text(value))


def _qc_add_text_list(
    document,
    label: str,
    values: object,
    *,
    empty: str,
) -> None:
    records = _qc_list(values)
    if not records:
        _qc_add_label(document, label, empty)
        return
    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run(f"{label}:")
    _qc_set_run_font(run, color=_QC_DARK_BLUE, bold=True)
    num_id = _qc_new_list_numbering(document, "bullet")
    for value in records:
        p = document.add_paragraph()
        _qc_apply_numbering(p, num_id)
        p.add_run(
            _qc_json(value) if isinstance(value, (dict, list)) else _qc_text(value)
        )


def _qc_table_geometry(
    table,
    widths_dxa: list[int],
    *,
    header: bool = True,
    indent_dxa: int = _QC_TABLE_INDENT_DXA,
) -> None:
    """Apply fixed DXA geometry and standard-business-brief table tokens."""
    if sum(widths_dxa) != _QC_CONTENT_WIDTH_DXA:
        raise ValueError("QC table column widths must total 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for tag in ("tblW", "tblInd", "tblLayout", "tblCellMar", "tblBorders"):
        old = tbl_pr.find(qn(f"w:{tag}"))
        if old is not None:
            tbl_pr.remove(old)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(_QC_CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    margins = OxmlElement("w:tblCellMar")
    for edge, amount in (
        ("top", 80),
        ("start", 120),
        ("bottom", 80),
        ("end", 120),
    ):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(amount))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), _QC_BORDER)
        borders.append(border)
    tbl_pr.append(borders)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        if header and row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for col_index, cell in enumerate(row.cells):
            width = widths_dxa[col_index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header and row_index == 0:
                shd = tc_pr.find(qn("w:shd"))
                if shd is None:
                    shd = OxmlElement("w:shd")
                    tc_pr.append(shd)
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), _QC_LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    _qc_set_run_font(
                        run,
                        size=10,
                        color=(
                            _QC_INK
                            if header and row_index == 0
                            else "000000"
                        ),
                        bold=True if header and row_index == 0 else None,
                    )


def _qc_add_table(
    document,
    headers: list[str],
    rows: list[list[object]],
    widths_dxa: list[int],
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = _qc_text(value, "")
    _qc_table_geometry(table, widths_dxa)
    return table


def _qc_version_note(qc_result: dict) -> str:
    version_index = qc_result.get("version_index")
    if isinstance(version_index, int):
        return f" of document version v{version_index + 1}"
    return ""


def _render_qc_closing(document, qc_result: dict, *, compact: bool) -> None:
    """The Final-QC section appended to the issued spec (compact form)."""
    document.add_page_break()
    _centered(document, "FINAL QC SUMMARY")
    model = str(qc_result.get("model") or "the QC model")
    finished = str(qc_result.get("finished_at") or "")
    document.add_paragraph(
        f"Final quality-control review{_qc_version_note(qc_result)} by "
        f"{model} ({finished}). Every finding below survived an adversarial "
        "verification pass. This summary is advisory and is not a substitute "
        "for review by a licensed design professional."
    )
    summary = str(qc_result.get("summary") or "").strip()
    if summary:
        document.add_paragraph(summary)

    findings = qc_result.get("findings") or []
    open_findings = [f for f in findings if f.get("status") == "open"]
    if open_findings:
        rows = [
            (
                str(f.get("severity", "")).upper(),
                (f"[{f.get('element_id')}] " if f.get("element_id") else "")
                + str(f.get("title", "")),
            )
            for f in _sorted_by_severity(open_findings)
        ]
        _schedule_table(document, rows, ("Severity", "Open finding"))
    else:
        document.add_paragraph("No surviving finding remains open.")

    if not compact:
        return
    applied = sum(1 for f in findings if f.get("status") == "applied")
    dismissed = sum(1 for f in findings if f.get("status") == "dismissed")
    refuted = len(qc_result.get("refuted") or [])
    inconclusive = len(qc_result.get("inconclusive") or [])
    document.add_paragraph(
        "Surviving-finding dispositions: "
        f"{len(open_findings)} open, {applied} applied, {dismissed} dismissed. "
        "Other candidate outcomes: "
        f"{refuted} substantively refuted, {inconclusive} "
        "infrastructure-inconclusive; those candidates are not shown in this "
        "compact closing but remain in the full audit report."
    )


def _sorted_by_severity(findings: list[dict]) -> list[dict]:
    rank = {s: i for i, s in enumerate(_SEVERITY_ORDER)}
    return sorted(findings, key=lambda f: rank.get(str(f.get("severity")), 99))


# The standalone report builder follows below. The compact closing above is
# intentionally unchanged for exports that append a short QC summary to the spec.


def _qc_add_masthead(
    document,
    qc_result: dict,
    section: SpecSection,
    *,
    stale: bool,
    execution_label: str,
) -> None:
    title = document.add_paragraph(style="Title")
    title.add_run("FINAL QC AUDIT REPORT")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run(
        f"Audit trail for Section {section.number or '[TBD]'} | "
        f"{section.title or '[TBD]'}"
    )

    profile = section.project_profile or {}
    where = ", ".join(
        str(value)
        for value in (
            profile.get("city"),
            profile.get("state_or_province"),
            profile.get("country"),
        )
        if value
    )
    project_name = str(
        profile.get("project_name")
        or profile.get("name")
        or profile.get("client_name")
        or "Not recorded"
    )
    rows = [
        ("To", "Project design and specification review team"),
        ("From", "Build-a-Spec Final QC"),
        ("Date", qc_result.get("finished_at") or "Not recorded"),
        (
            "Re",
            f"Section {section.number or '[TBD]'} | "
            f"{section.title or '[TBD]'}",
        ),
        ("Project", project_name),
    ]
    if where:
        rows.append(("Location", where))
    if profile.get("client_name"):
        rows.append(("Client", profile.get("client_name")))
    rows.extend(
        [
            ("Run ID", qc_result.get("run_id") or "Legacy result: not recorded"),
            ("Status", execution_label),
            ("Input state", "STALE" if stale else "Current at export"),
        ]
    )
    for label, value in rows:
        _qc_add_label(document, label, value, style="QC Metadata")

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(10)
    _qc_set_paragraph_border(rule, edge="bottom", color=_QC_INK, size=12, space=3)


def _qc_schema_version(qc_result: dict) -> int:
    try:
        return int(qc_result.get("schema_version", 1) or 1)
    except (TypeError, ValueError):
        return 1


def _qc_expected_lens_specs(qc_result: dict) -> list[dict[str, str]]:
    configuration = _qc_dict(
        _qc_dict(qc_result.get("input_manifest")).get("configuration")
    )
    configured = [
        item
        for item in _qc_list(configuration.get("lenses"))
        if isinstance(item, dict) and str(item.get("lens_id") or "").strip()
    ]
    if configured:
        return [
            {
                "lens_id": str(item.get("lens_id")).strip(),
                "title": str(item.get("title") or "").strip(),
                "brief": str(item.get("brief") or "").strip(),
            }
            for item in configured
        ]
    return [
        {
            "lens_id": lens_id,
            "title": lens_id.replace("_", " ").title(),
            "brief": "",
        }
        for lens_id in _QC_EXPECTED_LENS_IDS
    ]


def _qc_expected_panel_size(qc_result: dict, finding: dict) -> int:
    recorded = finding.get("verification_panel_size")
    try:
        if int(recorded or 0) > 0:
            return int(recorded)
    except (TypeError, ValueError):
        pass
    configuration = _qc_dict(
        _qc_dict(qc_result.get("input_manifest")).get("configuration")
    )
    severity = str(
        finding.get("original_severity") or finding.get("severity") or ""
    ).lower()
    key = (
        "verifiers_critical"
        if severity in {"critical", "high"}
        else "verifiers_standard"
    )
    try:
        configured = int(configuration.get(key) or 0)
    except (TypeError, ValueError):
        configured = 0
    if configured > 0:
        return configured
    return 3 if severity in {"critical", "high"} else 2


def _qc_execution_issues(qc_result: dict) -> list[str]:
    """Return observable coverage defects; never infer missing seats as passes."""
    issues: list[str] = []
    schema_version = _qc_schema_version(qc_result)
    if schema_version >= 2 and not str(
        qc_result.get("execution_status") or ""
    ).strip():
        issues.append("The schema v2 execution-status field is missing.")
    statuses = [
        item
        for item in _qc_list(qc_result.get("lens_statuses"))
        if isinstance(item, dict)
    ]
    if schema_version >= 2:
        expected_specs = _qc_expected_lens_specs(qc_result)
        expected_ids = [item["lens_id"] for item in expected_specs]
        recorded_ids = [str(item.get("lens_id") or "") for item in statuses]
        if len(set(expected_ids)) != len(expected_ids):
            issues.append("The input manifest contains duplicate configured lens IDs.")
        duplicate_ids = sorted(
            {lens_id for lens_id in recorded_ids if recorded_ids.count(lens_id) > 1}
        )
        if duplicate_ids:
            issues.append(
                "Duplicate persisted lens record(s): " + ", ".join(duplicate_ids) + "."
            )
        missing = [lens_id for lens_id in expected_ids if lens_id not in recorded_ids]
        unexpected = [
            lens_id for lens_id in recorded_ids if lens_id not in set(expected_ids)
        ]
        if missing:
            issues.append("Missing required lens record(s): " + ", ".join(missing) + ".")
        if unexpected:
            issues.append(
                "Unexpected lens record(s): " + ", ".join(unexpected) + "."
            )
        for lens in statuses:
            lens_id = str(lens.get("lens_id") or "unnamed lens")
            status = str(lens.get("status") or "missing").lower()
            if status != "completed":
                issues.append(f"Lens {lens_id} status is {status}, not completed.")
            elif not _qc_list(lens.get("reviewed_checks")):
                issues.append(
                    f"Lens {lens_id} has no required reviewed-check audit record."
                )
    elif not statuses:
        issues.append("No legacy lens execution records were persisted.")
    else:
        for lens in statuses:
            if str(lens.get("status") or "").lower() != "completed":
                issues.append(
                    f"Legacy lens {lens.get('lens_id') or 'unnamed'} did not complete."
                )

    raw_refuted = [
        item
        for item in _qc_list(qc_result.get("refuted"))
        if isinstance(item, dict)
    ]
    legacy_inconclusive = [
        item
        for item in raw_refuted
        if str(item.get("verification_outcome") or "").lower()
        in {"default_refuted", "inconclusive"}
    ]
    normalized_collections = (
        ("surviving", _qc_list(qc_result.get("findings"))),
        (
            "refuted",
            [item for item in raw_refuted if item not in legacy_inconclusive],
        ),
        (
            "infrastructure-inconclusive",
            [*_qc_list(qc_result.get("inconclusive")), *legacy_inconclusive],
        ),
    )
    for prefix, collection in normalized_collections:
        for finding in collection:
            if not isinstance(finding, dict):
                continue
            finding_id = str(finding.get("finding_id") or "unnamed finding")
            expected = _qc_expected_panel_size(qc_result, finding)
            verdicts = [
                item
                for item in _qc_list(finding.get("verdicts"))
                if isinstance(item, dict)
            ]
            if len(verdicts) != expected:
                issues.append(
                    f"{prefix.title()} finding {finding_id} has {len(verdicts)} of "
                    f"{expected} required verifier record(s)."
                )
            indexes: list[int] = []
            for verdict in verdicts:
                try:
                    indexes.append(int(verdict.get("reviewer_index", 0) or 0))
                except (TypeError, ValueError):
                    indexes.append(0)
            expected_indexes = set(range(1, expected + 1))
            if schema_version >= 2 and set(indexes) != expected_indexes:
                issues.append(
                    f"{prefix.title()} finding {finding_id} verifier seat indexes are "
                    f"{indexes}; expected 1 through {expected}."
                )
            elif schema_version < 2 and indexes and not (
                set(indexes) == expected_indexes or all(index == 0 for index in indexes)
            ):
                issues.append(
                    f"Legacy {prefix} finding {finding_id} has inconsistent verifier indexes."
                )
            incomplete = [
                str(verdict.get("status") or "completed")
                for verdict in verdicts
                if str(verdict.get("status") or "completed").lower() != "completed"
            ]
            if incomplete:
                issues.append(
                    f"{prefix.title()} finding {finding_id} has incomplete verifier "
                    f"status(es): {', '.join(incomplete)}."
                )
    return issues


def _qc_execution_label(qc_result: dict, stale: bool) -> str:
    if stale:
        return "STALE - RERUN REQUIRED"
    if _qc_export_control_issues(qc_result):
        return "BLOCKED - NOT ISSUE-READY"
    recorded = str(qc_result.get("execution_status") or "").strip().lower()
    if recorded and recorded not in {"complete", "completed", "success"}:
        return f"{recorded.upper()} - COVERAGE INCOMPLETE"
    if _qc_execution_issues(qc_result):
        return "PARTIAL - EXECUTION COVERAGE INCOMPLETE"
    return "COMPLETE"


def _qc_export_control_issues(qc_result: dict) -> list[str]:
    """Return export-time QC controls that forbid a complete sign-off.

    The report body can be a retained success while a newer attempt is still
    running or has failed.  Likewise, readiness is computed from the live
    session and can reject a body whose historical execution record is
    complete.  Treat these generated export facts as controlling so neither
    the masthead nor the signature page can promote an older success.
    """
    state = _qc_dict(qc_result.get("export_current_state"))
    if not state:
        return []

    issues: list[str] = []
    blocking_statuses = {"failed", "cancelled", "partial", "running"}
    attempt = _qc_dict(state.get("latest_attempt"))
    attempt_status = str(attempt.get("status") or "").strip().lower()
    if attempt_status in blocking_statuses:
        issues.append(
            "Latest attempt "
            f"{attempt.get('run_id') or 'with unrecorded run ID'} is "
            f"{attempt_status}."
        )

    runner = _qc_dict(state.get("runner"))
    runner_status = str(runner.get("status") or "").strip().lower()
    if runner_status in blocking_statuses and runner_status != attempt_status:
        issues.append(f"QC runner state is {runner_status}.")

    readiness = _qc_dict(state.get("readiness"))
    checks = [
        item
        for item in _qc_list(readiness.get("checks"))
        if isinstance(item, dict)
    ]
    for check_id in ("qc_current", "qc_audit_complete"):
        matching = next(
            (item for item in checks if str(item.get("id") or "") == check_id),
            None,
        )
        if matching is not None and matching.get("ok") is not True:
            detail = str(matching.get("detail") or "No detail recorded").strip()
            issues.append(f"Readiness check {check_id} is blocked: {detail}")
    return issues


def _qc_severity_counts(findings: list[dict]) -> list[int]:
    counts = [
        sum(
            1
            for item in findings
            if str(item.get("severity", "")).lower() == level
        )
        for level in _SEVERITY_ORDER
    ]
    return [*counts, len(findings)]


def _qc_render_executive_status(
    document,
    qc_result: dict,
    *,
    stale: bool,
    execution_label: str,
) -> None:
    _qc_heading(document, "Executive Status", 1)
    findings = [
        item
        for item in _qc_list(qc_result.get("findings"))
        if isinstance(item, dict)
    ]
    raw_refuted = [
        item
        for item in _qc_list(qc_result.get("refuted"))
        if isinstance(item, dict)
    ]
    legacy_inconclusive = [
        item
        for item in raw_refuted
        if str(item.get("verification_outcome") or "").lower()
        in {"default_refuted", "inconclusive"}
    ]
    refuted = [item for item in raw_refuted if item not in legacy_inconclusive]
    inconclusive = [
        item
        for item in _qc_list(qc_result.get("inconclusive"))
        if isinstance(item, dict)
    ]
    inconclusive.extend(legacy_inconclusive)
    open_findings = [
        item
        for item in findings
        if str(item.get("status") or "open") == "open"
    ]
    applied = [
        item for item in findings if str(item.get("status") or "") == "applied"
    ]
    dismissed = [
        item
        for item in findings
        if str(item.get("status") or "") == "dismissed"
    ]
    other = [
        item
        for item in findings
        if str(item.get("status") or "open")
        not in {"open", "applied", "dismissed"}
    ]
    statuses = [
        item
        for item in _qc_list(qc_result.get("lens_statuses"))
        if isinstance(item, dict)
    ]
    completed_lenses = sum(
        1 for item in statuses if item.get("status") == "completed"
    )
    failed_lenses = len(statuses) - completed_lenses

    if stale:
        accent = _QC_RISK
        message = (
            "One or more material QC inputs no longer match this run. The document, "
            "research profile, standards, module or discipline, model configuration, "
            "or source policy may have changed. Findings and dispositions belong to "
            "the recorded input only. Re-run Final QC before relying on this report."
        )
    elif _qc_export_control_issues(qc_result):
        accent = _QC_RISK
        message = (
            "The export-time latest attempt or required QC readiness gate blocks "
            "reliance on this report for issue. A retained successful record is "
            "historical evidence only and cannot override the controlling state."
        )
    elif execution_label != "COMPLETE":
        accent = _QC_CAUTION
        message = (
            "The saved result contains incomplete execution coverage. Read every "
            "failed or missing lens and verifier record below; absence of a finding "
            "from an incomplete reviewer is not a pass."
        )
    elif open_findings:
        accent = (
            _QC_RISK
            if any(
                str(item.get("severity")) in {"critical", "high"}
                for item in open_findings
            )
            else _QC_CAUTION
        )
        message = (
            f"{len(open_findings)} verified finding(s) remain open. The report "
            "records the evidence, verification votes, proposed operations, and "
            "current disposition for reviewer action."
        )
    else:
        accent = _QC_POSITIVE
        message = (
            "No verified finding is currently marked open. This means the saved "
            "QC findings are applied or dismissed; it is not a certification that "
            "the specification is error-free or professionally approved."
        )
    _qc_add_callout(document, execution_label, message, accent=accent)
    execution_issues = _qc_execution_issues(qc_result)
    if execution_issues:
        _qc_heading(document, "Execution Coverage Exceptions", 2)
        _qc_add_bullets(document, execution_issues)
    export_control_issues = _qc_export_control_issues(qc_result)
    if export_control_issues:
        _qc_heading(document, "Export-Time QC Control Exceptions", 2)
        _qc_add_bullets(document, export_control_issues)

    summary = str(qc_result.get("summary") or "").strip()
    _qc_add_label(
        document,
        "Run summary",
        summary or "No overall narrative summary was persisted.",
    )
    _qc_add_label(
        document,
        "Lens execution",
        f"{completed_lenses} completed; {failed_lenses} failed or incomplete; "
        f"{len(statuses)} recorded total",
    )

    rows = [
        ["Open survivors", *_qc_severity_counts(open_findings)],
        ["Applied survivors", *_qc_severity_counts(applied)],
        ["Dismissed survivors", *_qc_severity_counts(dismissed)],
    ]
    if other:
        rows.append(["Other status survivors", *_qc_severity_counts(other)])
    rows.extend(
        [
            ["All surviving findings", *_qc_severity_counts(findings)],
            ["Refuted candidates", *_qc_severity_counts(refuted)],
            [
                "Infrastructure-inconclusive candidates",
                *_qc_severity_counts(inconclusive),
            ],
            [
                "All recorded candidates",
                *_qc_severity_counts([*findings, *refuted, *inconclusive]),
            ],
        ]
    )
    _qc_add_table(
        document,
        ["Disposition", "Critical", "High", "Medium", "Low", "Total"],
        rows,
        [2360, 1400, 1400, 1400, 1400, 1400],
    )
    citation = document.add_paragraph(style="QC Table Citation")
    citation.add_run(
        "Counts are derived from the persisted report at export. "
        "Surviving, refuted, and infrastructure-inconclusive candidates are "
        "mutually exclusive."
    )


def _qc_render_manifest(
    document, manifest: object, *, title: str = "Input Manifest"
) -> None:
    _qc_heading(document, title, 2)
    if not isinstance(manifest, dict) or not manifest:
        document.add_paragraph(
            "No structured input manifest was persisted. This is expected for "
            "legacy schema results; document-only identity may still be available."
        )
        return
    for key, value in manifest.items():
        label = str(key).replace("_", " ").strip().title()
        if isinstance(value, dict):
            _qc_add_label(document, label, _qc_json(value))
        elif isinstance(value, list):
            if not value:
                _qc_add_label(document, label, "[]")
                continue
            heading = document.add_paragraph()
            head_run = heading.add_run(f"{label}:")
            _qc_set_run_font(head_run, color=_QC_DARK_BLUE, bold=True)
            num_id = _qc_new_list_numbering(document, "bullet")
            for entry in value:
                p = document.add_paragraph()
                _qc_apply_numbering(p, num_id)
                if isinstance(entry, str) and _qc_is_safe_http_url(entry):
                    _qc_add_hyperlink(p, entry, entry)
                else:
                    p.add_run(
                        _qc_json(entry)
                        if isinstance(entry, (dict, list))
                        else _qc_text(entry)
                    )
        else:
            _qc_add_label(document, label, value)


def _qc_render_identity(
    document,
    qc_result: dict,
    section: SpecSection,
    *,
    stale: bool,
    execution_label: str,
) -> None:
    _qc_heading(document, "Run and Input Identity", 1)
    legacy = _qc_legacy_schema(document)
    version = qc_result.get("version_index")
    version_text = (
        f"v{version + 1}" if isinstance(version, int) else "Not recorded"
    )
    duration = qc_result.get("duration_ms")
    duration_text = (
        f"{int(duration):,} ms ({float(duration) / 1000:,.3f} seconds)"
        if isinstance(duration, (int, float)) and not isinstance(duration, bool)
        else "Not recorded"
    )
    identities = [
        (
            "Report schema",
            qc_result.get("schema_version") or "Legacy schema: not recorded",
        ),
        (
            "QC protocol",
            qc_result.get("protocol_version") or "Legacy protocol: not recorded",
        ),
        ("Run ID", qc_result.get("run_id") or "Not recorded"),
        ("Execution status", execution_label),
        ("Started", qc_result.get("started_at") or "Not recorded"),
        ("Finished", qc_result.get("finished_at") or "Not recorded"),
        ("Recorded duration", "Not recorded" if legacy else duration_text),
        ("Reviewed document version", version_text),
        (
            "Document version fingerprint",
            qc_result.get("version_fingerprint") or "Not recorded",
        ),
        (
            "Complete input fingerprint",
            qc_result.get("input_fingerprint") or "Not recorded",
        ),
        ("Export-time input state", "STALE" if stale else "Current at export"),
        ("Section number", section.number or "[TBD]"),
        ("Section title", section.title or "[TBD]"),
        ("Model", qc_result.get("model") or "Not recorded"),
        (
            "Reasoning effort",
            "Not recorded" if legacy else qc_result.get("effort") or "Not recorded",
        ),
        (
            "Maximum output tokens",
            (
                "Not recorded"
                if legacy
                else qc_result.get("max_tokens") or "Not recorded"
            ),
        ),
        (
            "Research profile present",
            (
                "Yes"
                if qc_result.get("research_profile_present")
                else "No or not recorded"
            ),
        ),
    ]
    for label, value in identities:
        _qc_add_label(document, label, value)
    _qc_render_manifest(document, qc_result.get("input_manifest"))


def _qc_render_export_current_state(document, qc_result: dict) -> None:
    state = _qc_dict(qc_result.get("export_current_state"))
    if not state:
        return
    _qc_heading(document, "Export-Time State and Report History", 1)
    runner = _qc_dict(state.get("runner"))
    attempt = _qc_dict(state.get("latest_attempt"))
    report_run_id = str(qc_result.get("run_id") or "")
    attempt_run_id = str(attempt.get("run_id") or "")
    attempt_status = str(attempt.get("status") or "Not recorded")
    if attempt_run_id and attempt_run_id != report_run_id:
        _qc_add_callout(
            document,
            "LATEST ATTEMPT DIFFERS FROM THIS REPORT",
            (
                f"Latest attempt {attempt_run_id} is {attempt_status}. The main "
                f"report body identifies run {report_run_id or 'not recorded'}. "
                "Treat the latest-attempt status as controlling for readiness."
            ),
            accent=_QC_CAUTION,
        )
    elif attempt_status.lower() in {"failed", "cancelled", "partial", "running"}:
        _qc_add_callout(
            document,
            f"LATEST ATTEMPT: {attempt_status.upper()}",
            str(attempt.get("error") or "See the preserved execution records."),
            accent=_QC_CAUTION,
        )
    values = [
        ("Export generated", state.get("generated_at")),
        ("Selected report run ID", report_run_id or "Not recorded"),
        (
            "Selected report execution status",
            qc_result.get("execution_status") or "Not recorded",
        ),
        ("Active document version", state.get("document_version")),
        ("Active document fingerprint", state.get("document_fingerprint")),
        ("Active full-input fingerprint", state.get("current_input_fingerprint")),
        ("Report matches all active inputs", state.get("report_matches_current_inputs")),
        ("Report stale at export", state.get("stale")),
        ("Runner status", runner.get("status")),
        ("Runner error", runner.get("error") or "None recorded"),
        ("Latest attempt run ID", attempt_run_id or "Not recorded"),
        ("Latest attempt status", attempt_status),
        ("Latest attempt started", attempt.get("started_at")),
        ("Latest attempt finished", attempt.get("finished_at")),
        ("Latest attempt error", attempt.get("error") or "None recorded"),
        ("Latest attempt report available", attempt.get("report_available")),
    ]
    for label, value in values:
        _qc_add_label(document, label, value)

    retained = _qc_dict(state.get("last_successful_report"))
    if retained:
        retained_run_id = str(retained.get("run_id") or "Not recorded")
        _qc_heading(document, "Retained Prior Successful Report Identity", 2)
        _qc_add_callout(
            document,
            "HISTORICAL SUCCESS - DOES NOT CONTROL CURRENT READINESS",
            (
                f"Run {retained_run_id} is retained for traceability. The selected "
                f"report is run {report_run_id or 'not recorded'}, and the latest "
                "attempt plus export-time readiness checks control whether the "
                "specification may proceed."
            ),
            accent=_QC_CAUTION,
        )
        retained_values = [
            ("Retained successful run ID", retained_run_id),
            (
                "Retained successful execution status",
                retained.get("execution_status") or "Not recorded",
            ),
            ("Retained successful start", retained.get("started_at")),
            ("Retained successful finish", retained.get("finished_at")),
            ("Retained reviewed document version", retained.get("version_index")),
            (
                "Retained reviewed document fingerprint",
                retained.get("version_fingerprint"),
            ),
            (
                "Retained full-input fingerprint",
                retained.get("input_fingerprint"),
            ),
            ("Retained report summary", retained.get("summary")),
        ]
        for label, value in retained_values:
            _qc_add_label(document, label, value)

    readiness = _qc_dict(state.get("readiness"))
    _qc_add_label(document, "Issue readiness at export", readiness.get("ready"))
    checks = [
        item
        for item in _qc_list(readiness.get("checks"))
        if isinstance(item, dict)
    ]
    if checks:
        _qc_heading(document, "Readiness Checks at Export", 2)
        _qc_add_bullets(
            document,
            [
                f"{'PASS' if check.get('ok') else 'BLOCK'} - "
                f"{check.get('id') or 'unnamed'}: "
                f"{check.get('detail') or 'No detail recorded'}"
                for check in checks
            ],
        )
    _qc_render_manifest(
        document,
        state.get("current_input_manifest"),
        title="Current Input Manifest at Export",
    )


def _qc_render_methodology(document, qc_result: dict) -> None:
    _qc_heading(document, "Methodology and Interpretation", 1)
    document.add_paragraph(
        "This report presents the observable execution record persisted by "
        "Final QC. It does not expose hidden model reasoning. It records which "
        "reviewers ran, the checks and searches they reported, which evidence "
        "was retrieved or accepted, how candidate findings were challenged, "
        "and what disposition was saved."
    )
    _qc_add_numbered_steps(
        document,
        [
            (
                "Input fixation",
                "The reviewed document version and, for current schemas, the full "
                "material-input manifest are content-addressed by fingerprints.",
            ),
            (
                "Independent lens review",
                "Each configured lens records its own completion state, summary, "
                "reported checks, searches, retrieved sources, and candidate count.",
            ),
            (
                "Source grounding",
                "A source marked accepted passed the recorded grounding check. A "
                "source merely cited or retrieved is retained with that narrower label.",
            ),
            (
                "Adversarial verification",
                "Candidate findings are challenged by a severity-based reviewer "
                "panel. The report preserves completed, failed, and cancelled "
                "reviewer records when the active schema provides them.",
            ),
            (
                "Operation validation and disposition",
                "Proposed edits are recorded in full with their validation state. "
                "Open, applied, dismissed, substantively refuted, and "
                "infrastructure-inconclusive outcomes remain distinct.",
            ),
        ],
    )
    threshold_note = qc_result.get("protocol_version") or "the legacy protocol"
    _qc_add_callout(
        document,
        "Reading rule",
        f"A completed lens that reports zero findings is different from a failed "
        f"or unrecorded lens. Verification thresholds are shown per finding when "
        f"persisted by {threshold_note}; missing records are labeled, never inferred "
        "as successful work.",
        accent=_QC_BLUE,
    )


def _qc_render_usage_values(document, usage: object, *, label: str) -> None:
    values = _qc_dict(usage)
    if not values:
        _qc_add_label(document, label, "No usage counters were persisted.")
        return
    _qc_add_label(
        document,
        label,
        "; ".join(
            f"{str(key).replace('_', ' ')}={_qc_text(value)}"
            for key, value in values.items()
        ),
    )


def _qc_render_record(
    document,
    record: dict,
    *,
    ignored_keys: set[str] | None = None,
) -> None:
    ignored = ignored_keys or set()
    for key, value in record.items():
        if key in ignored:
            continue
        label = str(key).replace("_", " ").strip().title()
        if isinstance(value, list) and any(
            token in key.lower() for token in ("source", "url")
        ):
            _qc_add_source_list(document, label, value)
        elif isinstance(value, (dict, list)):
            _qc_add_label(document, label, _qc_json(value))
        else:
            _qc_add_label(document, label, value)


def _qc_render_lens(document, lens: dict, index: int, findings: list[dict]) -> None:
    lens_id = str(lens.get("lens_id") or f"lens-{index}")
    title = str(lens.get("title") or lens_id.replace("_", " ").title())
    status = str(lens.get("status") or "not_recorded").lower()
    _qc_heading(document, f"Lens {index}: {title}", 2)
    color = _QC_POSITIVE if status == "completed" else _QC_RISK
    _qc_add_label(document, "Lens ID", lens_id)
    _qc_add_label(document, "Execution status", status.upper(), color=color)
    if lens.get("error"):
        _qc_add_callout(
            document,
            "Recorded failure",
            str(lens.get("error")),
            accent=_QC_RISK,
        )

    lens_findings = [
        item for item in findings if str(item.get("lens_id")) == lens_id
    ]
    reported_count = lens.get("finding_count")
    grounded_count = lens.get("grounded_count")
    legacy = _qc_legacy_schema(document)
    checks = [
        entry
        for entry in _qc_list(lens.get("reviewed_checks"))
        if isinstance(entry, dict)
    ]
    contract_complete = status == "completed" and (legacy or bool(checks))
    _qc_add_label(
        document,
        "Candidate findings reported by lens",
        reported_count if reported_count is not None else "Not recorded",
    )
    _qc_add_label(
        document,
        "Grounded candidate count",
        grounded_count if grounded_count is not None else "Not recorded",
    )
    _qc_add_label(
        document,
        "Candidate records retained in this report",
        len(lens_findings),
    )
    if contract_complete and int(reported_count or 0) == 0:
        _qc_add_callout(
            document,
            "Zero-finding distinction",
            "This lens completed and reported zero candidate findings. That is a "
            "recorded review result, not the same as a lens that failed or did not run.",
            accent=_QC_POSITIVE,
        )
    elif not contract_complete:
        _qc_add_callout(
            document,
            "Coverage limitation",
            "This lens is failed, missing, or lacks the required reviewed-check "
            "ledger. No conclusion should be drawn from a zero or missing finding "
            "count for this lens.",
            accent=_QC_RISK,
        )

    _qc_add_label(document, "Review brief", lens.get("brief") or "Not recorded")
    _qc_add_label(document, "Lens summary", lens.get("summary") or "Not recorded")
    _qc_add_text_list(
        document,
        "Search queries",
        lens.get("search_queries"),
        empty="No search-query record was persisted.",
    )
    _qc_add_source_list(
        document,
        "Retrieved sources",
        lens.get("retrieved_sources"),
        empty="No retrieved-source record was persisted.",
    )
    _qc_add_text_list(
        document,
        "All billed-attempt search queries",
        lens.get("attempted_search_queries"),
        empty="No separate billed-attempt query record was persisted.",
    )
    _qc_add_source_list(
        document,
        "All billed-attempt sources",
        lens.get("attempted_sources"),
        empty="No separate billed-attempt source record was persisted.",
    )
    _qc_render_usage_values(document, lens.get("usage_totals"), label="Lens usage")
    _qc_add_label(
        document,
        "Estimated lens cost (USD)",
        lens.get("estimated_cost_usd", "Not recorded"),
    )
    _qc_add_label(
        document,
        "API requests",
        (
            "Not recorded"
            if legacy
            else lens.get("api_request_count")
            if "api_request_count" in lens
            else "Not recorded"
        ),
    )
    _qc_add_label(
        document,
        "Model responses",
        (
            "Not recorded"
            if legacy
            else lens.get("model_response_count")
            if "model_response_count" in lens
            else "Not recorded"
        ),
    )

    _qc_heading(document, "Reviewed Checks", 3)
    if not checks:
        if legacy:
            document.add_paragraph(
                "No per-check audit records were persisted. Legacy results can "
                "retain the lens outcome without its observable-work ledger."
            )
        else:
            _qc_add_callout(
                document,
                "Missing required audit record",
                "Schema v2 requires at least one reviewed check for every completed "
                "lens. This lens is incomplete even if its status says completed.",
                accent=_QC_RISK,
            )
    for check_index, check in enumerate(checks, start=1):
        outcome = str(check.get("outcome") or "not_recorded").upper()
        check_name = str(check.get("check") or "Unnamed check")
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"Check {check_index} | {outcome} | {check_name}")
        _qc_set_run_font(run, color=_QC_DARK_BLUE, bold=True)
        _qc_add_label(document, "Notes", check.get("notes") or "No note recorded")
        _qc_add_label(
            document,
            "Element IDs",
            ", ".join(
                str(value) for value in _qc_list(check.get("element_ids"))
            )
            or "None recorded",
        )
        _qc_add_source_list(
            document,
            "Check source URLs",
            check.get("source_urls"),
            empty="No source URL recorded for this check.",
        )
        source_checks = [
            entry
            for entry in _qc_list(check.get("source_checks"))
            if isinstance(entry, dict)
        ]
        if source_checks:
            _qc_render_source_checks(document, source_checks)


def _qc_render_lenses(document, qc_result: dict) -> None:
    document.add_page_break()
    _qc_heading(document, "Lens-by-Lens Audit Trail", 1)
    document.add_paragraph(
        "Every persisted lens record appears below, including failures and completed "
        "lenses that reported zero findings. Counts describe saved outputs; they do "
        "not silently fill gaps in failed execution."
    )
    recorded = [
        item
        for item in _qc_list(qc_result.get("lens_statuses"))
        if isinstance(item, dict)
    ]
    all_findings = [
        item
        for item in [
            *_qc_list(qc_result.get("findings")),
            *_qc_list(qc_result.get("refuted")),
            *_qc_list(qc_result.get("inconclusive")),
        ]
        if isinstance(item, dict)
    ]
    if not recorded:
        _qc_add_callout(
            document,
            "Missing lens ledger",
            "No per-lens execution records were persisted. Findings may still be "
            "reviewed, but coverage completeness cannot be established from this report.",
            accent=_QC_RISK,
        )
    pending = list(recorded)
    lenses: list[dict] = []
    for spec in _qc_expected_lens_specs(qc_result):
        match_index = next(
            (
                position
                for position, lens in enumerate(pending)
                if str(lens.get("lens_id") or "") == spec["lens_id"]
            ),
            None,
        )
        if match_index is None:
            lenses.append(
                {
                    **spec,
                    "status": "missing",
                    "error": "No execution record was persisted for this required lens.",
                }
            )
        else:
            lens = pending.pop(match_index)
            if not lens.get("brief") and spec.get("brief"):
                lens = {**lens, "brief": spec["brief"]}
            lenses.append(lens)
    lenses.extend(pending)
    for index, lens in enumerate(lenses, start=1):
        _qc_render_lens(document, lens, index, all_findings)


def _qc_vote_counts(verdicts: list[dict]) -> tuple[int, int, int, int]:
    uphold = refute = failed = 0
    for verdict in verdicts:
        status = str(verdict.get("status") or "completed").lower()
        if status != "completed":
            failed += 1
        elif verdict.get("upholds") is True:
            uphold += 1
        else:
            refute += 1
    return uphold, refute, failed, len(verdicts)


def _qc_render_verdict(document, verdict: dict, index: int) -> None:
    reviewer_index = verdict.get("reviewer_index")
    label = reviewer_index if reviewer_index not in (None, "", 0) else index
    status = str(verdict.get("status") or "completed").upper()
    _qc_heading(document, f"Verifier {label}: {status}", 3)
    if "upholds" in verdict:
        _qc_add_label(
            document,
            "Vote",
            "UPHOLD" if verdict.get("upholds") is True else "REFUTE",
        )
    else:
        _qc_add_label(document, "Vote", "Not recorded")
    _qc_add_label(
        document,
        "Revised severity",
        verdict.get("revised_severity") or "No revision recorded",
    )
    _qc_add_label(document, "Verdict note", verdict.get("note") or "Not recorded")
    if verdict.get("error"):
        _qc_add_callout(
            document,
            "Verifier error",
            str(verdict.get("error")),
            accent=_QC_RISK,
        )
    _qc_add_text_list(
        document,
        "Verifier search queries",
        verdict.get("search_queries"),
        empty="No verifier search-query record was persisted.",
    )
    _qc_add_source_list(
        document,
        "Verifier retrieved sources",
        verdict.get("retrieved_sources"),
        empty="No verifier retrieved-source record was persisted.",
    )
    _qc_add_text_list(
        document,
        "All billed-attempt verifier queries",
        verdict.get("attempted_search_queries"),
        empty="No separate billed-attempt verifier-query record was persisted.",
    )
    _qc_add_source_list(
        document,
        "All billed-attempt verifier sources",
        verdict.get("attempted_sources"),
        empty="No separate billed-attempt verifier-source record was persisted.",
    )
    _qc_render_usage_values(
        document, verdict.get("usage_totals"), label="Verifier usage"
    )
    _qc_add_label(
        document,
        "Estimated verifier cost (USD)",
        verdict.get("estimated_cost_usd", "Not recorded"),
    )
    legacy = _qc_legacy_schema(document)
    request_value = (
        "Not recorded"
        if legacy
        else
        verdict.get("api_request_count")
        if "api_request_count" in verdict
        else verdict.get("request_count", "Not recorded")
    )
    response_value = (
        "Not recorded"
        if legacy
        else
        verdict.get("model_response_count")
        if "model_response_count" in verdict
        else verdict.get("response_count", "Not recorded")
    )
    _qc_add_label(document, "API requests", request_value)
    _qc_add_label(document, "Model responses", response_value)


def _qc_render_source_checks(document, checks: object) -> None:
    records = [
        item for item in _qc_list(checks) if isinstance(item, dict)
    ]
    if not records:
        _qc_add_label(document, "Source-check records", "None persisted")
        return
    for index, record in enumerate(records, start=1):
        url = str(
            record.get("url") or record.get("source_url") or ""
        ).strip()
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"Source check {index}")
        _qc_set_run_font(run, color=_QC_DARK_BLUE, bold=True)
        if url:
            p.add_run(" | ")
            _qc_add_hyperlink(p, url, url)
        _qc_render_record(
            document,
            record,
            ignored_keys={"url", "source_url"},
        )


def _qc_render_ops(
    document, finding: dict, *, candidate_kind: str = "surviving"
) -> None:
    _qc_heading(document, "Proposed Operations and Validation", 3)
    operations = [
        item
        for item in _qc_list(finding.get("proposed_ops"))
        if isinstance(item, dict)
    ]
    if not operations:
        document.add_paragraph(
            "No mechanical edit operation was proposed. The finding is advisory or "
            "requires human-authored correction."
        )
        if candidate_kind == "inconclusive":
            validation = (
                "NOT EVALUATED - CANDIDATE INFRASTRUCTURE-INCONCLUSIVE"
            )
        elif candidate_kind == "refuted":
            validation = "NOT EVALUATED - CANDIDATE REFUTED"
        else:
            validation = "Not applicable - no operation proposed"
        _qc_add_label(document, "Operation validation", validation)
        return
    for index, operation in enumerate(operations, start=1):
        p = document.add_paragraph()
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"Operation {index} of {len(operations)}")
        _qc_set_run_font(run, color=_QC_DARK_BLUE, bold=True)
        op = document.add_paragraph(style="QC Operation")
        op.add_run(_qc_json(operation))
    valid = finding.get("ops_valid")
    if candidate_kind != "surviving":
        kind_label = (
            "INFRASTRUCTURE-INCONCLUSIVE"
            if candidate_kind == "inconclusive"
            else "REFUTED"
        )
        _qc_add_label(
            document,
            "Operation validation",
            f"NOT EVALUATED - CANDIDATE {kind_label}",
            color=_QC_MUTED,
        )
        _qc_add_label(
            document,
            "Validation detail",
            "The pipeline does not dry-run proposed operations for candidates "
            "that are substantively refuted or lack a complete verifier panel. "
            "A false ops_valid default is not an invalidity finding for these "
            "operations.",
        )
        return
    _qc_add_label(
        document,
        "Operation validation",
        "VALID" if valid is True else "INVALID OR NOT VALIDATED",
        color=_QC_POSITIVE if valid is True else _QC_RISK,
    )
    detail = finding.get("ops_invalid_reason")
    if not detail:
        detail = (
            "All proposed operations passed the saved validation."
            if valid
            else "No validation reason was persisted."
        )
    _qc_add_label(document, "Validation detail", detail)


def _qc_render_disposition(
    document, finding: dict, *, candidate_kind: str = "surviving"
) -> None:
    _qc_heading(document, "Disposition Record", 3)
    if candidate_kind != "surviving":
        outcome_label = (
            "INCONCLUSIVE - REQUIRED VERIFIER COVERAGE DID NOT COMPLETE"
            if candidate_kind == "inconclusive"
            else "REFUTED - NOT A SURVIVING FINDING"
        )
        _qc_add_label(document, "QC outcome", outcome_label)
        _qc_add_label(
            document,
            "Recorded finding status field",
            finding.get("status") or "Not recorded",
        )
    else:
        _qc_add_label(
            document,
            "Current disposition",
            str(finding.get("status") or "open").upper(),
        )
    _qc_add_label(
        document,
        "Dismissal reason",
        finding.get("dismiss_reason") or "None recorded",
    )
    events = [
        item
        for item in _qc_list(finding.get("disposition_events"))
        if isinstance(item, dict)
    ]
    if not events:
        document.add_paragraph(
            "No disposition-event history was persisted. The current status above is "
            "the only saved disposition state for this finding."
        )
        return
    for index, event in enumerate(events, start=1):
        _qc_add_label(document, f"Disposition event {index}", _qc_json(event))


def _render_memo_finding(
    document,
    finding: dict,
    *,
    ordinal: str = "",
    candidate_kind: str = "surviving",
) -> None:
    legacy = _qc_legacy_schema(document)
    severity = str(finding.get("severity") or "not recorded").upper()
    title = str(finding.get("title") or "Untitled finding")
    prefix = f"{ordinal} | " if ordinal else ""
    _qc_heading(document, f"{prefix}{severity} | {title}", 2)
    _qc_add_label(
        document, "Finding ID", finding.get("finding_id") or "Not recorded"
    )
    _qc_add_label(
        document,
        "Originating lens",
        finding.get("lens_id") or "Not recorded",
    )
    _qc_add_label(
        document,
        "Element reference",
        finding.get("element_id") or "Section-level finding",
    )
    _qc_add_label(
        document,
        "Reviewed reference",
        (
            "Not recorded"
            if legacy
            else finding.get("reviewed_ref") or "Not recorded"
        ),
    )
    if not legacy and "element_resolved" in finding:
        resolved = bool(finding.get("element_resolved"))
        _qc_add_label(
            document,
            "Reviewed anchor resolved",
            "Yes" if resolved else "No",
            color=_QC_POSITIVE if resolved else _QC_RISK,
        )
        if not resolved:
            _qc_add_callout(
                document,
                "Unresolved reviewed anchor",
                "The model-supplied element reference did not resolve against the "
                "reviewed input. Treat the finding as section-level unless a human "
                "reviewer identifies the intended provision from the issue text.",
                accent=_QC_RISK,
            )
    else:
        _qc_add_label(document, "Reviewed anchor resolved", "Not recorded")
    _qc_add_label(
        document,
        "Exact reviewed text",
        (
            "Not recorded"
            if legacy
            else finding.get("reviewed_text") or "Not recorded"
        ),
    )
    _qc_add_label(
        document,
        "Original severity",
        (
            f"Not separately recorded (current: {finding.get('severity') or 'unknown'})"
            if legacy
            else finding.get("original_severity")
            or finding.get("severity")
            or "Not recorded"
        ),
    )
    _qc_add_label(
        document, "Final severity", finding.get("severity") or "Not recorded"
    )
    _qc_add_label(document, "Issue", finding.get("issue") or "Not recorded")
    _qc_add_label(
        document, "Rationale", finding.get("rationale") or "Not recorded"
    )

    _qc_heading(document, "Evidence and Grounding", 3)
    _qc_add_label(
        document,
        "Grounded",
        "Yes" if finding.get("grounded") else "No",
    )
    _qc_add_source_list(
        document,
        "Accepted sources",
        finding.get("accepted_sources"),
        empty="No source was recorded as accepted for this finding.",
    )
    _qc_add_source_list(
        document,
        "Model-cited sources",
        finding.get("source_urls"),
        empty="No model-cited source URL was persisted.",
    )
    _qc_render_source_checks(document, finding.get("source_checks"))

    _qc_heading(document, "Verification Record", 3)
    persisted_outcome = str(finding.get("verification_outcome") or "").strip()
    if candidate_kind == "inconclusive":
        outcome = "infrastructure-inconclusive"
    elif candidate_kind == "refuted":
        outcome = "refuted"
    else:
        outcome = persisted_outcome or "upheld (derived from survivor collection)"
    threshold = finding.get("verification_threshold")
    _qc_add_label(document, "Verification outcome", outcome)
    if persisted_outcome and persisted_outcome.lower() != outcome.lower():
        _qc_add_label(
            document,
            "Persisted legacy outcome label",
            persisted_outcome,
        )
        document.add_paragraph(
            "The label above is retained for record fidelity. This report classifies "
            "the candidate from its collection and verifier-seat evidence; legacy "
            "default-refuted records with incomplete seats are infrastructure-"
            "inconclusive, not substantive refutations."
        )
    panel_size = finding.get("verification_panel_size")
    _qc_add_label(
        document,
        "Required panel size",
        panel_size if panel_size not in (None, "", 0) else "Not recorded",
    )
    _qc_add_label(
        document,
        "Required uphold threshold",
        threshold if threshold not in (None, "", 0) else "Not recorded",
    )
    verdicts = [
        item
        for item in _qc_list(finding.get("verdicts"))
        if isinstance(item, dict)
    ]
    uphold, refute_votes, failed, total = _qc_vote_counts(verdicts)
    _qc_add_label(
        document,
        "Persisted panel record",
        f"{uphold} uphold; {refute_votes} refute; {failed} failed or cancelled; "
        f"{total} reviewer record(s)",
    )
    if not verdicts:
        document.add_paragraph(
            "No individual verifier records were persisted. This is a legacy audit "
            "limitation; collection membership records the outcome but not each vote."
        )
    for index, verdict in enumerate(verdicts, start=1):
        _qc_render_verdict(document, verdict, index)

    _qc_render_ops(document, finding, candidate_kind=candidate_kind)
    _qc_render_disposition(document, finding, candidate_kind=candidate_kind)


def _qc_render_surviving_findings(document, qc_result: dict) -> None:
    document.add_page_break()
    _qc_heading(document, "Complete Surviving Findings Register", 1)
    findings = [
        item
        for item in _qc_list(qc_result.get("findings"))
        if isinstance(item, dict)
    ]
    if not findings:
        _qc_add_callout(
            document,
            "Zero surviving findings",
            "No candidate finding survived verification. This is distinct from lens "
            "or verifier failure; consult the execution records above before treating "
            "the result as complete.",
            accent=_QC_POSITIVE,
        )
        return
    document.add_paragraph(
        "Each surviving finding is reproduced with its identity, full issue and "
        "rationale, evidence status, source checks, verifier records, proposed "
        "operations, validation result, and complete saved disposition history."
    )
    counter = 0
    for severity in _SEVERITY_ORDER:
        band = [
            item
            for item in findings
            if str(item.get("severity") or "").lower() == severity
        ]
        for finding in band:
            counter += 1
            _render_memo_finding(
                document,
                finding,
                ordinal=f"SF-{counter:03d}",
                candidate_kind="surviving",
            )
    unranked = [
        item
        for item in findings
        if str(item.get("severity") or "").lower() not in _SEVERITY_ORDER
    ]
    for finding in unranked:
        counter += 1
        _render_memo_finding(
            document,
            finding,
            ordinal=f"SF-{counter:03d}",
            candidate_kind="surviving",
        )


def _qc_render_refuted_appendix(document, qc_result: dict) -> None:
    document.add_page_break()
    _qc_heading(document, "Appendix A: Complete Refuted Candidate Register", 1)
    refuted = [
        item
        for item in _qc_list(qc_result.get("refuted"))
        if isinstance(item, dict)
        and str(item.get("verification_outcome") or "").lower()
        not in {"default_refuted", "inconclusive"}
    ]
    if not refuted:
        document.add_paragraph(
            "No refuted candidate finding was persisted for this run."
        )
        return
    document.add_paragraph(
        "These candidates were raised during lens review but did not meet the saved "
        "verification outcome. They are retained in full for transparency and are "
        "not surviving QC findings. Proposed operations shown here must not be "
        "applied solely because they appear in this appendix."
    )
    for index, finding in enumerate(_sorted_by_severity(refuted), start=1):
        _render_memo_finding(
            document,
            finding,
            ordinal=f"RF-{index:03d}",
            candidate_kind="refuted",
        )


def _qc_render_inconclusive_appendix(document, qc_result: dict) -> None:
    document.add_page_break()
    _qc_heading(
        document,
        "Appendix A2: Infrastructure-Inconclusive Candidate Register",
        1,
    )
    candidates = [
        item
        for item in _qc_list(qc_result.get("inconclusive"))
        if isinstance(item, dict)
    ]
    # Legacy v2 records placed incomplete panels in ``refuted`` with a
    # default_refuted outcome. Surface those honestly even before a rerun.
    legacy_candidates = [
        item
        for item in _qc_list(qc_result.get("refuted"))
        if isinstance(item, dict)
        and str(item.get("verification_outcome") or "").lower()
        in {"default_refuted", "inconclusive"}
    ]
    candidates.extend(legacy_candidates)
    if not candidates:
        document.add_paragraph(
            "No infrastructure-inconclusive candidate was persisted for this run."
        )
        return
    document.add_paragraph(
        "These candidates did not receive every required completed verifier seat. "
        "They are neither surviving findings nor substantive refutations. Failed "
        "and cancelled seat records, evidence, and proposed operations are retained "
        "below so a reviewer can trace paid work and decide whether to rerun. Their "
        "operations were not validated and must not be applied from this record."
    )
    for index, finding in enumerate(_sorted_by_severity(candidates), start=1):
        _render_memo_finding(
            document,
            finding,
            ordinal=f"IC-{index:03d}",
            candidate_kind="inconclusive",
        )


def _qc_source_strings(value: object) -> list[str]:
    if isinstance(value, str):
        text = value.strip()
        return [text] if text else []
    if isinstance(value, list):
        return [item for entry in value for item in _qc_source_strings(entry)]
    if isinstance(value, dict):
        results: list[str] = []
        source_keys = {
            "url",
            "urls",
            "uri",
            "href",
            "link",
            "source",
            "source_url",
            "source_urls",
            "accepted_sources",
            "retrieved_sources",
        }
        for key, entry in value.items():
            if str(key).lower() in source_keys or "url" in str(key).lower():
                results.extend(_qc_source_strings(entry))
        return results
    return []


def _qc_evidence_register(qc_result: dict) -> list[dict[str, object]]:
    entries: dict[str, dict[str, object]] = {}

    def add(value: object, classification: str, reference: str) -> None:
        for source in _qc_source_strings(value):
            entry = entries.setdefault(
                source,
                {"source": source, "classes": [], "references": []},
            )
            classes = entry["classes"]
            references = entry["references"]
            if classification not in classes:
                classes.append(classification)
            if reference not in references:
                references.append(reference)

    for lens_index, lens in enumerate(
        _qc_list(qc_result.get("lens_statuses")), start=1
    ):
        if not isinstance(lens, dict):
            continue
        lens_label = str(lens.get("lens_id") or f"lens-{lens_index}")
        add(lens.get("retrieved_sources"), "retrieved", f"Lens {lens_label}")
        add(
            lens.get("attempted_sources"),
            "billed attempt",
            f"Lens {lens_label}",
        )
        for check_index, check in enumerate(
            _qc_list(lens.get("reviewed_checks")), start=1
        ):
            if not isinstance(check, dict):
                continue
            add(
                check.get("source_urls"),
                "check citation",
                f"{lens_label} check {check_index}",
            )
            add(
                check.get("source_checks"),
                "check grounding",
                f"{lens_label} check {check_index}",
            )

    for collection, prefix in (
        ("findings", "SF"),
        ("refuted", "RF"),
        ("inconclusive", "IC"),
    ):
        for index, finding in enumerate(
            _qc_list(qc_result.get(collection)), start=1
        ):
            if not isinstance(finding, dict):
                continue
            reference = (
                f"{prefix}-{index:03d} {finding.get('finding_id') or ''}".strip()
            )
            add(finding.get("accepted_sources"), "accepted", reference)
            add(finding.get("source_urls"), "model cited", reference)
            add(finding.get("source_checks"), "finding grounding", reference)
            for verdict_index, verdict in enumerate(
                _qc_list(finding.get("verdicts")), start=1
            ):
                if isinstance(verdict, dict):
                    add(
                        verdict.get("retrieved_sources"),
                        "verifier retrieved",
                        f"{reference} verifier {verdict_index}",
                    )
                    add(
                        verdict.get("attempted_sources"),
                        "verifier billed attempt",
                        f"{reference} verifier {verdict_index}",
                    )
    add(qc_result.get("input_manifest"), "input manifest", "Run input")
    return list(entries.values())


def _qc_render_evidence_register(document, qc_result: dict) -> None:
    document.add_page_break()
    _qc_heading(document, "Appendix B: Evidence Register", 1)
    document.add_paragraph(
        "This register deduplicates URLs from the input manifest, lens retrievals, "
        "reviewed checks, findings, source checks, and verifier retrievals. The class "
        "column preserves the narrowest recorded claim: retrieved or cited is not the "
        "same as accepted. Safe HTTP(S) targets are clickable."
    )
    entries = _qc_evidence_register(qc_result)
    if not entries:
        document.add_paragraph(
            "No external source URL was persisted. Findings may be based on internal "
            "consistency or discipline review rather than web evidence."
        )
        return

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, label in zip(
        table.rows[0].cells,
        ("ID", "Class", "Source", "Referenced by"),
    ):
        cell.text = label
    for index, entry in enumerate(entries, start=1):
        cells = table.add_row().cells
        cells[0].text = f"E-{index:03d}"
        cells[1].text = ", ".join(entry["classes"])
        source_p = cells[2].paragraphs[0]
        source_p.text = ""
        source = str(entry["source"])
        _qc_add_hyperlink(source_p, source, source)
        if not _qc_is_safe_http_url(source):
            note = source_p.add_run(" (not linked: failed HTTP(S) safety checks)")
            _qc_set_run_font(note, size=9, color=_QC_MUTED, italic=True)
        cells[3].text = "; ".join(entry["references"])
    _qc_table_geometry(table, [720, 1440, 4320, 2880])
    note = document.add_paragraph(style="QC Table Citation")
    note.add_run(
        "Acceptance records retrieval/grounding validation, not a legal or technical "
        "guarantee that the source proves every proposition in a finding."
    )


def _qc_flatten_cost(
    value: object, prefix: str = "Estimated cost"
) -> list[tuple[str, str]]:
    if isinstance(value, dict):
        rows: list[tuple[str, str]] = []
        for key, entry in value.items():
            rows.extend(
                _qc_flatten_cost(
                    entry,
                    f"{prefix} - {str(key).replace('_', ' ')}",
                )
            )
        return rows
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return [(prefix, f"${float(value):,.6f}")]
    return [(prefix, _qc_text(value))]


def _qc_render_usage_and_cost(document, qc_result: dict) -> None:
    _qc_heading(document, "Usage, Requests, and Estimated Cost", 1)
    legacy = _qc_legacy_schema(document)
    rows: list[list[object]] = []
    usage = _qc_dict(qc_result.get("usage_totals"))
    if usage:
        for key, value in usage.items():
            rows.append(
                [
                    str(key).replace("_", " ").title(),
                    _qc_text(value),
                    "Overall persisted run total",
                ]
            )
    else:
        rows.append(
            ["Usage totals", "Not recorded", "Legacy or incomplete result"]
        )
    rows.extend(
        [
            [
                "API request count",
                (
                    "Not recorded"
                    if legacy
                    else _qc_text(
                        qc_result.get("api_request_count"), "Not recorded"
                    )
                ),
                "Overall persisted run total",
            ],
            [
                "Model response count",
                (
                    "Not recorded"
                    if legacy
                    else _qc_text(
                        qc_result.get("model_response_count"), "Not recorded"
                    )
                ),
                "Overall persisted run total",
            ],
            [
                "Maximum output tokens",
                (
                    "Not recorded"
                    if legacy
                    else _qc_text(qc_result.get("max_tokens"), "Not recorded")
                ),
                "Configured ceiling",
            ],
            [
                "Recorded duration (ms)",
                (
                    "Not recorded"
                    if legacy
                    else _qc_text(qc_result.get("duration_ms"), "Not recorded")
                ),
                "Wall-clock execution record",
            ],
            [
                "Unique evidence records",
                len(_qc_evidence_register(qc_result)),
                "Deduplicated structured source fields",
            ],
        ]
    )
    cost = qc_result.get("estimated_cost_usd")
    if cost is None or legacy:
        rows.append(["Estimated cost", "Not recorded", "No saved estimate"])
    else:
        for label, value in _qc_flatten_cost(cost):
            rows.append([label, value, "Application pricing estimate"])
    _qc_add_table(
        document,
        ["Metric", "Value", "Meaning"],
        rows,
        [3000, 2100, 4260],
    )
    p = document.add_paragraph(style="QC Table Citation")
    p.add_run(
        "Cost is an application estimate based on persisted usage and configured "
        "pricing. The provider invoice is authoritative. Per-lens and per-verifier "
        "counters remain in their detailed records above."
    )
    cost_basis = _qc_dict(qc_result.get("cost_basis"))
    _qc_heading(document, "Saved Pricing Basis", 2)
    if cost_basis:
        for key, value in cost_basis.items():
            _qc_add_label(
                document,
                str(key).replace("_", " ").title(),
                _qc_json(value) if isinstance(value, (dict, list)) else value,
            )
    else:
        document.add_paragraph(
            "No pricing-rate snapshot was saved. The aggregate estimate cannot be "
            "independently recomputed from this legacy record alone."
        )


def _qc_signoff_recommendation(
    qc_result: dict, *, stale: bool
) -> tuple[str, str]:
    statuses = [
        item
        for item in _qc_list(qc_result.get("lens_statuses"))
        if isinstance(item, dict)
    ]
    incomplete = not statuses or any(
        item.get("status") != "completed" for item in statuses
    )
    findings = [
        item
        for item in _qc_list(qc_result.get("findings"))
        if isinstance(item, dict)
    ]
    open_findings = [
        item
        for item in findings
        if str(item.get("status") or "open") == "open"
    ]
    if stale:
        return (
            "HOLD - RERUN FINAL QC",
            "One or more material run inputs no longer match current state.",
        )
    export_control_issues = _qc_export_control_issues(qc_result)
    if export_control_issues:
        controls = "; ".join(
            issue.rstrip(".") for issue in export_control_issues
        )
        return (
            "HOLD - EXPORT-TIME QC CONTROL BLOCKED",
            f"The selected report is not issue-ready: {controls}.",
        )
    if incomplete or str(
        qc_result.get("execution_status") or "complete"
    ).lower() not in {"complete", "completed", "success"}:
        return (
            "HOLD - COMPLETE FAILED OR MISSING REVIEWERS",
            "Execution coverage is incomplete; missing results are not passes.",
        )
    if any(
        str(item.get("severity")) in {"critical", "high"}
        for item in open_findings
    ):
        return (
            "HOLD - RESOLVE HIGH-PRIORITY OPEN FINDINGS",
            "At least one verified critical or high finding remains open.",
        )
    if open_findings:
        return (
            "REVIEW REQUIRED - OPEN FINDINGS REMAIN",
            "Resolve or formally disposition every open finding before issue.",
        )
    return (
        "QC RECORD COMPLETE - PROFESSIONAL APPROVAL STILL REQUIRED",
        "No surviving finding is currently open in the saved result.",
    )


def _qc_render_limitations_and_signoff(
    document,
    qc_result: dict,
    *,
    stale: bool,
) -> None:
    _qc_heading(document, "Limitations", 1)
    limitations = [
        "This report is advisory. It is not a substitute for review, sealing, or approval by the licensed design professional of record.",
        "A zero-finding lens result means that reviewer reported no candidate defect; it does not prove that no defect exists.",
        "Failed, cancelled, missing, or legacy-unrecorded reviewer data are coverage limitations and must not be read as successful checks.",
        "Source acceptance records retrieval and grounding checks. It does not independently establish code applicability, legal authority, edition currency, or complete support for every statement.",
        "Validated proposed operations passed the saved structural validation only. Human review remains necessary for design intent, coordination, and downstream consequences.",
        "Cost is estimated from saved counters and configured rates. Billing records from the provider control if figures differ.",
        "Older report schemas may omit input manifests, per-check work, source decisions, verifier failures, usage, request counts, or disposition history; each omission is labeled in this report.",
    ]
    if stale:
        limitations.insert(
            0,
            "This report is stale because at least one material QC input no longer matches the recorded run. The changed input may be the document, research profile, standards, module or discipline, model configuration, or source policy.",
        )
    failed_lenses = [
        item
        for item in _qc_list(qc_result.get("lens_statuses"))
        if isinstance(item, dict) and item.get("status") != "completed"
    ]
    if failed_lenses:
        limitations.append(
            f"{len(failed_lenses)} lens record(s) are failed or incomplete; see "
            "the lens audit trail for the exact errors."
        )
    unresolved = [
        item
        for item in [
            *_qc_list(qc_result.get("findings")),
            *_qc_list(qc_result.get("refuted")),
            *_qc_list(qc_result.get("inconclusive")),
        ]
        if isinstance(item, dict)
        and "element_resolved" in item
        and not item.get("element_resolved")
    ]
    if unresolved:
        limitations.append(
            f"{len(unresolved)} candidate finding(s) have unresolved reviewed anchors. "
            "Their exact reviewed text is unavailable and they must be located by a "
            "human reviewer before any edit is made."
        )
    _qc_add_bullets(document, limitations)

    _qc_heading(document, "Reviewer Sign-off", 1)
    recommendation, basis = _qc_signoff_recommendation(qc_result, stale=stale)
    _qc_add_callout(
        document,
        "Recommended control state",
        f"{recommendation}. {basis}",
        accent=(
            _QC_RISK if recommendation.startswith("HOLD") else _QC_CAUTION
        ),
    )
    document.add_paragraph(
        "The undersigned reviewer should confirm the run identity, read every failed "
        "reviewer record, resolve or accept every surviving finding, inspect source "
        "applicability, and verify all applied edits in the current specification. "
        "Signing this page records human review; the software report itself does not "
        "approve the specification."
    )
    _qc_add_bullets(
        document,
        [
            "Run ID and reviewed-input fingerprint confirmed.",
            "All open findings and disposition events reviewed.",
            "Failed, cancelled, and missing reviewer coverage assessed.",
            "Evidence applicability and source authority independently checked.",
            "Current specification and any applied operations reviewed in context.",
        ],
    )
    document.add_paragraph(
        "Reviewer decision (select in signed copy): Accept for coordination | "
        "Return for correction | Re-run Final QC"
    )
    for label in ("Reviewer name and role", "Signature", "Date", "Comments"):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(f"{label}: ")
        _qc_set_run_font(run, bold=True)
        _qc_set_paragraph_border(
            p, edge="bottom", color=_QC_BORDER, size=4, space=2
        )


def _qc_configure_core_properties(
    document, qc_result: dict, section: SpecSection
) -> None:
    """Stamp trustworthy package metadata on the generated audit artifact."""
    section_ref = (
        f"Section {section.number or '[TBD]'} | {section.title or '[TBD]'}"
    )
    run_id = str(qc_result.get("run_id") or "legacy run ID not recorded")
    generated_at = datetime.now(timezone.utc)
    properties = document.core_properties
    properties.title = f"Final QC Audit Report - {section_ref}"[:255]
    properties.subject = (
        f"Build-a-Spec Final QC audit record for run {run_id} | {section_ref}"
    )[:255]
    properties.author = "Build-a-Spec"
    properties.last_modified_by = "Build-a-Spec"
    properties.created = generated_at
    properties.modified = generated_at


def build_qc_memo(qc_result: dict, section: SpecSection, *, stale: bool) -> bytes:
    """Build the standalone, audit-grade Final-QC report.

    The export preserves execution identity, every lens outcome, every
    surviving, refuted, and infrastructure-inconclusive candidate, evidence
    and grounding records, verifier votes, complete proposed operations and validation, dispositions,
    usage/cost, limitations, and human sign-off. Missing legacy fields are
    labeled as missing rather than treated as passes.
    """
    document = Document()
    _qc_configure_core_properties(document, qc_result, section)
    try:
        schema_version = int(qc_result.get("schema_version", 1) or 1)
    except (TypeError, ValueError):
        schema_version = 1
    setattr(document, "_qc_schema_version", schema_version)
    _style_base(document)
    _qc_configure_styles(document)
    _qc_setup_page(document, section)

    execution_label = _qc_execution_label(qc_result, stale)
    _qc_add_masthead(
        document,
        qc_result,
        section,
        stale=stale,
        execution_label=execution_label,
    )
    _qc_render_executive_status(
        document,
        qc_result,
        stale=stale,
        execution_label=execution_label,
    )
    _qc_render_identity(
        document,
        qc_result,
        section,
        stale=stale,
        execution_label=execution_label,
    )
    _qc_render_export_current_state(document, qc_result)
    _qc_render_methodology(document, qc_result)
    _qc_render_lenses(document, qc_result)
    _qc_render_surviving_findings(document, qc_result)
    _qc_render_refuted_appendix(document, qc_result)
    _qc_render_inconclusive_appendix(document, qc_result)
    _qc_render_evidence_register(document, qc_result)
    _qc_render_usage_and_cost(document, qc_result)
    _qc_render_limitations_and_signoff(document, qc_result, stale=stale)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
