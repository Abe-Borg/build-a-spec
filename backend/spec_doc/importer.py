"""Master-spec ``.docx`` import: parse an office master into the tree.

The extraction mechanics are ported from Claude-Spec-Critic
``src/input/extractor.py`` — specifically the pieces that took field
sessions to get right: the **Accept-All tracked-changes text resolution**
(python-docx ``Paragraph.text`` silently produces a hybrid that matches
neither Accept-All nor Reject-All when a redline is pending), pending-
revision *detection* across body and tables, and the drawing-heavy
**content-loss warning**. The SectionFormat tree builder on top is
Build-a-Spec-native: Spec Critic extracts flat text for review; drafting
needs the PART → article → paragraph hierarchy.

Parsing philosophy — **keep everything, warn loudly**: office masters vary
wildly, so unrecognized structure is never dropped. A paragraph that fits
no heading pattern becomes a level-0 paragraph under the current article; a
paragraph arriving before any article lands in a synthetic ``IMPORTED
CONTENT`` article; nesting deeper than SectionFormat's four levels clamps
to level four. Every such decision is recorded in ``ImportResult.warnings``
so the reviewer knows exactly where the parse guessed. Every imported block
enters with status ``imported`` (not yet reviewed for this project) and the
interview pivots to gap-and-adapt mode.

Both manual-label masters ("A. Provide...") and Word-auto-numbered masters
(labels live in ``w:numPr``, not the text) are handled: explicit text
labels win; otherwise the paragraph's numbering indent level (``ilvl``)
drives the depth — placed *relative to its own list*, because ``ilvl`` is
an indent level inside a numbering definition and not an absolute outline
depth (see :meth:`_TreeBuilder.numbered_paragraph`).

Auto-numbered PART and ARTICLE **headings** are recognized through the
numbering definitions themselves: a ``w:numPr`` paragraph whose level's
``lvlText`` renders the literal word ``PART`` is a part heading, and one
whose ``lvlText`` is two decimal placeholders joined by a dot (``%1.%2`` —
the auto-numbered form of the ``2.01 TITLE`` label) is an article heading.
That is the structural signal the visible text cannot carry — an
auto-numbered article's text is just ``SUMMARY`` — and it is provably safe
for the export/re-import round trip: the app's own numbering definitions
use only single-token ``lvlText``s (see ``word_numbering._LEVELS``), so
neither grammar can match a normalized export.

Numbering is resolved the way Word resolves it: a paragraph's own
``w:numPr`` first, else the ``w:numPr`` its paragraph STYLE carries (through
``w:basedOn`` chains). Office masters keep the whole outline on their
PRT/ART/PR1-PR4 styles and the paragraph carries only ``w:pStyle`` — read
direct numbering alone, every heading in such a master exposes its bare
title ("SUMMARY") and the file arrives as one flat blob. The CSI style NAMES
(PRT, ART, PR1..PR5) are the secondary signal, consulted only when the
resolved numbering promotes nothing: they are a decades-old convention, not
a localized free-form label.

The section's identity is decided in the FRONT MATTER — the body before its
first PART or article heading — and decided once: a ``SECTION 21 05 00``
line, a bare ``21 05 00 — TITLE`` line, or a cover page's ``Section Number:
21 05 00`` field. A SECTION-shaped line after structure has begun is a
provision citing a sibling section ("Section 09 90 00 – Painting and
Coating" in a Related Requirements list), never a header, and the first
header found is never overwritten by a later one. Everything else in the
front matter (cover page, revision history, table of contents) is recorded
but not modelled: the appearance-preserving export emits it verbatim, in
place, ahead of the section.
"""
from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from lxml.etree import XMLSyntaxError

from .model import (
    MAX_PARAGRAPH_DEPTH,
    Article,
    Paragraph,
    SpecSection,
    iter_paragraphs,
)
from .source_format import (
    FormatAnchor,
    HEADER_SOURCE_CHROME,
    HEADER_SOURCE_FRONT_MATTER,
    HEADER_SOURCE_LINE,
    SECTION_TITLE_UID,
    LABEL_AUTO,
    LABEL_MANUAL,
    LABEL_NONE,
    SourceFormatMap,
    build_format_map,
)
from .source_mapping import (
    SourceBodyMap,
    SourceParagraphBinding,
    bind_opaque_projection,
    bind_source_paragraph,
    build_source_body_map,
)

# ---------------------------------------------------------------------------
# Accept-All tracked-changes text resolution
# [PORT ≈verbatim: Spec Critic src/input/extractor.py]
# ---------------------------------------------------------------------------

_W_R = qn("w:r")
_W_HYPERLINK = qn("w:hyperlink")
_W_INS = qn("w:ins")
_W_DEL = qn("w:del")
_W_MOVE_FROM = qn("w:moveFrom")
_W_MOVE_TO = qn("w:moveTo")

_ACCEPTED_REVISION_WRAPPERS = frozenset({_W_INS, _W_MOVE_TO})
_W_TXBX_CONTENT = qn("w:txbxContent")
_MC_FALLBACK = (
    "{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback"
)


def _text_box_text(run_el) -> list[str]:
    """The text of every text box anchored in ``run_el``, one string per box.

    A cover page is routinely built from text boxes, and a text box's
    paragraphs live under ``w:txbxContent`` inside the drawing rather than
    among the paragraph's own runs — invisible to ``Paragraph.text``, so a
    whole cover page used to read as empty and be dropped. Word writes each
    DrawingML text box twice, a ``mc:Choice`` drawing plus a ``mc:Fallback``
    VML copy; the fallback is skipped or every box would read twice.
    """
    boxes: list[str] = []
    for box in run_el.iter(_W_TXBX_CONTENT):
        if any(ancestor.tag == _MC_FALLBACK for ancestor in box.iterancestors()):
            continue
        if any(
            ancestor.tag == _W_TXBX_CONTENT and ancestor is not box
            for ancestor in box.iterancestors()
        ):
            continue  # a box nested in a box is read by its outer box
        lines = [
            " ".join(_accept_all_paragraph_text(p_el).split())
            for p_el in box.iter(qn("w:p"))
        ]
        text = " ".join(line for line in lines if line)
        if text:
            boxes.append(text)
    return boxes
_REVISION_MARKER_TAGS = (_W_INS, _W_DEL, _W_MOVE_FROM, _W_MOVE_TO)


def _collect_accept_all_text(container, parts: list[str]) -> None:
    """Append the Accept-All run/hyperlink text under ``container``.

    Mirrors python-docx ``CT_P.text`` with one addition: descends through
    *accepted* revision wrappers (``<w:ins>`` / ``<w:moveTo>``) and skips
    ``<w:del>`` / ``<w:moveFrom>`` entirely — the text that remains once
    the redline is accepted, i.e. what will actually be issued. A document
    with no revision markup yields output byte-identical to
    ``Paragraph.text``.
    """
    for child in container:
        tag = child.tag
        if not isinstance(tag, str):
            continue  # comments / processing instructions carry no run text
        if tag == _W_R or tag == _W_HYPERLINK:
            parts.append(child.text or "")
            if tag == _W_R:
                boxed = _text_box_text(child)
                if boxed:
                    parts.append(" " + " ".join(boxed))
        elif tag in _ACCEPTED_REVISION_WRAPPERS:
            _collect_accept_all_text(child, parts)


def _accept_all_paragraph_text(p_el) -> str:
    parts: list[str] = []
    _collect_accept_all_text(p_el, parts)
    return "".join(parts)


def _element_has_tracked_changes(el) -> bool:
    return any(el.find(".//" + tag) is not None for tag in _REVISION_MARKER_TAGS)


# ---------------------------------------------------------------------------
# Content-loss warning [PORT: Spec Critic extractor._detect_content_loss_warning]
# ---------------------------------------------------------------------------

_CONTENT_LOSS_WARNING_THRESHOLD = 0.20


def _detect_content_loss_warning(body) -> str | None:
    drawing_qn = qn("w:drawing")
    pict_qn = qn("w:pict")
    object_qn = qn("w:object")
    sect_pr_qn = qn("w:sectPr")

    total = 0
    non_text = 0
    for child in body:
        if child.tag == sect_pr_qn:
            continue
        total += 1
        if (
            child.find(".//" + drawing_qn) is not None
            or child.find(".//" + pict_qn) is not None
            or child.find(".//" + object_qn) is not None
        ):
            non_text += 1
    if total == 0 or non_text == 0:
        return None
    proportion = non_text / total
    if proportion <= _CONTENT_LOSS_WARNING_THRESHOLD:
        return None
    return (
        f"The master contains {round(proportion * 100)}% non-text elements "
        "(drawings, pictures, or embedded objects). Some content may not "
        "have been imported — verify against the source visually."
    )


# ---------------------------------------------------------------------------
# Heading / label patterns (Build-a-Spec-native tree heuristics)
# ---------------------------------------------------------------------------

_SECTION_RE = re.compile(
    r"^SECTION\s+(\d{2})\s*(\d{2})\s*(\d{2})(?:\.(\d{2}))?\b\s*[-–—]?\s*(.*)$",
    re.IGNORECASE,
)
# PART 1-3 are SectionFormat's parts; 4/5 are accepted as spec structure and
# mapped (loudly) onto PART 3 rather than demoting the file to unstructured.
_PART_RE = re.compile(r"^PART\s*([1-5])\b", re.IGNORECASE)
_END_RE = re.compile(r"^END\s+OF\s+SECTION\b", re.IGNORECASE)
# "1.1 SUMMARY" / "1.01 SUMMARY" / "2.3 - PIPING" (part digit + article no.).
# Digits 4/5 are accepted alongside PART 4/5 above and remapped to PART 3 the
# same way — an accepted "PART 4" whose "4.01" articles were NOT accepted
# would shed its article titles into the synthetic IMPORTED CONTENT catch-all.
_ARTICLE_RE = re.compile(r"^([1-5])\.(\d{1,2})\.?\s+[-–—]?\s*(\S.*)$")
# A header line without the SECTION keyword ("23 05 48 — TITLE"). Consulted
# for the FIRST content line only — anywhere else six digits and a dash are
# far more likely to be a provision than a header.
_BARE_SECTION_RE = re.compile(
    r"^(\d{2})\s?(\d{2})\s?(\d{2})(?:\.(\d{2}))?\s*[-–—]\s*(\S.*)$"
)
# Manual paragraph labels by depth.
_LEVEL_RES = (
    re.compile(r"^([A-Z]{1,2})\.\s+(\S.*)$"),  # A.  (depth 0)
    re.compile(r"^(\d{1,2})\.\s+(\S.*)$"),  # 1.  (depth 1)
    re.compile(r"^([a-z]{1,2})\.\s+(\S.*)$"),  # a.  (depth 2)
    re.compile(r"^(\d{1,2})\)\s+(\S.*)$"),  # 1)  (depth 3)
)

# The label grammars that promote a ``w:numPr`` paragraph to a HEADING. A
# level whose rendered label carries the literal word PART is a part heading
# whatever its number format ("PART %1" under decimal or upperRoman alike);
# one that renders two decimal placeholders joined by a dot is the
# auto-numbered form of the "2.01 TITLE" article label. Everything else —
# including every ``lvlText`` the app's own normalized exports write
# ("%1.", "%2.", "%3.", "%4)") — stays a provision, which is what keeps the
# export/re-import round trip untouched by construction.
_PART_LVLTEXT_RE = re.compile(r"\bPART\b", re.IGNORECASE)
_ARTICLE_LVLTEXT_RE = re.compile(r"^%\d+\.%\d+\.?\s*$")
_ARTICLE_NUM_FMTS = frozenset({"", "decimal", "decimalZero"})

# A cover page's "Section Number: 21 05 00" field. Distinct from _SECTION_RE
# (which needs the digits right after the keyword), so a number carried this
# way is still found; the title then comes from the cover page's own title
# line beside it, or from the page header/footer.
_SECTION_NUMBER_FIELD_RE = re.compile(
    r"^SECTION\s+(?:NUMBER|NO\.?|#)\s*[:\-–—]?\s*"
    r"(\d{2})\s?(\d{2})\s?(\d{2})(?:\.(\d{2}))?\s*\.?$",
    re.IGNORECASE,
)
# Any MasterFormat-shaped number, for the page header/footer fallback.
_MASTERFORMAT_RE = re.compile(r"\b(\d{2})\s?(\d{2})\s?(\d{2})(?:\.(\d{2}))?\b")
# The CSI/MasterSpec paragraph-style convention (PRT = part, ART = article,
# PR1..PR5 = provision levels), matched on the style id or name after folding
# case and separators. A secondary signal only: consulted when the resolved
# numbering definition promoted nothing.
_CSI_STYLE_KINDS: dict[str, object] = {
    "PRT": "part",
    "ART": "article",
    "PR1": 0,
    "PR2": 1,
    "PR3": 2,
    "PR4": 3,
    "PR5": 3,
}
# Cover-page lines that are labels or metadata rather than the title.
_FRONT_MATTER_TITLE_STOP_RE = re.compile(
    r"^(?:specification|spec\b|prepared|issued|revision|rev\b|date\b|page\b|"
    r"sheet\b|for\b|by\b|table of contents|contents\b)",
    re.IGNORECASE,
)
def _clip(value: str, limit: int = 60) -> str:
    return value if len(value) <= limit else value[:limit] + "…"


def _title_like(text: str) -> bool:
    """Does a front-matter or header/footer line read like a section title?

    Two-plus mostly-alphabetic words, no MasterFormat number, no
    ``Label: value`` colon, and none of the cover-page label words. This is
    a fallback for a document that carries its identity only on a cover
    page or in its page furniture; whatever it picks is reported in the
    import notes as read-from-there, so the user checks it once.
    """
    if not text or len(text) > 120 or ":" in text:
        return False
    if (
        _MASTERFORMAT_RE.search(text)
        or _SECTION_RE.match(text)
        or _SECTION_NUMBER_FIELD_RE.match(text)
        or _FRONT_MATTER_TITLE_STOP_RE.match(text)
    ):
        return False
    words = text.split()
    if len(words) < 2 or len(words) > 14:
        return False
    non_space = sum(1 for ch in text if not ch.isspace())
    letters = sum(1 for ch in text if ch.isalpha())
    return non_space > 0 and letters / non_space >= 0.7


def _identity_from_chrome(lines: tuple[str, ...]) -> tuple[str, str]:
    """``(number, title)`` as far as the page headers/footers state them.

    A spec's running header conventionally carries the title and its footer
    the number (``21 05 00 - 1``); either may also carry a whole
    ``SECTION 21 05 00 - TITLE`` line. Consulted only when the body itself
    carries no header line, and always disclosed in the import notes.
    """
    number = ""
    title = ""
    for line in lines:
        section_match = _SECTION_RE.match(line)
        bare_match = _BARE_SECTION_RE.match(line) if not section_match else None
        if section_match or bare_match:
            g1, g2, g3, g4, remainder = (section_match or bare_match).groups()
            if not number:
                number = f"{g1} {g2} {g3}" + (f".{g4}" if g4 else "")
                if remainder.strip() and _title_like(remainder.strip()):
                    title = remainder.strip()
            continue
        found = _MASTERFORMAT_RE.search(line)
        if found and not number:
            number = f"{found.group(1)} {found.group(2)} {found.group(3)}" + (
                f".{found.group(4)}" if found.group(4) else ""
            )
    if not title:
        for line in lines:
            if _title_like(line):
                title = line
                break
    return number, title


# Said once, at the top of the warning list, when a file carries none of the
# three SectionFormat markers. The tree is still SectionFormat underneath (the
# whole editing/lint/diff/QC machinery is typed to it), so this states plainly
# what the app did rather than letting the panel imply the document always had
# a section header and three parts.
UNSTRUCTURED_IMPORT_WARNING = (
    "This file does not look like a SectionFormat spec section — no SECTION "
    "number, PART heading, or numbered article was found. Its content was "
    "imported in document order under a placeholder article, and the original "
    "file is retained exactly. Set a section number and title when you are "
    "ready to turn it into a spec section."
)


@dataclass
class ImportResult:
    """A parsed master: the tree plus the parse's honesty trail."""

    section: SpecSection
    warnings: list[str] = field(default_factory=list)
    tracked_changes_detected: bool = False
    imported_block_count: int = 0
    skipped_empty_count: int = 0
    # Immutable UID -> source-body anchors for P1 source-preserving export.
    # Kept outside SpecSection/version snapshots and never sent to the model.
    source_map: SourceBodyMap | None = None
    # Where every semantic element came from in the uploaded package, for the
    # appearance-preserving export. Like ``source_map`` it lives outside the
    # SpecSection/version snapshots and is never sent to the model.
    format_map: SourceFormatMap | None = None
    # True when the parse recognized at least one SectionFormat marker (a
    # SECTION line, a PART heading, or an ``N.M`` article). False means the
    # spec scaffolding around the content — section header, empty parts, the
    # synthetic article — is entirely the importer's, so the UI must not
    # present it as something the source document carried. Defaults True so a
    # caller that never sets it keeps the historical posture.
    spec_shape_detected: bool = True
    # Body blocks before the section's first PART/article heading — a cover
    # page, a revision history, a table of contents — kept for export exactly
    # as they are and NOT modelled in the tree. Readable lines, in order.
    front_matter: tuple[str, ...] = ()
    # Where the section number/title came from (one of the HEADER_SOURCE_*
    # values, or "" when nothing stated it).
    header_source: str = ""
    # Whether any paragraph's numbering was inherited from its style — a
    # diagnostics fact, so a bundle can say which parser path a master took.
    style_numbering_resolved: bool = False


class MasterImportError(ValueError):
    """The file could not be parsed as a master spec at all."""


def _direct_numbering(
    paragraph: DocxParagraph,
) -> tuple[int | None, int | None] | None:
    """The paragraph's OWN ``w:numPr`` as ``(numId, ilvl)``, or ``None``.

    Either member may be ``None`` when the element names only the other —
    a direct ``w:numPr`` that carries just an ``ilvl`` inherits the
    definition its style names, which :func:`_effective_numbering` resolves.
    """
    p_pr = paragraph._p.pPr
    if p_pr is None or p_pr.numPr is None:
        return None
    num_pr = p_pr.numPr
    num_id = None
    if num_pr.numId is not None and num_pr.numId.val is not None:
        num_id = int(num_pr.numId.val)
    ilvl = None
    if num_pr.ilvl is not None and num_pr.ilvl.val is not None:
        ilvl = int(num_pr.ilvl.val)
    return num_id, ilvl


def _paragraph_style_id(paragraph: DocxParagraph) -> str:
    p_pr = paragraph._p.pPr
    if p_pr is None or p_pr.pStyle is None:
        return ""
    return str(p_pr.pStyle.val or "")


def _effective_numbering(
    paragraph: DocxParagraph,
    style_numbering: dict[str, tuple[int, int]],
    default_style_id: str = "",
) -> tuple[int, int] | None:
    """``(numId, ilvl)`` the way Word resolves it, or ``None`` when unnumbered.

    Direct ``w:numPr`` wins. ``numId`` 0 is OOXML for "no numbering" and
    cancels whatever the style would have supplied. A direct element naming
    only a level inherits the style's definition; one naming only a
    definition starts at that definition's level 0. With no direct element
    at all the paragraph's style (or the document's default paragraph
    style) decides — which is where every office master keeps its outline.
    ``ilvl`` is an indent level within the definition, not an outline depth
    (:meth:`_TreeBuilder.numbered_paragraph` places it).
    """
    style_id = _paragraph_style_id(paragraph) or default_style_id
    inherited = style_numbering.get(style_id) if style_id else None
    direct = _direct_numbering(paragraph)
    if direct is None:
        return inherited
    num_id, ilvl = direct
    if num_id is not None and num_id <= 0:
        return None
    if num_id is None:
        if inherited is None:
            return None
        num_id = inherited[0]
        if ilvl is None:
            ilvl = inherited[1]
    if ilvl is None:
        ilvl = 0
    return num_id, ilvl


def _int_attr(element, attribute: str) -> int | None:
    value = element.get(qn(attribute)) if element is not None else None
    if value is None:
        return None
    try:
        return int(value)
    except ValueError:
        return None


def _styles_root(document):
    """The styles part's root, read through the relationship (never created)."""
    try:
        part = document.part.part_related_by(RELATIONSHIP_TYPE.STYLES)
        return part.element
    except (KeyError, AttributeError, ValueError):
        return None


def _default_paragraph_style_id(document) -> str:
    root = _styles_root(document)
    if root is None:
        return ""
    for style_el in root.findall(qn("w:style")):
        if (style_el.get(qn("w:type")) or "paragraph") != "paragraph":
            continue
        if style_el.get(qn("w:default")) in ("1", "true", "on"):
            return str(style_el.get(qn("w:styleId")) or "")
    return ""


def _load_style_numbering(document) -> dict[str, tuple[int, int]]:
    """``styleId -> (numId, ilvl)`` for every paragraph style that numbers.

    Resolved through ``w:basedOn`` chains once, the way Word resolves them,
    so a ``PR2`` based on ``PR1`` based on a numbered ``Body`` style reads
    the numbering it actually renders with. A style whose own ``w:numPr``
    names ``numId`` 0 cancels its parent's numbering. Read via the
    relationship directly, like the numbering catalog — an importer never
    mutates the package it inspects — and any malformation degrades to no
    style numbering at all, i.e. the direct-``w:numPr``-only behavior every
    earlier import had.
    """
    root = _styles_root(document)
    if root is None:
        return {}
    try:
        # Each style's OWN ``w:numPr`` as (numId, ilvl), either side None
        # when the element names only the other: a derived style routinely
        # states just a deeper ``ilvl`` and inherits the definition.
        own: dict[str, tuple[int | None, int | None]] = {}
        based_on: dict[str, str] = {}
        for style_el in root.findall(qn("w:style")):
            if (style_el.get(qn("w:type")) or "paragraph") != "paragraph":
                continue
            style_id = style_el.get(qn("w:styleId"))
            if not style_id:
                continue
            based_el = style_el.find(qn("w:basedOn"))
            parent = based_el.get(qn("w:val")) if based_el is not None else None
            if parent:
                based_on[style_id] = parent
            p_pr = style_el.find(qn("w:pPr"))
            num_pr = p_pr.find(qn("w:numPr")) if p_pr is not None else None
            if num_pr is None:
                continue
            own[style_id] = (
                _int_attr(num_pr.find(qn("w:numId")), "w:val"),
                _int_attr(num_pr.find(qn("w:ilvl")), "w:val"),
            )
        resolved: dict[str, tuple[int, int]] = {}
        for style_id in set(own) | set(based_on):
            # Property-by-property inheritance, the way Word merges a style
            # chain: the nearest style naming a definition supplies it, the
            # nearest naming a level supplies that, independently.
            seen: set[str] = set()
            cursor: str = style_id
            num_id: int | None = None
            ilvl: int | None = None
            while cursor and cursor not in seen:
                seen.add(cursor)
                raw = own.get(cursor)
                if raw is not None:
                    if num_id is None and raw[0] is not None:
                        num_id = raw[0]
                    if ilvl is None and raw[1] is not None:
                        ilvl = raw[1]
                    if num_id is not None:
                        break
                cursor = based_on.get(cursor, "")
            if num_id is not None and num_id > 0:
                resolved[style_id] = (num_id, ilvl if ilvl is not None else 0)
        return resolved
    except Exception:  # noqa: BLE001 - a malformed styles part is no catalog
        return {}


def _load_style_kinds(document) -> dict[str, object]:
    """``styleId -> CSI kind`` for paragraph styles following the convention.

    Matched on the style id and on its name (Word derives one from the
    other, but not always identically), after folding case and separators.
    """
    root = _styles_root(document)
    if root is None:
        return {}
    kinds: dict[str, object] = {}
    try:
        for style_el in root.findall(qn("w:style")):
            if (style_el.get(qn("w:type")) or "paragraph") != "paragraph":
                continue
            style_id = style_el.get(qn("w:styleId"))
            if not style_id:
                continue
            name_el = style_el.find(qn("w:name"))
            name = name_el.get(qn("w:val")) if name_el is not None else ""
            for candidate in (style_id, name or ""):
                key = re.sub(r"[\s_\-]+", "", candidate).upper()
                if key in _CSI_STYLE_KINDS:
                    kinds[style_id] = _CSI_STYLE_KINDS[key]
                    break
    except Exception:  # noqa: BLE001 - malformed styles carry no convention
        return {}
    return kinds


def _lvl_entry(lvl_el) -> tuple[int, tuple[str, str]] | None:
    """One ``w:lvl`` element as ``(ilvl, (numFmt, lvlText))``, or ``None``."""
    raw_ilvl = lvl_el.get(qn("w:ilvl"))
    if raw_ilvl is None:
        return None
    try:
        ilvl = int(raw_ilvl)
    except ValueError:
        return None
    fmt_el = lvl_el.find(qn("w:numFmt"))
    text_el = lvl_el.find(qn("w:lvlText"))
    num_fmt = fmt_el.get(qn("w:val")) if fmt_el is not None else ""
    lvl_text = text_el.get(qn("w:val")) if text_el is not None else ""
    return ilvl, (num_fmt or "", lvl_text or "")


def _load_numbering_catalog(document) -> dict[tuple[int, int], tuple[str, str]]:
    """``(numId, ilvl) -> (numFmt, lvlText)`` from the package's numbering part.

    Read via the relationship directly — python-docx's ``numbering_part``
    property CREATES an empty part when none exists, and an importer must
    never mutate the package it is inspecting. A per-``num`` ``lvlOverride``
    that redefines a level replaces the abstract's grammar for it, so the
    label the reader actually sees is the one that decides. Any absence or
    malformation degrades to an empty catalog, i.e. exactly the
    no-promotion behavior every pre-catalog import had.
    """
    try:
        part = document.part.part_related_by(RELATIONSHIP_TYPE.NUMBERING)
        root = part.element
    except (KeyError, AttributeError, ValueError):
        return {}
    try:
        abstract: dict[str, dict[int, tuple[str, str]]] = {}
        for abstract_el in root.findall(qn("w:abstractNum")):
            abstract_id = abstract_el.get(qn("w:abstractNumId"))
            if abstract_id is None:
                continue
            levels: dict[int, tuple[str, str]] = {}
            for lvl_el in abstract_el.findall(qn("w:lvl")):
                entry = _lvl_entry(lvl_el)
                if entry is not None:
                    levels[entry[0]] = entry[1]
            abstract[abstract_id] = levels
        catalog: dict[tuple[int, int], tuple[str, str]] = {}
        for num_el in root.findall(qn("w:num")):
            raw_num_id = num_el.get(qn("w:numId"))
            abstract_ref = num_el.find(qn("w:abstractNumId"))
            if raw_num_id is None or abstract_ref is None:
                continue
            try:
                num_id = int(raw_num_id)
            except ValueError:
                continue
            for ilvl, entry in abstract.get(
                abstract_ref.get(qn("w:val")), {}
            ).items():
                catalog[(num_id, ilvl)] = entry
            for override in num_el.findall(qn("w:lvlOverride")):
                lvl_el = override.find(qn("w:lvl"))
                if lvl_el is None:
                    continue
                entry = _lvl_entry(lvl_el)
                if entry is not None:
                    catalog[(num_id, entry[0])] = entry[1]
        return catalog
    except Exception:  # noqa: BLE001 - a malformed numbering part is no catalog
        return {}


def _promoted_heading_kind(
    catalog: dict[tuple[int, int], tuple[str, str]],
    num_id: int,
    ilvl: int,
) -> str:
    """``"part"`` / ``"article"`` when the numbering label says so, else ``""``."""
    if not catalog or num_id <= 0:
        return ""
    entry = catalog.get((num_id, ilvl))
    if entry is None:
        return ""
    num_fmt, lvl_text = entry
    if _PART_LVLTEXT_RE.search(lvl_text):
        return "part"
    if _ARTICLE_LVLTEXT_RE.match(lvl_text) and num_fmt in _ARTICLE_NUM_FMTS:
        return "article"
    return ""


@dataclass(frozen=True)
class _BodyTextEntry:
    text: str
    paragraph: DocxParagraph | None
    body_child_index: int
    source_element: object
    opaque_blocker: str = ""
    # Non-empty when the body child is preserved content the appearance-
    # preserving export emits verbatim (model.LOCK_REASONS). Distinct from
    # ``opaque_blocker``, which is the older byte-exact mode's vocabulary.
    lock_reason: str = ""
    # Pre-normalized text that must survive the main loop's whitespace fold —
    # a preserved table keeps one line per row so it stays readable.
    preformatted: bool = False


def _header_footer_text(document) -> tuple[str, ...]:
    """Readable lines from every header and footer, in part order.

    Captured so the app can NOTICE when a spec footer still carries the
    number of the master the section was adapted from. Headers and footers
    are never rewritten — that is the preservation contract — so this is
    evidence for a warning, never an editing surface.
    """
    lines: list[str] = []
    for section in document.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            # Only a part with its OWN definition is read. python-docx's
            # ``paragraphs`` on a header that has none CREATES one — a new
            # part plus a ``w:headerReference`` in the body's ``w:sectPr`` —
            # and an importer must never mutate the package it inspects:
            # the source map hashes that very ``w:sectPr``. A later section's
            # linked header repeats the earlier text, which is deduplicated
            # below anyway.
            if part is None or getattr(part, "is_linked_to_previous", True):
                continue
            try:
                paragraphs = part.paragraphs
            except (AttributeError, ValueError):  # pragma: no cover - defensive
                continue
            for paragraph in paragraphs:
                text = " ".join(paragraph.text.split())
                if text and text not in lines:
                    lines.append(text)
    return tuple(lines)


def _iter_body_texts(document) -> list[_BodyTextEntry]:
    """Body content in document order: (accept-all text, paragraph or None).

    A table is ONE entry, not one per row. Under the appearance-preserving
    contract a table is emitted back as the single Word block it is, so it
    has to be one thing the user can move or delete — a row-per-paragraph
    projection made "delete this schedule" an N-paragraph operation and let
    a row beginning "A." be read as a heading. Rows are joined with newlines
    and cells with `` | `` so the grid still reads in the panel and in the
    model's context.

    Paragraphs carrying a picture, an embedded object, or a content control
    are locked for the same reason: their markup is carried through verbatim,
    so their text is not ours to retype.
    """
    results: list[_BodyTextEntry] = []
    body = document.element.body
    source_index = 0
    drawing_qn = qn("w:drawing")
    pict_qn = qn("w:pict")
    object_qn = qn("w:object")
    fld_char_qn = qn("w:fldChar")
    instr_text_qn = qn("w:instrText")
    # Open complex-field frames, outermost first; True marks a TOC field. A
    # cached table of contents is ordinary paragraphs to python-docx, and
    # its entries ("1.1 SUMMARY ... 3") match the article grammar exactly.
    # They are a Word field result: carried through verbatim, never parsed.
    open_fields: list[bool] = []
    for child in body.iterchildren():
        if not isinstance(child.tag, str):
            continue
        body_child_index = source_index
        source_index += 1
        if child.tag == qn("w:p"):
            paragraph = DocxParagraph(child, document)
            in_toc = any(open_fields)
            for field_el in child.iter():
                if field_el.tag == fld_char_qn:
                    kind = field_el.get(qn("w:fldCharType"))
                    if kind == "begin":
                        open_fields.append(False)
                    elif kind == "end" and open_fields:
                        if open_fields.pop():
                            in_toc = True
                elif field_el.tag == instr_text_qn and open_fields:
                    if (field_el.text or "").strip().upper().startswith("TOC"):
                        open_fields[-1] = True
                        in_toc = True
            if child.find(".//" + object_qn) is not None:
                lock = "embedded_object"
            elif (
                child.find(".//" + drawing_qn) is not None
                or child.find(".//" + pict_qn) is not None
            ):
                lock = "image"
            elif in_toc or any(open_fields):
                lock = "field"
            else:
                lock = ""
            results.append(
                _BodyTextEntry(
                    _accept_all_paragraph_text(child),
                    paragraph,
                    body_child_index,
                    child,
                    "",
                    lock,
                )
            )
        elif child.tag == qn("w:tbl"):
            table = DocxTable(child, document)
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = " ".join(
                        _accept_all_paragraph_text(p._p) for p in cell.paragraphs
                    ).strip()
                    cells.append(" ".join(cell_text.split()))
                line = " | ".join(c for c in cells if c)
                if line:
                    rows.append(line)
            if rows:
                results.append(
                    _BodyTextEntry(
                        "\n".join(rows),
                        None,
                        body_child_index,
                        child,
                        "table_projection",
                        "table",
                        True,
                    )
                )
        elif child.tag == qn("w:sdt"):
            # Structured document tags were silently dropped before, which
            # under a rebuild-the-body export would have deleted them from
            # the user's file. Keep the readable text, lock the block.
            text = " ".join(
                " ".join(_accept_all_paragraph_text(p).split())
                for p in child.iter(qn("w:p"))
            ).strip()
            if text:
                results.append(
                    _BodyTextEntry(
                        text,
                        None,
                        body_child_index,
                        child,
                        "content_control_projection",
                        "content_control",
                    )
                )
    return results


@dataclass(frozen=True)
class ReferenceExtraction:
    """Plain readable text pulled out of an attachment for model context.

    The counterpart to :func:`parse_master_docx` for files that are *not*
    becoming the document: no tree, no source map, no structural heuristics —
    just the readable text in document order. For a ``.docx`` that means the
    body resolved through the same Accept-All revision handling and table
    flattening a master import uses, so a reference file reads the way it does
    in Word.

    The type is shared with the non-Word extractors in
    ``backend/reference_extract.py`` (PDF, text, XML, CSV), which set ``kind``
    and may report ``warnings`` about what the read left out — a truncated
    page range, undecodable pages, a non-UTF-8 encoding. ``tracked_changes``
    is meaningful only for Word.
    """

    text: str
    block_count: int
    tracked_changes: bool
    kind: str = "docx"
    warnings: tuple[str, ...] = ()


def extract_reference_text(filepath: str | Path) -> ReferenceExtraction:
    """Extract a reference ``.docx``'s body text. Raises
    :class:`MasterImportError` for anything unreadable."""
    filepath = Path(filepath)
    try:
        document = Document(str(filepath))
    except (
        PackageNotFoundError,
        zipfile.BadZipFile,
        XMLSyntaxError,
        KeyError,
        ValueError,
    ) as exc:
        raise MasterImportError(
            "That file is not a readable .docx document."
        ) from exc

    lines: list[str] = []
    for entry in _iter_body_texts(document):
        text = " ".join(entry.text.split())
        if text:
            lines.append(text)
    return ReferenceExtraction(
        text="\n".join(lines),
        block_count=len(lines),
        tracked_changes=_element_has_tracked_changes(document.element.body),
    )


def _entry_text(entry: _BodyTextEntry) -> str:
    """The parse's normalized text for one entry ("" for a blank one)."""
    return (
        entry.text.strip()
        if entry.preformatted
        else " ".join(entry.text.split())
    )


def _structural_heading_kind(
    entry: _BodyTextEntry,
    text: str,
    catalog: dict[tuple[int, int], tuple[str, str]],
    style_numbering: dict[str, tuple[int, int]],
    default_style_id: str,
    style_kinds: dict[str, object],
) -> str:
    """``"part"`` / ``"article"`` when ANY recognition path would promote
    this entry to a heading, else ``""``. The same order the parse applies:
    the numbering definition's label grammar, the CSI style name, then the
    text grammars — which a numbered paragraph never reaches, because its
    label lives in the definition rather than in the text."""
    if entry.lock_reason or not text:
        return ""
    paragraph = entry.paragraph
    if paragraph is not None:
        numbering = _effective_numbering(
            paragraph, style_numbering, default_style_id
        )
        if numbering is not None:
            kind = _promoted_heading_kind(catalog, *numbering)
            if kind:
                return kind
        style_kind = style_kinds.get(
            _paragraph_style_id(paragraph) or default_style_id
        )
        if style_kind in ("part", "article"):
            return str(style_kind)
        if numbering is not None:
            return ""
    if _PART_RE.match(text):
        return "part"
    if _ARTICLE_RE.match(text):
        return "article"
    return ""


def _first_structure_line(
    entries: list[_BodyTextEntry],
    texts_by_line: dict[int, str],
    catalog: dict[tuple[int, int], tuple[str, str]],
    style_numbering: dict[str, tuple[int, int]],
    default_style_id: str,
    style_kinds: dict[str, object],
) -> int | None:
    """The 1-based line of the first PART/article heading, or ``None``.

    ``None`` means the file has no structure — or ends (``END OF SECTION``)
    before any — and keeps the historical posture in which every block is
    content. Everything before the returned line is front matter.
    """
    for line_no, entry in enumerate(entries, start=1):
        text = texts_by_line[line_no]
        if not text:
            continue
        if _structural_heading_kind(
            entry, text, catalog, style_numbering, default_style_id, style_kinds
        ):
            return line_no
        if (
            not entry.lock_reason
            and entry.paragraph is not None
            and _effective_numbering(
                entry.paragraph, style_numbering, default_style_id
            )
            is None
            and _END_RE.match(text)
        ):
            return None
    return None


def _promoted_part_title(requested: int, mapped: int, text: str) -> str:
    """The Part title for a promoted heading whose visible text is just the
    name ("GENERAL"): the master's own wording under the canonical label, or
    "" (keep whatever the part has) when the part was remapped."""
    if requested != mapped or not text:
        return ""
    if _PART_RE.match(text):
        return text
    return f"PART {mapped} - {text}"


class _TreeBuilder:
    """Builds the SpecSection with the same uid/seq discipline as apply."""

    def __init__(self) -> None:
        self.section = SpecSection.empty()
        self.current_part = None
        self.current_article: Article | None = None
        # Paragraph stack by depth for nesting (index = depth).
        self.stack: list[Paragraph] = []
        # How far this article's Word numbering has been shifted shallower —
        # see numbered_paragraph. Resets with the stack it is relative to.
        self.numbering_offset = 0
        self.warnings: list[str] = []
        # warning index -> the uid of the element the warning is about, so
        # the parse can append a FINDABLE reference once the tree is built.
        # "Line N" alone counts body children (blank paragraphs and table
        # rows included) — a number the user can locate nowhere in Word or
        # in the panel.
        self.warning_uids: dict[int, str] = {}
        self.imported_count = 0

    def part(self, number: int, title: str = ""):
        self.current_part = self.section.parts[number - 1]
        # The master's own heading wording ("PART 1 - GENERAL REQUIREMENTS")
        # rather than the canonical default, so the panel shows what the
        # file says and the appearance-preserving export finds the heading
        # unchanged instead of rewriting it. First heading for a part wins;
        # a PART 4 remapped onto PART 3 never renames it.
        if title and not self.current_part.articles:
            self.current_part.title = title
        self.current_article = None
        self.stack = []
        self.numbering_offset = 0

    def article(self, part_number: int, title: str) -> None:
        self.part(part_number)
        part = self.current_part
        article = Article(uid=f"{part.uid}.a{part.next_seq}", title=title)
        part.next_seq += 1
        part.articles.append(article)
        self.current_article = article
        self.stack = []
        self.numbering_offset = 0

    def ensure_article(self, line_no: int) -> None:
        """Synthesize a container when content precedes any article."""
        if self.current_article is not None:
            return
        if self.current_part is None:
            self.part(1)
        part_title = self.current_part.title
        self.article(self.current_part.number, "IMPORTED CONTENT")
        self.warnings.append(
            f"Line {line_no}: content arrived before any article heading — "
            "kept under a synthetic 'IMPORTED CONTENT' article in "
            f"{part_title}."
        )
        self.warning_uids[len(self.warnings) - 1] = self.current_article.uid

    def numbered_paragraph(self, ilvl: int, text: str, line_no: int) -> Paragraph:
        """A Word-auto-numbered paragraph, placed relative to its own list.

        ``ilvl`` is an *indent level within a numbering definition*, not an
        absolute outline depth. A list that reserves level 0 and starts at
        level 1 is as ordinary as one that starts at 0, and a document can
        hold several definitions that disagree about where they begin.

        Reading ``ilvl`` as absolute depth corrupted such a list, and not by
        an off-by-one: its first paragraph had no parent to hang from, so it
        was pushed to the shallowest free depth — while every *sibling* at
        the same ``ilvl`` then found that stack sufficient and landed one
        deeper, i.e. as its **child**. A flat list of articles came back as
        one article owning all the others.

        The shift is therefore remembered for the rest of the article, so
        the siblings move with the paragraph that was already moved. That is
        the whole rule, and it is what makes it safe: the offset can only
        grow at the moment a paragraph was going to be pushed shallower
        anyway, so it never relocates content the old code left alone. In
        particular a list nested under a parent the stack does support — a
        restarted sub-list carrying its own ``numId`` at a deliberately deep
        ``ilvl`` — triggers nothing and keeps its depth, which a per-``numId``
        rebasing of the raw levels would have promoted to top level.

        It is deliberately scoped to the article rather than the document:
        the stack it is relative to resets there, and two articles whose
        lists begin at different levels each get their own answer. A
        document-wide minimum cannot do that, and it also cannot help the
        second list at all when the first one already uses level 0.
        """
        self.ensure_article(line_no)
        depth = max(0, ilvl - self.numbering_offset)
        if depth > len(self.stack):
            self.numbering_offset += depth - len(self.stack)
            depth = len(self.stack)
        return self.paragraph(depth, text, line_no)

    def paragraph(self, depth: int, text: str, line_no: int) -> Paragraph:
        self.ensure_article(line_no)
        pending_warning_indexes: list[int] = []
        if depth >= MAX_PARAGRAPH_DEPTH:
            self.warnings.append(
                f"Line {line_no}: nesting deeper than "
                f"{MAX_PARAGRAPH_DEPTH} levels — clamped to level "
                f"{MAX_PARAGRAPH_DEPTH}."
            )
            pending_warning_indexes.append(len(self.warnings) - 1)
            depth = MAX_PARAGRAPH_DEPTH - 1
        # A deeper level than the stack supports attaches at the deepest
        # available parent + 1 (a master can open with "1." under nothing).
        if depth > len(self.stack):
            self.warnings.append(
                f"Line {line_no}: paragraph level jumped deeper than its "
                f"context — attached at level {len(self.stack)}."
            )
            pending_warning_indexes.append(len(self.warnings) - 1)
            depth = len(self.stack)
        if depth == 0:
            owner = self.current_article
            siblings = owner.paragraphs
        else:
            owner = self.stack[depth - 1]
            siblings = owner.children
        paragraph = Paragraph(
            uid=f"{owner.uid}.p{owner.next_seq}",
            text=text,
            status="imported",
        )
        owner.next_seq += 1
        siblings.append(paragraph)
        self.stack = self.stack[:depth] + [paragraph]
        self.imported_count += 1
        for warning_index in pending_warning_indexes:
            self.warning_uids[warning_index] = paragraph.uid
        return paragraph


def parse_master_docx(filepath: str | Path) -> ImportResult:
    """Parse a master ``.docx`` into an all-``imported`` SectionFormat tree.

    Raises :class:`MasterImportError` for a file that isn't a readable
    ``.docx``; every structural surprise inside a readable file becomes a
    warning, never a drop.
    """
    filepath = Path(filepath)
    try:
        source_bytes = filepath.read_bytes()
    except OSError as exc:
        raise MasterImportError(
            "That file is not a readable .docx document."
        ) from exc
    try:
        document = Document(str(filepath))
    except (
        PackageNotFoundError,
        zipfile.BadZipFile,
        XMLSyntaxError,
        KeyError,
        ValueError,
    ) as exc:
        raise MasterImportError(
            "That file is not a readable .docx document."
        ) from exc

    builder = _TreeBuilder()
    tracked = _element_has_tracked_changes(document.element.body)
    loss_warning = _detect_content_loss_warning(document.element.body)
    if loss_warning:
        builder.warnings.append(loss_warning)
    if tracked:
        builder.warnings.append(
            "The master carries pending tracked changes; text was imported "
            "as the Accept-All-Changes view (insertions kept, deletions "
            "removed)."
        )

    skipped_empty = 0
    saw_table = False
    pending_title = False  # SECTION number seen; next line may be the title
    # Any one of these means the file carried real SectionFormat structure.
    # Deliberately the same signals the parse itself acts on, so the verdict
    # can never disagree with the tree that was built: a paragraph consumed
    # by the numbering branch below counts only when its numbering label
    # PROMOTED it to real structure (a part or article the tree actually
    # holds), never merely for being numbered.
    saw_spec_marker = False
    source_bindings: list[SourceParagraphBinding] = []
    # Where each semantic element came from, for the appearance-preserving
    # export. Paragraph anchors are appended by add_mapped_paragraph;
    # headings are recorded at the branch that creates them.
    format_anchors: list[FormatAnchor] = []
    # The numbering-definition label grammars, for heading promotion.
    numbering_catalog = _load_numbering_catalog(document)
    # Numbering the way Word resolves it: the paragraph's own ``w:numPr``,
    # else the definition its paragraph STYLE carries. Office masters keep
    # the whole outline on their PRT/ART/PR1-4 styles, so without this every
    # heading exposes only its bare title and the file lands as one blob.
    style_numbering = _load_style_numbering(document)
    default_style_id = _default_paragraph_style_id(document)
    style_kinds = _load_style_kinds(document)
    style_numbering_used = False
    # Promoted PART headings are numbered by order of appearance — the
    # rendered "PART 1/2/3" is a counter this parse never runs.
    promoted_part_count = 0
    # The bare-section header is consulted for the first content line only
    # in a document with no structure at all (see ``first_structure``).
    saw_any_content = False
    # PART 4/5 remaps warn once per out-of-range number, however the content
    # arrives — a heading line, an auto-numbered heading, or a bare "4.01"
    # article — instead of once per line about the same remap.
    warned_out_of_range_parts: set[int] = set()
    # Where the section identity came from, and the rule that it is decided
    # ONCE: the first header found wins, and a SECTION-shaped line after
    # structure has begun is a provision citing a sibling section.
    header_source = ""
    # Body blocks before the first PART/article heading: recorded for the
    # import notes, the lint and the model, never modelled in the tree. The
    # appearance-preserving export carries them through verbatim, in place.
    front_matter_lines: list[str] = []
    front_matter_count = 0

    def _remap_out_of_range_part(part_number: int, line_no: int) -> int:
        if part_number <= 3:
            return part_number
        if part_number not in warned_out_of_range_parts:
            warned_out_of_range_parts.add(part_number)
            builder.warnings.append(
                f"Line {line_no}: SectionFormat carries three parts; "
                f"'PART {part_number}' content was kept under PART 3 — "
                "review placement."
            )
        return 3
    # Where END OF SECTION stopped the parse (1-based), for the trailing-
    # content accounting below. None = the file never said it ended.
    end_of_section_index: int | None = None

    entries = _iter_body_texts(document)
    texts_by_line = {
        line_no: _entry_text(entry)
        for line_no, entry in enumerate(entries, start=1)
    }
    # The first PART or article heading by ANY recognition path. Everything
    # before it is front matter, which is where — and only where — the
    # section identity is read. ``None`` means the file has no structure and
    # keeps the historical posture: every block is content, and a header
    # line may sit anywhere (the first one still wins).
    first_structure = _first_structure_line(
        entries,
        texts_by_line,
        numbering_catalog,
        style_numbering,
        default_style_id,
        style_kinds,
    )

    def _cover_page_title(line_no: int) -> str:
        """The cover-page title beside a ``Section Number:`` line: the
        nearest non-empty line above it, else the one below it, when it
        reads like a title and is still inside the front matter."""
        candidates: list[str] = []
        for span in (
            range(line_no - 1, 0, -1),
            range(line_no + 1, len(entries) + 1),
        ):
            for other in span:
                if first_structure is not None and other >= first_structure:
                    break
                candidate = texts_by_line.get(other, "")
                if candidate:
                    candidates.append(candidate)
                    break
        for candidate in candidates:
            if _title_like(candidate):
                return candidate
        return ""

    for line_no, entry in enumerate(entries, start=1):
        raw_text = entry.text
        docx_paragraph = entry.paragraph
        text = texts_by_line[line_no]
        if not text:
            skipped_empty += 1
            continue
        first_content = not saw_any_content
        saw_any_content = True
        in_front_matter = first_structure is not None and line_no < first_structure
        # The identity is read in the front matter, or anywhere in a file
        # with no structure — and only until something states it.
        header_allowed = not header_source and (
            in_front_matter or first_structure is None
        )
        numbering = (
            _effective_numbering(docx_paragraph, style_numbering, default_style_id)
            if docx_paragraph is not None
            else None
        )
        style_kind = (
            style_kinds.get(_paragraph_style_id(docx_paragraph) or default_style_id)
            if docx_paragraph is not None
            else None
        )

        def add_mapped_paragraph(
            depth: int,
            semantic_text: str,
            *,
            numbered: bool = False,
            manual_label: bool = False,
        ) -> None:
            # `depth` is a raw ``w:numPr`` indent level for a numbered
            # paragraph and an absolute depth for a manual label ("A." is
            # depth 0 by definition), so only the former is rebased.
            place = (
                builder.numbered_paragraph if numbered else builder.paragraph
            )
            paragraph = place(depth, semantic_text, line_no)
            if entry.opaque_blocker:
                binding = bind_opaque_projection(
                    uid=paragraph.uid,
                    body_child_index=entry.body_child_index,
                    element=entry.source_element,
                    source_visible_text=raw_text,
                    baseline_text=semantic_text,
                    blocker=entry.opaque_blocker,
                )
            else:
                binding = bind_source_paragraph(
                    uid=paragraph.uid,
                    body_child_index=entry.body_child_index,
                    element=entry.source_element,
                    source_visible_text=raw_text,
                    baseline_text=semantic_text,
                )
            source_bindings.append(binding)
            if entry.lock_reason:
                paragraph.locked = entry.lock_reason
            # How this provision's label was expressed decides whether the
            # export writes one back. Word renders an auto-numbered label
            # itself; a stripped literal has to be regenerated positionally.
            if numbered:
                label_kind = LABEL_AUTO
            elif manual_label:
                label_kind = LABEL_MANUAL
            else:
                label_kind = LABEL_NONE
            format_anchors.append(
                FormatAnchor(
                    uid=paragraph.uid,
                    origin_index=entry.body_child_index,
                    label_kind=label_kind,
                    locked=entry.lock_reason,
                )
            )

        # --- The section identity: decided in the front matter, once. -----
        if header_allowed and not entry.lock_reason and numbering is None:
            section_match = _SECTION_RE.match(text)
            if section_match:
                saw_spec_marker = True
                header_source = HEADER_SOURCE_LINE
                g1, g2, g3, g4, remainder = section_match.groups()
                builder.section.number = f"{g1} {g2} {g3}" + (
                    f".{g4}" if g4 else ""
                )
                format_anchors.append(
                    FormatAnchor(
                        uid="sec",
                        origin_index=entry.body_child_index,
                        label_kind=LABEL_NONE,
                    )
                )
                if remainder.strip():
                    builder.section.title = remainder.strip()
                    pending_title = False
                else:
                    pending_title = True
                continue
            field_match = _SECTION_NUMBER_FIELD_RE.match(text)
            if field_match and in_front_matter:
                # A cover page's "Section Number: 21 05 00". The line itself
                # stays front matter — emitted verbatim in place, never
                # anchored — so the identity is recorded without a header
                # element for the export to reproduce.
                saw_spec_marker = True
                header_source = HEADER_SOURCE_FRONT_MATTER
                f1, f2, f3, f4 = field_match.groups()
                builder.section.number = f"{f1} {f2} {f3}" + (
                    f".{f4}" if f4 else ""
                )
                title = _cover_page_title(line_no)
                if title:
                    builder.section.title = title
                    builder.warnings.append(
                        f"The section title ('{_clip(title)}') was read from "
                        "the cover page beside its 'Section Number' line — "
                        "check it."
                    )
                front_matter_lines.append(text)
                front_matter_count += 1
                continue
            if in_front_matter or first_content:
                # A keyword-less header — "23 05 48 — COMMON WORK RESULTS FOR
                # HVAC" — is accepted anywhere in the front matter, and on
                # the first content line of a file with no structure. Later
                # in a section the same shape is a provision citing a
                # sibling section, so it never re-arms.
                bare_match = _BARE_SECTION_RE.match(text)
                if bare_match:
                    saw_spec_marker = True
                    header_source = HEADER_SOURCE_LINE
                    b1, b2, b3, b4, bare_title = bare_match.groups()
                    builder.section.number = f"{b1} {b2} {b3}" + (
                        f".{b4}" if b4 else ""
                    )
                    builder.section.title = bare_title.strip()
                    format_anchors.append(
                        FormatAnchor(
                            uid="sec",
                            origin_index=entry.body_child_index,
                            label_kind=LABEL_NONE,
                        )
                    )
                    pending_title = False
                    builder.warnings.append(
                        f"Line {line_no}: no 'SECTION' keyword — the section "
                        "number and title were read from this line."
                    )
                    continue
        if pending_title:
            pending_title = False
            if (
                not entry.lock_reason
                and numbering is None
                and not _structural_heading_kind(
                    entry,
                    text,
                    numbering_catalog,
                    style_numbering,
                    default_style_id,
                    style_kinds,
                )
                and not _SECTION_NUMBER_FIELD_RE.match(text)
            ):
                builder.section.title = text
                format_anchors.append(
                    FormatAnchor(
                        uid=SECTION_TITLE_UID,
                        origin_index=entry.body_child_index,
                        label_kind=LABEL_NONE,
                    )
                )
                continue
        if in_front_matter:
            # Cover page, revision history, table of contents: kept for
            # export exactly as it is, reported, and not modelled.
            front_matter_lines.append(text)
            front_matter_count += 1
            continue

        # --- The section body. --------------------------------------------
        if entry.lock_reason == "table" and not saw_table:
            saw_table = True
            builder.warnings.append(
                "The master contains tables. Each is preserved as one "
                "read-only block — the export carries its cells and "
                "formatting through exactly as they came in. You can move "
                "or delete a table, but not retype it here."
            )
        if entry.lock_reason:
            # A preserved block is content, never structure. Running a table
            # row or a caption through the heading grammars is how a row
            # beginning "A." used to become an article; and since the block
            # is emitted verbatim, its text was never a label to strip.
            #
            # It must also never become a HIERARCHY PARENT. Placing it at
            # depth 0 made it the current depth-0 node, so a following "1."
            # provision nested underneath it — and a table cannot own a
            # subparagraph in Word, in the export, or in any reading of the
            # document. Restoring the stack leaves the block a sibling and
            # sends the next nested provision to the real provision above it.
            # (Caught in review on PR #141, Codex: the renderer skipped those
            # children and the trailing sweep excluded them as anchored, so
            # an untouched export silently DELETED them.)
            stack_before = list(builder.stack)
            offset_before = builder.numbering_offset
            add_mapped_paragraph(0, text)
            builder.stack = stack_before
            builder.numbering_offset = offset_before
            continue

        # Word numbering — the paragraph's own or its style's — is structural
        # metadata, so it must win over text-pattern heuristics. Normalized
        # exports deliberately keep the generated A./1./a./1) marker out of
        # w:t; their semantic text may therefore begin with strings such as
        # "END OF SECTION", "PART 2", "1.2", or "A." without becoming a false
        # heading or manual label on re-import. A numbered line also cannot
        # serve as the pending section title.
        if numbering is not None:
            num_id, ilvl = numbering
            if _direct_numbering(docx_paragraph) is None:
                style_numbering_used = True
            pending_title = False
            # An auto-numbered heading's visible text is just its title
            # ("SUMMARY") — no text pattern can ever reach it. The numbering
            # definition's own label grammar is the structural signal:
            # promote what it says is a PART or an article. When the grammar
            # says nothing but the style is named for what it is (PRT / ART),
            # the convention decides. Everything else stays on the provision
            # path. A promoted heading gets no source binding, exactly like
            # a text-matched heading (headings live in the fixed projection,
            # not the editable body surface).
            kind = _promoted_heading_kind(numbering_catalog, num_id, ilvl)
            if not kind and style_kind in ("part", "article"):
                kind = str(style_kind)
            if kind == "part":
                saw_spec_marker = True
                promoted_part_count += 1
                part_number = _remap_out_of_range_part(
                    promoted_part_count, line_no
                )
                builder.part(
                    part_number,
                    _promoted_part_title(promoted_part_count, part_number, text),
                )
                if builder.current_part is not None:
                    format_anchors.append(
                        FormatAnchor(
                            uid=builder.current_part.uid,
                            origin_index=entry.body_child_index,
                            label_kind=LABEL_AUTO,
                        )
                    )
                continue
            if kind == "article":
                saw_spec_marker = True
                builder.article(
                    builder.current_part.number
                    if builder.current_part is not None
                    else 1,
                    text,
                )
                if builder.current_article is not None:
                    # Word renders this heading's "2.01" itself, so the
                    # export must not write one into the text as well.
                    format_anchors.append(
                        FormatAnchor(
                            uid=builder.current_article.uid,
                            origin_index=entry.body_child_index,
                            label_kind=LABEL_AUTO,
                        )
                    )
                continue
            # The raw indent level: numbered_paragraph places it against
            # its own list, never as an absolute depth.
            add_mapped_paragraph(max(0, ilvl), text, numbered=True)
            continue

        # A CSI-named style (PRT / ART) with no numbering of its own: the
        # name says what the paragraph IS. A typed label in the text still
        # wins below, because the reader sees it and the export must write
        # it back.
        if style_kind == "part" and not _PART_RE.match(text):
            saw_spec_marker = True
            promoted_part_count += 1
            part_number = _remap_out_of_range_part(promoted_part_count, line_no)
            builder.part(
                part_number,
                _promoted_part_title(promoted_part_count, part_number, text),
            )
            if builder.current_part is not None:
                format_anchors.append(
                    FormatAnchor(
                        uid=builder.current_part.uid,
                        origin_index=entry.body_child_index,
                        label_kind=LABEL_AUTO,
                    )
                )
            continue
        if style_kind == "article" and not _ARTICLE_RE.match(text):
            saw_spec_marker = True
            builder.article(
                builder.current_part.number
                if builder.current_part is not None
                else 1,
                text,
            )
            if builder.current_article is not None:
                format_anchors.append(
                    FormatAnchor(
                        uid=builder.current_article.uid,
                        origin_index=entry.body_child_index,
                        label_kind=LABEL_AUTO,
                    )
                )
            continue

        # The normalized renderer emits this exact line only to show that a
        # PART has no articles. It is presentation, not a semantic provision;
        # recognizing it here prevents an export/re-import round trip from
        # manufacturing an ``IMPORTED CONTENT`` article in an empty part.
        if (
            text.casefold() == "(not used.)"
            and builder.current_part is not None
            and builder.current_article is None
        ):
            continue

        if _END_RE.match(text):
            end_of_section_index = line_no
            break

        part_match = _PART_RE.match(text)
        if part_match:
            saw_spec_marker = True
            requested_part = int(part_match.group(1))
            part_number = _remap_out_of_range_part(requested_part, line_no)
            builder.part(
                part_number, text if part_number == requested_part else ""
            )
            if builder.current_part is not None:
                format_anchors.append(
                    FormatAnchor(
                        uid=builder.current_part.uid,
                        origin_index=entry.body_child_index,
                        label_kind=LABEL_MANUAL,
                    )
                )
            continue

        article_match = _ARTICLE_RE.match(text)
        if article_match:
            saw_spec_marker = True
            part_digit, _article_no, title = article_match.groups()
            builder.article(
                _remap_out_of_range_part(int(part_digit), line_no),
                title.strip(),
            )
            if builder.current_article is not None:
                format_anchors.append(
                    FormatAnchor(
                        uid=builder.current_article.uid,
                        origin_index=entry.body_child_index,
                        label_kind=LABEL_MANUAL,
                    )
                )
            continue

        # Manual paragraph labels, most-specific first (uppercase before
        # digit before lowercase ordering is inherent to the regexes).
        matched_level = None
        for depth, pattern in enumerate(_LEVEL_RES):
            match = pattern.match(text)
            if match:
                matched_level = (depth, match.group(2).strip())
                break
        if matched_level is not None:
            add_mapped_paragraph(
                matched_level[0], matched_level[1], manual_label=True
            )
            continue

        if isinstance(style_kind, int):
            # A PRn-styled provision with neither Word numbering nor a typed
            # label: the style says how deep it sits.
            add_mapped_paragraph(style_kind, text)
            continue

        # Unlabeled content: keep as a level-0 paragraph (never drop).
        add_mapped_paragraph(0, text)

    # END OF SECTION still stops the parse — the app's own normalized export
    # puts its assumptions/open-items schedules after that line, and
    # re-importing them as content would corrupt the round trip — but
    # stopping must never be SILENT: dropped content with no warning
    # violates this module's whole philosophy, and for a combined
    # multi-section master it quietly discarded every section but the
    # first. The one suppression is the app's own trailing schedules,
    # recognized by the exact heading the exporter writes first.
    if end_of_section_index is not None:
        trailing = [
            " ".join(item.text.split())
            for item in entries[end_of_section_index:]
        ]
        trailing = [item for item in trailing if item]
        if trailing and trailing[0] != "ASSUMPTIONS SCHEDULE":
            next_section = next(
                (item for item in trailing if _SECTION_RE.match(item)), None
            )

            if next_section is not None:
                builder.warnings.append(
                    "This file contains more than one SECTION (next: "
                    f"'{_clip(next_section)}'); only the first was imported "
                    f"— split the file to import another. {len(trailing)} "
                    "block(s) after 'END OF SECTION' were not imported."
                )
            else:
                builder.warnings.append(
                    f"{len(trailing)} block(s) after 'END OF SECTION' were "
                    f"not imported (beginning '{_clip(trailing[0])}'). "
                    "Build-a-Spec authors one section at a time."
                )

    if front_matter_count:
        builder.warnings.append(
            f"{front_matter_count} block(s) before the first PART heading — a "
            "cover page, revision history or table of contents — are kept "
            "for export exactly as they are and are not part of the "
            f"editable section (beginning '{_clip(front_matter_lines[0])}')."
        )

    # A structured section whose body never stated its own identity: the
    # page header/footer conventionally carries it. Consulted last, and
    # disclosed, so a stale template footer is checked once rather than
    # trusted silently. A file with no structure gets no invented header.
    chrome_lines = _header_footer_text(document)
    if first_structure is not None and (
        not builder.section.number or not builder.section.title
    ):
        chrome_number, chrome_title = _identity_from_chrome(chrome_lines)
        if not builder.section.number and chrome_number:
            builder.section.number = chrome_number
            header_source = HEADER_SOURCE_CHROME
            saw_spec_marker = True
            builder.warnings.append(
                f"The section number ({chrome_number}) was read from the "
                "page header/footer — the body carries no SECTION line. "
                "Check it."
            )
        if builder.section.number and not builder.section.title and chrome_title:
            builder.section.title = chrome_title
            if not header_source:
                header_source = HEADER_SOURCE_CHROME
            builder.warnings.append(
                f"The section title ('{_clip(chrome_title)}') was read from "
                "the page header/footer. Check it."
            )

    # Append a FINDABLE reference to every warning that is about a specific
    # element: the display ref the panel and export schedules use plus the
    # stable id. Runs before UNSTRUCTURED_IMPORT_WARNING's insert(0), which
    # would shift the recorded indexes. Refs derive from final positions, so
    # this must be a post-pass — mid-build numbering is not settled yet.
    if builder.warning_uids:
        refs = {
            p.uid: ref
            for _part, _article, p, _depth, ref in iter_paragraphs(
                builder.section
            )
        }
        for part in builder.section.parts:
            for article_index, article in enumerate(part.articles, start=1):
                refs.setdefault(
                    article.uid, f"{part.number}.{article_index}"
                )
        for warning_index, uid in builder.warning_uids.items():
            ref = refs.get(uid)
            where = f" (at {ref}, id {uid})" if ref else f" (id {uid})"
            builder.warnings[warning_index] += where

    if builder.imported_count == 0 and builder.section.is_empty():
        raise MasterImportError(
            "No importable content found — the document has no recognizable "
            "SectionFormat structure and no body text."
        )

    # Lead with the verdict so it reads before the per-line parse notes it
    # explains. Those notes are kept: they say *where* the content landed,
    # which stays useful once the user starts working with it.
    if not saw_spec_marker:
        builder.warnings.insert(0, UNSTRUCTURED_IMPORT_WARNING)

    try:
        source_map = build_source_body_map(
            source_bytes=source_bytes,
            document=document,
            section=builder.section,
            bindings=source_bindings,
        )
    except (TypeError, ValueError, zipfile.BadZipFile) as exc:
        raise MasterImportError(
            "The document body could not be mapped safely for editing."
        ) from exc

    format_map = build_format_map(
        source_bytes=source_bytes,
        # Body CHILDREN, not semantic entries: origin indexes are positions
        # in the source body, and blank paragraphs occupy one each.
        body_child_count=sum(
            1
            for child in document.element.body.iterchildren()
            if isinstance(child.tag, str)
        ),
        anchors=format_anchors,
        header_footer_text=chrome_lines,
        front_matter_text=tuple(front_matter_lines),
        header_source=header_source,
        section_number=builder.section.number,
        section_title=builder.section.title,
    )

    return ImportResult(
        section=builder.section,
        format_map=format_map,
        warnings=builder.warnings,
        tracked_changes_detected=tracked,
        imported_block_count=builder.imported_count,
        skipped_empty_count=skipped_empty,
        source_map=source_map,
        spec_shape_detected=saw_spec_marker,
        front_matter=tuple(front_matter_lines),
        header_source=header_source,
        style_numbering_resolved=style_numbering_used,
    )
