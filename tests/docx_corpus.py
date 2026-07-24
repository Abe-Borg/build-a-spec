"""Deterministic, sanitized DOCX fidelity corpus recipes.

Bundled recipes are intentionally synthetic. Producer-like recipe cases
exercise package shapes and metadata associated with common workflows, but
never claim to have been saved by Microsoft Word, LibreOffice, a binary
``.doc`` converter, or a real consultant system. The manifest may also name
immutable, independently sanitized producer fixtures with explicit provenance.
See ``docs/DOCX_FIDELITY_CORPUS.md``.
"""
from __future__ import annotations

import argparse
import hashlib
import importlib.metadata
import io
import json
import platform
import re
import struct
import tempfile
import zipfile
import zlib
from dataclasses import dataclass
from pathlib import Path, PurePosixPath, PureWindowsPath
from typing import Any, Callable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from tests.docx_fidelity_helpers import (
    TARGET_PARA_ID,
    _add_relationship,
    _append_page_field,
    _set_para_id,
    document_xml,
    make_fidelity_master,
    make_numbered_island_master,
    rewrite_zip_members,
)


_MANIFEST_DIR = Path(__file__).parent / "fixtures" / "docx_corpus"
_MANIFEST_PATH = _MANIFEST_DIR / "manifest.json"
_FIXED_ZIP_TIME = (2026, 1, 1, 0, 0, 0)
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_EP_NS = (
    "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
)
_CP_CORE_NS = (
    "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
)
_DC_NS = "http://purl.org/dc/elements/1.1/"
_COMMENTS_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
_FOOTNOTES_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
)
_ENDNOTES_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"
)
_COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
_FOOTNOTES_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
)
_ENDNOTES_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"
)
_NUMBERING_REL_SUFFIX = "/numbering"
_RELOCATED_NUMBERING_PART = "word/corpus/numbering-relocated.xml"
_ARCHIVE_COMMENT = b"Build-a-Spec sanitized synthetic corpus ZIP comment"
_PRIVATE_EXTRA = struct.pack("<HH4s", 0xCAFE, 4, b"BASP")
_SHA256_RE = re.compile(r"[0-9a-f]{64}")
_CASE_ID_RE = re.compile(r"[A-Za-z0-9][A-Za-z0-9_.-]*")
_WINDOWS_RESERVED_NAMES = {
    "con",
    "prn",
    "aux",
    "nul",
    *(f"com{index}" for index in range(1, 10)),
    *(f"lpt{index}" for index in range(1, 10)),
}
_EXTERNAL_PRODUCER_FIELDS = (
    "product",
    "version",
    "platform",
    "production_method",
)
_EXTERNAL_SANITIZATION_FIELDS = (
    "tool",
    "procedure",
    "privacy_review",
)


@dataclass(frozen=True)
class CorpusCase:
    case_id: str
    filename: str
    recipe: str | None
    fixture: str | None
    expected_sha256: str | None
    description: str
    categories: tuple[str, ...]
    provenance: dict[str, Any]
    expectations: dict[str, Any]


def load_manifest() -> dict[str, Any]:
    return json.loads(_MANIFEST_PATH.read_text(encoding="utf-8"))


def _is_nonempty_string(value: Any) -> bool:
    return isinstance(value, str) and bool(value.strip())


def _is_portable_windows_component(value: str) -> bool:
    """Reject Win32 device aliases, trimmed suffixes, and illegal characters."""
    if not value or value.endswith((" ", ".")):
        return False
    if any(ord(character) < 32 or character in '<>:"/\\|?*' for character in value):
        return False
    device_name = value.split(".", 1)[0].casefold()
    return device_name not in _WINDOWS_RESERVED_NAMES


def _validate_case_id(case_id: str) -> None:
    """Require one portable path-neutral identifier for recipe workspaces."""
    if (
        not _CASE_ID_RE.fullmatch(case_id)
        or not _is_portable_windows_component(case_id)
    ):
        raise ValueError(
            f"Corpus case id {case_id!r} must be a portable identifier"
        )


def _validate_output_filename(case_id: str, filename: str) -> None:
    """Require a portable leaf name before materializing into an output root."""
    posix = PurePosixPath(filename)
    windows = PureWindowsPath(filename)
    if (
        not _is_nonempty_string(filename)
        or posix.is_absolute()
        or windows.is_absolute()
        or posix.name != filename
        or windows.name != filename
        or posix.suffix.casefold() != ".docx"
        or not _is_portable_windows_component(filename)
    ):
        raise ValueError(
            f"Corpus case {case_id!r} needs a leaf-only .docx filename"
        )


def _validate_fixture_reference(case_id: str, fixture: str) -> None:
    posix = PurePosixPath(fixture)
    windows = PureWindowsPath(fixture)
    if (
        not _is_nonempty_string(fixture)
        or posix.is_absolute()
        or windows.is_absolute()
        or ".." in posix.parts
        or ".." in windows.parts
        or posix.suffix.casefold() != ".docx"
        or any(
            not _is_portable_windows_component(part)
            for part in windows.parts
        )
    ):
        raise ValueError(
            f"Corpus case {case_id!r} has an unsafe fixture reference"
        )


def _validate_external_provenance(case: CorpusCase) -> None:
    actual = case.provenance.get("actual_producer")
    sanitization = case.provenance.get("sanitization")
    if not isinstance(actual, dict) or any(
        not _is_nonempty_string(actual.get(field))
        for field in _EXTERNAL_PRODUCER_FIELDS
    ):
        raise ValueError(
            f"External corpus case {case.case_id!r} needs structured "
            "producer product/version/platform/production_method provenance"
        )
    if not isinstance(sanitization, dict) or any(
        not _is_nonempty_string(sanitization.get(field))
        for field in _EXTERNAL_SANITIZATION_FIELDS
    ):
        raise ValueError(
            f"External corpus case {case.case_id!r} needs structured "
            "sanitization tool/procedure/privacy_review provenance"
        )
    modified_parts = sanitization.get("modified_parts")
    if not isinstance(modified_parts, list) or any(
        not _is_nonempty_string(part) for part in modified_parts
    ):
        raise ValueError(
            f"External corpus case {case.case_id!r} needs a modified_parts list"
        )


def _validate_case(case: CorpusCase) -> None:
    _validate_case_id(case.case_id)
    _validate_output_filename(case.case_id, case.filename)
    if not _is_nonempty_string(case.description) or not case.categories:
        raise ValueError(
            f"Corpus case {case.case_id!r} needs description and categories"
        )
    if case.provenance.get("sanitized") is not True or not _is_nonempty_string(
        case.provenance.get("statement")
    ):
        raise ValueError(
            f"Corpus case {case.case_id!r} needs explicit sanitization provenance"
        )

    kind = case.provenance.get("kind")
    if kind == "synthetic":
        if case.provenance.get("actual_producer") is not None:
            raise ValueError(
                f"Synthetic corpus case {case.case_id!r} claims an actual producer"
            )
        if not case.recipe or case.fixture is not None:
            raise ValueError(
                f"Synthetic corpus case {case.case_id!r} needs exactly one recipe"
            )
    elif kind == "external_sanitized":
        if case.recipe is not None or not case.fixture:
            raise ValueError(
                f"External corpus case {case.case_id!r} needs exactly one fixture"
            )
        _validate_fixture_reference(case.case_id, case.fixture)
        if case.expected_sha256 is None or not _SHA256_RE.fullmatch(
            case.expected_sha256
        ):
            raise ValueError(
                f"External corpus case {case.case_id!r} needs a lowercase SHA-256"
            )
        _validate_external_provenance(case)
    else:
        raise ValueError(
            f"Corpus case {case.case_id!r} has unsupported provenance kind {kind!r}"
        )

    expectations = case.expectations
    if expectations.get("importable") is not True:
        raise ValueError(f"Corpus case {case.case_id!r} must be importable")
    if expectations.get("exact_noop") is not True:
        raise ValueError(f"Corpus case {case.case_id!r} must support exact no-op")
    if expectations.get("source_mode") not in {"ready", "pass_through_only"}:
        raise ValueError(f"Corpus case {case.case_id!r} has invalid source_mode")
    blockers = expectations.get("mutation_blockers")
    if not isinstance(blockers, list) or any(
        not _is_nonempty_string(blocker) for blocker in blockers
    ):
        raise ValueError(
            f"Corpus case {case.case_id!r} needs a mutation_blockers list"
        )


def corpus_cases() -> tuple[CorpusCase, ...]:
    manifest = load_manifest()
    cases: list[CorpusCase] = []
    for item in manifest["cases"]:
        case = CorpusCase(
            case_id=str(item["id"]),
            filename=str(item["filename"]),
            recipe=(str(item["recipe"]) if item.get("recipe") else None),
            fixture=(str(item["fixture"]) if item.get("fixture") else None),
            expected_sha256=(
                str(item["sha256"]) if item.get("sha256") else None
            ),
            description=str(item["description"]),
            categories=tuple(str(value) for value in item["categories"]),
            provenance=dict(item["provenance"]),
            expectations=dict(item["expectations"]),
        )
        _validate_case(case)
        cases.append(case)
    if len({case.case_id.casefold() for case in cases}) != len(cases):
        raise ValueError("DOCX corpus case ids must be unique")
    if len({case.filename.casefold() for case in cases}) != len(cases):
        raise ValueError("DOCX corpus output filenames must be unique")
    return tuple(cases)


def _clone_zip_info(source: zipfile.ZipInfo) -> zipfile.ZipInfo:
    info = zipfile.ZipInfo(source.filename, date_time=_FIXED_ZIP_TIME)
    info.compress_type = source.compress_type
    info.comment = source.comment
    info.extra = source.extra
    info.create_system = 3
    info.external_attr = (
        (0o40755 if source.is_dir() else 0o100600) << 16
    )
    return info


def _canonicalize_docx(payload: bytes) -> bytes:
    """Return byte-deterministic ZIP metadata while preserving member order."""
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(payload), "r") as source:
        with zipfile.ZipFile(output, "w") as target:
            for source_info in source.infolist():
                info = _clone_zip_info(source_info)
                data = source.read(source_info)
                kwargs: dict[str, Any] = {"compress_type": info.compress_type}
                if info.compress_type == zipfile.ZIP_DEFLATED:
                    kwargs["compresslevel"] = 9
                target.writestr(info, data, **kwargs)
            target.comment = source.comment
    return output.getvalue()


def _replace_xml_member(payload: bytes, name: str, root) -> bytes:
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        archive_comment = archive.comment
    rewritten = rewrite_zip_members(
        payload,
        replacements={
            name: etree.tostring(
                root,
                encoding="UTF-8",
                xml_declaration=True,
                standalone=True,
            )
        },
    )
    if not archive_comment:
        return rewritten
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(rewritten), "r") as source:
        with zipfile.ZipFile(output, "w") as target:
            for source_info in source.infolist():
                target.writestr(source_info, source.read(source_info))
            target.comment = archive_comment
    return output.getvalue()


def _set_application(payload: bytes, application: str, version: str) -> bytes:
    name = "docProps/app.xml"
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        root = etree.fromstring(archive.read(name))
    app = root.find(f"{{{_EP_NS}}}Application")
    if app is None:
        app = etree.SubElement(root, f"{{{_EP_NS}}}Application")
    app.text = application
    app_version = root.find(f"{{{_EP_NS}}}AppVersion")
    if app_version is None:
        app_version = etree.SubElement(root, f"{{{_EP_NS}}}AppVersion")
    app_version.text = version
    return _replace_xml_member(payload, name, root)


def _sanitize_core_properties(payload: bytes, case_id: str) -> bytes:
    name = "docProps/core.xml"
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        root = etree.fromstring(archive.read(name))
    for tag, text in (
        (f"{{{_DC_NS}}}creator", "Build-a-Spec Synthetic Corpus"),
        (f"{{{_CP_CORE_NS}}}lastModifiedBy", "Build-a-Spec Synthetic Corpus"),
        (f"{{{_DC_NS}}}title", f"Synthetic DOCX corpus: {case_id}"),
        (f"{{{_CP_CORE_NS}}}keywords", "sanitized synthetic fidelity fixture"),
    ):
        node = root.find(tag)
        if node is None:
            node = etree.SubElement(root, tag)
        node.text = text
    return _replace_xml_member(payload, name, root)


def _base_fidelity(workspace: Path, filename: str) -> bytes:
    return make_fidelity_master(workspace, filename=filename)


def _word_like_rich(workspace: Path) -> bytes:
    payload = _base_fidelity(workspace, "word-like-source.docx")
    return _set_application(payload, "Microsoft Office Word", "16.0000")


def _libreoffice_like_metadata(workspace: Path) -> bytes:
    payload = _base_fidelity(workspace, "libreoffice-like-source.docx")
    return _set_application(
        payload,
        "LibreOffice/24.2.0.3$Windows_X86_64 LibreOffice_project/synthetic",
        "24.2",
    )


def _older_conversion_like(workspace: Path) -> bytes:
    payload = _base_fidelity(workspace, "older-conversion-source.docx")
    payload = _set_application(payload, "Microsoft Office Word", "12.0000")
    name = "word/settings.xml"
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        settings = etree.fromstring(archive.read(name))
    compat = settings.find(f"{{{_W_NS}}}compat")
    if compat is None:
        compat = etree.SubElement(settings, f"{{{_W_NS}}}compat")
    compatibility_mode = None
    for setting in compat.findall(f"{{{_W_NS}}}compatSetting"):
        if setting.get(f"{{{_W_NS}}}name") == "compatibilityMode":
            compatibility_mode = setting
            break
    if compatibility_mode is None:
        compatibility_mode = etree.SubElement(
            compat, f"{{{_W_NS}}}compatSetting"
        )
    compatibility_mode.set(f"{{{_W_NS}}}name", "compatibilityMode")
    compatibility_mode.set(
        f"{{{_W_NS}}}uri",
        "http://schemas.microsoft.com/office/word",
    )
    compatibility_mode.set(f"{{{_W_NS}}}val", "12")
    return _replace_xml_member(payload, name, settings)


def _consultant_template(workspace: Path) -> bytes:
    payload = _base_fidelity(workspace, "consultant-template-source.docx")
    return _set_application(
        payload,
        "Build-a-Spec Synthetic Consultant Template",
        "1.0000",
    )


def _mixed_section_layout(workspace: Path) -> bytes:
    document = Document()
    document.settings.odd_and_even_pages_header_footer = True
    first = document.sections[0]
    first.top_margin = Inches(0.72)
    first.bottom_margin = Inches(0.81)
    first.left_margin = Inches(0.88)
    first.right_margin = Inches(0.94)
    first.header_distance = Inches(0.25)
    first.footer_distance = Inches(0.28)
    first.different_first_page_header_footer = True
    first.header.paragraphs[0].text = "PORTRAIT DEFAULT HEADER - KEEP"
    first.first_page_header.paragraphs[0].text = "PORTRAIT FIRST HEADER - KEEP"
    first.even_page_header.paragraphs[0].text = "PORTRAIT EVEN HEADER - KEEP"
    first_footer = first.footer.paragraphs[0]
    first_footer.add_run("PORTRAIT PAGE ")
    _append_page_field(first_footer)

    document.add_paragraph("SECTION 21 13 17")
    document.add_paragraph("MIXED SECTION LAYOUT FIXTURE")
    document.add_paragraph("PART 1 - GENERAL")
    document.add_paragraph("1.1 SUMMARY")
    target = document.add_paragraph("A. Preserve mixed section geometry.")
    _set_para_id(target, TARGET_PARA_ID)
    document.add_paragraph("END OF SECTION 21 13 17")

    second = document.add_section(WD_SECTION.NEW_PAGE)
    second.orientation = WD_ORIENT.LANDSCAPE
    second.page_width, second.page_height = first.page_height, first.page_width
    second.top_margin = Inches(0.55)
    second.bottom_margin = Inches(0.63)
    second.left_margin = Inches(1.15)
    second.right_margin = Inches(0.61)
    second.header_distance = Inches(0.4)
    second.footer_distance = Inches(0.44)
    second.header.is_linked_to_previous = False
    second.footer.is_linked_to_previous = False
    second.header.paragraphs[0].text = "LANDSCAPE APPENDIX HEADER - KEEP"
    second_footer = second.footer.paragraphs[0]
    second_footer.add_run("LANDSCAPE PAGE ")
    _append_page_field(second_footer)
    document.add_paragraph("OPAQUE LANDSCAPE APPENDIX - KEEP")
    table = document.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.text = f"R{row_index + 1}C{column_index + 1}"

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _custom_numbering(workspace: Path) -> bytes:
    return make_numbered_island_master(
        workspace,
        filename="custom-numbering-source.docx",
    )


def _relocated_numbering_opc(workspace: Path) -> bytes:
    payload = _custom_numbering(workspace)
    rels_name = "word/_rels/document.xml.rels"
    content_types_name = "[Content_Types].xml"
    source_part = "word/numbering.xml"
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        rels = etree.fromstring(archive.read(rels_name))
        content_types = etree.fromstring(archive.read(content_types_name))
        numbering = archive.read(source_part)
    relationships = [
        item
        for item in rels.findall(f"{{{_REL_NS}}}Relationship")
        if item.get("Type", "").endswith(_NUMBERING_REL_SUFFIX)
    ]
    if len(relationships) != 1:
        raise AssertionError("numbering relocation needs one relationship")
    relationships[0].set("Target", "corpus/numbering-relocated.xml")
    overrides = [
        item
        for item in content_types.findall(f"{{{_CT_NS}}}Override")
        if item.get("PartName", "").lstrip("/") == source_part
    ]
    if len(overrides) != 1:
        raise AssertionError("numbering relocation needs one content type")
    overrides[0].set("PartName", f"/{_RELOCATED_NUMBERING_PART}")
    return rewrite_zip_members(
        payload,
        replacements={
            rels_name: etree.tostring(
                rels, encoding="UTF-8", xml_declaration=True, standalone=True
            ),
            content_types_name: etree.tostring(
                content_types,
                encoding="UTF-8",
                xml_declaration=True,
                standalone=True,
            ),
        },
        additions=[(_RELOCATED_NUMBERING_PART, numbering)],
        omit={source_part},
    )


def _unusual_utf8_declaration(workspace: Path) -> bytes:
    payload = _base_fidelity(workspace, "unusual-utf8-source.docx")
    original = document_xml(payload)
    declaration_end = original.find(b"?>")
    if declaration_end < 0:
        raise AssertionError("document.xml has no XML declaration")
    root_bytes = original[declaration_end + 2 :].lstrip(b"\r\n")
    unusual = (
        b"\xef\xbb\xbf<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\r\n"
        b"<!-- Build-a-Spec sanitized lexical corpus marker -->\r\n"
        b"<?build-a-spec-corpus unusual-utf8?>\r\n"
        + root_bytes
    )
    return rewrite_zip_members(
        payload,
        replacements={"word/document.xml": unusual},
    )


def _utf16_pass_through(workspace: Path) -> bytes:
    payload = _base_fidelity(workspace, "utf16-source.docx")
    root = etree.fromstring(document_xml(payload))
    utf16 = etree.tostring(
        root,
        encoding="UTF-16",
        xml_declaration=True,
        standalone=True,
    )
    return rewrite_zip_members(
        payload,
        replacements={"word/document.xml": utf16},
    )


def _png_chunk(kind: bytes, data: bytes) -> bytes:
    checksum = zlib.crc32(kind)
    checksum = zlib.crc32(data, checksum) & 0xFFFFFFFF
    return (
        struct.pack(">I", len(data))
        + kind
        + data
        + struct.pack(">I", checksum)
    )


def _large_deterministic_png(width: int = 768, height: int = 768) -> bytes:
    pixel_bytes = hashlib.shake_256(
        b"Build-a-Spec sanitized large media corpus"
    ).digest(width * height * 4)
    rows = bytearray()
    stride = width * 4
    for row in range(height):
        rows.append(0)
        start = row * stride
        rows.extend(pixel_bytes[start : start + stride])
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 6, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + _png_chunk(b"IHDR", ihdr)
        + _png_chunk(b"IDAT", zlib.compress(bytes(rows), 0))
        + _png_chunk(b"IEND", b"")
    )


def _large_media(workspace: Path) -> bytes:
    payload = _base_fidelity(workspace, "large-media-source.docx")
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        media = sorted(
            name for name in archive.namelist() if name.startswith("word/media/")
        )
    if not media:
        raise AssertionError("large-media recipe needs a referenced image")
    return rewrite_zip_members(
        payload,
        replacements={media[0]: _large_deterministic_png()},
    )


def _zip_comment_and_extra(workspace: Path) -> bytes:
    payload = _base_fidelity(workspace, "zip-envelope-source.docx")
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(payload), "r") as source:
        with zipfile.ZipFile(output, "w") as target:
            for source_info in source.infolist():
                info = zipfile.ZipInfo(source_info.filename, source_info.date_time)
                info.compress_type = source_info.compress_type
                info.external_attr = source_info.external_attr
                info.create_system = source_info.create_system
                info.comment = source_info.comment
                info.extra = source_info.extra
                if source_info.filename == "word/document.xml":
                    info.extra += _PRIVATE_EXTRA
                target.writestr(info, source.read(source_info))
            target.comment = _ARCHIVE_COMMENT
    return output.getvalue()


def _manual_note_part(
    *,
    collection_name: str,
    note_name: str,
    reference_name: str,
    text: str,
):
    collection = etree.Element(
        f"{{{_W_NS}}}{collection_name}",
        nsmap={"w": _W_NS},
    )
    for note_id, note_type, marker_name in (
        ("-1", "separator", "separator"),
        ("0", "continuationSeparator", "continuationSeparator"),
    ):
        note = etree.SubElement(collection, f"{{{_W_NS}}}{note_name}")
        note.set(f"{{{_W_NS}}}id", note_id)
        note.set(f"{{{_W_NS}}}type", note_type)
        paragraph = etree.SubElement(note, f"{{{_W_NS}}}p")
        run = etree.SubElement(paragraph, f"{{{_W_NS}}}r")
        etree.SubElement(run, f"{{{_W_NS}}}{marker_name}")

    note = etree.SubElement(collection, f"{{{_W_NS}}}{note_name}")
    note.set(f"{{{_W_NS}}}id", "1")
    paragraph = etree.SubElement(note, f"{{{_W_NS}}}p")
    reference_run = etree.SubElement(paragraph, f"{{{_W_NS}}}r")
    etree.SubElement(reference_run, f"{{{_W_NS}}}{reference_name}")
    text_run = etree.SubElement(paragraph, f"{{{_W_NS}}}r")
    text_node = etree.SubElement(text_run, f"{{{_W_NS}}}t")
    text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = f" {text}"
    return collection


def _manual_notes_ooxml(workspace: Path) -> bytes:
    payload = _base_fidelity(workspace, "manual-notes-source.docx")
    document_name = "word/document.xml"
    rels_name = "word/_rels/document.xml.rels"
    content_types_name = "[Content_Types].xml"
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        document_root = etree.fromstring(archive.read(document_name))
        rels = etree.fromstring(archive.read(rels_name))
        content_types = etree.fromstring(archive.read(content_types_name))

    paragraphs = document_root.xpath(
        ".//w:p[.//w:t[contains(., 'OPAQUE CLIENT APPENDIX')]]",
        namespaces={"w": _W_NS},
    )
    if len(paragraphs) != 1:
        raise AssertionError("manual note target was not unique")
    paragraph = paragraphs[0]
    parent = paragraph.getparent()
    if parent is None:
        raise AssertionError("manual note target has no parent")
    anchor_paragraph = etree.Element(f"{{{_W_NS}}}p")
    parent.insert(parent.index(paragraph) + 1, anchor_paragraph)
    for label, reference_name in (
        ("Footnote anchor", "footnoteReference"),
        (". Endnote anchor", "endnoteReference"),
    ):
        label_run = etree.SubElement(anchor_paragraph, f"{{{_W_NS}}}r")
        label_text = etree.SubElement(label_run, f"{{{_W_NS}}}t")
        label_text.set(
            "{http://www.w3.org/XML/1998/namespace}space",
            "preserve",
        )
        label_text.text = label
        reference_run = etree.SubElement(anchor_paragraph, f"{{{_W_NS}}}r")
        reference = etree.SubElement(
            reference_run,
            f"{{{_W_NS}}}{reference_name}",
        )
        reference.set(f"{{{_W_NS}}}id", "1")
    final_run = etree.SubElement(anchor_paragraph, f"{{{_W_NS}}}r")
    final_text = etree.SubElement(final_run, f"{{{_W_NS}}}t")
    final_text.text = "."

    section_properties = document_root.find(
        f"{{{_W_NS}}}body/{{{_W_NS}}}sectPr"
    )
    if section_properties is None:
        raise AssertionError("manual notes fixture has no section properties")
    page_margins = section_properties.find(f"{{{_W_NS}}}pgMar")
    if page_margins is None:
        raise AssertionError("manual notes fixture has no page margins")
    page_margins.set(f"{{{_W_NS}}}bottom", "2880")

    _add_relationship(rels, "rIdCorpusFootnotes", _FOOTNOTES_REL, "footnotes.xml")
    _add_relationship(rels, "rIdCorpusEndnotes", _ENDNOTES_REL, "endnotes.xml")
    for part_name, content_type in (
        ("/word/footnotes.xml", _FOOTNOTES_CONTENT_TYPE),
        ("/word/endnotes.xml", _ENDNOTES_CONTENT_TYPE),
    ):
        override = etree.SubElement(content_types, f"{{{_CT_NS}}}Override")
        override.set("PartName", part_name)
        override.set("ContentType", content_type)

    footnotes = _manual_note_part(
        collection_name="footnotes",
        note_name="footnote",
        reference_name="footnoteRef",
        text="Sanitized synthetic footnote content.",
    )
    endnotes = _manual_note_part(
        collection_name="endnotes",
        note_name="endnote",
        reference_name="endnoteRef",
        text="Sanitized synthetic endnote content.",
    )

    def serialize(root) -> bytes:
        return etree.tostring(
            root,
            encoding="UTF-8",
            xml_declaration=True,
            standalone=True,
        )

    return rewrite_zip_members(
        payload,
        replacements={
            document_name: serialize(document_root),
            rels_name: serialize(rels),
            content_types_name: serialize(content_types),
        },
        additions=[
            ("word/footnotes.xml", serialize(footnotes)),
            ("word/endnotes.xml", serialize(endnotes)),
        ],
    )


def _manual_comments_ooxml(workspace: Path) -> bytes:
    payload = _base_fidelity(workspace, "manual-comments-source.docx")
    document_name = "word/document.xml"
    rels_name = "word/_rels/document.xml.rels"
    content_types_name = "[Content_Types].xml"
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        document_root = etree.fromstring(archive.read(document_name))
        rels = etree.fromstring(archive.read(rels_name))
        content_types = etree.fromstring(archive.read(content_types_name))

    paragraphs = document_root.xpath(
        ".//w:p[.//w:t[contains(., 'OPAQUE CLIENT APPENDIX')]]",
        namespaces={"w": _W_NS},
    )
    if len(paragraphs) != 1:
        raise AssertionError("manual comment target was not unique")
    paragraph = paragraphs[0]
    run = paragraph.find(f"{{{_W_NS}}}r")
    if run is None:
        raise AssertionError("manual comment target has no run")
    run_index = paragraph.index(run)
    range_start = etree.Element(f"{{{_W_NS}}}commentRangeStart")
    range_start.set(f"{{{_W_NS}}}id", "0")
    range_end = etree.Element(f"{{{_W_NS}}}commentRangeEnd")
    range_end.set(f"{{{_W_NS}}}id", "0")
    reference_run = etree.Element(f"{{{_W_NS}}}r")
    reference = etree.SubElement(reference_run, f"{{{_W_NS}}}commentReference")
    reference.set(f"{{{_W_NS}}}id", "0")
    paragraph.insert(run_index, range_start)
    paragraph.insert(run_index + 2, range_end)
    paragraph.insert(run_index + 3, reference_run)

    _add_relationship(rels, "rIdCorpusComments", _COMMENTS_REL, "comments.xml")
    override = etree.SubElement(content_types, f"{{{_CT_NS}}}Override")
    override.set("PartName", "/word/comments.xml")
    override.set("ContentType", _COMMENTS_CONTENT_TYPE)

    comments = etree.Element(
        f"{{{_W_NS}}}comments",
        nsmap={"w": _W_NS},
    )
    comment = etree.SubElement(comments, f"{{{_W_NS}}}comment")
    comment.set(f"{{{_W_NS}}}id", "0")
    comment.set(f"{{{_W_NS}}}author", "Build-a-Spec Synthetic Corpus")
    comment.set(f"{{{_W_NS}}}date", "2026-01-01T00:00:00Z")
    comment_paragraph = etree.SubElement(comment, f"{{{_W_NS}}}p")
    comment_run = etree.SubElement(comment_paragraph, f"{{{_W_NS}}}r")
    comment_text = etree.SubElement(comment_run, f"{{{_W_NS}}}t")
    comment_text.text = "Sanitized synthetic review marker."

    serialize = lambda root: etree.tostring(  # noqa: E731
        root,
        encoding="UTF-8",
        xml_declaration=True,
        standalone=True,
    )
    return rewrite_zip_members(
        payload,
        replacements={
            document_name: serialize(document_root),
            rels_name: serialize(rels),
            content_types_name: serialize(content_types),
        },
        additions=[("word/comments.xml", serialize(comments))],
    )


_RECIPES: dict[str, Callable[[Path], bytes]] = {
    "word_like_rich": _word_like_rich,
    "libreoffice_like_metadata": _libreoffice_like_metadata,
    "older_conversion_like": _older_conversion_like,
    "consultant_template": _consultant_template,
    "mixed_section_layout": _mixed_section_layout,
    "custom_numbering": _custom_numbering,
    "relocated_numbering_opc": _relocated_numbering_opc,
    "unusual_utf8_declaration": _unusual_utf8_declaration,
    "utf16_pass_through": _utf16_pass_through,
    "large_media": _large_media,
    "zip_comment_and_extra": _zip_comment_and_extra,
    "manual_notes_ooxml": _manual_notes_ooxml,
    "manual_comments_ooxml": _manual_comments_ooxml,
}


def build_case(case: CorpusCase, workspace: Path) -> bytes:
    # Defend this public helper even when callers construct a CorpusCase
    # directly instead of obtaining it through the validated manifest.
    _validate_case_id(case.case_id)
    if bool(case.recipe) == bool(case.fixture):
        raise ValueError(
            f"Corpus case {case.case_id!r} needs exactly one recipe or fixture"
        )
    if case.fixture is not None:
        fixture_path = (_MANIFEST_DIR / case.fixture).resolve()
        if _MANIFEST_DIR.resolve() not in fixture_path.parents:
            raise ValueError(
                f"Corpus fixture escapes its manifest directory: {case.fixture}"
            )
        payload = fixture_path.read_bytes()
        actual_sha256 = hashlib.sha256(payload).hexdigest()
        if case.expected_sha256 is None or actual_sha256 != case.expected_sha256:
            raise ValueError(
                f"Corpus fixture checksum mismatch for {case.case_id!r}"
            )
        # Preserve the reviewed fixture exactly as delivered. Sanitization may
        # already have changed or repacked the original producer envelope, so
        # provenance describes the full producer-and-sanitizer path explicitly.
        return payload

    assert case.recipe is not None
    recipe = _RECIPES.get(case.recipe)
    if recipe is None:
        raise ValueError(f"Unknown DOCX corpus recipe: {case.recipe}")
    case_workspace = workspace / case.case_id
    case_workspace.mkdir(parents=True, exist_ok=True)
    payload = recipe(case_workspace)
    payload = _sanitize_core_properties(payload, case.case_id)
    return _canonicalize_docx(payload)


def materialize_corpus(output_dir: Path) -> dict[str, Any]:
    """Write the generated corpus and a resolved manifest with checksums."""
    output_dir = output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    source_manifest = load_manifest()
    cases = corpus_cases()
    if len(cases) != len(source_manifest["cases"]):
        raise ValueError("DOCX corpus manifest changed while materializing")
    resolved_cases: list[dict[str, Any]] = []
    with tempfile.TemporaryDirectory(prefix="build-a-spec-docx-corpus-") as raw:
        workspace = Path(raw)
        for case, metadata in zip(cases, source_manifest["cases"]):
            payload = build_case(case, workspace)
            target = (output_dir / case.filename).resolve()
            if target.parent != output_dir:
                raise ValueError(
                    f"Corpus output escapes its destination: {case.filename}"
                )
            target.write_bytes(payload)
            resolved = dict(metadata)
            resolved["sha256"] = hashlib.sha256(payload).hexdigest()
            resolved["size_bytes"] = len(payload)
            resolved_cases.append(resolved)
    resolved_manifest = dict(source_manifest)
    resolved_manifest["cases"] = resolved_cases
    resolved_manifest["generator_toolchain"] = {
        "python": platform.python_version(),
        "python_docx": importlib.metadata.version("python-docx"),
        "lxml": ".".join(str(value) for value in etree.LXML_VERSION),
        "zlib": zlib.ZLIB_VERSION,
    }
    (output_dir / "manifest.json").write_text(
        json.dumps(resolved_manifest, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    return resolved_manifest


def _main() -> None:
    parser = argparse.ArgumentParser(
        description="Materialize the sanitized DOCX fidelity corpus."
    )
    parser.add_argument("output_dir", type=Path)
    args = parser.parse_args()
    resolved = materialize_corpus(args.output_dir.resolve())
    print(f"Wrote {len(resolved['cases'])} corpus cases to {args.output_dir.resolve()}")


if __name__ == "__main__":
    _main()


__all__ = [
    "CorpusCase",
    "build_case",
    "corpus_cases",
    "load_manifest",
    "materialize_corpus",
]
