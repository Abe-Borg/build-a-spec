"""High-fidelity DOCX fixtures and structural assertions for P1.

The fixture deliberately mixes content Build-a-Spec owns (one simple body
provision) with content it must treat as opaque: headers, footers, fields,
images, a table, an SDT, relationships, styles, numbering, and custom OPC
parts.  Tests compare package payloads and XML subtrees rather than relying on
``python-docx``'s lossy object model.
"""
from __future__ import annotations

import hashlib
import io
import json
import posixpath
import struct
import urllib.parse
import zipfile
import zlib
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from backend.spec_doc.source_package import inspect_docx_package


DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
TARGET_PARA_ID = "A1B2C3D4"
NUMBERED_PARA_ID = "B1C2D3E4"
OPAQUE_PARA_ID = "C1D2E3F4"
TARGET_SOURCE_TEXT = "A. Install system per NFPA 13-2019."
TARGET_MODEL_TEXT = "Install system per NFPA 13-2019."
TARGET_EDITED_TEXT = "Install system per NFPA 13-2022."
TARGET_EDITED_SOURCE_TEXT = "A. Install system per NFPA 13-2022."
NUMBERED_ISLAND_PARA_IDS = ("A1B2C3D4", "B2C3D4E5", "C3D4E5F6")
NUMBERED_ISLAND_TEXTS = (
    "Provide monitored control valves.",
    "Install seismic bracing at required intervals.",
    "Label system components permanently.",
)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_NUMBERING_REL_TYPES = {
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
    "http://purl.oclc.org/ooxml/officeDocument/relationships/numbering",
}
_CP_NS = (
    "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
)
_VT_NS = (
    "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"
)
_NS = {"w": _W_NS, "w14": _W14_NS}


@dataclass(frozen=True)
class ZipMemberSnapshot:
    """The meaningful state of one OPC member.

    ``data`` is the strong preservation invariant. ZIP metadata is retained so
    diagnostics can explain archive-level changes without confusing a
    recompressed-but-identical part with lost Word content.
    """

    data: bytes
    compress_type: int
    date_time: tuple[int, int, int, int, int, int]
    external_attr: int


def _png_bytes(width: int = 12, height: int = 12) -> bytes:
    """Return a deterministic, valid RGBA PNG without a Pillow dependency."""

    def chunk(kind: bytes, payload: bytes) -> bytes:
        checksum = zlib.crc32(kind)
        checksum = zlib.crc32(payload, checksum) & 0xFFFFFFFF
        return struct.pack(">I", len(payload)) + kind + payload + struct.pack(">I", checksum)

    # A small blue square makes header/body image preservation visible when
    # the optional render regression runs.
    scanline = b"\x00" + bytes((38, 92, 145, 255)) * width
    pixels = scanline * height
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 6, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(pixels, 9))
        + chunk(b"IEND", b"")
    )


def _set_para_id(paragraph, para_id: str) -> None:
    paragraph._p.set(qn("w14:paraId"), para_id)
    paragraph._p.set(qn("w14:textId"), para_id[::-1])


def _set_direct_numbering(paragraph, num_id: int, ilvl: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl_el, num_id_el))


def _append_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, display, end))


def _append_hyperlink(paragraph, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.extend((color, underline))
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.extend((run_props, text_el))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_opaque_sdt(document) -> None:
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "CLIENT-OPAQUE-CONTROL")
    sdt_pr.append(tag)
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    paragraph.set(qn("w14:paraId"), OPAQUE_PARA_ID)
    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = "OPAQUE CONTENT CONTROL - KEEP EXACTLY"
    run.append(text_el)
    paragraph.append(run)
    content.append(paragraph)
    sdt.extend((sdt_pr, content))
    # sectPr must remain the final body child.
    document.element.body.insert(len(document.element.body) - 1, sdt)


def _next_num_id(document) -> int:
    nums = document.part.numbering_part.element.findall(qn("w:num"))
    if not nums:  # pragma: no cover - the bundled Word template has lists
        raise AssertionError("The default DOCX template has no numbering definitions")
    return int(nums[0].get(qn("w:numId")))


def _serialize_xml(root) -> bytes:
    return etree.tostring(
        root,
        encoding="UTF-8",
        xml_declaration=True,
        standalone=True,
    )


def _add_relationship(root, rel_id: str, rel_type: str, target: str) -> None:
    relationship = etree.SubElement(root, f"{{{_REL_NS}}}Relationship")
    relationship.set("Id", rel_id)
    relationship.set("Type", rel_type)
    relationship.set("Target", target)


def rewrite_zip_members(
    payload: bytes,
    *,
    replacements: dict[str, bytes] | None = None,
    additions: list[tuple[str, bytes]] | None = None,
    omit: set[str] | None = None,
) -> bytes:
    """Rewrite a ZIP while retaining each copied member's ``ZipInfo``."""

    replacements = replacements or {}
    omit = omit or set()
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(payload), "r") as source:
        with zipfile.ZipFile(output, "w") as target:
            for info in source.infolist():
                if info.filename in omit:
                    continue
                data = replacements.get(info.filename, source.read(info))
                target.writestr(info, data)
            for name, data in additions or []:
                info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o600 << 16
                target.writestr(info, data)
    return output.getvalue()


def mark_zip_members_encrypted(payload: bytes) -> bytes:
    """Set the ZIP encryption flag in local and central headers.

    The stdlib cannot create encrypted archives, but a standards-level flag
    is sufficient to verify that validation rejects them before extraction.
    """
    mutated = bytearray(payload)
    for signature, flag_offset in ((b"PK\x03\x04", 6), (b"PK\x01\x02", 8)):
        cursor = 0
        while True:
            cursor = mutated.find(signature, cursor)
            if cursor < 0:
                break
            flags = struct.unpack_from("<H", mutated, cursor + flag_offset)[0]
            struct.pack_into("<H", mutated, cursor + flag_offset, flags | 0x1)
            cursor += len(signature)
    return bytes(mutated)


def _inject_custom_parts(payload: bytes) -> bytes:
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        root_rels = etree.fromstring(archive.read("_rels/.rels"))
        document_rels = etree.fromstring(
            archive.read("word/_rels/document.xml.rels")
        )
        content_types = etree.fromstring(archive.read("[Content_Types].xml"))

    _add_relationship(
        document_rels,
        "rIdFixtureCustomXml",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml",
        "../customXml/clientFixture.xml",
    )
    _add_relationship(
        root_rels,
        "rIdFixtureCustomProperties",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties",
        "docProps/custom.xml",
    )
    override = etree.SubElement(content_types, f"{{{_CT_NS}}}Override")
    override.set("PartName", "/docProps/custom.xml")
    override.set(
        "ContentType",
        "application/vnd.openxmlformats-officedocument.custom-properties+xml",
    )

    custom_xml = (
        b'<?xml version="1.0" encoding="UTF-8"?>'
        b'<clientData xmlns="urn:build-a-spec:test">KEEP-CUSTOM-XML</clientData>'
    )
    custom_props = etree.Element(
        f"{{{_CP_NS}}}Properties",
        nsmap={None: _CP_NS, "vt": _VT_NS},
    )
    prop = etree.SubElement(custom_props, f"{{{_CP_NS}}}property")
    prop.set("fmtid", "{D5CDD505-2E9C-101B-9397-08002B2CF9AE}")
    prop.set("pid", "2")
    prop.set("name", "ClientFixtureMarker")
    value = etree.SubElement(prop, f"{{{_VT_NS}}}lpwstr")
    value.text = "KEEP-CUSTOM-PROPERTY"

    return rewrite_zip_members(
        payload,
        replacements={
            "_rels/.rels": _serialize_xml(root_rels),
            "word/_rels/document.xml.rels": _serialize_xml(document_rels),
            "[Content_Types].xml": _serialize_xml(content_types),
        },
        additions=[
            ("customXml/clientFixture.xml", custom_xml),
            ("docProps/custom.xml", _serialize_xml(custom_props)),
        ],
    )


def make_fidelity_master(
    tmp_path: Path,
    *,
    split_target_runs: bool = False,
    filename: str = "client-fidelity-master.docx",
) -> bytes:
    """Build the source package used by preservation acceptance tests."""

    logo_path = tmp_path / "client-fixture-logo.png"
    logo_path.write_bytes(_png_bytes())

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.83)
    section.right_margin = Inches(0.91)
    section.header_distance = Inches(0.29)
    section.footer_distance = Inches(0.31)
    section.different_first_page_header_footer = True
    document.settings.odd_and_even_pages_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    provision_style = document.styles.add_style(
        "Client Provision", WD_STYLE_TYPE.PARAGRAPH
    )
    provision_style.font.name = "Arial"
    provision_style.font.size = Pt(10.5)
    provision_style.font.color.rgb = RGBColor(31, 78, 121)
    provision_style.paragraph_format.space_after = Pt(7)
    provision_style.paragraph_format.left_indent = Inches(0.35)

    header = section.header.paragraphs[0]
    header.add_run("CLIENT MASTER HEADER | KEEP EXACT | ")
    header.add_run().add_picture(str(logo_path), width=Inches(0.14))
    first_header = section.first_page_header.paragraphs[0]
    first_header.text = "CLIENT FIRST-PAGE HEADER | KEEP EXACT"
    even_header = section.even_page_header.paragraphs[0]
    even_header.text = "CLIENT EVEN-PAGE HEADER | KEEP EXACT"

    footer = section.footer.paragraphs[0]
    footer.add_run("CLIENT CONFIDENTIAL | PAGE ")
    _append_page_field(footer)
    section.first_page_footer.paragraphs[0].text = (
        "CLIENT FIRST-PAGE FOOTER | KEEP EXACT"
    )
    section.even_page_footer.paragraphs[0].text = (
        "CLIENT EVEN-PAGE FOOTER | KEEP EXACT"
    )

    document.add_paragraph("SECTION 21 13 13")
    document.add_paragraph("WET-PIPE SPRINKLER SYSTEMS")
    document.add_paragraph("PART 1 - GENERAL")
    document.add_paragraph("1.1 SUMMARY")

    target = document.add_paragraph(style="Client Provision")
    _set_para_id(target, TARGET_PARA_ID)
    if split_target_runs:
        target.add_run("A. Install system per ").bold = True
        target.add_run("NFPA 13-2019.").italic = True
    else:
        target.add_run(TARGET_SOURCE_TEXT).bold = True

    numbered = document.add_paragraph("Retain real Word numbering definitions.")
    _set_para_id(numbered, NUMBERED_PARA_ID)
    _set_direct_numbering(numbered, _next_num_id(document), ilvl=0)

    linked = document.add_paragraph("B. Unrelated linked provision: ")
    _append_hyperlink(linked, "client requirements", "https://example.invalid/spec")
    document.add_paragraph("END OF SECTION 21 13 13")

    # Everything below END OF SECTION is intentionally opaque to the semantic
    # importer but remains visible in Word and must survive a source edit.
    page_break = document.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("OPAQUE CLIENT APPENDIX - KEEP EXACT")
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Client key"
    table.cell(0, 1).text = "Client value"
    table.cell(1, 0).text = "K-factor"
    table.cell(1, 1).text = "5.6"
    document.add_picture(str(logo_path), width=Inches(0.35))
    _append_opaque_sdt(document)

    raw = io.BytesIO()
    document.save(raw)
    payload = _inject_custom_parts(raw.getvalue())
    # Exercise both product validation and python-docx reopening at fixture
    # construction time so later failures point at the patcher, not the test.
    inspect_docx_package(payload)
    Document(io.BytesIO(payload))
    (tmp_path / filename).write_bytes(payload)
    return payload


def make_numbered_island_master(
    tmp_path: Path,
    *,
    separator: str | None = None,
    mixed_num_id: bool = False,
    invalid_num_id: str | None = None,
    inconsistent_format: bool = False,
    ilvls: tuple[int, int, int] = (0, 0, 0),
    complex_middle: str | None = None,
    constant_level_text: str | None = None,
    level_number_format: str | None = None,
    filename: str = "client-numbered-island-master.docx",
) -> bytes:
    """Build a source master with a genuine Word-numbered safe island.

    The three provisions are direct ``w:body/w:p`` siblings using the same
    ``w:numPr`` by default.  Optional variants deliberately violate exactly
    one structural-safety precondition for fail-closed tests.
    """

    if separator not in {None, "empty", "empty_after_second", "sdt"}:
        raise ValueError(
            "separator must be None, 'empty', 'empty_after_second', or 'sdt'"
        )
    if complex_middle not in {None, "field", "hyperlink"}:
        raise ValueError("complex_middle must be None, 'field', or 'hyperlink'")
    if invalid_num_id not in {None, "zero", "dangling"}:
        raise ValueError("invalid_num_id must be None, 'zero', or 'dangling'")
    if len(ilvls) != 3:
        raise ValueError("ilvls must contain exactly three levels")

    logo_path = tmp_path / "numbered-island-logo.png"
    logo_path.write_bytes(_png_bytes(width=14, height=10))

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.71)
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(0.84)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.32)
    section.different_first_page_header_footer = True
    document.settings.odd_and_even_pages_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    provision_style = document.styles.add_style(
        "Client Auto Numbered Provision", WD_STYLE_TYPE.PARAGRAPH
    )
    provision_style.font.name = "Arial"
    provision_style.font.size = Pt(10.5)
    provision_style.font.color.rgb = RGBColor(31, 78, 121)
    provision_style.paragraph_format.space_after = Pt(7)
    provision_style.paragraph_format.left_indent = Inches(0.38)

    header = section.header.paragraphs[0]
    header.add_run("NUMBERED MASTER HEADER | KEEP EXACT | ")
    header.add_run().add_picture(str(logo_path), width=Inches(0.14))
    section.first_page_header.paragraphs[0].text = (
        "NUMBERED FIRST-PAGE HEADER | KEEP EXACT"
    )
    section.even_page_header.paragraphs[0].text = (
        "NUMBERED EVEN-PAGE HEADER | KEEP EXACT"
    )
    footer = section.footer.paragraphs[0]
    footer.add_run("NUMBERED MASTER FOOTER | PAGE ")
    _append_page_field(footer)
    section.first_page_footer.paragraphs[0].text = (
        "NUMBERED FIRST-PAGE FOOTER | KEEP EXACT"
    )
    section.even_page_footer.paragraphs[0].text = (
        "NUMBERED EVEN-PAGE FOOTER | KEEP EXACT"
    )

    document.add_paragraph("SECTION 21 13 15")
    document.add_paragraph("AUTOMATIC NUMBERING FIDELITY FIXTURE")
    document.add_paragraph("PART 1 - GENERAL")
    document.add_paragraph("1.1 NUMBERED REQUIREMENTS")

    numbering_root = document.part.numbering_part.element
    numbering_instances = numbering_root.findall(qn("w:num"))
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering_instances]
    if not num_ids:  # pragma: no cover - the bundled template has lists
        raise AssertionError("The default DOCX template has no numbering definitions")
    # Mint dedicated numbering instances for the fixture. Reusing one of the
    # bundled template's numIds would let unrelated list styles share the same
    # document-wide counter, which is intentionally outside the safe-island
    # contract under test.
    primary_num_id = max(num_ids) + 1
    primary_instance = deepcopy(numbering_instances[0])
    primary_instance.set(qn("w:numId"), str(primary_num_id))
    numbering_root.append(primary_instance)
    definition_num_id = primary_num_id
    secondary_num_id = primary_num_id + 1
    secondary_template = (
        numbering_instances[1]
        if len(numbering_instances) > 1
        else numbering_instances[0]
    )
    secondary_instance = deepcopy(secondary_template)
    secondary_instance.set(qn("w:numId"), str(secondary_num_id))
    numbering_root.append(secondary_instance)
    num_ids.extend((primary_num_id, secondary_num_id))
    if invalid_num_id == "zero":
        primary_num_id = 0
    elif invalid_num_id == "dangling":
        primary_num_id = max(num_ids) + 999

    if constant_level_text is not None or level_number_format is not None:
        numbering = document.part.numbering_part.element
        instances = [
            item
            for item in numbering.findall(qn("w:num"))
            if int(item.get(qn("w:numId"))) == definition_num_id
        ]
        assert len(instances) == 1
        abstract_ref = instances[0].find(qn("w:abstractNumId"))
        assert abstract_ref is not None
        abstract_id = abstract_ref.get(qn("w:val"))
        abstracts = [
            item
            for item in numbering.findall(qn("w:abstractNum"))
            if item.get(qn("w:abstractNumId")) == abstract_id
        ]
        assert len(abstracts) == 1
        levels = [
            item
            for item in abstracts[0].findall(qn("w:lvl"))
            if item.get(qn("w:ilvl")) == "0"
        ]
        assert len(levels) == 1
        number_format = levels[0].find(qn("w:numFmt"))
        level_text = levels[0].find(qn("w:lvlText"))
        assert number_format is not None and level_text is not None
        number_format.set(
            qn("w:val"),
            level_number_format or "upperLetter",
        )
        if constant_level_text is not None:
            level_text.set(qn("w:val"), constant_level_text)

    for index, (para_id, text, ilvl) in enumerate(
        zip(NUMBERED_ISLAND_PARA_IDS, NUMBERED_ISLAND_TEXTS, ilvls)
    ):
        paragraph = document.add_paragraph(style=provision_style)
        _set_para_id(paragraph, para_id)
        # These volatile attributes prove that a newly synthesized paragraph
        # does not clone source identity/revision-session metadata.
        paragraph._p.set(qn("w:rsidR"), f"00ABC{index:03X}")
        paragraph._p.set(qn("w:rsidRDefault"), f"00DEF{index:03X}")
        num_id = secondary_num_id if mixed_num_id and index == 1 else primary_num_id
        _set_direct_numbering(paragraph, num_id, ilvl=ilvl)
        if complex_middle == "hyperlink" and index == 1:
            paragraph.add_run("Install seismic bracing at ").bold = True
            _append_hyperlink(
                paragraph,
                "required intervals.",
                "https://example.invalid/numbered-island",
            )
        elif complex_middle == "field" and index == 1:
            paragraph.add_run("Install seismic bracing at required interval ").bold = True
            _append_page_field(paragraph)
        else:
            run = paragraph.add_run(text)
            run.bold = True
            if inconsistent_format and index == 1:
                run.italic = True
            run.font.color.rgb = RGBColor(31, 78, 121)
            run._r.set(qn("w:rsidR"), f"00123{index:03X}")

        if (
            index == 0
            and separator == "empty"
            or index == 1
            and separator == "empty_after_second"
        ):
            empty = document.add_paragraph()
            empty._p.set(qn("w14:paraId"), "D4E5F607")
        elif index == 0 and separator == "sdt":
            sdt = OxmlElement("w:sdt")
            sdt_pr = OxmlElement("w:sdtPr")
            tag = OxmlElement("w:tag")
            tag.set(qn("w:val"), "NUMBERED-ISLAND-BOUNDARY")
            sdt_pr.append(tag)
            content = OxmlElement("w:sdtContent")
            opaque_p = OxmlElement("w:p")
            opaque_r = OxmlElement("w:r")
            opaque_t = OxmlElement("w:t")
            opaque_t.text = "OPAQUE ISLAND BOUNDARY - KEEP EXACT"
            opaque_r.append(opaque_t)
            opaque_p.append(opaque_r)
            content.append(opaque_p)
            sdt.extend((sdt_pr, content))
            document.element.body.insert(len(document.element.body) - 1, sdt)

    document.add_paragraph("END OF SECTION 21 13 15")

    # Opaque content after END OF SECTION still exercises package fidelity,
    # but is never exposed as semantic content the app can mutate.
    page_break = document.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("NUMBERED OPAQUE APPENDIX - KEEP EXACT")
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Client key"
    table.cell(0, 1).text = "Client value"
    table.cell(1, 0).text = "Hazard"
    table.cell(1, 1).text = "Ordinary Group 1"
    document.add_picture(str(logo_path), width=Inches(0.35))
    _append_opaque_sdt(document)

    raw = io.BytesIO()
    document.save(raw)
    payload = _inject_custom_parts(raw.getvalue())
    inspect_docx_package(payload)
    Document(io.BytesIO(payload))
    (tmp_path / filename).write_bytes(payload)
    return payload


def remove_numbering_relationship(payload: bytes) -> bytes:
    """Leave an orphan numbering part that Word no longer wires to the body."""
    rels_name = "word/_rels/document.xml.rels"
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        rels = etree.fromstring(archive.read(rels_name))
    matches = [
        relationship
        for relationship in rels.findall(f"{{{_REL_NS}}}Relationship")
        if relationship.get("Type") in _NUMBERING_REL_TYPES
    ]
    assert len(matches) == 1
    rels.remove(matches[0])
    orphaned = rewrite_zip_members(
        payload,
        replacements={rels_name: _serialize_xml(rels)},
    )
    # The baseline upload validator intentionally checks only the required OPC
    # shell; the structural-numbering gate must reject this subtler orphan.
    inspect_docx_package(orphaned)
    return orphaned


def replace_numbering_content_type(payload: bytes, content_type: str) -> bytes:
    """Return a package whose numbering target has the wrong effective type."""
    part_name = "[Content_Types].xml"
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        content_types = etree.fromstring(archive.read(part_name))
    matches = [
        override
        for override in content_types.findall(f"{{{_CT_NS}}}Override")
        if override.get("PartName", "").lstrip("/") == "word/numbering.xml"
    ]
    assert len(matches) == 1
    matches[0].set("ContentType", content_type)
    changed = rewrite_zip_members(
        payload,
        replacements={part_name: _serialize_xml(content_types)},
    )
    inspect_docx_package(changed)
    return changed


def add_document_protection(payload: bytes) -> bytes:
    """Return the fixture with enforced read-only Word protection."""
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        settings = etree.fromstring(archive.read("word/settings.xml"))
    protection = etree.Element(f"{{{_W_NS}}}documentProtection")
    protection.set(f"{{{_W_NS}}}edit", "readOnly")
    protection.set(f"{{{_W_NS}}}enforcement", "1")
    settings.insert(0, protection)
    protected = rewrite_zip_members(
        payload,
        replacements={"word/settings.xml": _serialize_xml(settings)},
    )
    inspect_docx_package(protected)
    Document(io.BytesIO(protected))
    return protected


def add_tracked_change(payload: bytes) -> bytes:
    """Wrap the target run in a pending insertion revision."""
    root = etree.fromstring(document_xml(payload))
    targets = root.xpath(
        f'.//w:p[@w14:paraId="{TARGET_PARA_ID}"]', namespaces=_NS
    )
    assert len(targets) == 1
    target = targets[0]
    run = target.find(f"{{{_W_NS}}}r")
    assert run is not None
    index = target.index(run)
    target.remove(run)
    insertion = etree.Element(f"{{{_W_NS}}}ins")
    insertion.set(f"{{{_W_NS}}}id", "77")
    insertion.set(f"{{{_W_NS}}}author", "Fixture Reviewer")
    insertion.set(f"{{{_W_NS}}}date", "2026-01-01T00:00:00Z")
    insertion.append(run)
    target.insert(index, insertion)
    tracked = rewrite_zip_members(
        payload,
        replacements={"word/document.xml": _serialize_xml(root)},
    )
    inspect_docx_package(tracked)
    Document(io.BytesIO(tracked))
    return tracked


def add_paragraph_property_change(payload: bytes) -> bytes:
    """Add a pending ``w:pPrChange`` revision without wrapping body text.

    Property-change revisions are easy to miss when a detector only searches
    for ``w:ins``/``w:del``. They still make a source package pass-through-only.
    """

    root = etree.fromstring(document_xml(payload))
    targets = root.xpath(
        f'.//w:p[@w14:paraId="{TARGET_PARA_ID}"]', namespaces=_NS
    )
    assert len(targets) == 1
    target = targets[0]
    p_pr = target.find(f"{{{_W_NS}}}pPr")
    assert p_pr is not None
    change = etree.Element(f"{{{_W_NS}}}pPrChange")
    change.set(f"{{{_W_NS}}}id", "78")
    change.set(f"{{{_W_NS}}}author", "Fixture Reviewer")
    change.set(f"{{{_W_NS}}}date", "2026-01-02T00:00:00Z")
    prior = etree.SubElement(change, f"{{{_W_NS}}}pPr")
    prior_style = etree.SubElement(prior, f"{{{_W_NS}}}pStyle")
    prior_style.set(f"{{{_W_NS}}}val", "Normal")
    p_pr.append(change)
    tracked = rewrite_zip_members(
        payload,
        replacements={"word/document.xml": _serialize_xml(root)},
    )
    inspect_docx_package(tracked)
    Document(io.BytesIO(tracked))
    return tracked


def add_signature_origin_marker(payload: bytes) -> bytes:
    """Add the OPC signature-origin plumbing used by the mutation blocker."""
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        root_rels = etree.fromstring(archive.read("_rels/.rels"))
        content_types = etree.fromstring(archive.read("[Content_Types].xml"))
    _add_relationship(
        root_rels,
        "rIdFixtureSignatureOrigin",
        "http://schemas.openxmlformats.org/package/2006/relationships/digital-signature/origin",
        "/_xmlsignatures/origin.sigs",
    )
    override = etree.SubElement(content_types, f"{{{_CT_NS}}}Override")
    override.set("PartName", "/_xmlsignatures/origin.sigs")
    override.set(
        "ContentType",
        "application/vnd.openxmlformats-package.digital-signature-origin",
    )
    marked = rewrite_zip_members(
        payload,
        replacements={
            "_rels/.rels": _serialize_xml(root_rels),
            "[Content_Types].xml": _serialize_xml(content_types),
        },
        additions=[("_xmlsignatures/origin.sigs", b"")],
    )
    inspect_docx_package(marked)
    Document(io.BytesIO(marked))
    return marked


def add_active_content_marker(payload: bytes) -> bytes:
    """Add a related embedded-object part that must be pass-through-only."""
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        document_rels = etree.fromstring(
            archive.read("word/_rels/document.xml.rels")
        )
        content_types = etree.fromstring(archive.read("[Content_Types].xml"))
    _add_relationship(
        document_rels,
        "rIdFixtureEmbeddedObject",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject",
        "embeddings/clientFixture.bin",
    )
    override = etree.SubElement(content_types, f"{{{_CT_NS}}}Override")
    override.set("PartName", "/word/embeddings/clientFixture.bin")
    override.set(
        "ContentType",
        "application/vnd.openxmlformats-officedocument.oleObject",
    )
    marked = rewrite_zip_members(
        payload,
        replacements={
            "word/_rels/document.xml.rels": _serialize_xml(document_rels),
            "[Content_Types].xml": _serialize_xml(content_types),
        },
        additions=[("word/embeddings/clientFixture.bin", b"OPAQUE-OLE-FIXTURE")],
    )
    inspect_docx_package(marked)
    Document(io.BytesIO(marked))
    return marked


def make_table_projection_master(tmp_path: Path) -> bytes:
    """Build a master whose only provision is an opaque table-row projection."""
    document = Document()
    document.sections[0].header.paragraphs[0].text = "TABLE MASTER HEADER - KEEP"
    document.add_paragraph("SECTION 21 13 14")
    document.add_paragraph("TABLE PROJECTION FIXTURE")
    document.add_paragraph("PART 2 - PRODUCTS")
    document.add_paragraph("2.1 SCHEDULE")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "K-factor"
    table.cell(0, 1).text = "5.6"
    document.add_paragraph("END OF SECTION 21 13 14")
    raw = io.BytesIO()
    document.save(raw)
    payload = raw.getvalue()
    inspect_docx_package(payload)
    return payload


def zip_snapshot(payload: bytes) -> tuple[list[str], dict[str, ZipMemberSnapshot]]:
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        infos = archive.infolist()
        return [info.filename for info in infos], {
            info.filename: ZipMemberSnapshot(
                data=archive.read(info),
                compress_type=info.compress_type,
                date_time=info.date_time,
                external_attr=info.external_attr,
            )
            for info in infos
        }


def assert_untouched_parts_identical(
    before: bytes,
    after: bytes,
    *,
    changed_parts: set[str] | frozenset[str] = frozenset(
        {"word/document.xml"}
    ),
) -> None:
    before_order, before_parts = zip_snapshot(before)
    after_order, after_parts = zip_snapshot(after)
    assert after_order == before_order, "The source package member order changed"
    assert set(after_parts) == set(before_parts)
    for name in before_order:
        if name not in changed_parts:
            assert after_parts[name].data == before_parts[name].data, (
                f"Untouched OPC part changed: {name}"
            )


def _relationship_source_part(rel_part: str) -> str:
    if rel_part == "_rels/.rels":
        return ""
    marker = "/_rels/"
    if marker not in rel_part or not rel_part.endswith(".rels"):
        raise AssertionError(f"Malformed relationship part path: {rel_part}")
    prefix, rel_name = rel_part.split(marker, 1)
    return posixpath.join(prefix, rel_name[: -len(".rels")])


def _resolve_internal_target(rel_part: str, target: str) -> str:
    target = urllib.parse.unquote(target.split("#", 1)[0])
    if target.startswith("/"):
        resolved = posixpath.normpath(target.lstrip("/"))
    else:
        source_part = _relationship_source_part(rel_part)
        resolved = posixpath.normpath(
            posixpath.join(posixpath.dirname(source_part), target)
        )
    assert resolved and not resolved.startswith("../"), (
        f"Relationship escapes the package: {rel_part} -> {target}"
    )
    return resolved


def assert_valid_docx_package(payload: bytes) -> None:
    """Validate ZIP integrity, XML, OPC graph, content types, and reopening."""

    inspect_docx_package(payload)
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        assert archive.testzip() is None
        names = {info.filename for info in archive.infolist() if not info.is_dir()}
        parsed: dict[str, etree._Element] = {}
        for name in names:
            if name.endswith((".xml", ".rels")) or name == "[Content_Types].xml":
                parsed[name] = etree.fromstring(archive.read(name))

        for rel_name, root in parsed.items():
            if not rel_name.endswith(".rels"):
                continue
            for relationship in root.findall(f"{{{_REL_NS}}}Relationship"):
                if relationship.get("TargetMode", "Internal") == "External":
                    continue
                target = relationship.get("Target")
                assert target
                resolved = _resolve_internal_target(rel_name, target)
                assert resolved in names, (
                    f"Missing relationship target: {rel_name} -> {resolved}"
                )

        content_types = parsed["[Content_Types].xml"]
        defaults = {
            node.get("Extension", "").lower(): node.get("ContentType")
            for node in content_types.findall(f"{{{_CT_NS}}}Default")
        }
        overrides = {
            node.get("PartName", "").lstrip("/"): node.get("ContentType")
            for node in content_types.findall(f"{{{_CT_NS}}}Override")
        }
        for name in names - {"[Content_Types].xml"}:
            extension = name.rsplit(".", 1)[-1].lower() if "." in name else ""
            assert name in overrides or extension in defaults, (
                f"No content type declaration covers {name}"
            )

    reopened = Document(io.BytesIO(payload))
    assert reopened.sections


def document_xml(payload: bytes) -> bytes:
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        return archive.read("word/document.xml")


def _canonical(element) -> bytes:
    return etree.tostring(element, method="c14n", with_comments=True)


def _body_and_target(payload: bytes, para_id: str = TARGET_PARA_ID):
    root = etree.fromstring(document_xml(payload))
    body = root.find(f"{{{_W_NS}}}body")
    assert body is not None
    matches = body.xpath(
        f'./w:p[@w14:paraId="{para_id}"]', namespaces=_NS
    )
    assert len(matches) == 1, f"Expected one direct body paragraph {para_id}"
    target = matches[0]
    return root, body, target, list(body).index(target)


def paragraph_text(payload: bytes, para_id: str = TARGET_PARA_ID) -> str:
    _root, _body, target, _index = _body_and_target(payload, para_id)
    return "".join(target.xpath(".//w:t/text()", namespaces=_NS))


def assert_only_target_text_changed(
    before: bytes,
    after: bytes,
    *,
    para_id: str = TARGET_PARA_ID,
    expected_before: str = TARGET_SOURCE_TEXT,
    expected_after: str = TARGET_EDITED_SOURCE_TEXT,
) -> None:
    before_root, before_body, before_target, before_index = _body_and_target(
        before, para_id
    )
    after_root, after_body, after_target, after_index = _body_and_target(
        after, para_id
    )
    assert before_index == after_index
    assert len(before_body) == len(after_body)
    assert etree.QName(before_root).localname == etree.QName(after_root).localname

    for index, (old_child, new_child) in enumerate(zip(before_body, after_body)):
        if index != before_index:
            assert _canonical(old_child) == _canonical(new_child), (
                f"Untargeted body child {index} changed"
            )

    before_text_nodes = before_target.xpath(".//w:t", namespaces=_NS)
    after_text_nodes = after_target.xpath(".//w:t", namespaces=_NS)
    assert len(before_text_nodes) == len(after_text_nodes)
    differing = [
        index
        for index, (old, new) in enumerate(zip(before_text_nodes, after_text_nodes))
        if old.text != new.text
    ]
    assert differing == [0], "A P1a edit may change exactly one w:t node"
    assert "".join(node.text or "" for node in before_text_nodes) == expected_before
    assert "".join(node.text or "" for node in after_text_nodes) == expected_after

    # Replacing the changed text with the same sentinel must make the full
    # target paragraph canonical form identical. This covers pPr, rPr,
    # labels, attributes, bookmarks, and all other inline markup at once.
    old_skeleton = deepcopy(before_target)
    new_skeleton = deepcopy(after_target)
    old_skeleton.xpath(".//w:t", namespaces=_NS)[0].text = "__EDITED_TEXT__"
    new_skeleton.xpath(".//w:t", namespaces=_NS)[0].text = "__EDITED_TEXT__"
    assert _canonical(old_skeleton) == _canonical(new_skeleton)


def package_manifest(payload: bytes) -> dict:
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        return json.loads(archive.read("manifest.json"))


def sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()
