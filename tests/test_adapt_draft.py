"""The adapt-imported pass + the imported-review context line.

``POST /api/draft/adapt`` is ``/api/draft/full``'s counterpart for the other
on-ramp: a document that arrived FULL of imported starter content instead of
empty. Same posture — a thin endpoint handing back a server-owned directive
the frontend sends through the ordinary chat path — same prerequisite gate,
plus one guard of its own (nothing imported = nothing to adapt).

Also here: the IMPORTED-STARTER REVIEW line in PROJECT CONTEXT. Imported
blocks are deliberately not open items, so before this line the model had no
work-list driver at all on a fresh import — the strongest driver path A has
(a visible list shrinking each turn) simply did not exist in path B.
"""
from __future__ import annotations

import io
import json

from docx import Document
from fastapi.testclient import TestClient

from backend.app import create_app
from backend import sessions
from tests.fakes import FakeClient, request_context_text, text_turn, tool_turn


def _client() -> TestClient:
    return TestClient(create_app())


def _master_bytes() -> bytes:
    document = Document()
    for text in (
        "SECTION 23 05 48",
        "VIBRATION AND SEISMIC CONTROLS FOR HVAC",
        "PART 1 - GENERAL",
        "1.1 SECTION INCLUDES",
        "A. Vibration isolation devices for HVAC equipment.",
        "B. Seismic restraints for HVAC equipment.",
        "END OF SECTION 23 05 48",
    ):
        document.add_paragraph(text)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _import_starting_point(client: TestClient) -> dict:
    """Import a master as a starting point (detached — no sweep involved)."""
    response = client.post(
        "/api/import/master",
        files={
            "file": (
                "23 05 48.docx",
                _master_bytes(),
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document",
            )
        },
        data={"detach": "true"},
    )
    assert response.status_code == 200, response.text
    return response.json()


_PROJECT_TYPE_OP = {
    "action": "set_project_identity",
    "target_id": "sec",
    "project_type": "Hyperscale Data Center",
}
_COUNTRY_OP = {
    "action": "set_project_profile",
    "target_id": "sec",
    "country": "United States",
}


def _establish(client: TestClient, *ops: dict) -> None:
    resp = client.post("/api/doc/edit", json={"ops": list(ops)})
    assert resp.status_code == 200, resp.text


def _parse_sse(body: str) -> list[dict]:
    return [
        json.loads(line[len("data: "):])
        for line in body.splitlines()
        if line.startswith("data: ")
    ]


def _patch_client(monkeypatch, fake: FakeClient) -> None:
    monkeypatch.setattr("backend.llm.conversation.get_client", lambda: fake)


# ---------------------------------------------------------------------------
# Endpoint guards + payload
# ---------------------------------------------------------------------------


def test_draft_adapt_returns_the_directive_for_an_imported_starter():
    client = _client()
    _import_starting_point(client)
    # The master satisfied the section prerequisite itself; the other two
    # anchor facts are what "adapt to THIS project" is defined by.
    _establish(client, _PROJECT_TYPE_OP, _COUNTRY_OP)

    resp = client.post("/api/draft/adapt")
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["ok"] is True
    assert data["ready"] is True
    assert data["missing"] == []
    assert data["imported_count"] > 0
    message = data["message"]
    assert message.startswith("Walk the ENTIRE imported starter")
    # Anchored on the established facts, the full-draft pattern.
    assert "- SECTION: 23 05 48 VIBRATION AND SEISMIC CONTROLS FOR HVAC" in (
        message
    )
    assert "- PROJECT TYPE (facility/use): Hyperscale Data Center" in message
    assert "- COUNTRY: USA" in message


def test_draft_adapt_409_while_turn_active():
    sessions.get_session().turn_active = True
    try:
        resp = _client().post("/api/draft/adapt")
        assert resp.status_code == 409
        assert "streaming" in resp.json()["error"]
    finally:
        sessions.get_session().turn_active = False


def test_draft_adapt_409_while_research_running():
    sessions.get_session().research.status = "running"
    resp = _client().post("/api/draft/adapt")
    assert resp.status_code == 409
    assert "research" in resp.json()["error"].lower()


def test_draft_adapt_refuses_when_nothing_is_imported():
    """A blank page has nothing to adapt, and so does a fully reviewed one."""
    client = _client()
    resp = client.post("/api/draft/adapt")
    assert resp.status_code == 409
    assert "nothing left to adapt" in resp.json()["error"].lower()

    # Seed one imported block, then retire it: the guard reads live state.
    _establish(
        client,
        {"action": "add_article", "target_id": "pt1", "text": "SUMMARY"},
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "text": "Starter provision.",
            "status": "imported",
        },
    )
    assert client.post("/api/draft/adapt").status_code == 200
    doc = client.get("/api/doc").json()["doc"]
    uid = doc["parts"][0]["articles"][0]["paragraphs"][0]["id"]
    _establish(
        client,
        {"action": "set_status", "target_id": uid, "status": "confirmed"},
    )
    assert client.post("/api/draft/adapt").status_code == 409


def test_draft_adapt_collects_the_missing_facts_instead_of_adapting_blind():
    client = _client()
    _import_starting_point(client)

    resp = client.post("/api/draft/adapt")
    assert resp.status_code == 200
    data = resp.json()
    assert data["ok"] is True
    assert data["ready"] is False
    # The master's own header settled the section; the project facts didn't
    # ride in with the file.
    assert data["missing"] == ["project_type", "country"]

    message = data["message"]
    assert not message.startswith("Walk the ENTIRE imported starter")
    assert "adapt the imported document" in message
    assert "set_project_identity" in message
    assert "set_project_profile" in message
    assert "Do NOT start the adapt pass in this turn" in message
    assert "tell me the adapt pass is ready" in message
    # Settled facts named as settled.
    assert "- SECTION: 23 05 48 VIBRATION AND SEISMIC CONTROLS FOR HVAC" in (
        message
    )


def test_the_full_draft_collect_wording_is_byte_stable():
    """The adapt phrasing rides keyword parameters whose DEFAULTS must keep
    reproducing the full-draft wording exactly — the full-draft tests pin the
    substance; this pins the mechanism that keeps them independent."""
    from backend.llm.prompts import (
        draft_prerequisites,
        draft_prerequisites_directive,
    )

    prereqs = draft_prerequisites(project_type="Hospital", country="US")
    message = draft_prerequisites_directive(prereqs)
    assert "Before you draft the full section" in message
    assert "a whole-section draft anchors on it" in message
    assert "Do NOT draft the section in this turn" in message
    assert "tell me the full draft is ready and I will run it." in message


# ---------------------------------------------------------------------------
# Directive obligations (prompt snapshot)
# ---------------------------------------------------------------------------


def test_adapt_directive_carries_the_gap_and_adapt_obligations():
    from backend.llm.prompts import ADAPT_IMPORTED_DIRECTIVE

    text = ADAPT_IMPORTED_DIRECTIVE
    # Whole-document, in order, three dispositions per block.
    assert "PART by PART" in text
    assert "keep it" in text and "adapt it" in text and "delete" in text
    assert "set_status" in text
    # Edition hygiene + remnant hunt.
    assert "stale" in text
    assert "remnants" in text
    # Boundary awareness and the honest way out.
    assert "IMPORTED DOCX EDITING BOUNDARY" in text
    assert "Edit freely" in text
    # Provenance link + the close.
    assert "source_item_id" in text
    assert "follow-up questions" in text


# ---------------------------------------------------------------------------
# The imported-review context line (the work-list driver)
# ---------------------------------------------------------------------------


def test_the_context_names_the_imported_review_debt(monkeypatch):
    client = _client()
    _import_starting_point(client)

    fake = FakeClient([text_turn(["Noted."])])
    _patch_client(monkeypatch, fake)
    assert client.post("/api/chat", json={"message": "hi"}).status_code == 200

    context = request_context_text(fake.messages.requests[0])
    assert "IMPORTED-STARTER REVIEW:" in context
    assert "still carry status 'imported'" in context
    assert "confirmed, adapted, or deleted" in context


def test_a_document_with_no_imported_blocks_carries_no_review_line(
    monkeypatch,
):
    client = _client()
    fake = FakeClient([text_turn(["Noted."])])
    _patch_client(monkeypatch, fake)
    assert client.post("/api/chat", json={"message": "hi"}).status_code == 200
    context = request_context_text(fake.messages.requests[0])
    assert context
    assert "IMPORTED-STARTER REVIEW" not in context


def test_the_review_line_never_enters_the_stable_prompt():
    """The cache rule: session-varying counts live in dynamic context only."""
    from backend.llm.prompts import render_system_prompt
    from backend.spec_modules.registry import DEFAULT_MODULE

    prompt = render_system_prompt(DEFAULT_MODULE)
    assert "IMPORTED-STARTER REVIEW" not in prompt


# ---------------------------------------------------------------------------
# End-to-end: the directive drives an ordinary turn
# ---------------------------------------------------------------------------


def test_adapt_directive_drives_an_ordinary_turn(monkeypatch):
    client = _client()
    payload = _import_starting_point(client)
    _establish(client, _PROJECT_TYPE_OP, _COUNTRY_OP)

    first_uid = payload["doc"]["parts"][0]["articles"][0]["paragraphs"][0][
        "id"
    ]
    second_uid = payload["doc"]["parts"][0]["articles"][0]["paragraphs"][1][
        "id"
    ]
    plan = client.post("/api/draft/adapt").json()
    assert plan["ready"] is True

    fake = FakeClient(
        [
            tool_turn(
                ["Adapting PART 1. "],
                {
                    "edits": [
                        {
                            "action": "replace",
                            "target_id": first_uid,
                            "text": "Vibration isolation for this data "
                            "center's HVAC equipment.",
                            "status": "assumed",
                        },
                        {
                            "action": "set_status",
                            "target_id": second_uid,
                            "status": "confirmed",
                        },
                    ]
                },
                tool_id="toolu_a1",
            ),
            text_turn(["Adapt pass complete — kept 1, adapted 1."]),
        ]
    )
    _patch_client(monkeypatch, fake)

    resp = client.post("/api/chat", json={"message": plan["message"]})
    events = _parse_sse(resp.text)
    assert [e for e in events if e["type"] == "doc_patch"]
    assert events[-1]["type"] == "turn_complete"

    # The directive is the honest user turn in history.
    history = sessions.get_session().history
    assert history[0]["content"] == [
        {"type": "text", "text": plan["message"]}
    ]

    doc = client.get("/api/doc").json()["doc"]
    paragraphs = doc["parts"][0]["articles"][0]["paragraphs"]
    assert paragraphs[0]["status"] == "assumed"
    assert paragraphs[1]["status"] == "confirmed"
