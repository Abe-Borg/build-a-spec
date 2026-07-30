"""Backend API tests: health, key handling, SSE chat + tool loop, document
endpoints, export, and project save/resume — all against the scripted fake
streaming client in ``tests/fakes.py``."""
from __future__ import annotations

import io
import json
import logging
from types import SimpleNamespace

from docx import Document
from docx.oxml.ns import qn
from fastapi.testclient import TestClient

from backend.app import create_app
from backend import sessions
from tests.fakes import (
    FakeClient,
    chat_search_blocks,
    raw_turn,
    request_context_text,
    text_block,
    text_turn,
    thinking_block,
    tool_turn,
    tool_use_block,
)


def _client() -> TestClient:
    return TestClient(create_app())


def _parse_sse(body: str) -> list[dict]:
    events = []
    for line in body.splitlines():
        if line.startswith("data: "):
            events.append(json.loads(line[len("data: "):]))
    return events


def _patch_client(monkeypatch, fake: FakeClient) -> None:
    monkeypatch.setattr("backend.llm.conversation.get_client", lambda: fake)


_SEED_EDITS = {
    "edits": [
        {
            "action": "replace",
            "target_id": "sec",
            "text": "WET-PIPE SPRINKLER SYSTEMS",
            "numbering": "21 13 13",
        },
        {"action": "add_article", "target_id": "pt1", "text": "SUMMARY"},
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "text": "Section includes wet-pipe systems per NFPA 13-2025.",
            "status": "assumed",
        },
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "text": "Design density: [TBD: density] over remote area.",
            "status": "needs_input",
        },
    ]
}


def _seed_doc_via_chat(client: TestClient, monkeypatch) -> None:
    fake = FakeClient(
        [tool_turn(["Drafting."], _SEED_EDITS), text_turn(["Done."])]
    )
    _patch_client(monkeypatch, fake)
    resp = client.post("/api/chat", json={"message": "Start 21 13 13"})
    assert _parse_sse(resp.text)[-1]["type"] == "turn_complete"


# ---------------------------------------------------------------------------
# Phase 1 surface
# ---------------------------------------------------------------------------


def test_health_reports_model_and_key(monkeypatch):
    resp = _client().get("/api/health")
    assert resp.status_code == 200
    data = resp.json()
    assert data["status"] == "ok"
    assert data["model"]
    assert data["api_key_present"] is True  # conftest injects the env key


def test_chat_streams_deltas_and_updates_history(monkeypatch):
    fake = FakeClient([text_turn(["PART 1 ", "- GENERAL"])])
    _patch_client(monkeypatch, fake)

    resp = _client().post("/api/chat", json={"message": "Start 21 13 13"})
    assert resp.status_code == 200
    events = _parse_sse(resp.text)

    deltas = [e["text"] for e in events if e["type"] == "text_delta"]
    assert "".join(deltas) == "PART 1 - GENERAL"
    assert events[-1]["type"] == "turn_complete"
    assert events[-1]["stop_reason"] == "end_turn"

    history = sessions.get_session().history
    assert len(history) == 2
    assert history[0]["role"] == "user"
    assert history[1]["role"] == "assistant"
    assert history[1]["content"][0]["text"] == "PART 1 - GENERAL"

    # The request: cached stable system prompt (nothing else in system),
    # the PROJECT CONTEXT block riding the user message, the document +
    # web tools, and explicit adaptive thinking at the configured effort.
    request = fake.messages.last_request
    assert len(request["system"]) == 1
    assert request["system"][0]["cache_control"] == {"type": "ephemeral", "ttl": "1h"}
    context = request_context_text(request)
    assert "document" in context
    # Order is the cached-prefix contract: additions go on the end so the
    # existing tool bytes stay a stable prefix.
    assert [t["name"] for t in request["tools"]] == [
        "apply_spec_edits",
        "create_figure",
        "web_search",
        "web_fetch",
        "suggest_prompts",
        "read_reference_doc",
    ]
    # Both web tools invoke directly. The provider default for the
    # ``_20260209`` versions is a code-execution caller, whose pause_turn
    # continuations need a container id the chat loop does not send — and
    # which stops streaming the per-search inputs the 🔍 chips are built
    # from. Same builders, same mode, as research and Final QC.
    web_tools = [t for t in request["tools"] if t["name"] in ("web_search", "web_fetch")]
    assert [t["type"] for t in web_tools] == [
        "web_search_20260209",
        "web_fetch_20260209",
    ]
    assert all(t["allowed_callers"] == ["direct"] for t in web_tools)
    # No per-turn locale here: it would bust the cached tool prefix for the
    # whole session (see ``_chat_tools``).
    assert all("user_location" not in t for t in web_tools)
    # Adaptive thinking with the summarized-display opt-in (the "see what
    # the model is thinking" stream) at the configured effort.
    assert request["thinking"] == {"type": "adaptive", "display": "summarized"}
    assert request["output_config"] == {"effort": "high"}

    # Committed history stores the user's text ONLY — the context block
    # is per-request, never fossilized into history.
    assert history[0]["content"] == [
        {"type": "text", "text": "Start 21 13 13"}
    ]


def test_chat_error_leaves_history_clean(monkeypatch):
    def _boom():
        raise RuntimeError("kaput")

    monkeypatch.setattr("backend.llm.conversation.get_client", _boom)

    resp = _client().post("/api/chat", json={"message": "hello"})
    events = _parse_sse(resp.text)
    assert events == [
        {"type": "error", "message": "Unexpected error: kaput"}
    ]
    assert sessions.get_session().history == []


def test_empty_message_is_rejected(monkeypatch):
    fake = FakeClient([text_turn(["never used"])])
    _patch_client(monkeypatch, fake)
    resp = _client().post("/api/chat", json={"message": "   "})
    events = _parse_sse(resp.text)
    assert events[0]["type"] == "error"
    assert fake.messages.last_request is None


def test_session_reset_clears_history_and_document(monkeypatch):
    client = _client()
    _seed_doc_via_chat(client, monkeypatch)
    assert sessions.get_session().history
    assert not sessions.get_session().doc.doc.is_empty()

    resp = client.post("/api/session/reset")
    # Batch 10: the reset response reports the (kept) module + discipline. The
    # neutral default is now the generic module; project_context is echoed too.
    assert resp.json() == {
        "ok": True,
        "module_id": "generic",
        "module": sessions.get_session().module.display_name,
        "discipline": "",
        "project_context": "",
    }
    assert sessions.get_session().history == []
    assert sessions.get_session().doc.doc.is_empty()
    assert len(sessions.get_session().doc.versions) == 1


# ---------------------------------------------------------------------------
# Tool-use continuation loop
# ---------------------------------------------------------------------------


def test_tool_turn_patches_document_and_continues(monkeypatch):
    fake = FakeClient(
        [
            tool_turn(["Drafting the summary. "], _SEED_EDITS),
            text_turn(["Two questions next."]),
        ]
    )
    _patch_client(monkeypatch, fake)

    resp = _client().post("/api/chat", json={"message": "Start 21 13 13"})
    events = _parse_sse(resp.text)

    # Text from both rounds streamed out.
    text = "".join(e["text"] for e in events if e["type"] == "text_delta")
    assert text == "Drafting the summary. Two questions next."

    patches = [e for e in events if e["type"] == "doc_patch"]
    assert len(patches) == 1
    applied_ids = [op["id"] for op in patches[0]["ops"]]
    assert applied_ids == ["sec", "pt1.a1", "pt1.a1.p1", "pt1.a1.p2"]
    assert patches[0]["doc"]["section"]["number"] == "21 13 13"

    # Mid-turn patches carry the pre-commit version pointer; the committed
    # snapshot after the turn carries the real one.
    assert patches[0]["doc"]["version"] == {"index": 0, "count": 1}
    (snapshot_evt,) = [e for e in events if e["type"] == "doc_snapshot"]
    assert snapshot_evt["doc"]["version"] == {"index": 1, "count": 2}

    (open_evt,) = [e for e in events if e["type"] == "open_questions"]
    kinds = {i["kind"] for i in open_evt["items"]}
    assert kinds == {"tbd", "needs_input"}

    assert events[-1] == {
        "type": "turn_complete",
        "stop_reason": "end_turn",
        "usage": {},
    }

    # History: user, assistant(tool_use), user(tool_result), assistant.
    history = sessions.get_session().history
    assert [m["role"] for m in history] == ["user", "assistant", "user", "assistant"]
    tool_use = history[1]["content"][-1]
    assert tool_use["type"] == "tool_use" and tool_use["name"] == "apply_spec_edits"
    tool_result = history[2]["content"][0]
    assert tool_result["tool_use_id"] == tool_use["id"]
    assert "outline" in tool_result["content"]

    # The continuation request carried the tool result back.
    second_request = fake.messages.requests[1]
    assert second_request["messages"][-1]["content"][0]["type"] == "tool_result"

    # One committed version for the turn.
    store = sessions.get_session().doc
    assert len(store.versions) == 2 and store.index == 1


def test_invalid_edit_batch_becomes_tool_error_not_turn_failure(monkeypatch):
    fake = FakeClient(
        [
            tool_turn([], {"edits": [{"action": "delete", "target_id": "zzz"}]}),
            text_turn(["Let me fix that."]),
        ]
    )
    _patch_client(monkeypatch, fake)

    resp = _client().post("/api/chat", json={"message": "go"})
    events = _parse_sse(resp.text)

    assert [e for e in events if e["type"] == "doc_patch"] == []
    assert [e for e in events if e["type"] == "open_questions"] == []
    assert events[-1]["type"] == "turn_complete"

    history = sessions.get_session().history
    tool_result = history[2]["content"][0]
    assert tool_result["is_error"] is True
    assert "rejected" in tool_result["content"]

    store = sessions.get_session().doc
    assert store.doc.is_empty() and len(store.versions) == 1


def test_failure_mid_continuation_rolls_everything_back(monkeypatch):
    fake = FakeClient(
        [
            tool_turn(["Working. "], _SEED_EDITS),
            RuntimeError("kaput"),
        ]
    )
    _patch_client(monkeypatch, fake)

    resp = _client().post("/api/chat", json={"message": "go"})
    events = _parse_sse(resp.text)

    # The doc_patch streamed optimistically, but the turn failed…
    assert any(e["type"] == "doc_patch" for e in events)
    assert events[-1] == {"type": "error", "message": "Unexpected error: kaput"}

    # …so nothing stuck: history untouched, document rolled back.
    assert sessions.get_session().history == []
    store = sessions.get_session().doc
    assert store.doc.is_empty()
    assert len(store.versions) == 1 and store.index == 0


def test_client_disconnect_mid_turn_rolls_back(monkeypatch):
    from backend.llm.conversation import stream_user_turn

    fake = FakeClient(
        [tool_turn(["Working. "], _SEED_EDITS), text_turn(["never reached"])]
    )
    _patch_client(monkeypatch, fake)
    session = sessions.get_session()

    gen = stream_user_turn(session, "Start 21 13 13")
    saw_patch = False
    for event in gen:
        if event["type"] == "doc_patch":
            saw_patch = True
            break
    assert saw_patch
    # The SSE consumer goes away (browser reload / fetch abort): the
    # generator is closed at the yield, which except-clauses cannot see.
    gen.close()

    assert session.history == []
    assert session.doc.doc.is_empty()
    assert len(session.doc.versions) == 1
    # A fresh turn starts from clean state (no orphaned backup adopted).
    assert session.doc._turn_backup is None


def test_session_reset_mid_turn_discards_zombie_turn(monkeypatch):
    from backend.llm.conversation import stream_user_turn

    fake = FakeClient(
        [tool_turn(["Round one. "], _SEED_EDITS), text_turn(["Round two."])]
    )
    _patch_client(monkeypatch, fake)
    session = sessions.get_session()

    events = []
    gen = stream_user_turn(session, "Start 21 13 13")
    for event in gen:
        events.append(event)
        if event["type"] == "doc_patch":
            # "New session" lands between continuation rounds.
            sessions.reset_session()
    assert events[-1]["type"] == "error"
    assert "reset" in events[-1]["message"]

    # The fresh session stayed exactly fresh.
    assert session.history == []
    assert session.doc.doc.is_empty()
    assert len(session.doc.versions) == 1 and session.doc.index == 0


def test_max_tokens_mid_tool_use_does_not_wedge_history(monkeypatch):
    fake = FakeClient(
        [tool_turn(["Partial draft"], _SEED_EDITS, stop_reason="max_tokens")]
    )
    _patch_client(monkeypatch, fake)
    client = _client()

    resp = client.post("/api/chat", json={"message": "go"})
    events = _parse_sse(resp.text)
    assert events[-1] == {
        "type": "turn_complete",
        "stop_reason": "max_tokens",
        "usage": {},
    }
    # The unexecuted tool call never touched the doc and is not in history
    # (a dangling tool_use would invalidate every later request).
    assert sessions.get_session().doc.doc.is_empty()
    history = sessions.get_session().history
    assert [m["role"] for m in history] == ["user", "assistant"]
    assert all(b["type"] == "text" for b in history[1]["content"])

    # The next turn goes through cleanly on the committed history.
    fake2 = FakeClient([text_turn(["Continuing."])])
    _patch_client(monkeypatch, fake2)
    resp = client.post("/api/chat", json={"message": "continue"})
    assert _parse_sse(resp.text)[-1]["type"] == "turn_complete"


def test_tool_round_exhaustion_is_a_safe_failure(monkeypatch):
    from backend.llm.conversation import MAX_TOOL_ROUNDS

    turns = [
        tool_turn(
            [f"round {i} "],
            {"edits": [{"action": "add_article", "target_id": "pt1", "text": f"A{i}"}]},
            tool_id=f"toolu_round_{i}",
        )
        for i in range(MAX_TOOL_ROUNDS)
    ]
    fake = FakeClient(turns)
    _patch_client(monkeypatch, fake)

    resp = _client().post("/api/chat", json={"message": "go"})
    events = _parse_sse(resp.text)

    # Every round patched optimistically, then the turn failed as a unit.
    assert len([e for e in events if e["type"] == "doc_patch"]) == MAX_TOOL_ROUNDS
    assert events[-1]["type"] == "error"
    assert "tool rounds" in events[-1]["message"]
    assert sessions.get_session().history == []
    store = sessions.get_session().doc
    assert store.doc.is_empty() and len(store.versions) == 1


# ---------------------------------------------------------------------------
# Document endpoints
# ---------------------------------------------------------------------------


def test_doc_snapshot_undo_redo_endpoints(monkeypatch):
    client = _client()

    empty = client.get("/api/doc").json()
    assert empty["doc"]["version"] == {"index": 0, "count": 1}
    assert empty["open_questions"] == []

    _seed_doc_via_chat(client, monkeypatch)
    payload = client.get("/api/doc").json()
    assert payload["doc"]["section"]["number"] == "21 13 13"
    assert payload["doc"]["version"] == {"index": 1, "count": 2}
    assert len(payload["open_questions"]) == 2

    undone = client.post("/api/doc/undo")
    assert undone.status_code == 200
    assert undone.json()["doc"]["section"]["number"] == ""
    assert undone.json()["open_questions"] == []

    assert client.post("/api/doc/undo").status_code == 409

    redone = client.post("/api/doc/redo")
    assert redone.status_code == 200
    assert redone.json()["doc"]["section"]["number"] == "21 13 13"
    assert client.post("/api/doc/redo").status_code == 409


def test_docx_export_smoke(monkeypatch):
    client = _client()
    _seed_doc_via_chat(client, monkeypatch)

    resp = client.get("/api/export/docx")
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml"
    )
    assert "SECTION 21 13 13" in resp.headers["content-disposition"]

    document = Document(io.BytesIO(resp.content))
    texts = [p.text for p in document.paragraphs]
    assert "SECTION 21 13 13" in texts
    assert "WET-PIPE SPRINKLER SYSTEMS" in texts
    provision = next(
        p
        for p in document.paragraphs
        if p.text == "Section includes wet-pipe systems per NFPA 13-2025."
    )
    assert provision._p.pPr is not None
    num_pr = provision._p.pPr.find(qn("w:numPr"))
    assert num_pr is not None
    assert num_pr.find(qn("w:ilvl")).get(qn("w:val")) == "0"
    assert int(num_pr.find(qn("w:numId")).get(qn("w:val"))) > 0
    assert "ASSUMPTIONS SCHEDULE" in texts

    # The assumed block is scheduled with its numbering; the TBD is an
    # open item.
    tables = document.tables
    assert len(tables) == 2
    assumed_rows = [
        (row.cells[0].text, row.cells[1].text) for row in tables[0].rows[1:]
    ]
    assert assumed_rows == [
        ("1.1.A", "Section includes wet-pipe systems per NFPA 13-2025.")
    ]
    open_rows = [row.cells[1].text for row in tables[1].rows[1:]]
    assert any("density" in t for t in open_rows)


# ---------------------------------------------------------------------------
# Project save / resume
# ---------------------------------------------------------------------------


def test_project_save_and_resume_round_trip(monkeypatch):
    client = _client()
    _seed_doc_via_chat(client, monkeypatch)

    saved = client.get("/api/project/save")
    assert saved.status_code == 200
    assert "attachment" in saved.headers["content-disposition"]
    assert ".baspec" in saved.headers["content-disposition"].lower()
    # Keep direct JSON-load coverage for legacy/P0 project files. The primary
    # download above is now the source-capable binary .baspec container.
    project = json.loads(json.dumps(sessions.project_payload(sessions.get_session())))
    assert project["kind"] == "buildaspec-project"

    client.post("/api/session/reset")
    assert sessions.get_session().doc.doc.is_empty()

    loaded = client.post("/api/project/load", json=project)
    assert loaded.status_code == 200
    data = loaded.json()
    assert data["doc"]["section"]["number"] == "21 13 13"
    assert data["doc"]["version"] == {"index": 1, "count": 2}
    assert len(data["open_questions"]) == 2
    # The transcript shows only text turns (no tool plumbing).
    assert [m["role"] for m in data["chat"]] == ["user", "assistant"]
    assert data["chat"][1]["text"] == "Drafting.\n\nDone."

    # Undo still works across the resume (full version history restored).
    assert client.post("/api/doc/undo").status_code == 200

    # History resumed in API shape (tool_use/tool_result intact).
    history = sessions.get_session().history
    assert [m["role"] for m in history] == ["user", "assistant", "user", "assistant"]


def test_a_stopped_web_turn_cannot_poison_the_saved_project(monkeypatch):
    """Save/resume after stopping mid-search, end to end over HTTP.

    A dangling ``server_tool_use`` written to a project file used to be
    permanent: every request built from that history is a 400, so the file
    was unopenable-in-practice forever after.
    """
    fake = FakeClient(
        [
            raw_turn(
                [
                    text_block("Looking that up."),
                    # A search that never returned — the shape a stop
                    # while the UI says "Searching the web…" produces.
                    SimpleNamespace(
                        type="server_tool_use",
                        id="srvtoolu_interrupted",
                        name="web_search",
                        input={"query": "NFPA 13"},
                    ),
                ],
                stop_reason="max_tokens",
            )
        ]
    )
    _patch_client(monkeypatch, fake)
    client = _client()
    assert client.post("/api/chat", json={"message": "check"}).status_code == 200

    project = json.loads(json.dumps(sessions.project_payload(sessions.get_session())))
    assert "srvtoolu_interrupted" not in json.dumps(project)

    client.post("/api/session/reset")
    assert client.post("/api/project/load", json=project).status_code == 200

    fake2 = FakeClient([text_turn(["Resumed cleanly."])])
    _patch_client(monkeypatch, fake2)
    assert client.post("/api/chat", json={"message": "carry on"}).status_code == 200
    sent = fake2.messages.requests[0]["messages"]
    assert not [
        b
        for m in sent
        for b in (m.get("content") or [])
        if isinstance(b, dict) and b.get("type") == "server_tool_use"
    ]


def test_loading_a_legacy_poisoned_project_repairs_it_and_says_so(
    monkeypatch, caplog
):
    """The other half: a file already on disk from before the fix.

    Commit-time scrubbing cannot reach it, so the load boundary repairs the
    history in memory — loudly, and without rewriting the user's file until
    they next save normally.
    """
    client = _client()
    _seed_doc_via_chat(client, monkeypatch)
    project = json.loads(json.dumps(sessions.project_payload(sessions.get_session())))

    # Hand-build the poison a pre-fix build would have written.
    project["history"].append(
        {
            "role": "assistant",
            "content": [
                {"type": "text", "text": "Let me check."},
                {
                    "type": "server_tool_use",
                    "id": "srvtoolu_legacy_dangling",
                    "name": "web_search",
                    "input": {"query": "obstruction"},
                },
            ],
        }
    )
    client.post("/api/session/reset")

    with caplog.at_level(logging.WARNING, logger="buildaspec.project"):
        assert client.post("/api/project/load", json=project).status_code == 200
    assert any(
        "unpaired server tool call" in record.getMessage()
        for record in caplog.records
    ), "the repair must never be silent"

    history = sessions.get_session().history
    assert "srvtoolu_legacy_dangling" not in json.dumps(history)
    # The rest of the turn survived — repair, not deletion.
    assert history[-1]["content"][0]["text"] == "Let me check."

    # And the next turn's outgoing request is valid, which is the whole point.
    fake = FakeClient([text_turn(["Recovered."])])
    _patch_client(monkeypatch, fake)
    assert client.post("/api/chat", json={"message": "again"}).status_code == 200
    sent = fake.messages.requests[0]["messages"]
    assert "srvtoolu_legacy_dangling" not in json.dumps(sent, default=str)


def test_project_load_rejects_garbage(monkeypatch):
    client = _client()
    resp = client.post("/api/project/load", json={"kind": "not-a-project"})
    assert resp.status_code == 400
    assert "project file" in resp.json()["error"]
    # Session untouched.
    assert sessions.get_session().history == []


# ---------------------------------------------------------------------------
# Phase 3: lint + standards over the API surface
# ---------------------------------------------------------------------------

_OVERRIDE_EDITS = {
    "edits": [
        {
            "action": "replace",
            "target_id": "sec",
            "text": "WET-PIPE SPRINKLER SYSTEMS",
            "numbering": "21 13 13",
        },
        {"action": "add_article", "target_id": "pt1", "text": "REFERENCES"},
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "text": "Comply with NFPA 13-2025 throughout.",
            "status": "confirmed",
        },
        {
            "action": "set_standard_edition",
            "target_id": "sec",
            "standard": "NFPA 13",
            "edition": "2019",
            "basis": "2021 VCC per user (Loudoun County, VA)",
        },
    ]
}


def test_health_reports_module(monkeypatch):
    client = _client()
    # The neutral default is the generic module; select fire to assert the
    # curated module's health fields.
    client.post("/api/session/reset", json={"module_id": "hyperscale_fire"})
    data = client.get("/api/health").json()
    assert data["module_id"] == "hyperscale_fire"
    assert "Fire Suppression" in data["module"]


def test_chat_turn_emits_lint_event_and_payloads_carry_standards(monkeypatch):
    fake = FakeClient(
        [tool_turn(["Recording."], _OVERRIDE_EDITS), text_turn(["Done."])]
    )
    _patch_client(monkeypatch, fake)
    client = _client()
    # NFPA 13 is a fire-module PIN, so the override records as is_override; the
    # neutral default is the unpinned generic module, so select fire first.
    client.post("/api/session/reset", json={"module_id": "hyperscale_fire"})

    resp = client.post("/api/chat", json={"message": "The AHJ is on 2021 VCC"})
    events = _parse_sse(resp.text)
    assert events[-1]["type"] == "turn_complete"

    (lint_evt,) = [e for e in events if e["type"] == "lint"]
    # The 2025 citation contradicts the just-recorded 2019 override.
    stale = [i for i in lint_evt["items"] if i["rule"] == "stale_edition"]
    assert len(stale) == 1
    assert "edition in effect is 2019" in stale[0]["message"]
    nfpa13 = next(s for s in lint_evt["standards"] if s["name"] == "NFPA 13")
    assert nfpa13["is_override"] is True and nfpa13["edition"] == "2019"

    # The override op streamed as a doc_patch touching "sec".
    patches = [e for e in events if e["type"] == "doc_patch"]
    assert any(
        op["action"] == "set_standard_edition" and op["id"] == "sec"
        for p in patches
        for op in p["ops"]
    )

    # REST snapshot carries the same lint + standards shape.
    payload = client.get("/api/doc").json()
    assert [i["rule"] for i in payload["lint"]] == ["stale_edition"]
    assert any(s["is_override"] for s in payload["standards"])
    # Every row carries the full flag set for the standards manager.
    for s in payload["standards"]:
        assert {"is_added", "is_suppressed", "reason"} <= s.keys()
    # The 2019 override is an edition change, not a user-added standard, and
    # is in effect (not excluded).
    nfpa13 = next(s for s in payload["standards"] if s["name"] == "NFPA 13")
    assert nfpa13["is_added"] is False and nfpa13["is_suppressed"] is False


def test_override_missing_basis_is_tool_error_not_turn_failure(monkeypatch):
    fake = FakeClient(
        [
            tool_turn(
                [],
                {
                    "edits": [
                        {
                            "action": "set_standard_edition",
                            "target_id": "sec",
                            "standard": "NFPA 13",
                            "edition": "2019",
                        }
                    ]
                },
            ),
            text_turn(["Let me include the basis."]),
        ]
    )
    _patch_client(monkeypatch, fake)

    resp = _client().post("/api/chat", json={"message": "record it"})
    events = _parse_sse(resp.text)
    assert events[-1]["type"] == "turn_complete"
    tool_result = sessions.get_session().history[2]["content"][0]
    assert tool_result["is_error"] is True
    assert "basis" in tool_result["content"]
    assert sessions.get_session().doc.doc.edition_overrides == {}


def test_override_survives_project_round_trip(monkeypatch):
    fake = FakeClient(
        [tool_turn(["Recording."], _OVERRIDE_EDITS), text_turn(["Done."])]
    )
    _patch_client(monkeypatch, fake)
    client = _client()
    # NFPA 13 is a fire-module PIN (so the override is is_override, not
    # is_added); the neutral default is the unpinned generic module.
    client.post("/api/session/reset", json={"module_id": "hyperscale_fire"})
    client.post("/api/chat", json={"message": "go"})

    project = json.loads(json.dumps(sessions.project_payload(sessions.get_session())))
    assert project["module_id"] == "hyperscale_fire"

    client.post("/api/session/reset")
    loaded = client.post("/api/project/load", json=project).json()
    assert loaded["ok"] is True
    nfpa13 = next(s for s in loaded["standards"] if s["name"] == "NFPA 13")
    assert nfpa13["edition"] == "2019" and nfpa13["is_override"]
    assert [i["rule"] for i in loaded["lint"]] == ["stale_edition"]

    # The PROJECT CONTEXT block reflects the loaded override on the next turn.
    fake2 = FakeClient([text_turn(["Continuing."])])
    _patch_client(monkeypatch, fake2)
    client.post("/api/chat", json={"message": "continue"})
    context = request_context_text(fake2.messages.last_request)
    assert "jurisdiction-adopted override" in context
    assert "2021 VCC per user" in context


def test_undo_rolls_back_override_and_lint(monkeypatch):
    fake = FakeClient(
        [tool_turn(["Recording."], _OVERRIDE_EDITS), text_turn(["Done."])]
    )
    _patch_client(monkeypatch, fake)
    client = _client()
    client.post("/api/chat", json={"message": "go"})

    undone = client.post("/api/doc/undo").json()
    assert undone["ok"] is True
    assert undone["lint"] == []
    assert all(not s["is_override"] for s in undone["standards"])


def test_stable_system_prompt_is_cached_and_module_rendered(monkeypatch):
    fake = FakeClient([text_turn(["ok"])])
    _patch_client(monkeypatch, fake)
    client = _client()
    # This pins the CURATED module's rendered content (its catalog + NFPA
    # pins); the neutral default is now the generic module.
    client.post("/api/session/reset", json={"module_id": "hyperscale_fire"})
    client.post("/api/chat", json={"message": "hello"})
    request = fake.messages.last_request
    # The system prompt is ONLY the stable module block — everything
    # session-varying rides the PROJECT CONTEXT block in the user message.
    assert len(request["system"]) == 1
    stable = request["system"][0]
    assert stable["cache_control"] == {"type": "ephemeral", "ttl": "1h"}
    assert "21 13 13 Wet-Pipe Sprinkler Systems" in stable["text"]
    assert "Standards editions in effect" not in stable["text"]
    context = request_context_text(request)
    assert "Standards editions in effect" in context
    assert "NFPA 13: 2025" in context


# ---------------------------------------------------------------------------
# Sonnet unleashed: context splice/strip, thinking, pause_turn, caching
# ---------------------------------------------------------------------------


def test_context_block_never_fossilizes_into_history(monkeypatch):
    client = _client()
    _seed_doc_via_chat(client, monkeypatch)

    history = sessions.get_session().history
    assert history[0]["content"] == [
        {"type": "text", "text": "Start 21 13 13"}
    ]
    assert "PROJECT CONTEXT" not in json.dumps(history)

    # The next turn's request carries exactly one, current, context block
    # (with the seeded document's full text in it).
    fake2 = FakeClient([text_turn(["Continuing."])])
    _patch_client(monkeypatch, fake2)
    client.post("/api/chat", json={"message": "continue"})
    request = fake2.messages.last_request
    contexts = [
        b["text"]
        for m in request["messages"]
        if isinstance(m.get("content"), list)
        for b in m["content"]
        if isinstance(b, dict) and "PROJECT CONTEXT" in b.get("text", "")
    ]
    assert len(contexts) == 1
    assert "WET-PIPE SPRINKLER SYSTEMS" in contexts[0]
    # Full text, not the 160-char truncation: the whole seeded paragraph.
    assert "Section includes wet-pipe systems per NFPA 13-2025." in contexts[0]
    # The lint/open-item feedback loop reaches the model too.
    assert "OPEN ITEMS" in contexts[0]


def test_thinking_blocks_preserved_mid_turn_and_stripped_at_commit(monkeypatch):
    fake = FakeClient(
        [
            raw_turn(
                [
                    thinking_block(),
                    text_block("Drafting."),
                    tool_use_block("toolu_t1", "apply_spec_edits", _SEED_EDITS),
                ],
                stop_reason="tool_use",
                chunks=["Drafting."],
            ),
            text_turn(["Done."]),
        ]
    )
    _patch_client(monkeypatch, fake)

    resp = _client().post("/api/chat", json={"message": "go"})
    assert _parse_sse(resp.text)[-1]["type"] == "turn_complete"

    # The continuation request re-sent the thinking block verbatim (the
    # adaptive-thinking tool-use contract)…
    continuation = fake.messages.requests[1]
    assistant = [
        m for m in continuation["messages"] if m["role"] == "assistant"
    ][-1]
    assert assistant["content"][0]["type"] == "thinking"
    assert assistant["content"][0]["signature"] == "sig-fake"

    # …and committed history dropped it (only required within the turn).
    history = sessions.get_session().history
    assert "thinking" not in {
        b["type"] for m in history for b in m["content"]
    }
    assert not sessions.get_session().doc.doc.is_empty()


def test_pause_turn_resumes_and_emits_web_activity(monkeypatch):
    fake = FakeClient(
        [
            raw_turn(
                chat_search_blocks(
                    "NFPA 13 2025 obstruction rules", ["https://nfpa.org"]
                ),
                stop_reason="pause_turn",
            ),
            text_turn(["Verified against nfpa.org."]),
        ]
    )
    _patch_client(monkeypatch, fake)

    resp = _client().post("/api/chat", json={"message": "check that"})
    events = _parse_sse(resp.text)
    assert events[-1]["type"] == "turn_complete"

    # The search surfaced as a UI activity event.
    (search_evt,) = [e for e in events if e["type"] == "web_search"]
    assert search_evt["query"] == "NFPA 13 2025 obstruction rules"

    # The resume followed the pause_turn contract: assistant content
    # re-sent, no synthetic user turn, no tool_result.
    resumed = fake.messages.requests[1]["messages"]
    assert resumed[-1]["role"] == "assistant"
    assert resumed[-1]["content"][0]["type"] == "server_tool_use"
    # A pause that carried no container does not make one up.
    assert all("container" not in r for r in fake.messages.requests)

    # The server-tool blocks survive into committed history (they carry
    # the retrieval record), unlike thinking blocks.
    history = sessions.get_session().history
    types = {b["type"] for m in history for b in m["content"]}
    assert "server_tool_use" in types and "web_search_tool_result" in types


def test_chat_carries_the_container_through_a_turn_and_drops_it_next_turn(
    monkeypatch,
):
    """The chat loop's half of the continuation-container contract.

    Direct callers mean none is expected in practice; this is the
    defense-in-depth read, and the interesting part is the scope. The
    container belongs to ONE turn: every later round of it — a pause_turn
    resume *and* a continuation after a client tool_result — reuses the id,
    and the next user turn, which is a new conversation, starts clean.
    """
    fake = FakeClient(
        [
            raw_turn(
                chat_search_blocks("NFPA 13 obstruction rules", ["https://nfpa.org"]),
                stop_reason="pause_turn",
                container="cont_chat_1",
            ),
            tool_turn(
                [],
                {
                    "edits": [
                        {
                            "action": "replace",
                            "target_id": "sec",
                            "text": "WET-PIPE SPRINKLER SYSTEMS",
                            "numbering": "21 13 13",
                        }
                    ]
                },
            ),
            text_turn(["Recorded the section header."]),
        ]
    )
    _patch_client(monkeypatch, fake)
    client = _client()
    assert client.post("/api/chat", json={"message": "check that"}).status_code == 200

    requests = fake.messages.requests
    assert len(requests) == 3
    # Round 0 has nothing to know about yet.
    assert "container" not in requests[0]
    # Round 1 resumes the pause inside the container it started in.
    assert requests[1]["container"] == "cont_chat_1"
    # Round 2 follows a client tool_result — still the same turn, and
    # neither the tool_turn nor the tool dispatch supplied a container, so
    # this only passes if the id is retained rather than re-read per round.
    assert requests[2]["messages"][-1]["content"][0]["type"] == "tool_result"
    assert requests[2]["container"] == "cont_chat_1"

    # A second user turn is a new conversation: no inherited container.
    fake2 = FakeClient([text_turn(["Anything else?"])])
    _patch_client(monkeypatch, fake2)
    assert client.post("/api/chat", json={"message": "next"}).status_code == 200
    assert "container" not in fake2.messages.requests[0]

    # The id is a request argument and nothing more: it never reaches the
    # cached prefix, the conversation, committed history, or the saved
    # project file.
    for request in requests:
        assert "cont_chat_1" not in str(request["system"])
        assert "cont_chat_1" not in str(request["tools"])
        assert "cont_chat_1" not in str(request["messages"])
    assert "cont_chat_1" not in str(sessions.get_session().history)
    saved = client.get("/api/project/save")
    assert saved.status_code == 200
    assert "cont_chat_1" not in saved.text


# ---------------------------------------------------------------------------
# Rolling committed-history cache breakpoint (Chunk 4.2)
# ---------------------------------------------------------------------------
#
# A tail breakpoint alone cannot cache across turns: its entry is keyed on a
# prefix ending in that turn's PROJECT CONTEXT, and commit strips exactly
# those bytes. The committed-history boundary is keyed on the stripped form
# every later turn re-sends, so each turn's entry is a byte-prefix of the
# next turn's request.


def _one_turn_request(client, monkeypatch, message: str) -> dict:
    """Send one chat turn; return the request the model actually received."""
    fake = FakeClient([text_turn(["ok"])])
    _patch_client(monkeypatch, fake)
    client.post("/api/chat", json={"message": message})
    return fake.messages.last_request


def _marked_messages(request: dict) -> list[int]:
    """Indexes of messages carrying a cache breakpoint, in order."""
    return sorted(
        index
        for index, message in enumerate(request["messages"])
        for block in (message.get("content") or [])
        if isinstance(block, dict) and "cache_control" in block
    )


def _cache_controls(request: dict) -> list[dict]:
    """Every breakpoint value in the request, system block included."""
    blocks = [
        block
        for message in request["messages"]
        for block in (message.get("content") or [])
    ] + list(request["system"])
    return [
        block["cache_control"]
        for block in blocks
        if isinstance(block, dict) and "cache_control" in block
    ]


def _unannotated(messages: list) -> str:
    """Message bytes with every breakpoint annotation removed."""
    stripped = json.loads(json.dumps(messages))
    for message in stripped:
        for block in message.get("content") or []:
            if isinstance(block, dict):
                block.pop("cache_control", None)
    return json.dumps(stripped, sort_keys=True)


def test_the_committed_history_breakpoint_rolls_forward_each_turn(monkeypatch):
    client = _client()

    first = _one_turn_request(client, monkeypatch, "one")
    # Nothing is committed yet, so the tail is the only breakpoint there is
    # to place — one message, marked once.
    assert len(first["messages"]) == 1
    assert _marked_messages(first) == [0]

    second = _one_turn_request(client, monkeypatch, "two")
    # History is now [user, assistant]: the boundary marks the assistant
    # reply that closes it, and the tail marks the new user message.
    assert len(second["messages"]) == 3
    assert _marked_messages(second) == [1, 2]

    third = _one_turn_request(client, monkeypatch, "three")
    assert len(third["messages"]) == 5
    assert _marked_messages(third) == [3, 4]

    # Three breakpoints per request (two here + the system block), inside
    # the provider's limit of four and leaving no room wasted on a
    # separate tool breakpoint — the system block already closes tools.
    assert len(_cache_controls(third)) == 3

    # Each marks the LAST block of its message; a breakpoint anywhere else
    # would cache a partial message.
    for index in _marked_messages(third):
        content = third["messages"][index]["content"]
        assert "cache_control" in content[-1]
        assert not any("cache_control" in block for block in content[:-1])


def test_a_turns_cached_prefix_is_a_byte_prefix_of_the_next_request(
    monkeypatch,
):
    """The cache-read condition, asserted directly.

    An entry is only readable if its exact bytes lead the next request. If
    this fails, every turn silently re-bills the whole conversation as
    fresh input — which is the regression this chunk exists to fix.
    """
    client = _client()
    _one_turn_request(client, monkeypatch, "one")
    second = _one_turn_request(client, monkeypatch, "two")
    third = _one_turn_request(client, monkeypatch, "three")

    def cached_prefix(request: dict) -> list:
        boundary = _marked_messages(request)[0]
        return request["messages"][: boundary + 1]

    # What turn 2's boundary wrote, and the same span of turn 3's request.
    written = _unannotated(cached_prefix(second))
    assert written == _unannotated(third["messages"][:2])

    # And turn 3's own boundary strictly extends it — the increment is the
    # newest exchange, not the whole conversation.
    extended = _unannotated(cached_prefix(third))
    assert extended != written
    assert _unannotated(third["messages"][:2]) == written


def test_every_breakpoint_in_a_request_shares_one_ttl(monkeypatch):
    """Mixed TTLs are a nonretryable 400, not a degraded cache.

    The provider requires longer-lived entries to precede shorter-lived
    ones in tools -> system -> messages order. The fakes accept any request
    dict, so nothing but this assertion catches a mixed-TTL request before
    a provider does.
    """
    client = _client()
    _one_turn_request(client, monkeypatch, "one")
    request = _one_turn_request(client, monkeypatch, "two")

    controls = _cache_controls(request)
    assert controls, "a chat request must carry breakpoints"
    assert all(control == {"type": "ephemeral", "ttl": "1h"} for control in controls)
    assert len({json.dumps(c, sort_keys=True) for c in controls}) == 1


def test_continuation_rounds_keep_their_own_tail_breakpoint(monkeypatch):
    """A tool round extends the previous round's entry rather than rewriting."""
    client = _client()
    fake = FakeClient(
        [
            tool_turn(
                ["Drafting. "],
                {
                    "edits": [
                        {
                            "action": "add_article",
                            "target_id": "pt1",
                            "text": "SUMMARY",
                        }
                    ]
                },
            ),
            text_turn(["Done."]),
        ]
    )
    _patch_client(monkeypatch, fake)
    client.post("/api/chat", json={"message": "draft it"})

    # The second round's request ends with the tool_result user message,
    # and that message carries the tail breakpoint.
    request = fake.messages.last_request
    assert _marked_messages(request)[-1] == len(request["messages"]) - 1
    assert all(
        control == {"type": "ephemeral", "ttl": "1h"}
        for control in _cache_controls(request)
    )


def test_no_breakpoint_survives_into_history_or_a_saved_project(monkeypatch):
    client = _client()
    _seed_doc_via_chat(client, monkeypatch)
    _one_turn_request(client, monkeypatch, "another")

    assert "cache_control" not in json.dumps(sessions.get_session().history)
    saved = client.get("/api/project/save")
    assert saved.status_code == 200
    assert "cache_control" not in saved.text


def test_the_cache_ttl_setting_validates_and_degrades_loudly(monkeypatch, caplog):
    """An unsupported TTL is a 400 on every request, so it must not pass through."""
    from backend.settings import _cache_ttl_env

    monkeypatch.delenv("BUILD_A_SPEC_CHAT_CACHE_TTL", raising=False)
    assert _cache_ttl_env("BUILD_A_SPEC_CHAT_CACHE_TTL", "1h") == "1h"

    for supported in ("5m", "1h"):
        monkeypatch.setenv("BUILD_A_SPEC_CHAT_CACHE_TTL", supported)
        assert _cache_ttl_env("BUILD_A_SPEC_CHAT_CACHE_TTL", "1h") == supported

    monkeypatch.setenv("BUILD_A_SPEC_CHAT_CACHE_TTL", "7d")
    with caplog.at_level(logging.WARNING, logger="buildaspec.settings"):
        assert _cache_ttl_env("BUILD_A_SPEC_CHAT_CACHE_TTL", "1h") == "1h"
    # Silent degradation would leave an operator believing their override
    # took effect; the warning is the whole point of the fallback.
    assert "7d" in caplog.text


def test_the_history_boundary_fails_safe_if_sanitizing_ever_moves_messages():
    """Index arithmetic is checked, not trusted.

    Sanitization replaces messages positionally today. If that ever
    changes, dropping the extra breakpoint costs one cache read; guessing
    an index would annotate the wrong message and split the prefix in the
    wrong place.
    """
    from backend.llm.conversation import _committed_history_boundary

    assert _committed_history_boundary(2, 3, 3) == 1
    assert _committed_history_boundary(4, 5, 5) == 3
    # No committed history yet -> tail only.
    assert _committed_history_boundary(0, 1, 1) == -1
    # Count changed under us -> refuse to guess.
    assert _committed_history_boundary(2, 3, 2) == -1
    assert _committed_history_boundary(2, 3, 4) == -1


def test_sanitizing_a_request_never_adds_or_drops_a_message():
    """The invariant the committed-history boundary index rests on."""
    from backend.research.resend_sanitizer import sanitize_messages_for_resend

    cases = [
        [{"role": "user", "content": [{"type": "text", "text": "hi"}]}],
        [
            {"role": "user", "content": [{"type": "text", "text": "hi"}]},
            {"role": "assistant", "content": [{"type": "text", "text": "yes"}]},
            {"role": "user", "content": [{"type": "text", "text": "more"}]},
        ],
        # An assistant message whose only block is an unpaired server tool
        # use: emptied, then refilled with a placeholder — never removed.
        [
            {"role": "user", "content": [{"type": "text", "text": "search"}]},
            {
                "role": "assistant",
                "content": [
                    {
                        "type": "server_tool_use",
                        "id": "srvtoolu_orphan",
                        "name": "web_search",
                        "input": {},
                    }
                ],
            },
            {"role": "user", "content": [{"type": "text", "text": "and?"}]},
        ],
    ]
    for messages in cases:
        assert len(sanitize_messages_for_resend(messages)) == len(messages)


def test_tail_cache_breakpoint_rides_requests_not_history(monkeypatch):
    client = _client()
    _seed_doc_via_chat(client, monkeypatch)

    fake2 = FakeClient([text_turn(["ok"])])
    _patch_client(monkeypatch, fake2)
    client.post("/api/chat", json={"message": "continue"})

    # The request's final content block carries the incremental breakpoint…
    request = fake2.messages.last_request
    tail = request["messages"][-1]["content"][-1]
    assert tail["cache_control"] == {"type": "ephemeral", "ttl": "1h"}
    # …the stable system block carries the other…
    assert request["system"][0]["cache_control"] == {"type": "ephemeral", "ttl": "1h"}
    # …and stored history carries none (breakpoints are per-request).
    assert "cache_control" not in json.dumps(sessions.get_session().history)
