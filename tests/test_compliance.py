"""Compliance audit tests: the controlling-set trust model, payload
normalization, retry/failure paths, and the API lifecycle."""
from __future__ import annotations

import json
import time

from fastapi.testclient import TestClient

from backend.app import create_app
from backend import sessions
from backend.compliance.checker import (
    build_audit_system_prompt,
    build_audit_user_message,
    render_profile_block,
)
from backend.research import RequirementsProfile
from backend.spec_modules import DEFAULT_MODULE
from tests.fakes import (
    FakeClient,
    SequencedFakeClient,
    research_response,
    text_turn,
    tool_turn,
)
from tests.test_research_api import (
    _parse_sse,
    _patch_chat_client,
    _patch_research_client,
    _record_profile,
    _wait_terminal,
)
from tests.test_research_engine import _item, _scripts


def _client() -> TestClient:
    return TestClient(create_app())


def _profile_with(items: list[dict]) -> RequirementsProfile:
    return RequirementsProfile.from_dict(
        {
            "items": items,
            "dimension_statuses": [
                {"dimension_id": "governing_codes", "status": "completed"}
            ],
            "research_date": "2026-07-21",
            "project": {
                "city": "Ashburn",
                "state_or_province": "VA",
                "country": "US",
                "client_name": "ExampleCo",
            },
        }
    )


_GROUNDED = {
    "item_id": "r-aaa111aaa111",
    "dimension_id": "governing_codes",
    "topic": "code",
    "category": "governing_code",
    "requirement": "NFPA 13-2019 governs sprinkler installation.",
    "accepted_sources": ["https://a.gov"],
    "grounded": True,
    "confidence": 0.9,
    "actionability": "spec_requirement",
}
_UNVERIFIED = {
    "item_id": "r-bbb222bbb222",
    "dimension_id": "governing_codes",
    "topic": "owner",
    "category": "client_standard",
    "requirement": "Owner reportedly requires nitrogen inerting.",
    "grounded": False,
    "confidence": 0.3,
    "actionability": "spec_requirement",
}
_PROCESS = {
    "item_id": "r-ccc333ccc333",
    "dimension_id": "governing_codes",
    "topic": "fees",
    "category": "ahj_requirement",
    "requirement": "Plan review takes three weeks.",
    "accepted_sources": ["https://b.gov"],
    "grounded": True,
    "confidence": 0.8,
    "actionability": "process_advisory",
}


def test_profile_block_separates_controlling_from_unverified():
    block = render_profile_block(_profile_with([_GROUNDED, _UNVERIFIED, _PROCESS]))
    assert "CONTROLLING REQUIREMENTS" in block
    assert "r-aaa111aaa111" in block
    assert "NOT INDEPENDENTLY VERIFIED" in block
    assert "r-bbb222bbb222" in block
    # Process advisories are excluded entirely.
    assert "r-ccc333ccc333" not in block


def test_prompts_carry_persona_and_draft_ids():
    profile = _profile_with([_GROUNDED])
    system = build_audit_system_prompt(DEFAULT_MODULE)
    assert DEFAULT_MODULE.compliance_persona in system
    assert "submit_compliance_audit" in system

    session = sessions.get_session()
    session.doc.begin_turn()
    session.doc.apply_edits(
        [
            {"action": "add_article", "target_id": "pt1", "text": "SUMMARY"},
            {
                "action": "add_paragraph",
                "target_id": "pt1.a1",
                "text": "Comply with NFPA 13-2019.",
            },
        ]
    )
    session.doc.commit_turn()
    message = build_audit_user_message(session.doc.doc, profile)
    assert "[id: pt1.a1.p1]" in message
    assert "Comply with NFPA 13-2019." in message


def _audit_payload(coverage=None, findings=None) -> dict:
    return {
        "summary": "Looks broadly compliant.",
        "coverage": coverage or [],
        "findings": findings or [],
    }


def test_audit_normalization_enforces_the_controlling_set(monkeypatch):
    """Coverage for unknown/unverified ids drops; skipped controlling ids
    become 'unclear' — an unaudited requirement never looks audited."""
    from types import SimpleNamespace

    from backend.compliance.checker import run_compliance_audit
    from tests.fakes import tool_use_block, usage

    payload = _audit_payload(
        coverage=[
            {
                "requirement_id": "r-aaa111aaa111",
                "status": "represented",
                "evidence_quote": "Comply with NFPA 13-2019.",
                "element_id": "pt1.a1.p1",
                "note": "Edition matches.",
            },
            {  # Unverified item — must be dropped.
                "requirement_id": "r-bbb222bbb222",
                "status": "missing",
                "evidence_quote": None,
                "element_id": None,
                "note": "should not appear",
            },
            {  # Unknown id — dropped.
                "requirement_id": "r-zzz999zzz999",
                "status": "contradicted",
                "evidence_quote": None,
                "element_id": None,
                "note": "bogus",
            },
        ],
        findings=[
            {
                "severity": "HIGH",  # normalizes to lowercase
                "requirement_id": "r-bbb222bbb222",
                "element_id": None,
                "issue": "Confirm nitrogen inerting with the owner.",
                "suggestion": "Submit an RFI.",
            },
            {"severity": "medium", "issue": "", "requirement_id": None,
             "element_id": None, "suggestion": None},  # empty issue drops
        ],
    )

    class _OneShotClient:
        def __init__(self):
            self.messages = self

        def stream(self, **_request):
            response = SimpleNamespace(
                content=[
                    tool_use_block(
                        "toolu_audit", "submit_compliance_audit", payload
                    )
                ],
                stop_reason="tool_use",
                usage=usage(),
            )

            class _Ctx:
                def __enter__(self_inner):
                    return self_inner

                def __exit__(self_inner, *exc):
                    return False

                def get_final_message(self_inner):
                    return response

            return _Ctx()

    profile = _profile_with([_GROUNDED, _UNVERIFIED, _PROCESS])
    # Two controlling ids: the grounded spec_requirement only (process is
    # excluded, unverified is not controlling).
    session = sessions.get_session()
    session.doc.begin_turn()
    session.doc.apply_edits(
        [{"action": "add_article", "target_id": "pt1", "text": "SUMMARY"}]
    )
    session.doc.commit_turn()

    result = run_compliance_audit(
        session.doc.doc,
        profile,
        DEFAULT_MODULE,
        _OneShotClient(),
        model="claude-sonnet-5",
        max_tokens=2048,
    )
    assert [c["requirement_id"] for c in result["coverage"]] == [
        "r-aaa111aaa111"
    ]
    assert result["coverage"][0]["status"] == "represented"
    assert len(result["findings"]) == 1
    assert result["findings"][0]["severity"] == "high"
    assert result["parse_source"] == "structured"


def test_audit_requires_controlling_items():
    import pytest

    from backend.compliance.checker import (
        ComplianceAuditError,
        run_compliance_audit,
    )

    profile = _profile_with([_UNVERIFIED, _PROCESS])  # nothing controlling
    with pytest.raises(ComplianceAuditError, match="no grounded"):
        run_compliance_audit(
            sessions.get_session().doc.doc,
            profile,
            DEFAULT_MODULE,
            object(),
            model="claude-sonnet-5",
            max_tokens=2048,
        )


def _seed_doc_and_research(client: TestClient, monkeypatch) -> None:
    _record_profile(client, monkeypatch)
    draft = FakeClient(
        [
            tool_turn(
                ["Drafting."],
                {
                    "edits": [
                        {
                            "action": "add_article",
                            "target_id": "pt1",
                            "text": "SUMMARY",
                        },
                        {
                            "action": "add_paragraph",
                            "target_id": "pt1.a1",
                            "text": "Comply with the 2021 VCC.",
                        },
                    ]
                },
            ),
            text_turn(["Done."]),
        ]
    )
    _patch_chat_client(monkeypatch, draft)
    client.post("/api/chat", json={"message": "draft"})
    _patch_research_client(
        monkeypatch,
        SequencedFakeClient(
            _scripts(
                governing_codes=[
                    research_response(
                        items=[_item("2021 VCC governs.", ["https://a.gov"])],
                        searched_urls=["https://a.gov"],
                    )
                ]
            )
        ),
    )
    client.post("/api/research/start")
    _wait_terminal(client)


def test_audit_api_lifecycle_and_export_and_round_trip(monkeypatch):
    client = _client()

    # Gates: no research yet.
    assert client.post("/api/audit/start").status_code == 400

    _seed_doc_and_research(client, monkeypatch)

    from types import SimpleNamespace

    from tests.fakes import tool_use_block, usage

    grounded_id = None
    snapshot = client.get("/api/research/status").json()
    grounded_id = snapshot["profile"]["items"][0]["item_id"]

    payload = _audit_payload(
        coverage=[
            {
                "requirement_id": grounded_id,
                "status": "represented",
                "evidence_quote": "Comply with the 2021 VCC.",
                "element_id": "pt1.a1.p1",
                "note": "Stated in the summary.",
            }
        ]
    )

    class _AuditClient:
        def __init__(self):
            self.messages = self

        def stream(self, **_request):
            response = SimpleNamespace(
                content=[
                    tool_use_block(
                        "toolu_audit", "submit_compliance_audit", payload
                    )
                ],
                stop_reason="tool_use",
                usage=usage(),
            )

            class _Ctx:
                def __enter__(self_inner):
                    return self_inner

                def __exit__(self_inner, *exc):
                    return False

                def get_final_message(self_inner):
                    return response

            return _Ctx()

    monkeypatch.setattr("backend.app.get_client", lambda: _AuditClient())
    assert client.post("/api/audit/start").json()["ok"] is True

    deadline = time.monotonic() + 5
    while time.monotonic() < deadline:
        status = client.get("/api/audit/status").json()
        if status["status"] in ("complete", "failed"):
            break
        time.sleep(0.05)
    assert status["status"] == "complete"
    assert status["result"]["coverage"][0]["status"] == "represented"
    assert status["result"]["version_index"] == session_version(client)

    # Export carries the compliance closing section.
    import io

    from docx import Document

    resp = client.get("/api/export/docx")
    texts = [p.text for p in Document(io.BytesIO(resp.content)).paragraphs]
    assert "COMPLIANCE AUDIT SUMMARY" in texts

    # Project round-trip restores the audit.
    project = json.loads(client.get("/api/project/save").content)
    assert project["audit_result"]["coverage"]
    client.post("/api/session/reset")
    assert client.get("/api/audit/status").json()["status"] == "idle"
    client.post("/api/project/load", json=project)
    restored = client.get("/api/audit/status").json()
    assert restored["status"] == "complete"
    assert restored["result"]["coverage"][0]["requirement_id"] == grounded_id


def session_version(client) -> int:
    return client.get("/api/doc").json()["doc"]["version"]["index"]


def test_audit_user_message_carries_discipline_only_when_stated():
    # Batch 9: the deprecated audit path threads the session discipline so
    # the generic compliance_persona's "discipline stated for the session"
    # reference isn't left dangling. Curated audits stay byte-identical.
    from backend.compliance.checker import build_audit_user_message
    from backend.research import RequirementsProfile
    from backend.spec_doc.model import SpecSection

    section = SpecSection()
    profile = RequirementsProfile(items=[])

    without = build_audit_user_message(section, profile)
    assert "<project_discipline>" not in without
    assert build_audit_user_message(section, profile, "") == without

    with_d = build_audit_user_message(section, profile, "Electrical")
    assert "<project_discipline>\nElectrical\n</project_discipline>" in with_d
