"""Developer-tools diagnostics: logging, crash capture, endpoints, bundle.

Hermetic like everything else: ``conftest.py`` sets ``BUILD_A_SPEC_LOG=0``
and ``BUILD_A_SPEC_TRACE=0`` for the suite; these tests opt back in with
tmp directories through the two fixtures below and tear the process-global
logging state down again (``diagnostics.reset_for_tests``).
"""
from __future__ import annotations

import hashlib
import json
import logging
import os
import time
import zipfile
from io import BytesIO

import pytest
from fastapi.testclient import TestClient

from backend import diagnostics, sessions
from backend.app import create_app
from backend.tracing import capture, config as trace_config
from backend.tracing import recorder as recorder_module
from backend.tracing.recorder import set_recorder
from backend.tracing.redaction import scrub_data


@pytest.fixture
def log_env(monkeypatch, tmp_path):
    """Opt logging back in against a tmp dir; fully torn down after."""
    monkeypatch.setenv(diagnostics.ENV_LOG, "1")
    monkeypatch.delenv(diagnostics.ENV_LOG_LEVEL, raising=False)
    log_dir = tmp_path / "logs"
    monkeypatch.setenv(diagnostics.ENV_LOG_DIR, str(log_dir))
    yield log_dir
    diagnostics.reset_for_tests()


@pytest.fixture
def trace_env(monkeypatch, tmp_path):
    """Opt tracing back in against a tmp dir (the test_tracing pattern)."""
    monkeypatch.setenv(trace_config.ENV_TRACE, "1")
    trace_dir = tmp_path / "traces"
    monkeypatch.setenv(trace_config.ENV_TRACE_DIR, str(trace_dir))
    set_recorder(None)
    yield trace_dir
    rec = recorder_module.get_recorder()
    if rec is not None:
        rec.stop()
    set_recorder(None)


def _wait_events(predicate, timeout=3.0):
    """Poll the current recorder's events.jsonl until predicate holds.

    Per-line flush makes the file readable while the writer thread is
    live, but the queue drains asynchronously — hence the poll.
    """
    rec = recorder_module.get_recorder()
    assert rec is not None, "expected a live recorder"
    path = rec.trace_dir / "events.jsonl"
    deadline = time.monotonic() + timeout
    events: list[dict] = []
    while time.monotonic() < deadline:
        if path.exists():
            events = [
                json.loads(line)
                for line in path.read_text(encoding="utf-8").splitlines()
                if line.strip()
            ]
            if predicate(events):
                return events
        time.sleep(0.02)
    return events


def _make_log_run(
    root,
    token: str,
    *,
    timestamp: float,
    pid: int,
    clean: bool,
    payload_bytes: int = 16,
):
    run_id = "process-" + token * 32
    run_dir = root / run_id
    run_dir.mkdir(parents=True)
    marker = {
        "run_id": run_id,
        "started_at": timestamp - 1,
        "updated_at": timestamp,
        "pid": pid,
        "clean": clean,
        "state": "clean_shutdown" if clean else "active",
    }
    if clean:
        marker["ended_at"] = timestamp
    marker_path = run_dir / diagnostics.RUN_MARKER_FILENAME
    marker_path.write_text(json.dumps(marker), encoding="utf-8")
    log_path = run_dir / diagnostics.LOG_FILENAME
    log_path.write_bytes(b"x" * payload_bytes)
    os.utime(marker_path, (timestamp, timestamp))
    os.utime(log_path, (timestamp, timestamp))
    os.utime(run_dir, (timestamp, timestamp))
    return run_dir


# --- logging bootstrap ------------------------------------------------------


def test_init_logging_writes_to_the_configured_dir_and_is_idempotent(log_env):
    path = diagnostics.init_logging(force=True)
    assert path is not None and path.parent.parent == log_env
    assert path.parent == diagnostics.current_log_dir()
    assert path.parent.name.startswith("process-")
    logging.getLogger("buildaspec.test").info("hello file")
    assert path.exists()
    written = path.read_text(encoding="utf-8")
    assert "hello file" in written
    assert f"pid={os.getpid()}" in written
    assert "version=" in written
    assert "trace=" in written

    handler_count = sum(
        1
        for h in logging.getLogger().handlers
        if isinstance(h, logging.handlers.RotatingFileHandler)
    )
    assert diagnostics.init_logging() == path  # idempotent
    assert (
        sum(
            1
            for h in logging.getLogger().handlers
            if isinstance(h, logging.handlers.RotatingFileHandler)
        )
        == handler_count
    )
    health = diagnostics.snapshot()["logging"]
    assert health["initialization_attempted"] is True
    assert health["initialization_succeeded"] is True
    assert health["handler_attached"] is True
    assert health["capture_active"] is True


def test_logging_startup_failure_is_visible_without_exposing_message(
    log_env, monkeypatch
):
    class SensitiveFailureHandler:
        def __init__(self, *args, **kwargs):
            raise PermissionError("private-path-should-not-be-reported")

    monkeypatch.setattr(
        logging.handlers, "RotatingFileHandler", SensitiveFailureHandler
    )
    assert diagnostics.init_logging(force=True) is None

    health = diagnostics.snapshot()["logging"]
    assert health["initialization_attempted"] is True
    assert health["initialization_succeeded"] is False
    assert health["handler_attached"] is False
    assert health["capture_active"] is False
    assert health["initialization_failure_type"] == "PermissionError"
    assert health["initialization_failure_count"] == 1
    assert "private-path" not in json.dumps(health)


def test_file_log_redacts_credentials_in_messages_and_tracebacks(log_env):
    path = diagnostics.init_logging(force=True)
    assert path is not None
    secret = "sk-ant-abcdefghijklmnop"
    logger = logging.getLogger("buildaspec.redaction-test")
    logger.error("pasted credential: %s", secret)
    try:
        raise RuntimeError(f"provider rejected {secret}")
    except RuntimeError:
        logger.exception("request failed with %s", secret)

    written = path.read_text(encoding="utf-8")
    assert secret not in written
    assert written.count("<redacted>") >= 3
    assert "RuntimeError: provider rejected <redacted>" in written


def test_file_log_preserves_lone_surrogates_as_visible_escapes(log_env):
    path = diagnostics.init_logging(force=True)
    assert path is not None
    logging.getLogger("buildaspec.unicode-test").error("bad %s marker", "\ud800")

    written = path.read_text(encoding="utf-8")
    assert "bad \\ud800 marker" in written


def test_log_env_off_and_level_knobs(monkeypatch, tmp_path):
    try:
        monkeypatch.setenv(diagnostics.ENV_LOG, "0")
        assert diagnostics.init_logging(force=True) is None
        assert diagnostics.current_log_file() is None
        assert not (tmp_path / "logs").exists()

        monkeypatch.setenv(diagnostics.ENV_LOG, "1")
        monkeypatch.setenv(diagnostics.ENV_LOG_DIR, str(tmp_path / "logs"))
        monkeypatch.setenv(diagnostics.ENV_LOG_LEVEL, "INFO")
        assert diagnostics.configured_level() == logging.INFO
        monkeypatch.setenv(diagnostics.ENV_LOG_LEVEL, "garbage")
        assert diagnostics.configured_level() == logging.DEBUG
    finally:
        diagnostics.reset_for_tests()


def test_log_retention_policy_env_defaults_zero_and_invalid(monkeypatch):
    for name in (
        diagnostics.ENV_LOG_MAX_RUNS,
        diagnostics.ENV_LOG_MAX_AGE_DAYS,
        diagnostics.ENV_LOG_MAX_MIB,
    ):
        monkeypatch.delenv(name, raising=False)
    assert diagnostics.log_retention_policy().to_dict() == {
        "max_runs": diagnostics.DEFAULT_LOG_MAX_RUNS,
        "max_age_days": diagnostics.DEFAULT_LOG_MAX_AGE_DAYS,
        "max_bytes": diagnostics.DEFAULT_LOG_MAX_MIB * 1024 * 1024,
    }

    monkeypatch.setenv(diagnostics.ENV_LOG_MAX_RUNS, "0")
    monkeypatch.setenv(diagnostics.ENV_LOG_MAX_AGE_DAYS, "12")
    monkeypatch.setenv(diagnostics.ENV_LOG_MAX_MIB, "3")
    assert diagnostics.log_retention_policy().to_dict() == {
        "max_runs": 0,
        "max_age_days": 12,
        "max_bytes": 3 * 1024 * 1024,
    }

    monkeypatch.setenv(diagnostics.ENV_LOG_MAX_RUNS, "invalid")
    monkeypatch.setenv(diagnostics.ENV_LOG_MAX_AGE_DAYS, "-1")
    monkeypatch.setenv(diagnostics.ENV_LOG_MAX_MIB, "-2")
    assert diagnostics.log_retention_policy().to_dict() == {
        "max_runs": diagnostics.DEFAULT_LOG_MAX_RUNS,
        "max_age_days": diagnostics.DEFAULT_LOG_MAX_AGE_DAYS,
        "max_bytes": diagnostics.DEFAULT_LOG_MAX_MIB * 1024 * 1024,
    }


def test_live_sibling_is_not_misclassified_as_a_crash(
    log_env, monkeypatch, caplog
):
    live_pid = 515151
    live_dir = _make_log_run(
        log_env,
        "a",
        timestamp=time.time() - 60,
        pid=live_pid,
        clean=False,
    )
    monkeypatch.setattr(
        diagnostics,
        "_process_is_alive",
        lambda pid: pid in {os.getpid(), live_pid},
    )

    with caplog.at_level(logging.WARNING, logger="buildaspec"):
        diagnostics.init_logging(force=True)

    assert "did not shut down cleanly" not in caplog.text
    assert live_dir.exists()
    process = diagnostics.snapshot()["process"]
    assert process["previous_run_marker"]["present"] is False
    retention = diagnostics.snapshot()["logging"]["last_retention"]
    assert retention["protected_runs"] >= 2


def test_unclean_shutdown_marker_round_trip(log_env, caplog, monkeypatch):
    diagnostics.init_logging(force=True)
    marker_path = diagnostics.current_log_dir() / diagnostics.RUN_MARKER_FILENAME
    marker = json.loads(marker_path.read_text(encoding="utf-8"))
    assert marker["clean"] is False and marker["pid"] == os.getpid()
    assert marker["state"] == "active" and marker["run_id"].startswith(
        "process-"
    )
    process = diagnostics.snapshot()["process"]
    assert process["capture_state"] == "active_at_capture"
    assert process["current_run_marker"]["shutdown_state"] == (
        "active_at_capture"
    )
    assert process["current_run_marker"]["unclean_shutdown_evidence"] is False

    diagnostics.mark_clean_shutdown()
    marker = json.loads(marker_path.read_text(encoding="utf-8"))
    assert marker["clean"] is True and "ended_at" in marker

    # Simulate a distinct crashed launch. The current launch keeps its own
    # marker, so another live instance can never overwrite or impersonate it.
    prior_run_id = "process-" + "d" * 32
    prior_dir = log_env / prior_run_id
    prior_dir.mkdir()
    (prior_dir / diagnostics.RUN_MARKER_FILENAME).write_text(
        json.dumps(
            {
                "run_id": prior_run_id,
                "clean": False,
                "pid": 4242,
                "started_at": 1000.0,
                "updated_at": 1000.0,
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(
        diagnostics,
        "_process_is_alive",
        lambda pid: pid == os.getpid(),
    )
    with caplog.at_level(logging.WARNING, logger="buildaspec"):
        diagnostics.init_logging(force=True)
    assert "did not shut down cleanly" in caplog.text
    assert "4242" in caplog.text
    previous = diagnostics.snapshot()["process"]["previous_run_marker"]
    assert previous["pid"] == 4242
    assert previous["shutdown_state"] == "unclean_shutdown"
    assert previous["unclean_shutdown_evidence"] is True


def test_log_retention_count_protects_current_live_and_recent_crash(
    tmp_path, monkeypatch
):
    root = tmp_path / "logs"
    now = time.time()
    current = _make_log_run(
        root, "1", timestamp=now - 1000, pid=1001, clean=True
    )
    live = _make_log_run(
        root, "2", timestamp=now - 900, pid=1002, clean=True
    )
    recent_crash = _make_log_run(
        root, "3", timestamp=now - 60, pid=1003, clean=False
    )
    oldest_eligible = _make_log_run(
        root, "4", timestamp=now - 800, pid=1004, clean=True
    )
    newest_eligible = _make_log_run(
        root, "5", timestamp=now - 100, pid=1005, clean=True
    )
    legacy = root / diagnostics.LOG_FILENAME
    legacy.write_text("legacy flat log", encoding="utf-8")
    unrelated = root / "do-not-delete"
    unrelated.mkdir()
    monkeypatch.setattr(
        diagnostics, "_process_is_alive", lambda pid: pid == 1002
    )

    result = diagnostics.prune_log_runs(
        root,
        current_run_id=current.name,
        policy=diagnostics.LogRetentionPolicy(
            max_runs=3, max_age_days=0, max_bytes=0
        ),
        now=now,
    )

    assert current.exists() and live.exists() and recent_crash.exists()
    assert not oldest_eligible.exists() and not newest_eligible.exists()
    assert legacy.read_text(encoding="utf-8") == "legacy flat log"
    assert unrelated.exists()
    assert result["removed_by_reason"] == {"count": 2}
    assert result["protected_runs"] == 3
    assert result["remaining_runs"] == 3
    assert result["limits_satisfied"] is True


def test_log_retention_age_removes_dead_clean_and_stale_unclean_runs(
    tmp_path, monkeypatch
):
    root = tmp_path / "logs"
    now = time.time()
    old_clean = _make_log_run(
        root,
        "6",
        timestamp=now - 40 * 24 * 60 * 60,
        pid=2001,
        clean=True,
    )
    recent_clean = _make_log_run(
        root,
        "7",
        timestamp=now - 24 * 60 * 60,
        pid=2002,
        clean=True,
    )
    stale_crash = _make_log_run(
        root,
        "8",
        timestamp=now - 8 * 24 * 60 * 60,
        pid=2003,
        clean=False,
    )
    recent_crash = _make_log_run(
        root,
        "9",
        timestamp=now - 24 * 60 * 60,
        pid=2004,
        clean=False,
    )
    monkeypatch.setattr(diagnostics, "_process_is_alive", lambda _pid: False)

    result = diagnostics.prune_log_runs(
        root,
        current_run_id="process-" + "f" * 32,
        policy=diagnostics.LogRetentionPolicy(
            max_runs=0, max_age_days=7, max_bytes=0
        ),
        now=now,
    )

    assert not old_clean.exists() and not stale_crash.exists()
    assert recent_clean.exists() and recent_crash.exists()
    assert result["removed_by_reason"] == {"age": 2}
    assert result["protected_runs"] == 1
    assert result["limits_satisfied"] is True


def test_log_retention_byte_limit_removes_oldest_first(tmp_path, monkeypatch):
    root = tmp_path / "logs"
    now = time.time()
    oldest = _make_log_run(
        root, "a", timestamp=now - 30, pid=3001, clean=True, payload_bytes=300
    )
    middle = _make_log_run(
        root, "b", timestamp=now - 20, pid=3002, clean=True, payload_bytes=200
    )
    newest = _make_log_run(
        root, "c", timestamp=now - 10, pid=3003, clean=True, payload_bytes=100
    )
    monkeypatch.setattr(diagnostics, "_process_is_alive", lambda _pid: False)
    infos, _ignored, _scan_failed = diagnostics._scan_log_runs(root)
    assert _scan_failed is False
    sizes = {info.path: info.size_bytes for info in infos}
    budget = sum(sizes.values()) - sizes[oldest]

    result = diagnostics.prune_log_runs(
        root,
        current_run_id="process-" + "f" * 32,
        policy=diagnostics.LogRetentionPolicy(
            max_runs=0, max_age_days=0, max_bytes=budget
        ),
        now=now,
    )

    assert not oldest.exists()
    assert middle.exists() and newest.exists()
    assert result["removed_by_reason"] == {"bytes": 1}
    assert result["remaining_bytes"] <= budget
    assert result["limits_satisfied"] is True


def test_usage_count_keys_survive_the_scrub():
    """The ported bare-``token`` key pattern redacted every usage count out
    of every span — counts are not secrets; auth tokens still are."""
    scrubbed = scrub_data(
        {
            "usage": {"input_tokens": 10, "output_tokens": 5, "thinking_tokens": 2},
            "auth_token": "abc",
            "api_key": "whatever",
        }
    )
    assert scrubbed["usage"] == {
        "input_tokens": 10,
        "output_tokens": 5,
        "thinking_tokens": 2,
    }
    assert scrubbed["auth_token"] == "<redacted>"
    assert scrubbed["api_key"] == "<redacted>"


def test_credential_shaped_mapping_keys_are_redacted_without_collision():
    first = "sk-ant-dynamicmapcredential111"
    second = "sk-ant-dynamicmapcredential222"
    scrubbed = scrub_data({first: "one", second: "two"})

    assert first not in scrubbed and second not in scrubbed
    assert set(scrubbed.values()) == {"one", "two"}
    assert len(scrubbed) == 2
    assert all(str(key).startswith("<redacted-key:") for key in scrubbed)


# --- request middleware + exception handlers --------------------------------


def test_request_middleware_logs_and_traces_requests(trace_env, caplog):
    client = TestClient(create_app())
    with caplog.at_level(logging.DEBUG, logger="buildaspec.api"):
        client.get("/api/doc")
        client.get("/api/health")
    assert "GET /api/doc -> 200" in caplog.text

    events = _wait_events(
        lambda evs: all(
            any(
                e["type"] == "api_request" and e["path"] == path
                for e in evs
            )
            for path in ("/api/doc", "/api/health")
        )
    )
    doc_requests = [
        e
        for e in events
        if e["type"] == "api_request" and e["path"] == "/api/doc"
    ]
    assert doc_requests and doc_requests[0]["status"] == 200
    assert doc_requests[0]["outcome_code"] == "success"
    assert isinstance(doc_requests[0]["ms"], int)
    # Poll endpoints are log-quiet but still counted in trace coverage.
    health_requests = [
        e
        for e in events
        if e["type"] == "api_request" and e["path"] == "/api/health"
    ]
    assert health_requests and all(
        e["outcome_code"] == "success"
        for e in health_requests
    )
    assert not any(
        e["type"] == "api_request" and e["path"] == "/api/health"
        and e["status"] != 200
        for e in events
    )
    # Boot itself is recorded.
    assert any(e["type"] == "server_started" for e in events)


def test_catch_all_handler_returns_the_error_idiom_and_logs(
    monkeypatch, caplog
):
    # Flag-gated so the conftest teardown (which also calls get_session
    # while this patch is still active) is unaffected after the test body.
    state = {"boom": True}
    real_get_session = sessions.get_session

    def maybe_boom():
        if state["boom"]:
            raise RuntimeError("kaboom")
        return real_get_session()

    monkeypatch.setattr("backend.sessions.get_session", maybe_boom)
    client = TestClient(create_app(), raise_server_exceptions=False)
    with caplog.at_level(logging.WARNING, logger="buildaspec.api"):
        # Any route that reaches ``get_session`` will do — this one is
        # deliberately trivial, so the test stays about the catch-all
        # handler rather than about how a busier route acquires its state.
        resp = client.get("/api/figures")
    state["boom"] = False
    assert resp.status_code == 500
    data = resp.json()
    assert data["ok"] is False
    assert data["code"] == "internal_error"
    assert "kaboom" not in data["error"]
    assert data["request_id"] == resp.headers["X-Request-ID"]
    # The traceback is in the log, not just the response.
    assert "Unhandled error on GET /api/figures" in caplog.text
    assert "RuntimeError: kaboom" in caplog.text


def test_422_translates_to_the_error_idiom_without_echoing_input(caplog):
    client = TestClient(create_app())
    with caplog.at_level(logging.DEBUG, logger="buildaspec.api"):
        resp = client.post(
            "/api/key",
            json={"api_key": {"v": "sk-ant-fake-abcdef1234567890"}},
        )
    assert resp.status_code == 422
    data = resp.json()
    assert data["ok"] is False and data["code"] == "validation_error"
    # pydantic v2 errors() carry the submitted input — the handler must
    # build its message from loc/msg only, or /api/key echoes the key.
    assert "sk-ant-fake" not in resp.text
    assert "sk-ant-fake" not in caplog.text


# --- diagnostics endpoints --------------------------------------------------


def test_diagnostics_snapshot_shape_and_no_key_material():
    client = TestClient(create_app())
    resp = client.get("/api/diagnostics")
    assert resp.status_code == 200
    data = resp.json()
    assert data["ok"] is True
    assert data["schema_version"] == 1
    for block in ("app", "tracing", "logging", "key", "workspace", "session", "usage"):
        assert block in data
    assert data["app"]["name"] and data["app"]["version"]
    assert data["key"]["present"] is True and data["key"]["source"] == "env"
    assert data["key"]["masked"].startswith("…")
    assert "test-key-hermetic" not in resp.text
    assert data["workspace"]["scope"] == "original"
    assert data["session"]["history_len"] == 0
    assert data["session"]["doc_empty"] is True
    assert data["process"]["pid"] == os.getpid()
    assert data["process"]["uptime_seconds"] >= 0
    assert data["process"]["capture_state"] == "active_at_capture"
    assert data["session"]["document_shape"] == {
        "parts": 3,
        "articles": 0,
        "paragraphs": 0,
        "maximum_paragraph_depth": 0,
    }
    assert data["session"]["research"]["worker_alive"] is False
    assert data["session"]["audit"]["status"] == "idle"
    assert data["session"]["qc"]["worker_settled"] is True
    assert data["session"]["import"]["present"] is False
    assert data["usage"]["turns"] == 0

    # Suite defaults: logging and tracing are off — grace, not errors.
    assert client.get("/api/diagnostics/log").json()["enabled"] is False
    assert client.get("/api/diagnostics/activity").json()["enabled"] is False


def test_import_warning_snapshot_uses_codes_and_fingerprints_not_prose():
    secret_title = "PRIVATE PROJECT SECTION TITLE"
    facts = diagnostics._import_report_facts(
        {
            "warnings": [
                "This file does not look like a SectionFormat master; "
                + secret_title,
                "Unrecognized warning mentioning " + secret_title,
            ]
        }
    )
    assert facts["warning_count"] == 2
    assert facts["warning_code_counts"] == {
        "other": 1,
        "unstructured_document": 1,
    }
    assert len(facts["warning_evidence"]) == 2
    assert all(len(item["fingerprint"]) == 64 for item in facts["warning_evidence"])
    assert secret_title not in json.dumps(facts)
    assert "warnings" not in facts


def test_import_warning_snapshot_counts_beyond_the_evidence_cap():
    facts = diagnostics._import_report_facts(
        {"warnings": [f"Unrecognized warning {index}" for index in range(1100)]}
    )
    assert facts["warning_count"] == 1100
    assert facts["warning_code_counts"] == {"other": 1100}
    assert len(facts["warning_evidence"]) == 64
    assert facts["warning_evidence_truncated"] == 1036


def test_the_snapshot_says_why_an_imported_document_is_read_only():
    """A bundle must answer "everything says READ-ONLY" on its own.

    It recorded that a source package was retained but nothing about the
    permissions derived from it, so a package-wide blocker disabling the
    whole document was indistinguishable from ordinary per-element ones
    without asking the user to hover a badge.
    """
    from tests.conftest import settle_capability_sweep
    from tests.test_importer import _numbered

    from docx import Document

    def master(protected: bool) -> bytes:
        document = Document()
        document.add_paragraph("SECTION 23 05 48")
        document.add_paragraph("VIBRATION AND SEISMIC CONTROLS FOR HVAC")
        document.add_paragraph("PART 1 - GENERAL")
        _numbered(document, "SECTION INCLUDES", 0)
        _numbered(document, "Vibration isolation for HVAC equipment.", 1)
        buffer = BytesIO()
        document.save(buffer)
        raw = buffer.getvalue()
        if not protected:
            return raw
        # What many office masters ship with, and a package-wide blocker.
        source = zipfile.ZipFile(BytesIO(raw))
        out = BytesIO()
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename == "word/settings.xml":
                    cut = data.index(b">", data.index(b"<w:settings")) + 1
                    data = (
                        data[:cut]
                        + b'<w:documentProtection w:edit="readOnly" '
                        b'w:enforcement="1"/>'
                        + data[cut:]
                    )
                target.writestr(item, data)
        return out.getvalue()

    def imported_source_facts(protected: bool) -> dict:
        # An import refuses a non-empty document, and the session is a
        # process singleton — start each case from a blank one.
        sessions.get_session().reset()
        client = TestClient(create_app())
        resp = client.post(
            "/api/import/master",
            files={
                "file": (
                    "230548.docx",
                    master(protected),
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document",
                )
            },
        )
        assert resp.status_code == 200, resp.text
        settle_capability_sweep()
        session_facts = client.get("/api/diagnostics").json()["session"]
        imported = session_facts["import"]
        assert imported["present"] is True
        assert imported["filename"] == "230548.docx"
        assert imported["imported_block_count"] > 0
        assert imported["warning_count"] == sum(
            imported["warning_code_counts"].values()
        )
        assert imported["warning_count"] == len(imported["warning_evidence"])
        return session_facts["source"]

    # No source document at all: nothing to explain, and no empty scaffolding
    # implying there was.
    fresh = TestClient(create_app()).get("/api/diagnostics").json()
    assert fresh["session"]["source"]["capabilities_status"] is None

    ordinary = imported_source_facts(protected=False)
    assert ordinary["retained"] is True
    assert ordinary["capabilities_status"] == "ready"
    # Headings are never source-patched, so a clean master still records
    # blockers — the point is that they are per-element and named.
    assert "heading_change" in ordinary["edit_blockers"]
    assert "document_protection" not in ordinary["edit_blockers"]
    assert ordinary["global_edit_blockers"]["causes"] == []
    assert "heading_change" in ordinary["per_operation_edit_blockers"]
    assert ordinary["capability_operation_counts"]["denied"] > 0

    protected = imported_source_facts(protected=True)
    assert protected["capabilities_status"] == "pass_through_only"
    assert "document_protection" in protected["edit_blockers"]
    assert protected["global_edit_blockers"]["causes"] == [
        "document_protection"
    ]
    assert protected["global_edit_blockers"]["denial_counts"][
        "document_protection"
    ] == protected["edit_blockers"]["document_protection"]
    assert "document_protection" not in protected["per_operation_edit_blockers"]

    # The closed blocker vocabulary travels; provision text never does.
    for facts in (ordinary, protected):
        assert all(
            key.replace("_", "").isalnum() for key in facts["edit_blockers"]
        )
        assert "Vibration isolation" not in json.dumps(facts)


def test_log_tail_bounds_and_missing_file_grace(log_env, monkeypatch):
    diagnostics.init_logging(force=True)
    log = logging.getLogger("buildaspec.test")
    for i in range(20):
        log.info("line %d", i)

    out = diagnostics.tail_log(5)
    assert out["enabled"] is True and len(out["lines"]) == 5
    assert out["lines"][-1].endswith("line 19")
    assert out["size_bytes"] > 0

    assert len(diagnostics.tail_log(0)["lines"]) == 1  # clamps low
    assert len(diagnostics.tail_log(10**9)["lines"]) >= 20  # clamps high
    assert diagnostics.tail_log("garbage")["lines"]  # type-lenient

    # Missing file: an empty log dir (delay=True — no emit, no file).
    monkeypatch.setenv(diagnostics.ENV_LOG_DIR, str(log_env.parent / "empty"))
    out = diagnostics.tail_log(5)
    assert out["enabled"] is True
    assert out["lines"] == [] and out["size_bytes"] == 0


def test_traces_list_newest_first_with_current_flag(trace_env):
    older = trace_env / "session-aaaa1111-1000"
    older.mkdir(parents=True)
    (older / "run.json").write_text(
        json.dumps({"run_id": older.name, "started_at": 1000.0, "ended_at": 1500.0})
    )
    (older / "events.jsonl").write_text('{"ts": 1000.0}\n')
    newer = trace_env / "session-bbbb2222-2000"
    newer.mkdir()
    (newer / "run.json").write_text(
        json.dumps({"run_id": newer.name, "started_at": 2000.0, "ended_at": None})
    )
    os.utime(older, (1000, 1000))
    os.utime(newer, (2000, 2000))

    out = diagnostics.list_trace_runs()
    assert out["root"] == str(trace_env)
    assert [r["run_id"] for r in out["runs"]] == [newer.name, older.name]
    assert out["runs"][1]["files"]["events.jsonl"] > 0
    assert out["runs"][1]["started_at"] == 1000.0
    assert not any(r["current"] for r in out["runs"])

    # A live recorder marks its own run as current.
    capture.app_event("server_started")
    out = diagnostics.list_trace_runs()
    current = [r for r in out["runs"] if r["current"]]
    assert len(current) == 1
    assert current[0]["run_id"] == recorder_module.get_recorder().run_id


def test_bundle_excludes_trace_owned_by_another_live_process(log_env, trace_env):
    now = time.time()
    live_id = "session-live-sibling-3000"
    dead_id = "session-dead-prior-2000"
    for run_id, stamp, pid in (
        (live_id, now - 10, os.getpid()),
        (dead_id, now - 20, None),
    ):
        run_dir = trace_env / run_id
        run_dir.mkdir(parents=True)
        metadata = {
            "run_id": run_id,
            "started_at": stamp - 5,
            "ended_at": stamp,
        }
        if pid is not None:
            metadata["environment"] = {"pid": pid}
        (run_dir / "run.json").write_text(
            json.dumps(metadata), encoding="utf-8"
        )
        (run_dir / "events.jsonl").write_text(
            json.dumps({"ts": stamp, "type": "error", "status": 500}) + "\n",
            encoding="utf-8",
        )
        os.utime(run_dir, (stamp, stamp))

    diagnostics.init_logging(force=True)
    capture.app_event("server_started")
    inventory = diagnostics.list_trace_runs(limit=100)["runs"]
    live = next(run for run in inventory if run["run_id"] == live_id)
    assert live["owner_pid"] == os.getpid()
    assert live["owner_process_alive"] is True

    bundle_path, _filename = diagnostics.build_bundle()
    try:
        with zipfile.ZipFile(bundle_path) as zf:
            names = set(zf.namelist())
            manifest = json.loads(zf.read("bundle-manifest.json"))
    finally:
        diagnostics.unlink_quietly(bundle_path)

    assert not any(name.startswith(f"traces/{live_id}/") for name in names)
    assert any(name.startswith(f"traces/{dead_id}/") for name in names)
    assert live_id in manifest["scope"]["excluded_live_trace_run_ids"]
    assert live_id not in manifest["included_run_ids"]
    assert dead_id in manifest["included_run_ids"]


def test_recent_trace_incident_limit_uses_timestamps_not_append_order():
    current = {"ts": 10_000.0, "run_id": "current", "type": "error"}
    historical = [
        {"ts": float(index), "run_id": "prior", "type": "error"}
        for index in range(100)
    ]

    recent = diagnostics._recent_trace_incidents([current, *historical])

    assert len(recent) == diagnostics._INCIDENT_LIMIT
    assert current in recent


def test_activity_endpoint_returns_recent_events_and_open_spans(trace_env):
    capture.app_event("session_reset", module_id="generic", had_content=False)
    handle = capture.turn_start(model="claude-sonnet-5", history_len=0)
    try:
        _wait_events(
            lambda evs: any(e["type"] == "session_reset" for e in evs)
        )
        out = diagnostics.read_recent_trace_events(tail=50)
        assert out["enabled"] is True
        assert any(e["type"] == "session_reset" for e in out["events"])
        assert any(s["kind"] == "turn" for s in out["spans"])
    finally:
        capture.turn_end(handle)


def test_snapshot_surfaces_bounded_current_trace_summary_and_health(trace_env):
    capture.app_event("session_reset", module_id="generic", had_content=False)
    TestClient(create_app()).get("/api/doc")
    data = diagnostics.snapshot()["tracing"]
    assert data["run_id"] == recorder_module.get_recorder().run_id
    assert data["metadata_flush_complete"] is True
    assert data["metadata_size_bytes"] <= diagnostics._RUN_META_READ_CAP
    assert data["trace_schema_version"] == 2
    assert data["summary"]["events_total"] >= 1
    assert data["summary"]["events_by_type"]["session_reset"] >= 1
    assert data["summary"]["request_outcome_counts"]["success"] >= 1
    assert data["recorder_health"]["state"] == "running"
    assert data["recorder_health"]["thread_alive"] is True
    assert data["recorder_health"]["dropped_records"] == 0
    assert data["retention_policy"] == {
        "max_runs": trace_config.DEFAULT_TRACE_MAX_RUNS,
        "max_age_days": trace_config.DEFAULT_TRACE_MAX_AGE_DAYS,
        "max_bytes": trace_config.DEFAULT_TRACE_MAX_MIB * 1024 * 1024,
    }


def test_bundle_zip_contains_snapshot_log_and_trace_and_never_the_key(
    log_env, trace_env
):
    import tempfile
    from pathlib import Path

    diagnostics.init_logging(force=True)
    boot_nonce = "diagnostic-bootstrap-nonce-must-never-be-persisted"
    nonce_fingerprint = hashlib.sha256(boot_nonce.encode("utf-8")).hexdigest()
    diagnostics.set_runtime_server_identity(
        host="127.0.0.1", port=43123, boot_nonce=boot_nonce
    )
    identity = diagnostics.runtime_server_identity()
    assert identity["boot_nonce_fingerprint"] == nonce_fingerprint
    assert "boot_nonce" not in identity
    current_marker_path = (
        diagnostics.current_log_dir() / diagnostics.RUN_MARKER_FILENAME
    )
    current_marker_text = current_marker_path.read_text(encoding="utf-8")
    assert boot_nonce not in current_marker_text
    assert nonce_fingerprint in current_marker_text
    # Simulate a rotation produced by an older release, before file-log
    # formatting scrubbed credential-shaped substrings. Bundle assembly is a
    # second redaction boundary for those historical files.
    legacy_secret = "sk-ant-legacycredential1234"
    legacy_name = f"{diagnostics.LOG_FILENAME}.1"
    (log_env / legacy_name).write_text(
        f"2026-01-01 ERROR legacy provider failure {legacy_secret}\n",
        encoding="utf-8",
    )
    # The legacy flat marker may come from the version that persisted the raw
    # bootstrap nonce. Bundle migration sanitizes it without rewriting source.
    legacy_marker = log_env / diagnostics.RUN_MARKER_FILENAME
    legacy_marker.write_text(
        json.dumps(
            {
                "run_id": "legacy-run",
                "pid": 4444,
                "clean": True,
                "server": {"boot_nonce": boot_nonce},
            }
        ),
        encoding="utf-8",
    )
    client = TestClient(create_app())
    client.get("/api/doc")  # an api_request event + log lines

    tempdir = Path(tempfile.gettempdir())
    before = set(tempdir.glob("buildaspec-diagnostics-*.zip"))
    resp = client.get("/api/diagnostics/bundle")
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/zip"
    assert "buildaspec-diagnostics-" in resp.headers.get(
        "content-disposition", ""
    )
    assert resp.headers.get("cache-control") == "no-store"
    # Streamed from a temp file, deleted by the background task after the
    # response finishes (TestClient runs background tasks synchronously).
    assert set(tempdir.glob("buildaspec-diagnostics-*.zip")) == before

    zf = zipfile.ZipFile(BytesIO(resp.content))
    names = zf.namelist()
    assert "snapshot.json" in names
    assert "bundle-manifest.json" in names
    assert "incident-index.json" in names
    current_log_member = (
        f"logs/{diagnostics._PROCESS_RUN_ID}/{diagnostics.LOG_FILENAME}"
    )
    assert current_log_member in names
    archived_log = zf.read(current_log_member).decode("utf-8")
    assert f"trace={recorder_module.get_recorder().run_id}" in archived_log
    archived_legacy = zf.read(f"logs/{legacy_name}").decode("utf-8")
    assert legacy_secret not in archived_legacy
    assert "<redacted>" in archived_legacy
    archived_legacy_marker = zf.read(
        f"logs/{diagnostics.RUN_MARKER_FILENAME}"
    ).decode("utf-8")
    assert boot_nonce not in archived_legacy_marker
    assert nonce_fingerprint in archived_legacy_marker
    assert boot_nonce in legacy_marker.read_text(encoding="utf-8")
    assert any(
        n.startswith("traces/") and n.endswith("events.jsonl") for n in names
    )
    # The pre-copy flush barrier guarantees the request THIS test just made
    # is in the archived events — a bundle grabbed right after an incident
    # must contain the incident.
    events_member = next(
        n for n in names if n.startswith("traces/") and n.endswith("events.jsonl")
    )
    archived = [
        json.loads(line)
        for line in zf.read(events_member).decode("utf-8").splitlines()
        if line.strip()
    ]
    assert any(
        e.get("type") == "api_request" and e.get("path") == "/api/doc"
        for e in archived
    )
    # The hermetic fake key must not appear in ANY member, decompressed.
    for name in names:
        assert b"test-key-hermetic" not in zf.read(name), name
        assert boot_nonce.encode("utf-8") not in zf.read(name), name
    snapshot = json.loads(zf.read("snapshot.json"))
    assert boot_nonce not in json.dumps(snapshot)
    assert snapshot["server"]["boot_nonce_fingerprint"] == nonce_fingerprint
    assert snapshot["app"]["version"]
    manifest = json.loads(zf.read("bundle-manifest.json"))
    incidents = json.loads(zf.read("incident-index.json"))
    assert manifest["schema_version"] == 1
    assert manifest["capture_id"] == incidents["capture_id"]
    assert manifest["current_trace_run_id"] in manifest["included_run_ids"]
    assert manifest["scope"]["current_trace"] == "through_pre_copy_flush"
    assert manifest["current_trace_flush_complete"] is True
    legacy_meta = next(
        item
        for item in manifest["artifacts"]
        if item["path"] == f"logs/{legacy_name}"
    )
    assert legacy_meta["transformation"] == "credential_substring_redaction"
    assert legacy_meta["truncated"] is False
    assert incidents["capture_state"] == "active_at_capture"
    assert snapshot["key"]["masked"].startswith("…")


def test_bundle_adds_bounded_prior_context_and_an_incident_index(
    log_env, trace_env
):
    prior_id = "session-prior-incident-1000"
    recent_prior = time.time() - 60
    prior = trace_env / prior_id
    prior.mkdir(parents=True)
    (prior / "run.json").write_text(
        json.dumps(
            {
                "run_id": prior_id,
                "started_at": recent_prior - 10,
                "ended_at": recent_prior,
            }
        ),
        encoding="utf-8",
    )
    # One oversized historical record proves the bundle keeps a bounded,
    # line-aligned tail while retaining the incident after it.
    oversized = json.dumps({"type": "noise", "payload": "x" * 600_000})
    legacy_trace_secret = "sk-ant-legacybundlecredential123"
    failed_record = {
        "type": "api_request",
        "path": "/api/qc/export",
        "status": 500,
        "code": "export_failed",
        "request_id": "request-correlation-1",
        "outcome_code": "export_failed",
        "exception_type": "ExportError",
        "process_instance_id": "process-instance-1",
        "record_seq": 42,
        "workspace_id_before": 7,
        "generation_before": 3,
        "workspace_id_after": 7,
        "generation_after": 3,
        legacy_trace_secret: "non-secret diagnostic value",
    }
    failed = json.dumps(failed_record)
    (prior / "events.jsonl").write_text(
        oversized + "\n" + failed + "\n", encoding="utf-8"
    )
    (prior / "spans.jsonl").write_text(
        json.dumps(
            {
                "kind": "turn",
                "status": "failed",
                "error_kind": "provider_error",
            }
        )
        + "\n",
        encoding="utf-8",
    )
    os.utime(prior, (recent_prior, recent_prior))

    diagnostics.init_logging(force=True)
    logging.getLogger("buildaspec.test").error("synthetic export incident")
    client = TestClient(create_app())
    client.get("/api/doc")
    inventory = diagnostics.list_trace_runs(limit=10)
    assert prior_id in [item["run_id"] for item in inventory["runs"]], inventory
    resp = client.get("/api/diagnostics/bundle")
    assert resp.status_code == 200

    zf = zipfile.ZipFile(BytesIO(resp.content))
    names = set(zf.namelist())
    events_tail = f"traces/{prior_id}/events-tail.jsonl"
    spans_tail = f"traces/{prior_id}/spans-tail.jsonl"
    assert events_tail in names and spans_tail in names, sorted(names)
    archived_events = [
        json.loads(line)
        for line in zf.read(events_tail).decode("utf-8").splitlines()
    ]
    assert archived_events == [scrub_data(failed_record)]
    assert legacy_trace_secret.encode("utf-8") not in zf.read(events_tail)

    manifest = json.loads(zf.read("bundle-manifest.json"))
    assert prior_id in manifest["included_run_ids"]
    scope = next(
        item
        for item in manifest["scope"]["prior_traces"]
        if item["run_id"] == prior_id
    )
    events_meta = scope["artifacts"]["events-tail.jsonl"]
    assert events_meta["truncated"] is True
    assert events_meta["included_bytes"] < events_meta["source_size_bytes"]
    assert events_meta["transformation"] == (
        "current_recursive_credential_redaction"
    )

    incidents = json.loads(zf.read("incident-index.json"))
    assert any(
        "synthetic export incident" in item["line"]
        for item in incidents["recent_log_errors"]
    )
    assert any(
        item.get("run_id") == prior_id
        and item.get("path") == "/api/qc/export"
        and item.get("status") == 500
        and item.get("request_id") == "request-correlation-1"
        and item.get("outcome_code") == "export_failed"
        and item.get("record_seq") == 42
        and item.get("generation_before") == 3
        and item.get("generation_after") == 3
        for item in incidents["recent_trace_incidents"]
    )


def test_bundle_text_artifacts_use_a_manifested_line_aligned_tail(
    tmp_path, monkeypatch
):
    source = tmp_path / "crash.log"
    source.write_text("old-stack-" * 100 + "\nrecent-stack\n", encoding="utf-8")
    monkeypatch.setattr(diagnostics, "_BUNDLE_LOG_TEXT_BYTES", 64)
    archive = BytesIO()

    with zipfile.ZipFile(archive, "w") as zf:
        metadata = diagnostics._add_redacted_text_file_if_present(
            zf, source, "logs/crash.log"
        )

    assert metadata is not None
    assert metadata["coverage"] == "bounded_recent_tail"
    assert metadata["truncated"] is True
    with zipfile.ZipFile(BytesIO(archive.getvalue())) as zf:
        bundled = zf.read("logs/crash.log").decode("utf-8")
    assert bundled == "recent-stack\n"


def test_client_event_collector_logs_traces_and_bounds(
    log_env, trace_env, caplog
):
    diagnostics.init_logging(force=True)
    client = TestClient(create_app())

    with caplog.at_level(logging.ERROR, logger="buildaspec.client"):
        resp = client.post(
            "/api/diagnostics/client-event",
            json={
                "kind": "error",
                "message": "TypeError: x is undefined",
                "stack": "at App.tsx:1:1",
                "source": "index.js:10",
            },
        )
    assert resp.status_code == 200 and resp.json()["ok"] is True
    assert "TypeError: x is undefined" in caplog.text
    events = _wait_events(
        lambda evs: any(e["type"] == "client_error" for e in evs)
    )
    client_events = [e for e in events if e["type"] == "client_error"]
    assert client_events and client_events[0]["kind"] == "error"

    # Oversized → 400 in the error idiom.
    resp = client.post(
        "/api/diagnostics/client-event",
        json={"kind": "error", "message": "x" * 33_000},
    )
    assert resp.status_code == 400 and resp.json()["ok"] is False

    # Unknown kind coerces rather than rejects (a reporter must not fail).
    resp = client.post(
        "/api/diagnostics/client-event",
        json={"kind": "weird", "message": "m"},
    )
    assert resp.status_code == 200

    # A malformed payload gets the translated 422 idiom.
    resp = client.post("/api/diagnostics/client-event", json={"kind": "error"})
    assert resp.status_code == 422 and resp.json()["code"] == "validation_error"


# --- capture call sites -----------------------------------------------------


def test_capture_sites_leave_events(trace_env):
    client = TestClient(create_app())
    resp = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {"action": "add_article", "target_id": "pt1", "text": "SCOPE"}
            ]
        },
    )
    assert resp.status_code == 200
    client.post("/api/doc/undo")
    client.get("/api/project/save")
    client.post("/api/session/reset")

    events = _wait_events(
        lambda evs: {
            "doc_edit",
            "doc_history",
            "project_save",
            "session_reset",
        }
        <= {e["type"] for e in evs}
    )
    types = {e["type"] for e in events}
    assert {"doc_edit", "doc_history", "project_save", "session_reset"} <= types

    edit = next(e for e in events if e["type"] == "doc_edit")
    assert edit["ops"] == 1 and edit["ok"] is True
    assert edit["actions"] == ["add_article"]
    undo = next(e for e in events if e["type"] == "doc_history")
    assert undo["action"] == "undo" and undo["ok"] is True
    save = next(e for e in events if e["type"] == "project_save")
    assert save["bytes"] > 0 and save["filename"].startswith("buildaspec-")
    reset = next(e for e in events if e["type"] == "session_reset")
    assert reset["had_content"] is False  # the undo returned it to empty


def test_round_end_and_prompt_refs_are_recorded_per_turn(
    monkeypatch, trace_env
):
    from tests.fakes import FakeClient, text_turn, tool_turn

    fake = FakeClient(
        [
            tool_turn(
                ["Drafting."],
                {
                    "edits": [
                        {
                            "action": "add_article",
                            "target_id": "pt1",
                            "text": "SUMMARY",
                        }
                    ]
                },
            ),
            text_turn(["Done."]),
            text_turn(["Second turn."]),
        ]
    )
    monkeypatch.setattr("backend.llm.conversation.get_client", lambda: fake)
    client = TestClient(create_app())
    client.post("/api/chat", json={"message": "go"})
    client.post("/api/chat", json={"message": "again"})

    events = _wait_events(
        lambda evs: sum(1 for e in evs if e["type"] == "round_end") >= 3
        and sum(1 for e in evs if e["type"] == "prompt_refs") >= 2
    )
    rounds = [e for e in events if e["type"] == "round_end"]
    # Turn 1 = two rounds (tool_use then end_turn); turn 2 = one round.
    assert len(rounds) == 3
    assert rounds[0]["round"] == 0 and rounds[0]["stop_reason"] == "tool_use"
    assert rounds[0]["tool_uses"] == 1
    assert rounds[1]["round"] == 1 and rounds[1]["stop_reason"] == "end_turn"
    assert isinstance(rounds[0]["ms"], int)

    prompt_events = [e for e in events if e["type"] == "prompt_refs"]
    assert len(prompt_events) == 2
    assert "ref" in prompt_events[0]["system"]

    rec = recorder_module.get_recorder()
    prompts = [
        json.loads(line)
        for line in (rec.trace_dir / "prompts.jsonl")
        .read_text(encoding="utf-8")
        .splitlines()
        if line.strip()
    ]
    kinds = [p["kind"] for p in prompts]
    assert {"system", "project_context", "user"} <= set(kinds)
    # The stable system block is hash-deduped: one entry across two turns.
    assert kinds.count("system") == 1


def test_a_stopped_rounds_trace_event_discloses_its_estimated_output(
    monkeypatch, trace_env
):
    """A support bundle has to be able to tell an exact provider count from
    an estimated one — otherwise the round_end record silently asserts a
    precision it does not have."""
    from tests.fakes import text_block, token_usage
    from tests.test_stop import _stopped_turn_usage

    TestClient(create_app())  # start the app so the recorder is live
    usage = _stopped_turn_usage(
        [text_block("word " * 400)],
        snapshot_usage=token_usage(input=1000, output=4),
    )
    assert usage["estimated_output_tokens"] == 496

    events = _wait_events(
        lambda evs: any(e["type"] == "round_end" for e in evs)
    )
    rounds = [e for e in events if e["type"] == "round_end"]
    assert len(rounds) == 1
    recorded = rounds[0]["usage"]
    assert recorded["output_tokens"] == 4  # provider-reported
    assert recorded["estimated_output_tokens"] == 496
    assert recorded["usage_estimated"] is True


def test_a_normal_rounds_trace_event_claims_no_estimate(monkeypatch, trace_env):
    from tests.fakes import FakeClient, text_turn, token_usage

    fake = FakeClient(
        [text_turn(["Done."], usage=token_usage(input=100, output=25))]
    )
    monkeypatch.setattr("backend.llm.conversation.get_client", lambda: fake)
    client = TestClient(create_app())
    client.post("/api/chat", json={"message": "go"})

    events = _wait_events(
        lambda evs: any(e["type"] == "round_end" for e in evs)
    )
    recorded = next(e for e in events if e["type"] == "round_end")["usage"]
    assert recorded["output_tokens"] == 25
    assert "estimated_output_tokens" not in recorded
    assert "usage_estimated" not in recorded


def test_workspace_conflict_event_from_the_lease_middleware(
    monkeypatch, trace_env
):
    client = TestClient(create_app())

    def raiser(expected_workspace_id=None):
        raise sessions.WorkspaceConflictError("the workspace moved")

    monkeypatch.setattr("backend.sessions.active_write", raiser)
    resp = client.post(
        "/api/qc/dismiss", json={"finding_id": "x", "reason": "r"}
    )
    assert resp.status_code == 409
    assert resp.json()["code"] == "stale_workspace"
    events = _wait_events(
        lambda evs: any(e["type"] == "workspace_conflict" for e in evs)
    )
    conflict = next(e for e in events if e["type"] == "workspace_conflict")
    assert conflict["path"] == "/api/qc/dismiss"


def test_the_snapshot_names_which_research_coverage_never_completed():
    """A support bundle has to answer "which coverage failed" without opening
    the project — a partial run still reports ``complete``, so the status
    alone says nothing, and the raw provider message must not travel.
    """
    from backend.research.engine import DimensionStatus, RequirementsProfile

    session = sessions.get_workspace().session
    session.research.status = "complete"
    session.research.profile_result = RequirementsProfile(
        items=[],
        dimension_statuses=[
            DimensionStatus(
                dimension_id="governing_codes",
                status="completed",
                title="Governing codes",
            ),
            DimensionStatus(
                dimension_id="ahj_requirements",
                status="failed",
                title="AHJ requirements",
                error="BadRequestError: {'message': 'container_id required'}",
                error_kind="invalid_request",
            ),
        ],
        research_date="2026-07-21",
        rounds=[],
    )

    client = TestClient(create_app())
    resp = client.get("/api/diagnostics")
    assert resp.status_code == 200
    research = resp.json()["session"]["research"]
    assert research["status"] == "complete"
    assert research["dimension_count"] == 2
    assert research["incomplete_dimensions"] == [
        {
            "dimension_id": "ahj_requirements",
            "title": "AHJ requirements",
            "error_kind": "invalid_request",
        }
    ]
    # Sanitized: the kind travels, the provider's own payload does not.
    assert "container_id" not in resp.text
    assert "BadRequestError" not in resp.text


def test_a_crafted_project_kind_never_reaches_the_diagnostics_response():
    """The end of the path the review named: project file → snapshot → bundle.

    `.baspec` files are shared between people, so a hostile or hand-edited
    one must not be able to place text of its choosing in a support bundle
    through a field documented as a closed vocabulary.
    """
    from backend.research.engine import RequirementsProfile

    hostile = "BadRequestError: {'message': 'leaked payload'}"
    session = sessions.get_workspace().session
    session.research.status = "complete"
    session.research.profile_result = RequirementsProfile.from_dict(
        {
            "items": [],
            "dimension_statuses": [
                {
                    "dimension_id": "site_environment",
                    "status": "failed",
                    "title": "Site and environment",
                    "error_kind": hostile,
                }
            ],
            "research_date": "2026-07-21",
        }
    )

    resp = TestClient(create_app()).get("/api/diagnostics")
    assert resp.status_code == 200
    assert resp.json()["session"]["research"]["incomplete_dimensions"] == [
        {
            "dimension_id": "site_environment",
            "title": "Site and environment",
            "error_kind": "unrecognized",
        }
    ]
    assert "leaked payload" not in resp.text


def test_a_session_that_never_researched_reports_empty_coverage():
    client = TestClient(create_app())
    research = client.get("/api/diagnostics").json()["session"]["research"]
    assert research["status"] == "idle"
    assert research["rounds"] == 0
    assert research["dimension_count"] == 0
    assert research["incomplete_dimensions"] == []
