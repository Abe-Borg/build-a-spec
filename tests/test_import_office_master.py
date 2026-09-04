"""Importing a real office master: the shape that came back as one blob.

The reported document (a hyperscale client's 21 05 00 master) had a cover
page — "Specification for", the title, "Section Number: 21 05 00", a site
code, a location, a revision-history table — followed by PART and ARTICLE
headings whose numbering lived on the PRT/ART/PR1-PR4 STYLES rather than on
the paragraphs, and a Related Requirements list citing "Section 09 90 00 –
Painting and Coating". The import rendered the section as "SECTION 09 90 00
/ PAINTING AND COATING" and dumped every heading as a flat provision under a
synthetic IMPORTED CONTENT article. Three parser rules were wrong at once:

* numbering was read only from the paragraph's own ``w:numPr``, never from
  its style, so every heading exposed its bare title and matched nothing;
* ``SECTION …`` matched any body line and a LATER match overwrote an earlier
  one, so a cross-reference became the header — and appended a second
  ``sec`` anchor the format map refused, which made project Save a 500;
* a cover page before PART 1 was modelled as provisions of that synthetic
  article, and on export the renderer invented the article heading.

These tests build that document shape in memory and pin the fixes: style
numbering resolved through ``w:basedOn``, the CSI style-name convention as
the fallback, the section identity decided in the front matter and once,
the front matter kept for export verbatim (and never modelled), the page
header/footer as the last-resort identity, a cached table of contents
treated as a Word field, and text boxes read rather than dropped.
"""
from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastapi.testclient import TestClient
from lxml import etree

from backend import sessions
from backend.app import create_app
from backend.spec_doc.importer import (
    _iter_body_texts,
    parse_master_docx,
)
from backend.spec_doc.model import iter_paragraphs
from backend.spec_doc.source_format import (
    HEADER_SOURCE_CHROME,
    HEADER_SOURCE_FRONT_MATTER,
    HEADER_SOURCE_LINE,
    SourceFormatMap,
)
from backend.spec_doc.source_render import render_preserving_docx
from tests.docx_fidelity_helpers import DOCX_MEDIA_TYPE, _png_bytes
from tests.test_importer import _define_numbering

TITLE = "COMMON WORK RESULTS FOR FIRE SUPPRESSION"
CROSS_REFERENCE = "Section 09 90 00 – Painting and Coating"
_VML_NS = "urn:schemas-microsoft-com:vml"

# ilvl -> (numFmt, lvlText): the outline every MasterSpec-derived master
# carries. Level 0 renders "PART 1 - ", level 1 renders "1.01".
_OUTLINE_LEVELS = {
    0: ("decimal", "PART %1 - "),
    1: ("decimalZero", "%1.%2"),
    2: ("upperLetter", "%3."),
    3: ("decimal", "%4."),
    4: ("lowerLetter", "%5."),
}
_STYLE_LEVELS = {"PRT": 0, "ART": 1, "PR1": 2, "PR2": 3, "PR3": 4}
_NUM_ID = 70


@pytest.fixture
def client():
    sessions.reset_session()
    with TestClient(create_app()) as api:
        yield api
    sessions.reset_session()


def _add_style(
    document,
    name: str,
    *,
    num_id: int | None,
    ilvl: int | None,
    based_on=None,
):
    """A paragraph style, optionally carrying (part of) the outline numbering.

    ``num_id`` None with an ``ilvl`` writes a level-only ``w:numPr`` — the
    shape a derived style uses to sit one level deeper while inheriting the
    definition from the style it is based on.
    """
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if based_on is not None:
        style.base_style = based_on
    if num_id is not None or ilvl is not None:
        p_pr = style.element.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        if ilvl is not None:
            ilvl_el = OxmlElement("w:ilvl")
            ilvl_el.set(qn("w:val"), str(ilvl))
            num_pr.append(ilvl_el)
        if num_id is not None:
            num_id_el = OxmlElement("w:numId")
            num_id_el.set(qn("w:val"), str(num_id))
            num_pr.append(num_id_el)
        p_pr.append(num_pr)
    return style


def _text_box_paragraph(document, text: str):
    """A paragraph whose only content is a VML text box saying ``text``."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    pict = OxmlElement("w:pict")
    shape = etree.SubElement(pict, f"{{{_VML_NS}}}shape")
    textbox = etree.SubElement(shape, f"{{{_VML_NS}}}textbox")
    content = etree.SubElement(textbox, qn("w:txbxContent"))
    inner_p = etree.SubElement(content, qn("w:p"))
    inner_r = etree.SubElement(inner_p, qn("w:r"))
    inner_t = etree.SubElement(inner_r, qn("w:t"))
    inner_t.text = text
    run._r.append(pict)
    return paragraph


def _toc_paragraphs(document, entries: list[str]) -> None:
    """A cached table of contents: one complex field spanning paragraphs."""
    def _fld_char(kind: str):
        run = OxmlElement("w:r")
        char = OxmlElement("w:fldChar")
        char.set(qn("w:fldCharType"), kind)
        run.append(char)
        return run

    first = document.add_paragraph()
    first._p.append(_fld_char("begin"))
    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    instr_run.append(instr)
    first._p.append(instr_run)
    first._p.append(_fld_char("separate"))
    first.add_run(entries[0])
    for entry in entries[1:]:
        document.add_paragraph(entry)
    last = document.add_paragraph()
    last._p.append(_fld_char("end"))


def _office_master(
    *,
    cover: bool = True,
    number_line: bool = True,
    style_numbering: bool = True,
    chrome: bool = True,
    header_line: bool = False,
    cross_reference: bool = True,
    toc: bool = False,
    text_box: bool = False,
    number_in_text_box: bool = False,
    picture_between: bool = False,
) -> bytes:
    """The reported master's shape, with the knobs the tests need."""
    document = Document()
    if chrome:
        section = document.sections[0]
        section.header.paragraphs[0].text = TITLE
        section.footer.paragraphs[0].text = "21 05 00 - 1"
    if style_numbering:
        _define_numbering(document, _NUM_ID, _OUTLINE_LEVELS)
    body = _add_style(document, "SpecBody", num_id=None, ilvl=None)
    styles = {}
    for name, ilvl in _STYLE_LEVELS.items():
        if name == "PR3":
            continue
        styles[name] = _add_style(
            document,
            name,
            num_id=_NUM_ID if style_numbering else None,
            ilvl=ilvl if style_numbering else None,
            based_on=body,
        )
    # PR3 states only its level and inherits the DEFINITION through PR2 —
    # the basedOn chain a real styles part uses for its deeper levels.
    styles["PR3"] = _add_style(
        document,
        "PR3",
        num_id=None,
        ilvl=_STYLE_LEVELS["PR3"] if style_numbering else None,
        based_on=styles["PR2"],
    )

    if header_line:
        document.add_paragraph("SECTION 21 05 00")
        document.add_paragraph(TITLE)
    if cover:
        document.add_paragraph("Specification for")
        if text_box:
            _text_box_paragraph(document, TITLE)
        else:
            document.add_paragraph(TITLE)
        if number_line and number_in_text_box:
            _text_box_paragraph(document, "Section Number: 21 05 00")
        elif number_line:
            document.add_paragraph("Section Number: 21 05 00")
        document.add_paragraph("CLT11")
        document.add_paragraph("Maiden, NC")
        document.add_paragraph("PROJECT REVISION HISTORY")
        table = document.add_table(rows=3, cols=3)
        rows = [
            ("Rev", "Date", "Description / Justification"),
            ("0", "Initial draft, based on master revision A.2.0", ""),
            ("1", "2026-07-24", "30% IFR Package"),
        ]
        for row, values in zip(table.rows, rows):
            for cell, value in zip(row.cells, values):
                cell.text = value
        document.add_page_break()
    if toc:
        _toc_paragraphs(document, ["1.1 SUMMARY 3", "1.2 RELATED REQUIREMENTS 3"])

    def heading(text: str, style: str):
        document.add_paragraph(text, style=styles[style])

    heading("GENERAL", "PRT")
    heading("SUMMARY", "ART")
    heading("Section Includes:", "PR1")
    heading(
        "Piping materials and installation instructions common to most "
        "piping systems.",
        "PR2",
    )
    heading("Provide listed components throughout.", "PR3")
    heading("RELATED REQUIREMENTS", "ART")
    if cross_reference:
        heading(CROSS_REFERENCE, "PR1")
    heading("Section 21 13 13 – Wet-Pipe Sprinkler Systems", "PR1")
    heading("PRODUCTS", "PRT")
    heading("PIPE AND FITTINGS", "ART")
    heading("Steel pipe: ASTM A53.", "PR1")
    if picture_between:
        document.add_picture(io.BytesIO(_png_bytes()))
    heading("Fittings: ASME B16.9.", "PR1")
    heading("EXECUTION", "PRT")
    heading("INSTALLATION", "ART")
    heading("Install per NFPA 13.", "PR1")
    document.add_paragraph("END OF SECTION")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _parse(payload: bytes, tmp_path):
    path = tmp_path / "master.docx"
    path.write_bytes(payload)
    return parse_master_docx(path)


def _import(client: TestClient, source: bytes) -> dict:
    response = client.post(
        "/api/import/master",
        files={"file": ("21 05 00.docx", source, DOCX_MEDIA_TYPE)},
        data={"detach": "true"},
    )
    assert response.status_code == 200, response.text
    return response.json()


def _body_children(payload: bytes) -> list:
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        xml = archive.read("word/document.xml")
    body = etree.fromstring(xml).find(qn("w:body"))
    return [child for child in body.iterchildren() if isinstance(child.tag, str)]


def _texts(payload: bytes) -> list[str]:
    return [
        paragraph.text
        for paragraph in Document(io.BytesIO(payload)).paragraphs
        if paragraph.text.strip()
    ]


def _paragraph_texts(section) -> list[str]:
    return [p.text for _part, _article, p, _depth, _ref in iter_paragraphs(section)]


# ---------------------------------------------------------------------------
# 1. The outline lives on the styles
# ---------------------------------------------------------------------------


def test_a_style_numbered_master_arrives_with_its_real_structure(tmp_path):
    """Headings numbered by their STYLE are headings, not provisions.

    Direct ``w:numPr`` alone never fires for this master: every paragraph
    carries only ``w:pStyle``. Resolving the style chain (PR3 inherits its
    numbering through PR2) is what turns "GENERAL" into PART 1 and "SUMMARY"
    into article 1.1 instead of provisions A. and B. of IMPORTED CONTENT.
    """
    result = _parse(_office_master(), tmp_path)
    section = result.section

    assert [a.title for a in section.parts[0].articles] == [
        "SUMMARY",
        "RELATED REQUIREMENTS",
    ]
    assert [a.title for a in section.parts[1].articles] == ["PIPE AND FITTINGS"]
    assert [a.title for a in section.parts[2].articles] == ["INSTALLATION"]
    assert not any(
        article.title == "IMPORTED CONTENT"
        for part in section.parts
        for article in part.articles
    )
    summary = section.parts[0].articles[0]
    assert summary.paragraphs[0].text == "Section Includes:"
    child = summary.paragraphs[0].children[0]
    assert child.text.startswith("Piping materials")
    assert child.children[0].text == "Provide listed components throughout."
    assert result.style_numbering_resolved is True
    assert result.spec_shape_detected is True
    # The master's own PART wording, under the canonical label.
    assert section.parts[0].title == "PART 1 - GENERAL"
    assert section.parts[1].title == "PART 2 - PRODUCTS"


def test_csi_style_names_promote_headings_without_any_numbering(tmp_path):
    """PRT / ART / PR1.. by NAME, when neither paragraph nor style numbers.

    Some masters carry the convention's style names with no numbering
    definition behind them (the labels typed, or none at all). The name is
    still the structural signal; a typed text label would win over it.
    """
    result = _parse(_office_master(style_numbering=False), tmp_path)
    section = result.section

    assert [a.title for a in section.parts[0].articles] == [
        "SUMMARY",
        "RELATED REQUIREMENTS",
    ]
    assert section.parts[2].articles[0].title == "INSTALLATION"
    summary = section.parts[0].articles[0]
    # PR1 -> depth 0, PR2 -> depth 1, PR3 -> depth 2: the style says so.
    assert summary.paragraphs[0].text == "Section Includes:"
    assert summary.paragraphs[0].children[0].text.startswith("Piping materials")
    assert (
        summary.paragraphs[0].children[0].children[0].text
        == "Provide listed components throughout."
    )
    assert result.style_numbering_resolved is False


# ---------------------------------------------------------------------------
# 2. The section identity is decided in the front matter, once
# ---------------------------------------------------------------------------


def test_a_cross_reference_never_becomes_the_header(tmp_path):
    """The reported symptom: "SECTION 09 90 00 / PAINTING AND COATING".

    In a style-numbered master the Related Requirements entry's visible text
    BEGINS with "Section", so the old any-line, last-wins rule read it as
    the header. Identity is now read only in the front matter, and the
    cross-reference is the provision it always was.
    """
    result = _parse(_office_master(), tmp_path)

    assert result.section.number == "21 05 00"
    assert result.section.title == TITLE
    related = result.section.parts[0].articles[1]
    assert [p.text for p in related.paragraphs] == [
        CROSS_REFERENCE,
        "Section 21 13 13 – Wet-Pipe Sprinkler Systems",
    ]
    assert result.header_source == HEADER_SOURCE_FRONT_MATTER
    assert any("read from the cover page" in w for w in result.warnings)


def test_the_first_header_line_wins_and_is_anchored_once(tmp_path):
    """Two SECTION-shaped lines before END OF SECTION: one header, one anchor.

    The second one used to overwrite the first AND append a second ``sec``
    anchor, which ``SourceFormatMap.from_dict`` refuses — so the import
    succeeded and every later project Save failed.
    """
    document = Document()
    for line in (
        "SECTION 21 05 00",
        TITLE,
        "SECTION 21 05 00",  # a cover page repeating its own header
        "PART 1 - GENERAL",
        "1.1 RELATED REQUIREMENTS",
        "SECTION 09 90 00 - PAINTING AND COATING",
        "END OF SECTION",
    ):
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)

    result = _parse(buffer.getvalue(), tmp_path)
    assert result.section.number == "21 05 00"
    assert result.section.title == TITLE
    assert result.header_source == HEADER_SOURCE_LINE
    uids = [anchor.uid for anchor in result.format_map.anchors]
    assert uids.count("sec") == 1
    assert len(uids) == len(set(uids))
    # The body-level SECTION line is a provision of the article it sits in.
    assert _paragraph_texts(result.section) == [
        "SECTION 09 90 00 - PAINTING AND COATING"
    ]
    # The repeated cover line is front matter, not a provision.
    assert result.front_matter == ("SECTION 21 05 00",)
    # And the map round-trips, which is what Save depends on.
    SourceFormatMap.from_dict(result.format_map.to_dict())


def test_a_document_with_no_structure_keeps_the_first_header_it_finds(tmp_path):
    """No PART/article at all: a header line may sit anywhere, first wins."""
    document = Document()
    for line in (
        "Some memo text.",
        "SECTION 21 05 00 - COMMON WORK RESULTS",
        "More text.",
        "SECTION 09 90 00 - PAINTING AND COATING",
    ):
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)

    result = _parse(buffer.getvalue(), tmp_path)
    assert result.section.number == "21 05 00"
    assert result.section.title == "COMMON WORK RESULTS"
    assert result.front_matter == ()
    assert "SECTION 09 90 00 - PAINTING AND COATING" in _paragraph_texts(
        result.section
    )


# ---------------------------------------------------------------------------
# 3. The cover page is front matter
# ---------------------------------------------------------------------------


def test_the_cover_page_is_front_matter_not_provisions(tmp_path):
    result = _parse(_office_master(), tmp_path)

    assert "Specification for" in result.front_matter
    assert "CLT11" in result.front_matter
    assert "Maiden, NC" in result.front_matter
    assert "Section Number: 21 05 00" in result.front_matter
    assert any("30% IFR Package" in line for line in result.front_matter)
    texts = _paragraph_texts(result.section)
    assert "Specification for" not in texts
    assert "CLT11" not in texts
    assert not any("IFR Package" in text for text in texts)
    assert any("before the first PART heading" in w for w in result.warnings)
    assert result.format_map.front_matter_text == result.front_matter
    assert result.format_map.header_source == HEADER_SOURCE_FRONT_MATTER


def test_the_front_matter_rides_the_import_report_and_the_payload(client):
    data = _import(client, _office_master())
    report = data["import_report"]
    assert report["front_matter"]["count"] >= 6
    assert "CLT11" in report["front_matter"]["lines"]
    assert any("before the first PART heading" in w for w in data["warnings"])
    assert data["doc"]["section"] == {"number": "21 05 00", "title": TITLE}
    # No synthetic article, and the headings are real structure.
    titles = [
        article["title"]
        for part in data["doc"]["parts"]
        for article in part["articles"]
    ]
    assert "IMPORTED CONTENT" not in titles
    assert "SUMMARY" in titles
    # The project saves (the duplicate-anchor 500 is the regression here).
    assert client.get("/api/project/save").status_code == 200


def test_identity_falls_back_to_the_page_header_and_footer(tmp_path):
    """No SECTION line and no Section Number field: the page furniture says."""
    result = _parse(_office_master(number_line=False), tmp_path)

    assert result.section.number == "21 05 00"
    assert result.section.title == TITLE
    assert result.header_source == HEADER_SOURCE_CHROME
    assert any("page header/footer" in w for w in result.warnings)


def test_a_structureless_file_gets_no_invented_identity_from_its_footer(tmp_path):
    document = Document()
    document.sections[0].footer.paragraphs[0].text = "21 05 00 - 1"
    for line in ("A memo.", "About nothing in particular."):
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)

    result = _parse(buffer.getvalue(), tmp_path)
    assert result.section.number == ""
    assert result.spec_shape_detected is False


def test_a_cached_table_of_contents_is_a_field_not_articles(tmp_path):
    """TOC entries read "1.1 SUMMARY 3" — exactly the article grammar."""
    payload = _office_master(toc=True)
    result = _parse(payload, tmp_path)

    titles = [a.title for part in result.section.parts for a in part.articles]
    assert "SUMMARY 3" not in titles
    assert "RELATED REQUIREMENTS 3" not in titles
    assert any(line.startswith("1.1 SUMMARY") for line in result.front_matter)
    entries = _iter_body_texts(Document(io.BytesIO(payload)))
    toc_entries = [e for e in entries if e.text.startswith("1.")]
    assert toc_entries and all(e.lock_reason == "field" for e in toc_entries)


def test_text_box_content_is_read_not_dropped(tmp_path):
    """A cover title drawn in a text box is invisible to ``Paragraph.text``."""
    result = _parse(_office_master(text_box=True), tmp_path)

    assert TITLE in result.front_matter
    assert result.section.title == TITLE


def test_identity_drawn_in_a_text_box_is_still_read(tmp_path):
    """Both the number field AND the title in text boxes (Codex, PR #145).

    A locked ``image`` paragraph used to be excluded from identity matching
    altogether, so a cover page built entirely from text boxes — the common
    template shape — left the section blank even though the exact line had
    been read out of the box. The block stays front matter, is never
    anchored (rewriting it on a rename would delete the box), and the export
    still carries it verbatim.
    """
    source = _office_master(text_box=True, number_in_text_box=True, chrome=False)
    result = _parse(source, tmp_path)

    assert result.section.number == "21 05 00"
    assert result.section.title == TITLE
    assert result.header_source == HEADER_SOURCE_FRONT_MATTER
    assert "Section Number: 21 05 00" in result.front_matter
    assert not any(anchor.uid == "sec" for anchor in result.format_map.anchors)
    exported = render_preserving_docx(
        source_bytes=source, format_map=result.format_map, current=result.section
    )
    before = _body_children(source)
    after = _body_children(exported)
    assert len(before) == len(after)


def test_a_section_line_drawn_in_a_text_box_is_the_identity_not_a_provision(tmp_path):
    document = Document()
    _text_box_paragraph(document, "SECTION 21 05 00")
    _text_box_paragraph(document, TITLE)
    document.add_paragraph("PART 1 - GENERAL")
    document.add_paragraph("1.1 SUMMARY")
    document.add_paragraph("A. Section includes piping.")
    document.add_paragraph("END OF SECTION")
    buffer = io.BytesIO()
    document.save(buffer)

    result = _parse(buffer.getvalue(), tmp_path)
    assert result.section.number == "21 05 00"
    assert result.section.title == TITLE
    assert result.header_source == HEADER_SOURCE_FRONT_MATTER
    assert result.front_matter == ("SECTION 21 05 00", TITLE)
    assert _paragraph_texts(result.section) == ["Section includes piping."]


# ---------------------------------------------------------------------------
# 4. The export keeps the front matter where it was
# ---------------------------------------------------------------------------


def test_front_matter_is_exported_verbatim_ahead_of_the_section(client):
    source = _office_master()
    data = _import(client, source)
    target = None
    for part in data["doc"]["parts"]:
        for article in part["articles"]:
            for paragraph in article["paragraphs"]:
                if paragraph["text"].startswith("Steel pipe"):
                    target = paragraph["id"]
    assert target
    edit = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "replace",
                    "target_id": target,
                    "text": "Steel pipe: ASTM A53, Schedule 40.",
                }
            ]
        },
    )
    assert edit.status_code == 200, edit.text

    exported = client.get("/api/export/docx", params={"mode": "preserved"})
    assert exported.status_code == 200, exported.text
    before = _body_children(source)
    after = _body_children(exported.content)
    # Every front-matter element — cover lines, the revision table, the page
    # break — is byte-identical and still ahead of the first heading.
    first_heading = next(
        index
        for index, child in enumerate(before)
        if "".join(child.itertext()).strip() == "GENERAL"
    )
    for index in range(first_heading):
        assert etree.tostring(before[index], method="c14n") == etree.tostring(
            after[index], method="c14n"
        )
    texts = _texts(exported.content)
    assert texts.index("Specification for") < texts.index("GENERAL")
    assert "Steel pipe: ASTM A53, Schedule 40." in texts
    assert "Steel pipe: ASTM A53." not in texts


def test_a_front_matter_identity_synthesizes_no_header(client):
    source = _office_master()
    _import(client, source)
    exported = client.get("/api/export/docx", params={"mode": "preserved"})
    assert exported.status_code == 200, exported.text
    texts = _texts(exported.content)
    assert texts.count("Section Number: 21 05 00") == 1
    assert not any(text.startswith("SECTION 21 05 00") for text in texts)
    assert texts.count(TITLE) == 1


def test_renaming_the_section_flags_the_cover_page(client):
    data = _import(client, _office_master())
    edit = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "replace",
                    "target_id": "sec",
                    "text": TITLE,
                    "numbering": "21 05 13",
                }
            ]
        },
    )
    assert edit.status_code == 200, edit.text
    findings = [
        issue
        for issue in edit.json()["lint"]
        if issue["rule"] == "stale_document_identifier"
    ]
    assert findings, edit.json()["lint"]
    assert "cover page" in findings[0]["message"]
    assert "21 05 00" in findings[0]["message"]
    # The cover page itself is untouched by the rename.
    exported = client.get("/api/export/docx", params={"mode": "preserved"})
    assert exported.status_code == 200
    assert "Section Number: 21 05 00" in _texts(exported.content)
    assert data["doc"]["section"]["number"] == "21 05 00"


def test_a_picture_only_paragraph_keeps_its_place_in_the_export(tmp_path):
    """Unmodelled content travels with the element below it, never to the end."""
    source = _office_master(picture_between=True)
    result = _parse(source, tmp_path)
    exported = render_preserving_docx(
        source_bytes=source, format_map=result.format_map, current=result.section
    )
    children = _body_children(exported.content if hasattr(exported, "content") else exported)
    texts = ["".join(child.itertext()).strip() for child in children]
    drawing = qn("w:drawing")
    picture_index = next(
        index for index, child in enumerate(children) if child.find(f".//{drawing}") is not None
    )
    assert texts[picture_index - 1] == "Steel pipe: ASTM A53."
    assert texts[picture_index + 1] == "Fittings: ASME B16.9."


def test_an_untouched_office_master_round_trips_element_for_element(tmp_path):
    source = _office_master(toc=True, text_box=True)
    result = _parse(source, tmp_path)
    exported = render_preserving_docx(
        source_bytes=source, format_map=result.format_map, current=result.section
    )
    before = _body_children(source)
    after = _body_children(exported)
    assert len(before) == len(after)
    for left, right in zip(before, after):
        assert etree.tostring(left, method="c14n") == etree.tostring(
            right, method="c14n"
        )
