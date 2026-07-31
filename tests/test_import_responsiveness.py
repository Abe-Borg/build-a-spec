"""The import/open paths must not freeze the rest of the app.

Parsing and lexically indexing a master specification is seconds of CPU on a
real section. It used to run inline in an ``async def`` handler, i.e. directly
on the asyncio event loop, so for the whole duration the server answered
nothing: the document panel sat still and a chat turn could not deliver a
single SSE frame until the upload finished. These tests pin the fix — the
blocking half runs on a worker thread, so concurrent requests keep flowing —
and the lexical-index memoization that made the blocking half far shorter.
"""
from __future__ import annotations

import io
import threading
import time

from docx import Document
from fastapi.testclient import TestClient

from backend import app as app_module, sessions
from backend.llm.conversation import (
    SessionState,
    _source_editing_boundary_block,
    _turn_context_text,
)
from tests.conftest import settle_capability_sweep

# How long a stalled request is allowed to hold the server before the test
# gives up on it. A blocked event loop cannot honour an asyncio timeout — the
# loop is the thing that is stuck — so every wait here is a real thread wait.
_BLOCK_SECONDS = 5.0
# A request served off the blocked path returns in milliseconds; anything near
# ``_BLOCK_SECONDS`` means it queued behind the import.
_RESPONSIVE_SECONDS = 1.5

_DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _master_bytes(articles: int = 3, paragraphs: int = 3) -> bytes:
    document = Document()
    document.add_paragraph("SECTION 21 13 13")
    document.add_paragraph("WET-PIPE SPRINKLER SYSTEMS")
    for part_index, part in enumerate(
        ["PART 1 - GENERAL", "PART 2 - PRODUCTS", "PART 3 - EXECUTION"], 1
    ):
        document.add_paragraph(part)
        for article in range(articles):
            document.add_paragraph(
                f"{part_index}.{article + 1} ARTICLE {article + 1}"
            )
            for paragraph in range(paragraphs):
                document.add_paragraph(
                    f"{chr(65 + paragraph)}. Provision {article + 1}."
                    f"{paragraph + 1}: install per NFPA 13."
                )
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _upload(source: bytes) -> dict:
    return {"file": ("master.docx", source, _DOCX_MEDIA_TYPE)}


def test_import_does_not_block_concurrent_requests(monkeypatch):
    """A slow import must not stall an unrelated request on the event loop."""
    real_prepare = app_module._prepare_master_import
    entered = threading.Event()
    release = threading.Event()

    def blocking_prepare(source_bytes: bytes, safe_filename: str):
        entered.set()
        # Stands in for the seconds of real parsing/indexing CPU. Blocking a
        # thread is the point: on the event loop it freezes every other
        # request, exactly like the parse it replaces used to.
        release.wait(_BLOCK_SECONDS)
        return real_prepare(source_bytes, safe_filename)

    monkeypatch.setattr(app_module, "_prepare_master_import", blocking_prepare)
    source = _master_bytes()

    with TestClient(app_module.create_app()) as client:
        def do_import():
            return client.post("/api/import/master", files=_upload(source))

        worker_outcome: dict = {}

        def run_import() -> None:
            worker_outcome["response"] = do_import()

        worker = threading.Thread(target=run_import, daemon=True)
        worker.start()
        assert entered.wait(_BLOCK_SECONDS + 5), "import never started"
        started = time.perf_counter()
        health = client.get("/api/health")
        elapsed = time.perf_counter() - started
        release.set()
        worker.join(_BLOCK_SECONDS + 10)

    assert health.status_code == 200
    assert worker_outcome["response"].status_code == 200
    assert elapsed < _RESPONSIVE_SECONDS, (
        f"an unrelated request waited {elapsed:.2f}s for the import — the "
        "blocking work is back on the event loop"
    )


def test_reference_upload_does_not_block_concurrent_requests(monkeypatch):
    """Attaching a reference document is the third ``async def`` upload path.

    Extracting text from a long document is the same kind of multi-second CPU
    as a master parse, so it is subject to the same rule: it belongs on a
    worker thread, never inline on the loop.
    """
    real_prepare = app_module._prepare_reference_upload
    entered = threading.Event()
    release = threading.Event()

    def blocking_prepare(source_bytes: bytes, *, filename: str):
        entered.set()
        release.wait(_BLOCK_SECONDS)
        return real_prepare(source_bytes, filename=filename)

    monkeypatch.setattr(
        app_module, "_prepare_reference_upload", blocking_prepare
    )
    source = _master_bytes()

    with TestClient(app_module.create_app()) as client:
        worker_outcome: dict = {}

        def run_upload() -> None:
            worker_outcome["response"] = client.post(
                "/api/reference/upload",
                files={"file": ("standard.docx", source, _DOCX_MEDIA_TYPE)},
            )

        worker = threading.Thread(target=run_upload, daemon=True)
        worker.start()
        assert entered.wait(_BLOCK_SECONDS + 5), "upload never started"
        started = time.perf_counter()
        health = client.get("/api/health")
        elapsed = time.perf_counter() - started
        release.set()
        worker.join(_BLOCK_SECONDS + 10)

    assert health.status_code == 200
    assert worker_outcome["response"].status_code == 200
    assert elapsed < _RESPONSIVE_SECONDS, (
        f"an unrelated request waited {elapsed:.2f}s for a reference upload — "
        "the extraction is back on the event loop"
    )


def test_template_import_does_not_block_concurrent_requests(monkeypatch, tmp_path):
    """Importing a template is the fourth ``async def`` upload path.

    Up to 16 MiB of JSON parsed, validated and atomically written under the
    catalog lock — the same seconds-of-CPU shape as a master parse, and
    subject to the same rule: it belongs on a worker thread, never inline on
    the loop.
    """
    from pathlib import Path

    from backend.templates import TemplateCatalog, template_bytes

    curated_root = (
        Path(__file__).resolve().parents[1] / "backend" / "templates" / "curated"
    )
    catalog = TemplateCatalog(
        personal_root=tmp_path / "personal", curated_root=curated_root
    )
    curated, _source = catalog.get("curated:complete-section-starter")
    payload = template_bytes(curated)

    entered = threading.Event()
    release = threading.Event()
    real_import = catalog.import_bytes

    def blocking_import(data: bytes):
        entered.set()
        release.wait(_BLOCK_SECONDS)
        return real_import(data)

    catalog.import_bytes = blocking_import
    monkeypatch.setattr(app_module, "get_template_catalog", lambda: catalog)

    with TestClient(app_module.create_app()) as client:
        worker_outcome: dict = {}

        def run_import() -> None:
            worker_outcome["response"] = client.post(
                "/api/templates/import",
                files={
                    "file": (
                        "starter.bastemplate",
                        payload,
                        "application/vnd.buildaspec.template+json",
                    )
                },
            )

        worker = threading.Thread(target=run_import, daemon=True)
        worker.start()
        assert entered.wait(_BLOCK_SECONDS + 5), "template import never started"
        started = time.perf_counter()
        health = client.get("/api/health")
        elapsed = time.perf_counter() - started
        release.set()
        worker.join(_BLOCK_SECONDS + 10)

    assert health.status_code == 200
    assert worker_outcome["response"].status_code == 200
    assert elapsed < _RESPONSIVE_SECONDS, (
        f"an unrelated request waited {elapsed:.2f}s for a template import — "
        "the parse and atomic write are back on the event loop"
    )


def test_project_load_does_not_block_concurrent_requests(monkeypatch):
    """The same guarantee for opening a project (it re-parses the master)."""
    real_stage = app_module._stage_project_load
    entered = threading.Event()
    release = threading.Event()

    def blocking_stage(payload: bytes):
        entered.set()
        release.wait(_BLOCK_SECONDS)
        return real_stage(payload)

    monkeypatch.setattr(app_module, "_stage_project_load", blocking_stage)

    with TestClient(app_module.create_app()) as client:
        saved = client.get("/api/project/save")
        assert saved.status_code == 200
        worker_outcome: dict = {}

        def run_load() -> None:
            worker_outcome["response"] = client.post(
                "/api/project/load-file",
                files={
                    "file": ("session.baspec", saved.content, "application/zip")
                },
            )

        worker = threading.Thread(target=run_load, daemon=True)
        worker.start()
        assert entered.wait(_BLOCK_SECONDS + 5), "load never started"
        started = time.perf_counter()
        health = client.get("/api/health")
        elapsed = time.perf_counter() - started
        release.set()
        worker.join(_BLOCK_SECONDS + 10)

    assert health.status_code == 200
    assert worker_outcome["response"].status_code == 200
    assert elapsed < _RESPONSIVE_SECONDS, (
        f"an unrelated request waited {elapsed:.2f}s for the project load — "
        "the blocking work is back on the event loop"
    )


def _document_xml(source: bytes) -> bytes:
    import zipfile

    with zipfile.ZipFile(io.BytesIO(source)) as archive:
        return archive.read("word/document.xml")


def test_import_refuses_to_land_in_a_session_started_while_it_was_read(
    monkeypatch,
):
    """A "New session" during the read must not receive the old master.

    Serving requests during the read is the whole point of the offload, and
    ``/api/session/reset`` is one of them. A fresh session is blank, so the
    handler's body-content check would wave this master straight into a
    session the user deliberately started over — possibly on a different
    module or discipline. The generation stamp is what catches it.
    """
    real_prepare = app_module._prepare_master_import
    entered = threading.Event()
    release = threading.Event()

    def blocking_prepare(source_bytes: bytes, safe_filename: str):
        entered.set()
        release.wait(_BLOCK_SECONDS)
        return real_prepare(source_bytes, safe_filename)

    monkeypatch.setattr(app_module, "_prepare_master_import", blocking_prepare)
    source = _master_bytes()

    with TestClient(app_module.create_app()) as client:
        outcome: dict = {}

        def run_import() -> None:
            outcome["response"] = client.post(
                "/api/import/master", files=_upload(source)
            )

        worker = threading.Thread(target=run_import, daemon=True)
        worker.start()
        assert entered.wait(_BLOCK_SECONDS + 5), "import never started"
        assert client.post("/api/session/reset").status_code == 200
        release.set()
        worker.join(_BLOCK_SECONDS + 10)

    response = outcome["response"]
    assert response.status_code == 409
    assert "session was replaced" in response.json()["error"]
    # The new session is untouched: no master, no imported content.
    session = sessions.get_session()
    assert session.source_docx_bytes is None
    assert session.import_report is None
    assert not session.doc.doc.has_body_content()


def test_project_load_refuses_to_overwrite_a_session_started_while_it_was_read(
    monkeypatch,
):
    """Same guarantee for opening a project — it replaces the whole session."""
    real_stage = app_module._stage_project_load
    entered = threading.Event()
    release = threading.Event()

    def blocking_stage(payload: bytes):
        entered.set()
        release.wait(_BLOCK_SECONDS)
        return real_stage(payload)

    with TestClient(app_module.create_app()) as client:
        saved = client.get("/api/project/save")
        assert saved.status_code == 200
        monkeypatch.setattr(app_module, "_stage_project_load", blocking_stage)

        outcome: dict = {}

        def run_load() -> None:
            outcome["response"] = client.post(
                "/api/project/load-file",
                files={
                    "file": ("session.baspec", saved.content, "application/zip")
                },
            )

        worker = threading.Thread(target=run_load, daemon=True)
        worker.start()
        assert entered.wait(_BLOCK_SECONDS + 5), "load never started"
        # The reviewer's scenario exactly: start a new session, then do work
        # in it, while the earlier open is still being read.
        assert client.post("/api/session/reset").status_code == 200
        edit = client.post(
            "/api/doc/edit",
            json={
                "ops": [
                    {
                        "action": "replace",
                        "target_id": "sec",
                        "text": "WORK DONE IN THE NEW SESSION",
                        "numbering": "23 05 00",
                    }
                ]
            },
        )
        assert edit.status_code == 200
        release.set()
        worker.join(_BLOCK_SECONDS + 10)

    response = outcome["response"]
    assert response.status_code == 409
    assert "session was replaced" in response.json()["error"]
    # The work done in the new session survives — the stale load never replayed.
    assert sessions.get_session().doc.doc.title == "WORK DONE IN THE NEW SESSION"


def test_a_real_import_keeps_the_server_answering():
    """End-to-end: no monkeypatching, a genuine multi-hundred-block master.

    The offload has to cover the *whole* handler, not just the parse. Building
    the response payload runs the first source-capability sweep, which on an
    imported master is the most expensive thing the app does — left on the
    event loop it would freeze the server exactly like the parse used to.
    """
    # Sized so the response-building phase (the capability sweep) is several
    # seconds on its own — small enough to stay a reasonable test, large
    # enough that leaving any phase on the loop is unmissable.
    source = _master_bytes(articles=6, paragraphs=8)
    health_times: list[float] = []
    done = threading.Event()

    with TestClient(app_module.create_app()) as client:
        outcome: dict = {}

        def run_import() -> None:
            try:
                outcome["response"] = client.post(
                    "/api/import/master", files=_upload(source)
                )
            finally:
                done.set()

        worker = threading.Thread(target=run_import, daemon=True)
        worker.start()
        # Poll health for as long as the import runs. Every sample must come
        # back promptly; one slow sample means the loop was blocked.
        while not done.wait(0.05):
            started = time.perf_counter()
            assert client.get("/api/health").status_code == 200
            health_times.append(time.perf_counter() - started)
            if len(health_times) > 400:
                break
        worker.join(120)

    assert outcome["response"].status_code == 200
    assert health_times, "the import finished too fast to sample — raise the fixture size"
    assert max(health_times) < _RESPONSIVE_SECONDS, (
        f"slowest health check during a real import was {max(health_times):.2f}s "
        f"over {len(health_times)} samples — the server stalls while importing"
    )


def test_import_reports_pending_permissions_without_sweeping_inline(monkeypatch):
    """The import response must not wait on the per-element permission sweep.

    That sweep probes every element against the real source gate — O(document)
    work per probe, ~5 probes per paragraph — so waiting for it made the
    upload take minutes on a real master. It now runs in the background and
    the response says so.

    The sweep is held for the duration of the assertions rather than raced
    against: this fixture is small enough that a fast runner can finish
    sweeping before the response is even read, which made the ``pending``
    assertion fail intermittently (green on 3.11, red on 3.12, same commit).
    Blocking the worker is the technique the tests above already use, and it
    sharpens what is being tested — with the sweep provably unfinished, a
    response reporting anything but ``pending`` can only mean it waited.
    """
    source = _master_bytes(articles=2, paragraphs=3)
    release = threading.Event()
    real_sweep = SessionState._sweep_and_publish

    # ``*args`` deliberately: the sweep is reached from a blocking caller and
    # from the warm thread, and its parameter list has already changed once.
    # A fixed arity here fails inside a background thread, which surfaces as
    # a confusing assertion on the *next* line rather than as itself.
    def held_sweep(self, *args, **kwargs):
        release.wait(_BLOCK_SECONDS)
        return real_sweep(self, *args, **kwargs)

    monkeypatch.setattr(SessionState, "_sweep_and_publish", held_sweep)

    with TestClient(app_module.create_app()) as client:
        response = client.post("/api/import/master", files=_upload(source))
        assert response.status_code == 200, response.text
        capabilities = response.json()["source_capabilities"]
        assert capabilities is not None
        assert capabilities["status"] == "pending"
        # Pending is fail-closed: an undecided permission is denied, never
        # granted. Workspace-only metadata stays available because it does
        # not touch the retained Word body.
        operations = capabilities["elements"]["pt1.a1.p1"]
        assert operations["replace_text"]["allowed"] is False
        assert operations["replace_text"]["blocker"] == "capabilities_pending"
        assert operations["delete"]["allowed"] is False
        assert operations["set_status"]["allowed"] is True

        # ...and once released, the background sweep settles into a real
        # report, which the panel picks up from the dedicated polling route.
        release.set()
        settle_capability_sweep()
        settled = client.get("/api/doc/capabilities")
        assert settled.status_code == 200, settled.text
        report = settled.json()["source_capabilities"]
        assert report["status"] == "ready"
        assert report["elements"]["pt1.a1.p1"]["replace_text"]["allowed"] is True


def test_a_pending_report_still_refuses_a_forged_body_edit():
    """Fail-closed survives the async window.

    The report is guidance, never authorization, so the point is not that
    ``pending`` denies the operation in the payload — it is that the real
    gate refuses the edit anyway. Here the edit is one the gate rejects on
    its own merits; the capability report never gets a vote either way.
    """
    source = _master_bytes(articles=4, paragraphs=4)

    with TestClient(app_module.create_app()) as client:
        assert client.post(
            "/api/import/master", files=_upload(source)
        ).status_code == 200
        # A heading change is outside the source-preserving boundary.
        forged = client.post(
            "/api/doc/edit",
            json={
                "ops": [
                    {
                        "action": "replace",
                        "target_id": "sec",
                        "text": "FORGED SECTION HEADING",
                    }
                ]
            },
        )
        assert forged.status_code == 400, forged.text
        assert "source-backed edit rejected" in forged.json()["error"].lower()


def test_the_model_is_told_permissions_are_pending_not_read_only():
    """The boundary block must not describe a pending sweep as read-only.

    Every operation is denied while the sweep runs, so rendering the report
    the ordinary way would tell the model the whole imported document is
    permanently read-only — and the model would repeat that to the user.
    """
    source = _master_bytes(articles=6, paragraphs=8)

    with TestClient(app_module.create_app()) as client:
        assert client.post(
            "/api/import/master", files=_upload(source)
        ).status_code == 200
        session = sessions.get_session()
        assert not session.capability_warm_settled(0.0), (
            "the fixture is too small to still be sweeping — raise its size"
        )
        pending_block = _source_editing_boundary_block(session)
        assert pending_block is not None
        assert "still being derived" in pending_block
        # The claim that would be false and that the model would repeat.
        assert "All other imported body IDs are read-only" not in pending_block
        assert "Text-editable IDs" not in pending_block

        settle_capability_sweep()
        settled_block = _source_editing_boundary_block(session)

    assert settled_block is not None
    assert "still being derived" not in settled_block
    assert "Text-editable IDs" in settled_block


def test_a_chat_turn_starts_promptly_on_a_freshly_imported_master():
    """The reported symptom: chat froze after importing a spec.

    Every turn renders the imported-source editing boundary into its PROJECT
    CONTEXT, and that used to derive the permission report inline — before
    ``stream_user_turn`` yielded a single SSE frame. On a freshly imported
    master the memo is cold, so the user watched an empty chat for as long as
    the sweep took. The turn must now start while the sweep is still running.
    """
    source = _master_bytes(articles=6, paragraphs=8)

    with TestClient(app_module.create_app()) as client:
        assert client.post(
            "/api/import/master", files=_upload(source)
        ).status_code == 200

        session = sessions.get_session()
        # Deliberately do NOT settle: the sweep for this state is in flight,
        # which is exactly the window the freeze used to live in.
        assert not session.capability_warm_settled(0.0), (
            "the fixture is too small to still be sweeping — raise its size"
        )

        started = time.perf_counter()
        context = _turn_context_text(session)
        elapsed = time.perf_counter() - started

    assert elapsed < _RESPONSIVE_SECONDS, (
        f"building a turn's PROJECT CONTEXT took {elapsed:.2f}s on a freshly "
        "imported master — the turn is waiting on the capability sweep again"
    )
    # The model is still told the boundary exists; it just says the exact
    # per-element map is still being derived. The real gate remains authority.
    assert "IMPORTED DOCX EDITING BOUNDARY" in context


def test_anchor_lookups_never_rescan_the_element_table():
    """Anchor resolution must be O(1), not a scan of every element.

    Binding a master's anchors calls these lookups once per body child. When
    each call walked the whole ``elements`` tuple the import was quadratic —
    ~12s of blocking CPU for a 6k-paragraph section. Counting iterations pins
    the complexity class directly, with no dependence on machine speed.
    """
    from backend.spec_doc.xml_lexical import build_source_xml_index

    index = build_source_xml_index(_document_xml(_master_bytes(articles=6)))
    snapshot = list(index.elements)
    assert len(snapshot) > 50, "fixture too small to be meaningful"

    scans = {"count": 0}

    class CountingTuple(tuple):
        def __iter__(self):
            scans["count"] += 1
            return tuple.__iter__(self)

    object.__setattr__(index, "elements", CountingTuple(snapshot))

    for element in snapshot:
        index.element_for_span(element.element_span)
        index.direct_children(element)

    # One pass to build each of the two element-keyed maps, and never again —
    # regardless of how many lookups follow.
    assert scans["count"] <= 2, (
        f"{scans['count']} full scans of the element table for "
        f"{len(snapshot) * 2} lookups — the memoized maps are not being used"
    )


def test_memoized_lookups_match_a_linear_scan():
    """The memoized anchor maps return exactly what the old scans returned."""
    import pytest

    from backend.spec_doc.xml_lexical import (
        XmlLexicalError,
        build_source_xml_index,
    )

    index = build_source_xml_index(
        _document_xml(_master_bytes(articles=4, paragraphs=3))
    )

    for child in index.body_children:
        assert index.body_child(child.body_child_index) is child
    for node in index.word_text_nodes:
        assert (
            index.word_text(node.body_child_index, node.text_node_ordinal)
            is node
        )
    for element in index.elements:
        assert index.element_for_span(element.element_span) is element
        expected = tuple(
            candidate
            for candidate in index.elements
            if candidate.parent_start == element.element_span.start
        )
        assert index.direct_children(element) == expected

    # Misses still raise the same typed blockers rather than returning None.
    with pytest.raises(XmlLexicalError):
        index.body_child(10_000)
    with pytest.raises(XmlLexicalError):
        index.word_text(10_000, 0)


def test_docx_export_does_not_hold_the_turn_lock_while_it_renders(monkeypatch):
    """A DOCX render must not block the lock a chat turn has to claim.

    Export used to hold ``session_state_guard`` — the turn-state lock —
    through ZIP rebuild, XML reparse, source-plan validation and the
    python-docx render, which is seconds on a real section. For that whole
    time no turn could start and no stop could be processed, on the one
    action a user is most likely to take right after asking for a draft.

    Coherence never needed the lock held, only the inputs captured
    together, so export now snapshots under the guard and renders without
    it. ``/api/doc`` is the probe because it takes exactly that lock.
    """
    real_build = app_module.build_docx
    entered = threading.Event()
    release = threading.Event()

    def blocking_build(*args, **kwargs):
        entered.set()
        # Stands in for the seconds of real render CPU.
        release.wait(_BLOCK_SECONDS)
        return real_build(*args, **kwargs)

    monkeypatch.setattr(app_module, "build_docx", blocking_build)

    with TestClient(app_module.create_app()) as client:
        assert client.post(
            "/api/doc/edit",
            json={
                "ops": [
                    {"action": "add_article", "target_id": "pt1", "text": "SUMMARY"}
                ]
            },
        ).status_code == 200

        worker_outcome: dict = {}

        def run_export() -> None:
            worker_outcome["response"] = client.get("/api/export/docx")

        worker = threading.Thread(target=run_export, daemon=True)
        worker.start()
        assert entered.wait(_BLOCK_SECONDS + 5), "the render never started"
        started = time.perf_counter()
        doc = client.get("/api/doc")
        elapsed = time.perf_counter() - started
        release.set()
        worker.join(_BLOCK_SECONDS + 10)

    assert doc.status_code == 200
    assert worker_outcome["response"].status_code == 200
    assert elapsed < _RESPONSIVE_SECONDS, (
        f"a request needing the turn-state lock waited {elapsed:.2f}s for a "
        "DOCX render — the export is holding the lock across it again"
    )


def test_source_preserving_export_does_not_hold_the_turn_lock_either(monkeypatch):
    """The same guarantee for the source-preserving branch.

    That path is the heavier of the two — a lexical index, a source-plan
    validation and a raw-ZIP rebuild — so it is the one that held the lock
    longest. Rendering from a captured snapshot means the imported bytes,
    map and baseline are all bound before the guard releases.
    """
    from backend.spec_doc import source_patch as source_patch_module

    real_build = app_module.build_source_preserving_docx
    entered = threading.Event()
    release = threading.Event()

    def blocking_build(**kwargs):
        entered.set()
        release.wait(_BLOCK_SECONDS)
        return real_build(**kwargs)

    monkeypatch.setattr(app_module, "build_source_preserving_docx", blocking_build)
    assert source_patch_module.build_source_preserving_docx is not blocking_build

    with TestClient(app_module.create_app()) as client:
        assert client.post(
            "/api/import/master", files=_upload(_master_bytes())
        ).status_code == 200
        settle_capability_sweep()

        worker_outcome: dict = {}

        def run_export() -> None:
            worker_outcome["response"] = client.get(
                "/api/export/docx", params={"mode": "source"}
            )

        worker = threading.Thread(target=run_export, daemon=True)
        worker.start()
        assert entered.wait(_BLOCK_SECONDS + 5), "the source render never started"
        started = time.perf_counter()
        doc = client.get("/api/doc")
        elapsed = time.perf_counter() - started
        release.set()
        worker.join(_BLOCK_SECONDS + 10)

    assert doc.status_code == 200
    assert worker_outcome["response"].status_code == 200
    assert elapsed < _RESPONSIVE_SECONDS, (
        f"a request needing the turn-state lock waited {elapsed:.2f}s for a "
        "source-preserving render"
    )
