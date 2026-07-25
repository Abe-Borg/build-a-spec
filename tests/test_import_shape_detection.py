"""Spec-shape detection and honest framing for imports that are not specs.

Import always builds a SectionFormat tree — the edit, lint, diff, and QC
machinery is typed to it, and that is not changing. What these tests pin is
that the app stops *asserting* the parts of that scaffolding the imported file
never had: a section header, empty parts, and a "you forgot the section
number" lint finding.

The other half of the contract is that a genuine spec master is untouched by
any of it, so every gate here is also exercised from the spec side.
"""
from __future__ import annotations

import io

from docx import Document
from fastapi.testclient import TestClient

from backend import sessions
from backend.app import create_app
from backend.llm.conversation import _turn_context_text
from backend.spec_doc.importer import (
    UNSTRUCTURED_IMPORT_WARNING,
    parse_master_docx,
)
from backend.spec_doc.linting import RULE_MISSING_SECTION_HEADER, lint_document
from backend.spec_doc.source_package import sanitize_import_report
from backend.spec_modules import get_module

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _docx_bytes(lines: list[str]) -> bytes:
    document = Document()
    for text in lines:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _write(tmp_path, name: str, lines: list[str]):
    path = tmp_path / name
    path.write_bytes(_docx_bytes(lines))
    return path


MEMO = [
    "MEMORANDUM",
    "TO: Project Team",
    "RE: Chilled water plant capacity review",
    "The plant has adequate capacity for the new classroom wing.",
]

SPEC = [
    "SECTION 21 13 13 - WET-PIPE SPRINKLER SYSTEMS",
    "PART 1 - GENERAL",
    "1.1 SUMMARY",
    "A. Section includes wet-pipe sprinkler systems.",
    "END OF SECTION 21 13 13",
]


# ---------------------------------------------------------------------------
# Detection
# ---------------------------------------------------------------------------


def test_a_document_with_no_spec_markers_is_flagged_unstructured(tmp_path):
    result = parse_master_docx(_write(tmp_path, "memo.docx", MEMO))

    assert result.spec_shape_detected is False
    # The verdict leads, so it reads before the parse notes it explains.
    assert result.warnings[0] == UNSTRUCTURED_IMPORT_WARNING
    # Nothing is dropped to earn that verdict.
    assert result.imported_block_count == len(MEMO)


def test_a_real_spec_master_is_untouched_by_detection(tmp_path):
    result = parse_master_docx(_write(tmp_path, "spec.docx", SPEC))

    assert result.spec_shape_detected is True
    assert UNSTRUCTURED_IMPORT_WARNING not in result.warnings
    assert result.section.number == "21 13 13"


def test_any_single_marker_is_enough_to_count_as_a_spec(tmp_path):
    """A partial master must not be demoted for missing the other markers."""
    cases = {
        "section-line-only": ["SECTION 23 05 00 - COMMON WORK", "Some body text."],
        "part-heading-only": ["PART 2 - PRODUCTS", "A. Provide listed devices."],
        "article-only": ["1.1 SUMMARY", "A. Section includes the work."],
    }
    for name, lines in cases.items():
        result = parse_master_docx(_write(tmp_path, f"{name}.docx", lines))
        assert result.spec_shape_detected is True, name
        assert UNSTRUCTURED_IMPORT_WARNING not in result.warnings, name


def test_detection_agrees_with_the_tree_that_was_built(tmp_path):
    """A numbered paragraph never reaches the heading patterns, so text that
    merely *looks* like a PART inside a Word list must not count as a marker —
    otherwise the verdict would claim structure the parse did not act on."""
    document = Document()
    paragraph = document.add_paragraph("PART 2 - PRODUCTS")
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = 1
    buffer = io.BytesIO()
    document.save(buffer)
    path = tmp_path / "numbered.docx"
    path.write_bytes(buffer.getvalue())

    result = parse_master_docx(path)

    assert result.spec_shape_detected is False
    # It landed as a provision, not as a part — the verdict matches the tree.
    assert not any(part.articles for part in result.section.parts[1:])


# ---------------------------------------------------------------------------
# Lint gating
# ---------------------------------------------------------------------------


def test_missing_section_header_is_suppressed_for_an_unstructured_import(tmp_path):
    result = parse_master_docx(_write(tmp_path, "memo.docx", MEMO))
    module = get_module(None)

    ungated = lint_document(result.section, module)
    gated = lint_document(result.section, module, unstructured_import=True)

    assert RULE_MISSING_SECTION_HEADER in {i["rule"] for i in ungated}
    assert RULE_MISSING_SECTION_HEADER not in {i["rule"] for i in gated}


def test_the_gate_changes_nothing_else_about_lint(tmp_path):
    """Only the header rule is suppressed; every other finding still stands."""
    lines = [
        "MEMORANDUM",
        "Comply with NFPA 13-2019 for the sprinkler work.",
        "TODO: confirm the hazard classification.",
    ]
    result = parse_master_docx(_write(tmp_path, "markers.docx", lines))
    module = get_module(None)

    gated = lint_document(result.section, module, unstructured_import=True)
    rules = {i["rule"] for i in gated}

    assert RULE_MISSING_SECTION_HEADER not in rules
    # A questionable edition citation and a template marker are real findings
    # in any prose. Which edition rule fires depends on whether the active
    # module pins editions, and that is not what this test is about.
    assert rules & {"stale_edition", "unrecorded_edition"}
    assert "template_marker" in rules


def test_the_gate_is_a_no_op_for_a_spec_master(tmp_path):
    result = parse_master_docx(_write(tmp_path, "spec.docx", SPEC))
    module = get_module(None)

    assert lint_document(result.section, module) == lint_document(
        result.section, module, unstructured_import=False
    )


# ---------------------------------------------------------------------------
# Persistence
# ---------------------------------------------------------------------------


def test_the_flag_survives_a_project_round_trip():
    report = {
        "filename": "memo.docx",
        "sha256": "a" * 64,
        "size_bytes": 10,
        "zip_member_count": 3,
        "zip_uncompressed_bytes": 100,
        "imported_block_count": 4,
        "skipped_empty_count": 0,
        "warnings": [UNSTRUCTURED_IMPORT_WARNING],
        "tracked_changes_detected": False,
        "spec_shape_detected": False,
    }

    restored = sanitize_import_report(report)

    assert restored is not None
    assert restored["spec_shape_detected"] is False


def test_a_legacy_report_without_the_flag_still_loads_as_spec_shaped():
    """Projects written before detection existed must not start suppressing
    chrome, and must certainly not fail to load."""
    legacy = {
        "filename": "master.docx",
        "sha256": "b" * 64,
        "size_bytes": 10,
        "zip_member_count": 3,
        "zip_uncompressed_bytes": 100,
        "imported_block_count": 4,
        "skipped_empty_count": 0,
        "warnings": [],
        "tracked_changes_detected": False,
    }

    restored = sanitize_import_report(legacy)

    assert restored is not None
    assert restored["spec_shape_detected"] is True


def test_a_malformed_flag_degrades_to_spec_shaped():
    forged = {
        "filename": "master.docx",
        "sha256": "c" * 64,
        "size_bytes": 10,
        "zip_member_count": 3,
        "zip_uncompressed_bytes": 100,
        "imported_block_count": 4,
        "skipped_empty_count": 0,
        "warnings": [],
        "tracked_changes_detected": False,
        "spec_shape_detected": "no",
    }

    restored = sanitize_import_report(forged)

    assert restored is not None
    assert restored["spec_shape_detected"] is True


# ---------------------------------------------------------------------------
# End to end through the API + the model's context
# ---------------------------------------------------------------------------


def _import(client: TestClient, name: str, lines: list[str]):
    return client.post(
        "/api/import/master",
        files={"file": (name, _docx_bytes(lines), DOCX_MEDIA_TYPE)},
    )


def test_importing_a_non_spec_file_reports_honestly_and_lints_clean():
    client = TestClient(create_app())

    body = _import(client, "memo.docx", MEMO).json()

    assert body["ok"] is True
    assert body["import_report"]["spec_shape_detected"] is False
    assert body["warnings"][0] == UNSTRUCTURED_IMPORT_WARNING
    assert RULE_MISSING_SECTION_HEADER not in {i["rule"] for i in body["lint"]}
    # The content is all there; only the false assertions are gone.
    assert body["import_report"]["imported_block_count"] == len(MEMO)


def test_importing_a_spec_master_keeps_the_spec_posture():
    client = TestClient(create_app())

    body = _import(client, "spec.docx", SPEC).json()

    assert body["import_report"]["spec_shape_detected"] is True
    assert UNSTRUCTURED_IMPORT_WARNING not in body["warnings"]


def test_an_unstructured_import_does_not_break_exact_source_export():
    """The honesty work must not disturb the preservation contract."""
    client = TestClient(create_app())
    payload = _docx_bytes(MEMO)

    client.post(
        "/api/import/master",
        files={"file": ("memo.docx", payload, DOCX_MEDIA_TYPE)},
    )
    response = client.get("/api/export/docx?mode=source")

    assert response.status_code == 200
    assert response.content == payload


def test_the_model_is_told_the_scaffolding_is_not_the_users_document():
    client = TestClient(create_app())
    _import(client, "memo.docx", MEMO)
    session = sessions.get_session()

    assert session.import_is_unstructured() is True
    context = _turn_context_text(session)
    assert "IMPORTED DOCUMENT IS NOT A SPEC SECTION" in context
    # It must say what to do instead of silently inventing a header.
    assert "do not invent a section number" in context


def test_a_spec_import_carries_no_such_block():
    client = TestClient(create_app())
    _import(client, "spec.docx", SPEC)
    session = sessions.get_session()

    assert session.import_is_unstructured() is False
    assert "IMPORTED DOCUMENT IS NOT A SPEC SECTION" not in _turn_context_text(
        session
    )


def test_a_session_with_no_import_is_never_unstructured():
    """A from-scratch draft still wants its [TBD] header placeholders."""
    assert sessions.get_session().import_is_unstructured() is False


def test_the_framing_lives_in_project_context_not_the_cached_prompt():
    """The cache rule: nothing session-varying may enter the stable block."""
    from backend.llm.prompts import render_system_prompt

    client = TestClient(create_app())
    _import(client, "memo.docx", MEMO)
    session = sessions.get_session()

    stable = render_system_prompt(session.module)
    assert "IMPORTED DOCUMENT IS NOT A SPEC SECTION" not in stable
