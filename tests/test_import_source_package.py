"""Safety, accounting, and recovery tests for imported DOCX packages.

P0 established the bounded validation and honesty boundary. P1 keeps those
guarantees while making a retained source the default export base; normalized
reconstruction remains an explicit mode.
"""
from __future__ import annotations

import asyncio
import base64
import hashlib
import io
import json
import struct
import warnings
import zipfile

import pytest
from docx import Document
from fastapi import UploadFile
from fastapi.testclient import TestClient

from backend import sessions
from backend.app import create_app
from backend.spec_doc.source_package import (
    SourcePackageError,
    UploadTooLargeError,
    inspect_docx_package,
    read_upload_bounded,
)


DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
REQUIRED_DOCX_MEMBERS = (
    "[Content_Types].xml",
    "_rels/.rels",
    "word/document.xml",
)
IMPORT_REPORT_KEYS = {
    "filename",
    "sha256",
    "size_bytes",
    "zip_member_count",
    "zip_uncompressed_bytes",
    "imported_block_count",
    "skipped_empty_count",
    "warnings",
    "tracked_changes_detected",
    "fidelity_notice",
}


def _client() -> TestClient:
    return TestClient(create_app())


def _legacy_project_payload() -> dict:
    """Return the source-less inner JSON shape used before .baspec files."""
    project = json.loads(json.dumps(sessions.project_payload(sessions.get_session())))
    for key in ("source_map", "source_body_map", "source_docx"):
        project.pop(key, None)
    return project


def _master_docx_bytes(*, header_marker: str = "SOURCE HEADER â€” KEEP") -> bytes:
    """Small valid master with one intentionally empty body paragraph."""
    document = Document()
    document.sections[0].header.paragraphs[0].text = header_marker
    for text in (
        "SECTION 21 13 13",
        "WET-PIPE SPRINKLER SYSTEMS",
        "PART 1 - GENERAL",
        "1.1 SUMMARY",
        "",
        "A. Section includes wet-pipe sprinkler systems.",
        "END OF SECTION 21 13 13",
    ):
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _empty_docx_bytes() -> bytes:
    document = Document()
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _post_master(
    client: TestClient,
    payload: bytes,
    *,
    filename: str = "Client Master 21 13 13.docx",
):
    return client.post(
        "/api/import/master",
        files={"file": (filename, payload, DOCX_MEDIA_TYPE)},
    )


def _rewrite_package(
    payload: bytes,
    *,
    omit: set[str] | None = None,
    additions: list[tuple[str, bytes]] | None = None,
) -> bytes:
    """Copy a package, optionally dropping or appending ZIP members."""
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(payload), "r") as source:
        with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as dest:
            for info in source.infolist():
                if info.filename not in (omit or set()):
                    dest.writestr(info.filename, source.read(info))
            for name, data in additions or []:
                # Duplicate names are intentional in one safety fixture.
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore", UserWarning)
                    dest.writestr(name, data)
    return output.getvalue()


def _mark_zip_members_encrypted(payload: bytes) -> bytes:
    """Set the encryption bit in local and central ZIP headers.

    Python's stdlib cannot create an encrypted archive, but inspection only
    needs a standards-compliant flag to reject one before extraction.  CRCs
    and member data stay untouched.
    """
    mutated = bytearray(payload)
    for signature, flag_offset in ((b"PK\x03\x04", 6), (b"PK\x01\x02", 8)):
        cursor = 0
        while True:
            cursor = mutated.find(signature, cursor)
            if cursor < 0:
                break
            flags = struct.unpack_from("<H", mutated, cursor + flag_offset)[0]
            struct.pack_into("<H", mutated, cursor + flag_offset, flags | 0x1)
            cursor += len(signature)
    return bytes(mutated)


def _assert_clean_import_rejection(response, *, status_code: int = 400) -> None:
    assert response.status_code == status_code
    body = response.json()
    assert body["ok"] is False
    assert isinstance(body.get("error"), str) and body["error"].strip()
    # Public failures must not leak a traceback or parser internals.
    assert "traceback" not in body["error"].lower()


def test_bounded_upload_reader_accepts_boundary_and_rejects_next_byte():
    exact = UploadFile(filename="exact.docx", file=io.BytesIO(b"12345678"))
    assert (
        asyncio.run(read_upload_bounded(exact, max_bytes=8, chunk_bytes=3))
        == b"12345678"
    )

    over = UploadFile(filename="over.docx", file=io.BytesIO(b"123456789"))
    with pytest.raises(UploadTooLargeError):
        asyncio.run(read_upload_bounded(over, max_bytes=8, chunk_bytes=3))


def test_import_endpoint_uses_bounded_upload_reader(monkeypatch):
    """Exercise endpoint wiring without allocating the production 25 MiB."""
    from backend.spec_doc.source_package import read_upload_bounded as real_reader

    async def tiny_reader(upload, **_kwargs):
        return await real_reader(upload, max_bytes=64, chunk_bytes=17)

    monkeypatch.setattr("backend.app.read_upload_bounded", tiny_reader)
    client = _client()
    response = _post_master(client, _master_docx_bytes())
    _assert_clean_import_rejection(response, status_code=413)
    assert "large" in response.json()["error"].lower()
    assert client.get("/api/import/original").status_code == 404


def test_package_inspection_reports_size_and_has_configurable_limits():
    payload = _master_docx_bytes()
    info = inspect_docx_package(payload)
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        members = package.infolist()
        expected_uncompressed = sum(member.file_size for member in members)
        largest_member = max(member.file_size for member in members)

    assert info.member_count == len(members)
    assert info.uncompressed_bytes == expected_uncompressed

    with pytest.raises(SourcePackageError):
        inspect_docx_package(payload, max_members=info.member_count - 1)
    with pytest.raises(SourcePackageError):
        inspect_docx_package(
            payload, max_total_uncompressed=info.uncompressed_bytes - 1
        )
    with pytest.raises(SourcePackageError):
        inspect_docx_package(
            payload, max_member_uncompressed=largest_member - 1
        )


@pytest.mark.parametrize("required_member", REQUIRED_DOCX_MEMBERS)
def test_package_inspection_rejects_each_missing_required_part(required_member):
    payload = _rewrite_package(
        _master_docx_bytes(),
        omit={required_member},
    )
    with pytest.raises(SourcePackageError):
        inspect_docx_package(payload)


def test_package_inspection_rejects_traversal_duplicate_and_encryption():
    original = _master_docx_bytes()
    with zipfile.ZipFile(io.BytesIO(original)) as package:
        document_xml = package.read("word/document.xml")

    traversal = _rewrite_package(original, additions=[("../escape.xml", b"no")])
    duplicate = _rewrite_package(
        original, additions=[("word/document.xml", document_xml)]
    )
    encrypted = _mark_zip_members_encrypted(original)

    for unsafe in (traversal, duplicate, encrypted):
        with pytest.raises(SourcePackageError):
            inspect_docx_package(unsafe)


@pytest.mark.parametrize(
    "payload",
    [
        b"not a zip package",
        pytest.param(
            _rewrite_package(_master_docx_bytes(), omit={"word/document.xml"}),
            id="missing-document-part",
        ),
        pytest.param(
            _rewrite_package(
                _master_docx_bytes(), additions=[("../escape.xml", b"no")]
            ),
            id="path-traversal",
        ),
        pytest.param(
            _rewrite_package(
                _master_docx_bytes(),
                additions=[("word/document.xml", b"duplicate")],
            ),
            id="duplicate-member",
        ),
        pytest.param(
            _mark_zip_members_encrypted(_master_docx_bytes()),
            id="encrypted-member",
        ),
    ],
)
def test_import_endpoint_rejects_unsafe_packages_without_retaining_them(payload):
    client = _client()
    response = _post_master(client, payload)
    _assert_clean_import_rejection(response)

    doc = client.get("/api/doc").json()
    assert doc["doc"]["section"]["number"] == ""
    assert doc["import_report"] is None
    assert doc["source_available"] is False
    assert client.get("/api/import/original").status_code == 404


@pytest.mark.parametrize(
    ("payload", "passes_package_inspection"),
    [
        pytest.param(
            _empty_docx_bytes(),
            True,
            id="no-importable-body-content",
        ),
        pytest.param(
            _rewrite_package(
                _master_docx_bytes(),
                omit={"word/document.xml"},
                additions=[("word/document.xml", b"<malformed-document-xml")],
            ),
            False,
            id="malformed-document-xml",
        ),
    ],
)
def test_validation_or_parse_failure_is_atomic_and_retains_nothing(
    payload, passes_package_inspection
):
    """Failures at either side of the python-docx boundary stay atomic."""
    client = _client()
    if passes_package_inspection:
        inspect_docx_package(payload)  # this fixture reaches python-docx
    else:
        with pytest.raises(SourcePackageError):
            inspect_docx_package(payload)

    response = _post_master(client, payload)
    _assert_clean_import_rejection(response)
    snapshot = client.get("/api/doc").json()
    assert snapshot["doc"]["section"]["number"] == ""
    assert snapshot["import_report"] is None
    assert snapshot["source_available"] is False
    assert client.get("/api/import/original").status_code == 404


def test_successful_import_reports_honestly_and_retains_exact_source():
    client = _client()
    header_marker = "SOURCE HEADER â€” EXACT ORIGINAL ONLY"
    source = _master_docx_bytes(header_marker=header_marker)
    package_info = inspect_docx_package(source)

    response = _post_master(client, source)
    assert response.status_code == 200
    body = response.json()
    assert body["ok"] is True
    assert body["skipped_empty_count"] == 1
    assert body["source_available"] is True

    report = body["import_report"]
    assert set(report) == IMPORT_REPORT_KEYS
    assert report["filename"] == "Client Master 21 13 13.docx"
    assert report["sha256"] == hashlib.sha256(source).hexdigest()
    assert report["size_bytes"] == len(source)
    assert report["zip_member_count"] == package_info.member_count
    assert report["zip_uncompressed_bytes"] == package_info.uncompressed_bytes
    assert report["imported_block_count"] == body["imported_block_count"] == 1
    assert report["skipped_empty_count"] == body["skipped_empty_count"] == 1
    assert report["warnings"] == body["warnings"]
    assert report["tracked_changes_detected"] is False
    assert isinstance(report["fidelity_notice"], str)
    assert report["fidelity_notice"].strip()
    assert "source_available" not in report

    snapshot = client.get("/api/doc").json()
    assert snapshot["import_report"] == report
    assert snapshot["source_available"] is True

    original = client.get("/api/import/original")
    assert original.status_code == 200
    assert original.content == source
    assert original.headers["content-type"].startswith(DOCX_MEDIA_TYPE)
    assert original.headers["cache-control"] == "no-store"
    assert original.headers["x-content-type-options"] == "nosniff"
    assert "attachment" in original.headers["content-disposition"]
    assert "Client%20Master%2021%2013%2013.docx" in original.headers[
        "content-disposition"
    ]

    # P1 makes the source-preserving path the imported-document default. The
    # old semantic reconstruction still exists, but only by explicit request.
    preserving = client.get("/api/export/docx")
    assert preserving.status_code == 200
    assert preserving.content == source
    preserved_doc = Document(io.BytesIO(preserving.content))
    preserved_headers = "\n".join(
        paragraph.text
        for section in preserved_doc.sections
        for paragraph in section.header.paragraphs
    )
    assert header_marker in preserved_headers

    normalized = client.get("/api/export/docx", params={"mode": "normalized"})
    assert normalized.status_code == 200
    assert normalized.content != source
    normalized_doc = Document(io.BytesIO(normalized.content))
    normalized_headers = "\n".join(
        paragraph.text
        for section in normalized_doc.sections
        for paragraph in section.header.paragraphs
    )
    assert header_marker not in normalized_headers


def test_reset_clears_retained_source_and_import_report():
    client = _client()
    assert _post_master(client, _master_docx_bytes()).status_code == 200
    assert client.get("/api/import/original").status_code == 200

    assert client.post("/api/session/reset").status_code == 200
    snapshot = client.get("/api/doc").json()
    assert snapshot["import_report"] is None
    assert snapshot["source_available"] is False
    assert client.get("/api/import/original").status_code == 404


def test_legacy_json_project_persists_report_but_never_source_bytes():
    client = _client()
    source = _master_docx_bytes()
    imported = _post_master(client, source).json()
    report = imported["import_report"]

    project = _legacy_project_payload()
    assert project["import_report"] == report
    serialized = json.dumps(project)
    assert "source_docx_bytes" not in serialized
    assert "source_available" not in serialized
    assert base64.b64encode(source).decode("ascii") not in serialized

    client.post("/api/session/reset")
    loaded_response = client.post("/api/project/load", json=project)
    assert loaded_response.status_code == 200
    loaded = loaded_response.json()
    assert loaded["import_report"] == report
    assert loaded["source_available"] is False

    unavailable = client.get("/api/import/original")
    assert unavailable.status_code == 409
    assert unavailable.json()["ok"] is False

    saved_again = _legacy_project_payload()
    assert saved_again["import_report"] == report


def test_project_load_sanitizes_report_and_clears_an_active_source():
    client = _client()
    imported = _post_master(client, _master_docx_bytes()).json()
    canonical_report = imported["import_report"]
    project = _legacy_project_payload()
    project["import_report"].update(
        {
            "filename": "../../evil.docx",
            "sha256": project["import_report"]["sha256"].upper(),
            "warnings": ["  keep this warning  ", 123],
            "fidelity_notice": "Perfect fidelity is guaranteed!",
            "source_bytes": base64.b64encode(_master_docx_bytes()).decode("ascii"),
            "source_available": True,
            "unexpected": "drop me",
        }
    )

    # Load directly over the live import: this must discard the active bytes,
    # retain only the safe report allowlist, and derive availability from the
    # session instead of trusting project JSON.
    loaded_response = client.post("/api/project/load", json=project)
    assert loaded_response.status_code == 200
    loaded = loaded_response.json()
    report = loaded["import_report"]
    assert set(report) == IMPORT_REPORT_KEYS
    assert report["filename"] == "evil.docx"
    assert report["sha256"] == canonical_report["sha256"]
    assert report["warnings"] == ["keep this warning"]
    assert report["fidelity_notice"] == canonical_report["fidelity_notice"]
    assert loaded["source_available"] is False
    assert client.get("/api/import/original").status_code == 409


def test_malformed_optional_import_report_does_not_break_project_load():
    client = _client()
    project = _legacy_project_payload()
    project["import_report"] = {
        "filename": "master.docx",
        "sha256": "not-a-sha256",
    }

    loaded_response = client.post("/api/project/load", json=project)
    assert loaded_response.status_code == 200
    loaded = loaded_response.json()
    assert loaded["import_report"] is None
    assert loaded["source_available"] is False


def test_legacy_project_without_import_report_still_loads_and_clears_source():
    client = _client()
    legacy_project = _legacy_project_payload()
    legacy_project.pop("import_report", None)

    # Loading any project replaces the active session.  A pre-P0 project has
    # no report and must not accidentally inherit a prior import's bytes.
    assert _post_master(client, _master_docx_bytes()).status_code == 200
    assert client.get("/api/import/original").status_code == 200

    loaded_response = client.post("/api/project/load", json=legacy_project)
    assert loaded_response.status_code == 200
    loaded = loaded_response.json()
    assert loaded["import_report"] is None
    assert loaded["source_available"] is False
    assert client.get("/api/import/original").status_code == 404
