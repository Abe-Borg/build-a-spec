"""Master-import tests: round-trip goldens, tracked changes, structure
heuristics — docx fixtures built in-memory with python-docx."""
from __future__ import annotations

import io

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from backend.spec_doc.docx_export import build_docx
from backend.spec_doc.importer import (
    ImportResult,
    MasterImportError,
    parse_master_docx,
)
from backend.spec_doc.model import DocumentStore, SpecSection


def _write_docx(tmp_path, paragraphs: list[str], name="master.docx"):
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    path = tmp_path / name
    document.save(str(path))
    return path


_MASTER_LINES = [
    "SECTION 21 13 13",
    "WET-PIPE SPRINKLER SYSTEMS",
    "PART 1 - GENERAL",
    "1.1 SUMMARY",
    "A. Section includes wet-pipe sprinkler systems.",
    "B. Related sections: 21 13 19.",
    "1.2 REFERENCES",
    "A. NFPA 13 - Standard for the Installation of Sprinkler Systems.",
    "PART 2 - PRODUCTS",
    "2.1 PIPE AND FITTINGS",
    "A. Steel pipe: ASTM A53.",
    "1. Schedule 40 for 2 inches and smaller.",
    "a. Threaded joints.",
    "1) Approved thread sealant.",
    "PART 3 - EXECUTION",
    "3.1 INSTALLATION",
    "A. Install per NFPA 13.",
    "END OF SECTION 21 13 13",
    "This trailing text is after END OF SECTION and must not import.",
]


def test_import_builds_the_tree_with_imported_status(tmp_path):
    result = parse_master_docx(_write_docx(tmp_path, _MASTER_LINES))
    section = result.section
    assert section.number == "21 13 13"
    assert section.title == "WET-PIPE SPRINKLER SYSTEMS"
    assert [a.title for a in section.parts[0].articles] == [
        "SUMMARY",
        "REFERENCES",
    ]
    assert [a.title for a in section.parts[1].articles] == ["PIPE AND FITTINGS"]
    assert [a.title for a in section.parts[2].articles] == ["INSTALLATION"]

    # Nesting: A. -> 1. -> a. -> 1)
    piping = section.parts[1].articles[0]
    level0 = piping.paragraphs[0]
    assert level0.text == "Steel pipe: ASTM A53."
    assert level0.children[0].text.startswith("Schedule 40")
    assert level0.children[0].children[0].text == "Threaded joints."
    assert (
        level0.children[0].children[0].children[0].text
        == "Approved thread sealant."
    )

    # Every block imported; nothing after END OF SECTION.
    statuses = set()

    def collect(paragraphs):
        for p in paragraphs:
            statuses.add(p.status)
            collect(p.children)

    for part in section.parts:
        for article in part.articles:
            collect(article.paragraphs)
    assert statuses == {"imported"}
    assert result.imported_block_count == 8
    # The break at END OF SECTION stands (nothing after it entered the
    # tree), but the drop is no longer SILENT: keep-everything-warn-loudly
    # owes the reviewer a count and a first line for what stayed behind.
    assert not any(
        "trailing" in p.text.lower()
        for part in section.parts
        for article in part.articles
        for p in article.paragraphs
    )
    dropped = [w for w in result.warnings if "END OF SECTION" in w]
    assert len(dropped) == 1
    assert "1 block(s)" in dropped[0]
    assert "This trailing text is after END OF SECTION" in dropped[0]


def test_import_round_trip_through_store_and_export(tmp_path):
    result = parse_master_docx(_write_docx(tmp_path, _MASTER_LINES))
    store = DocumentStore()
    store.adopt_imported(result.section)
    # One version snapshot; undo returns to blank.
    assert store.index == 1 and len(store.versions) == 2
    assert store.undo() and store.doc.is_empty()
    assert store.redo()

    # Serialization round-trip (the integrity check accepts the built ids).
    restored = SpecSection.from_dict(store.doc.to_dict())
    assert restored.parts[1].articles[0].paragraphs[0].status == "imported"

    # Export carries the imported-provisions schedule.
    payload = build_docx(store.doc)
    document = Document(io.BytesIO(payload))
    texts = [p.text for p in document.paragraphs]
    assert "IMPORTED PROVISIONS NOT YET REVIEWED" in texts


def test_import_re_parse_golden(tmp_path):
    """docx → tree → export docx → re-import ≈ same tree shape."""
    first = parse_master_docx(_write_docx(tmp_path, _MASTER_LINES))
    exported = build_docx(first.section)
    exported_path = tmp_path / "exported.docx"
    exported_path.write_bytes(exported)
    second = parse_master_docx(exported_path)
    assert second.section.number == first.section.number
    assert [a.title for a in second.section.parts[0].articles][:2] == [
        "SUMMARY",
        "REFERENCES",
    ]
    # Paragraph counts survive (schedules after END OF SECTION are ignored).
    def count(section):
        total = 0

        def walk(paragraphs):
            nonlocal total
            for p in paragraphs:
                total += 1
                walk(p.children)

        for part in section.parts:
            for article in part.articles:
                walk(article.paragraphs)
        return total

    assert count(second.section) == count(first.section)


def test_unlabeled_and_orphan_content_is_kept_with_warnings(tmp_path):
    lines = [
        "PART 1 - GENERAL",
        "Orphan text before any article.",
        "1.1 SUMMARY",
        "A. Labeled provision.",
        "Continuation prose with no label.",
    ]
    result = parse_master_docx(_write_docx(tmp_path, lines))
    part1 = result.section.parts[0]
    assert [a.title for a in part1.articles] == ["IMPORTED CONTENT", "SUMMARY"]
    assert part1.articles[0].paragraphs[0].text == "Orphan text before any article."
    assert [p.text for p in part1.articles[1].paragraphs] == [
        "Labeled provision.",
        "Continuation prose with no label.",
    ]
    assert any("before any article" in w for w in result.warnings)


def test_tables_flatten_with_warning(tmp_path):
    document = Document()
    document.add_paragraph("PART 2 - PRODUCTS")
    document.add_paragraph("2.1 SCHEDULE")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "K5.6"
    table.rows[0].cells[1].text = "Quick response"
    path = tmp_path / "table.docx"
    document.save(str(path))

    result = parse_master_docx(path)
    schedule = result.section.parts[1].articles[0]
    assert schedule.paragraphs[0].text == "K5.6 | Quick response"
    assert any("tables" in w for w in result.warnings)


def test_tracked_changes_import_accept_all_view(tmp_path):
    document = Document()
    document.add_paragraph("PART 1 - GENERAL")
    document.add_paragraph("1.1 SUMMARY")
    paragraph = document.add_paragraph("A. Comply with ")
    # <w:ins><w:r>NFPA 13-2025</w:r></w:ins> — inserted text (kept).
    ins = OxmlElement("w:ins")
    ins_run = OxmlElement("w:r")
    ins_text = OxmlElement("w:t")
    ins_text.text = "NFPA 13-2025"
    ins_run.append(ins_text)
    ins.append(ins_run)
    paragraph._p.append(ins)
    # <w:del><w:r><w:delText>NFPA 13-2016</w:delText></w:r></w:del> (dropped).
    deleted = OxmlElement("w:del")
    del_run = OxmlElement("w:r")
    del_text = OxmlElement("w:delText")
    del_text.text = "NFPA 13-2016"
    del_run.append(del_text)
    deleted.append(del_run)
    paragraph._p.append(deleted)
    tail = paragraph.add_run(" throughout.")
    # Move the tail run after the revision wrappers (document order).
    paragraph._p.append(tail._r)

    path = tmp_path / "redline.docx"
    document.save(str(path))
    result = parse_master_docx(path)
    text = result.section.parts[0].articles[0].paragraphs[0].text
    assert text == "Comply with NFPA 13-2025 throughout."
    assert "NFPA 13-2016" not in text
    assert result.tracked_changes_detected is True
    assert any("tracked changes" in w.lower() for w in result.warnings)


def test_auto_numbered_master_uses_ilvl(tmp_path):
    document = Document()
    document.add_paragraph("PART 3 - EXECUTION")
    document.add_paragraph("3.1 INSTALLATION")

    def numbered(text: str, ilvl: int):
        paragraph = document.add_paragraph(text)
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(ilvl))
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "1")
        num_pr.append(ilvl_el)
        num_pr.append(num_id)
        p_pr.append(num_pr)

    numbered("Install per the working plans.", 0)
    numbered("Support piping per NFPA 13.", 1)
    path = tmp_path / "autonum.docx"
    document.save(str(path))

    result = parse_master_docx(path)
    article = result.section.parts[2].articles[0]
    top = article.paragraphs[0]
    assert top.text == "Install per the working plans."
    assert top.children[0].text == "Support piping per NFPA 13."


def _numbered(document, text: str, ilvl: int, num_id: str = "1"):
    """A body paragraph carrying real Word auto-numbering at ``ilvl``."""
    paragraph = document.add_paragraph(text)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    return paragraph


def _outline_master(tmp_path, base: int, name="outline.docx"):
    """A master whose numbered outline starts at ``base`` rather than 0.

    Three sibling headings at the outline's top level, each with one child —
    the shape every office master uses, and the one that came back
    mis-parented when ``ilvl`` was read as an absolute depth.
    """
    document = Document()
    document.add_paragraph("SECTION 23 05 48")
    document.add_paragraph("VIBRATION AND SEISMIC CONTROLS FOR HVAC")
    document.add_paragraph("PART 1 - GENERAL")
    _numbered(document, "SECTION INCLUDES", base)
    _numbered(document, "Vibration isolation for HVAC equipment.", base + 1)
    _numbered(document, "REFERENCE STANDARDS", base)
    _numbered(document, "ASHRAE Handbook - HVAC Applications.", base + 1)
    _numbered(document, "QUALITY ASSURANCE", base)
    _numbered(document, "Minimum three years documented experience.", base + 1)
    path = tmp_path / name
    document.save(str(path))
    return path


def test_numbering_depth_is_relative_to_the_documents_own_top_level(tmp_path):
    """An outline starting at ``ilvl`` 1 keeps siblings as siblings.

    ``ilvl`` is an indent level inside a numbering definition, not an
    absolute outline depth. Reading it as one clamped the first numbered
    paragraph of a part to depth 0 and then hung its same-level siblings
    off it as children — a flat list of headings came back as one heading
    owning all the others.
    """
    result = parse_master_docx(_outline_master(tmp_path, base=1))
    article = result.section.parts[0].articles[0]

    assert [p.text for p in article.paragraphs] == [
        "SECTION INCLUDES",
        "REFERENCE STANDARDS",
        "QUALITY ASSURANCE",
    ]
    assert [
        [child.text for child in p.children] for p in article.paragraphs
    ] == [
        ["Vibration isolation for HVAC equipment."],
        ["ASHRAE Handbook - HVAC Applications."],
        ["Minimum three years documented experience."],
    ]
    # The old clamp announced itself; nothing jumped once depth is relative.
    assert not any("jumped deeper" in w for w in result.warnings)


def test_an_outline_starting_at_level_one_matches_one_starting_at_zero(
    tmp_path,
):
    """The same outline is the same tree wherever its levels are numbered."""
    from_zero = parse_master_docx(
        _outline_master(tmp_path, base=0, name="zero.docx")
    )
    from_one = parse_master_docx(
        _outline_master(tmp_path, base=1, name="one.docx")
    )
    assert from_one.section.to_dict() == from_zero.section.to_dict()


def test_a_second_list_that_begins_deeper_than_the_first_is_not_corrupted(
    tmp_path,
):
    """Each article's list answers for itself, not for the document.

    A master can hold several numbering definitions that disagree about
    where they begin. A document-wide minimum is 0 as soon as *any* list
    uses level 0, which leaves every other list to reproduce the original
    sibling-becomes-child corruption.
    """
    document = Document()
    document.add_paragraph("PART 1 - GENERAL")
    document.add_paragraph("1.1 SUMMARY")
    _numbered(document, "Section includes seismic restraint.", 0, num_id="1")
    _numbered(document, "Related: Section 23 05 29.", 0, num_id="1")
    document.add_paragraph("1.2 REFERENCE STANDARDS")
    _numbered(document, "ASHRAE Handbook - HVAC Applications.", 1, num_id="2")
    _numbered(document, "SMACNA Seismic Restraint Manual.", 1, num_id="2")
    path = tmp_path / "twolists.docx"
    document.save(str(path))

    part = parse_master_docx(path).section.parts[0]
    first, second = part.articles[0], part.articles[1]
    assert [p.text for p in first.paragraphs] == [
        "Section includes seismic restraint.",
        "Related: Section 23 05 29.",
    ]
    # The list that begins at level 1 must not adopt its own sibling.
    assert [p.text for p in second.paragraphs] == [
        "ASHRAE Handbook - HVAC Applications.",
        "SMACNA Seismic Restraint Manual.",
    ]
    assert all(p.children == [] for p in second.paragraphs)


def test_a_restarted_sub_list_keeps_the_depth_its_parent_supports(tmp_path):
    """A deliberately deep list under a real parent is left alone.

    Word mints a fresh ``numId`` whenever a list restarts, so a nested
    sub-list routinely carries its own definition at a deep ``ilvl``.
    Rebasing each definition by its own minimum level would read that as
    "this list starts at its top level" and promote it out of its parent.
    Nothing shifts unless a paragraph was going to be pushed shallower
    anyway.
    """
    document = Document()
    document.add_paragraph("PART 2 - PRODUCTS")
    document.add_paragraph("2.1 PIPE AND FITTINGS")
    _numbered(document, "Steel pipe: ASTM A53.", 0, num_id="1")
    _numbered(document, "Schedule 40 for 2 inches and smaller.", 1, num_id="1")
    # A restarted a)/b) sub-list: its own numId, used only at level 2.
    _numbered(document, "Threaded joints.", 2, num_id="7")
    _numbered(document, "Grooved joints.", 2, num_id="7")
    path = tmp_path / "restarted.docx"
    document.save(str(path))

    article = parse_master_docx(path).section.parts[1].articles[0]
    top = article.paragraphs[0]
    assert [p.text for p in article.paragraphs] == ["Steel pipe: ASTM A53."]
    nested = top.children[0]
    assert nested.text == "Schedule 40 for 2 inches and smaller."
    assert [p.text for p in nested.children] == [
        "Threaded joints.",
        "Grooved joints.",
    ]


def test_a_document_that_uses_level_zero_is_parsed_exactly_as_before(tmp_path):
    """The normalization is a no-op for any document reaching level 0.

    Every Build-a-Spec normalized export writes provisions at ``w:ilvl`` 0,
    so this is what keeps the export/re-import round trip byte-identical:
    the base is 0 and no depth moves. A deeper sibling elsewhere in the
    document must not shift the paragraphs that do use level 0.
    """
    document = Document()
    document.add_paragraph("PART 3 - EXECUTION")
    document.add_paragraph("3.1 INSTALLATION")
    _numbered(document, "Install per the working plans.", 0)
    _numbered(document, "Support piping per NFPA 13.", 1)
    _numbered(document, "Brace per the seismic drawings.", 2)
    path = tmp_path / "haslevelzero.docx"
    document.save(str(path))

    article = parse_master_docx(path).section.parts[2].articles[0]
    top = article.paragraphs[0]
    assert top.text == "Install per the working plans."
    assert top.children[0].text == "Support piping per NFPA 13."
    assert top.children[0].children[0].text == "Brace per the seismic drawings."


def test_unreadable_and_empty_files_error(tmp_path):
    bogus = tmp_path / "not_a_docx.docx"
    bogus.write_bytes(b"this is not a zip")
    with pytest.raises(MasterImportError, match="not a readable"):
        parse_master_docx(bogus)

    empty = _write_docx(tmp_path, [], name="empty.docx")
    with pytest.raises(MasterImportError, match="No importable content"):
        parse_master_docx(empty)


def test_import_result_is_integrity_clean(tmp_path):
    result = parse_master_docx(_write_docx(tmp_path, _MASTER_LINES))
    # from_dict runs the full id/seq integrity check — must not raise.
    SpecSection.from_dict(result.section.to_dict())
    assert isinstance(result, ImportResult)


# ---------------------------------------------------------------------------
# API surface
# ---------------------------------------------------------------------------


def test_import_endpoint_gates_and_gap_adapt_context(tmp_path, monkeypatch):
    from fastapi.testclient import TestClient

    from backend import sessions
    from backend.app import create_app
    from tests.fakes import FakeClient, request_context_text, text_turn

    client = TestClient(create_app())
    path = _write_docx(tmp_path, _MASTER_LINES)
    lease_before = client.get("/api/health").json()

    # Wrong extension refused.
    resp = client.post(
        "/api/import/master",
        files={"file": ("master.txt", b"nope", "text/plain")},
    )
    assert resp.status_code == 400

    with open(path, "rb") as fh:
        resp = client.post(
            "/api/import/master",
            files={
                "file": (
                    "master.docx",
                    fh.read(),
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document",
                )
            },
        )
    data = resp.json()
    assert data["ok"] is True
    assert data["imported_block_count"] == 8
    assert data["doc"]["section"]["number"] == "21 13 13"
    assert data["doc"]["version"] == {"index": 1, "count": 2}
    lease_after = client.get("/api/health").json()
    assert data["workspace_id"] == lease_before["workspace_id"]
    assert data["workspace_scope"] == "original"
    assert data["generation"] > lease_before["generation"]
    assert data["generation"] == lease_after["generation"]
    statuses = {
        p["status"]
        for part in data["doc"]["parts"]
        for a in part["articles"]
        for p in a["paragraphs"]
    }
    assert statuses == {"imported"}

    # Non-empty doc → 409 (import is a starting point).
    with open(path, "rb") as fh:
        resp = client.post(
            "/api/import/master",
            files={"file": ("master.docx", fh.read(), "application/zip")},
        )
    assert resp.status_code == 409

    # The next turn's outline shows imported statuses (gap-and-adapt fuel)
    # and the stable prompt carries the policy.
    fake = FakeClient([text_turn(["Walking the master."])])
    monkeypatch.setattr("backend.llm.conversation.get_client", lambda: fake)
    client.post("/api/chat", json={"message": "adapt it"})
    request = fake.messages.last_request
    assert "(imported)" in request_context_text(request)
    assert "Gap-and-adapt" in request["system"][0]["text"]

    # Undo steps back to the blank page.
    undone = client.post("/api/doc/undo").json()
    assert undone["doc"]["section"]["number"] == ""


def test_import_allowed_over_project_setup_and_preserves_it(tmp_path):
    """Project setup entered before importing must not block the import.

    Regression: the import button gates on *body* content (frontend
    ``hasContent``), but the backend used to reject any non-empty document —
    and ``is_empty`` also counts project profile / edition overrides /
    excluded standards. A user who filled in the project profile first got an
    enabled button but a silent 409, so nothing landed in the document panel.
    Importing over body-less project setup now succeeds and carries that
    setup onto the imported tree.
    """
    from fastapi.testclient import TestClient

    from backend.app import create_app

    client = TestClient(create_app())

    # Body-less project setup: profile + an edition override + a suppression.
    assert client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_project_profile",
                    "target_id": "sec",
                    "city": "Phoenix",
                    "state": "Arizona",
                    "country": "USA",
                    "client": "Acme",
                }
            ]
        },
    ).status_code == 200
    assert client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_standard_edition",
                    "target_id": "sec",
                    "standard": "NFPA 13",
                    "edition": "2019",
                    "basis": "local amendment",
                    "title": "Installation of Sprinkler Systems",
                }
            ]
        },
    ).status_code == 200
    assert client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_standard_suppressed",
                    "target_id": "sec",
                    "standard": "NFPA 291",
                    "suppressed": True,
                    "basis": "out of scope",
                }
            ]
        },
    ).status_code == 200

    path = _write_docx(tmp_path, _MASTER_LINES)
    with open(path, "rb") as fh:
        resp = client.post(
            "/api/import/master",
            files={
                "file": (
                    "master.docx",
                    fh.read(),
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document",
                )
            },
        )
    assert resp.status_code == 200, resp.json()
    doc = resp.json()["doc"]
    # The master landed...
    assert doc["section"]["number"] == "21 13 13"
    assert any(part["articles"] for part in doc["parts"])
    # ...and the pre-import project setup survived the tree replacement.
    assert doc["project_profile"]["city"] == "Phoenix"
    assert doc["project_profile"]["client_name"] == "Acme"
    assert {k.lower() for k in doc["edition_overrides"]} == {"nfpa 13"}
    assert doc["suppressed_standards"]  # NFPA 291 excluded, still excluded

    # Undo returns to the pre-import setup (not a blank page), and re-import
    # is now blocked because real body content exists.
    undone = client.post("/api/doc/undo").json()
    assert undone["doc"]["section"]["number"] == ""
    assert undone["doc"]["project_profile"]["city"] == "Phoenix"

    client.post("/api/doc/redo")
    with open(path, "rb") as fh:
        blocked = client.post(
            "/api/import/master",
            files={
                "file": (
                    "master.docx",
                    fh.read(),
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document",
                )
            },
        )
    assert blocked.status_code == 409


# ---------------------------------------------------------------------------
# Auto-numbered heading promotion (the numbering-definition label grammar)
# ---------------------------------------------------------------------------


def _define_numbering(
    document,
    num_id: int,
    levels: dict[int, tuple[str, str]],
    overrides: dict[int, tuple[str, str]] | None = None,
):
    """Add a REAL numbering definition (abstractNum + num) to the package.

    ``levels``/``overrides`` map ``ilvl -> (numFmt, lvlText)``. Fixture
    num_ids should stay clear of 1-9 (the python-docx template's own).
    """
    root = document.part.numbering_part.element
    abstract_id = str(900 + num_id)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), abstract_id)
    for ilvl, (fmt, lvl_text) in levels.items():
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        fmt_el = OxmlElement("w:numFmt")
        fmt_el.set(qn("w:val"), fmt)
        text_el = OxmlElement("w:lvlText")
        text_el.set(qn("w:val"), lvl_text)
        lvl.append(fmt_el)
        lvl.append(text_el)
        abstract.append(lvl)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), abstract_id)
    num.append(ref)
    for ilvl, (fmt, lvl_text) in (overrides or {}).items():
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), str(ilvl))
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        fmt_el = OxmlElement("w:numFmt")
        fmt_el.set(qn("w:val"), fmt)
        text_el = OxmlElement("w:lvlText")
        text_el.set(qn("w:val"), lvl_text)
        lvl.append(fmt_el)
        lvl.append(text_el)
        override.append(lvl)
        num.append(override)
    first_num = root.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstract)
    else:
        root.append(abstract)
    root.append(num)


def test_auto_numbered_article_headings_become_real_articles(tmp_path):
    """The reported 23 05 48 symptom, fixed at its root.

    A Word-auto-numbered ARTICLE heading's visible text is just its title
    ("SECTION INCLUDES") — the "1.1" is generated — so no text pattern can
    ever recognize it, and every heading used to land as a provision inside
    one synthetic IMPORTED CONTENT article with the "not a spec section"
    banner on a real spec. The numbering definition's own label grammar
    ("%1.%2") is the structural signal the text cannot carry.
    """
    document = Document()
    document.add_paragraph("SECTION 23 05 48")
    document.add_paragraph("VIBRATION AND SEISMIC CONTROLS FOR HVAC")
    document.add_paragraph("PART 1 - GENERAL")
    _define_numbering(
        document,
        50,
        {1: ("decimal", "%1.%2"), 2: ("upperLetter", "%3.")},
    )
    _numbered(document, "SECTION INCLUDES", 1, num_id="50")
    _numbered(document, "Vibration isolation for HVAC equipment.", 2, "50")
    _numbered(document, "REFERENCE STANDARDS", 1, num_id="50")
    _numbered(document, "ASHRAE Handbook - HVAC Applications.", 2, "50")
    _numbered(document, "QUALITY ASSURANCE", 1, num_id="50")
    _numbered(document, "Minimum three years documented experience.", 2, "50")
    path = tmp_path / "auto-articles.docx"
    document.save(str(path))

    result = parse_master_docx(path)
    part = result.section.parts[0]
    assert [a.title for a in part.articles] == [
        "SECTION INCLUDES",
        "REFERENCE STANDARDS",
        "QUALITY ASSURANCE",
    ]
    assert [[p.text for p in a.paragraphs] for a in part.articles] == [
        ["Vibration isolation for HVAC equipment."],
        ["ASHRAE Handbook - HVAC Applications."],
        ["Minimum three years documented experience."],
    ]
    # The RELEASE_WINDOWS claim, now executable: no invented article, no
    # unstructured banner, the shape verdict agrees with the real tree.
    assert not any("IMPORTED CONTENT" in a.title for a in part.articles)
    assert result.spec_shape_detected is True
    assert not any("before any article" in w for w in result.warnings)


def test_auto_numbered_part_headings_map_by_order_of_appearance(tmp_path):
    document = Document()
    document.add_paragraph("SECTION 23 05 48")
    document.add_paragraph("VIBRATION AND SEISMIC CONTROLS FOR HVAC")
    _define_numbering(
        document,
        51,
        {
            0: ("decimal", "PART %1"),
            1: ("decimal", "%1.%2"),
            2: ("upperLetter", "%3."),
        },
    )
    for part_title, article_title, provision in (
        ("GENERAL", "SUMMARY", "Section includes seismic restraint."),
        ("PRODUCTS", "PIPE AND FITTINGS", "Steel pipe: ASTM A53."),
        ("EXECUTION", "INSTALLATION", "Install per the working plans."),
    ):
        _numbered(document, part_title, 0, num_id="51")
        _numbered(document, article_title, 1, num_id="51")
        _numbered(document, provision, 2, num_id="51")
    path = tmp_path / "auto-parts.docx"
    document.save(str(path))

    result = parse_master_docx(path)
    for index, article_title in enumerate(
        ("SUMMARY", "PIPE AND FITTINGS", "INSTALLATION")
    ):
        articles = result.section.parts[index].articles
        assert [a.title for a in articles] == [article_title]
        assert len(articles[0].paragraphs) == 1
    assert result.spec_shape_detected is True


def test_a_fourth_auto_numbered_part_lands_under_part_three_loudly(tmp_path):
    document = Document()
    _define_numbering(
        document, 52, {0: ("decimal", "PART %1"), 1: ("decimal", "%1.%2")}
    )
    for part_title in ("GENERAL", "PRODUCTS", "EXECUTION", "COMMISSIONING"):
        _numbered(document, part_title, 0, num_id="52")
        _numbered(document, f"{part_title} ARTICLE", 1, num_id="52")
    path = tmp_path / "four-parts.docx"
    document.save(str(path))

    result = parse_master_docx(path)
    assert [a.title for a in result.section.parts[2].articles] == [
        "EXECUTION ARTICLE",
        "COMMISSIONING ARTICLE",
    ]
    assert any("kept under PART 3" in w for w in result.warnings)


def test_a_dangling_num_id_stays_on_the_provision_path(tmp_path):
    """No definition = no label grammar = exactly the pre-catalog tree."""
    document = Document()
    document.add_paragraph("PART 1 - GENERAL")
    document.add_paragraph("1.1 SUMMARY")
    _numbered(document, "Section includes seismic restraint.", 0, "77")
    _numbered(document, "Related: Section 23 05 29.", 0, num_id="77")
    path = tmp_path / "dangling.docx"
    document.save(str(path))

    article = parse_master_docx(path).section.parts[0].articles[0]
    assert article.title == "SUMMARY"
    assert [p.text for p in article.paragraphs] == [
        "Section includes seismic restraint.",
        "Related: Section 23 05 29.",
    ]


def test_a_num_level_override_decides_the_label_grammar(tmp_path):
    """The label the reader actually SEES is the one that promotes."""
    document = Document()
    document.add_paragraph("PART 1 - GENERAL")
    _define_numbering(
        document,
        53,
        {1: ("decimal", "%2.")},
        overrides={1: ("decimal", "%1.%2")},
    )
    _numbered(document, "SECTION INCLUDES", 1, num_id="53")
    path = tmp_path / "override.docx"
    document.save(str(path))

    part = parse_master_docx(path).section.parts[0]
    assert [a.title for a in part.articles] == ["SECTION INCLUDES"]


def test_the_apps_own_label_grammar_never_promotes(tmp_path):
    """word_numbering's single-token lvlTexts are provision labels.

    This is the round-trip safety property stated directly: a normalized
    export's numbering can never match the PART or article grammar, so
    export -> re-import is untouched by promotion by construction.
    """
    document = Document()
    document.add_paragraph("PART 3 - EXECUTION")
    document.add_paragraph("3.1 INSTALLATION")
    _define_numbering(
        document,
        54,
        {
            0: ("upperLetter", "%1."),
            1: ("decimal", "%2."),
            2: ("lowerLetter", "%3."),
            3: ("decimal", "%4)"),
        },
    )
    _numbered(document, "Install per the working plans.", 0, num_id="54")
    _numbered(document, "Support piping per NFPA 13.", 1, num_id="54")
    path = tmp_path / "own-grammar.docx"
    document.save(str(path))

    article = parse_master_docx(path).section.parts[2].articles[0]
    assert article.title == "INSTALLATION"
    top = article.paragraphs[0]
    assert top.text == "Install per the working plans."
    assert top.children[0].text == "Support piping per NFPA 13."


# ---------------------------------------------------------------------------
# END OF SECTION is loud; header tolerance
# ---------------------------------------------------------------------------


def test_a_second_section_is_named_not_silently_discarded(tmp_path):
    path = _write_docx(
        tmp_path,
        [
            *_MASTER_LINES[:-1],
            "SECTION 23 05 93",
            "TESTING, ADJUSTING, AND BALANCING FOR HVAC",
            "PART 1 - GENERAL",
            "1.1 SUMMARY",
            "A. TAB of hydronic and air systems.",
        ],
        name="combined.docx",
    )
    result = parse_master_docx(path)
    assert result.section.number == "21 13 13"
    dropped = [w for w in result.warnings if "more than one SECTION" in w]
    assert len(dropped) == 1
    assert "SECTION 23 05 93" in dropped[0]
    assert "split the file" in dropped[0]


def test_reimporting_the_apps_own_export_stays_end_of_section_quiet(tmp_path):
    """The exporter's own schedules follow END OF SECTION by design.

    Warning about them on every export -> re-import round trip would train
    users to ignore the warning that matters; the exporter's first schedule
    heading is the recognized suppression.
    """
    first = parse_master_docx(_write_docx(tmp_path, _MASTER_LINES))
    exported = tmp_path / "own-export.docx"
    exported.write_bytes(build_docx(first.section))
    second = parse_master_docx(exported)
    assert not any("END OF SECTION" in w for w in second.warnings)


def test_a_bare_first_line_header_sets_the_section(tmp_path):
    path = _write_docx(
        tmp_path,
        [
            "23 05 48 — VIBRATION AND SEISMIC CONTROLS FOR HVAC",
            "PART 1 - GENERAL",
            "1.1 SUMMARY",
            "A. Section includes seismic restraint.",
        ],
        name="bare-header.docx",
    )
    result = parse_master_docx(path)
    assert result.section.number == "23 05 48"
    assert result.section.title == "VIBRATION AND SEISMIC CONTROLS FOR HVAC"
    assert result.spec_shape_detected is True
    assert any("no 'SECTION' keyword" in w for w in result.warnings)


def test_a_bare_header_shape_later_in_the_document_is_a_provision(tmp_path):
    """Six digits and a dash mid-document cite a sibling section; only the
    FIRST content line may be a keyword-less header."""
    path = _write_docx(
        tmp_path,
        [
            "PART 1 - GENERAL",
            "1.1 RELATED WORK",
            "23 05 29 — HANGERS AND SUPPORTS FOR HVAC",
        ],
        name="citation.docx",
    )
    result = parse_master_docx(path)
    assert result.section.number == ""
    article = result.section.parts[0].articles[0]
    assert [p.text for p in article.paragraphs] == [
        "23 05 29 — HANGERS AND SUPPORTS FOR HVAC"
    ]


def test_part_four_text_heading_lands_under_part_three_loudly(tmp_path):
    path = _write_docx(
        tmp_path,
        [
            "PART 1 - GENERAL",
            "1.1 SUMMARY",
            "A. Section includes commissioning of HVAC systems.",
            "PART 4 - COMMISSIONING",
            "A. Verify installation per the checklists.",
        ],
        name="part-four.docx",
    )
    result = parse_master_docx(path)
    part3 = result.section.parts[2]
    assert part3.articles, "PART 4 content must land under PART 3"
    assert any("PART 4" in w and "PART 3" in w for w in result.warnings)


def test_warnings_cite_a_findable_reference_not_just_a_line_ordinal(tmp_path):
    """"Line N" counts body children — blank paragraphs and table rows
    included — which corresponds to nothing a user can locate in Word or in
    the panel. Element-specific warnings now also carry the display ref the
    panel and export schedules use, plus the stable id."""
    lines = [
        "PART 1 - GENERAL",
        "Orphan text before any article.",
        "1.1 SUMMARY",
        "A. Labeled provision.",
        "1. Nested provision.",
        "a. Deeper provision.",
        "1) Deepest provision.",
    ]
    result = parse_master_docx(_write_docx(tmp_path, lines, name="refs.docx"))

    orphan = next(w for w in result.warnings if "before any article" in w)
    # The synthetic article's display number and id are both named.
    assert "(at 1.1, id pt1.a1)" in orphan
