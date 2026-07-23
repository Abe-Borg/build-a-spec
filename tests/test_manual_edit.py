"""WI2: direct manual editing via POST /api/doc/edit, and the set_status op."""
from __future__ import annotations

import io
import json

from docx import Document
from fastapi.testclient import TestClient

from backend.app import create_app
from backend import sessions
from backend.spec_doc.model import APPLY_SPEC_EDITS_TOOL
from tests.fakes import FakeClient, text_turn, tool_turn


def _client() -> TestClient:
    return TestClient(create_app())


def _parse_sse(body: str) -> list[dict]:
    return [
        json.loads(line[len("data: "):])
        for line in body.splitlines()
        if line.startswith("data: ")
    ]


_SEED_EDITS = {
    "edits": [
        {
            "action": "replace",
            "target_id": "sec",
            "text": "WET-PIPE SPRINKLER SYSTEMS",
            "numbering": "21 13 13",
        },
        {"action": "add_article", "target_id": "pt1", "text": "SUMMARY"},
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "text": "Section includes wet-pipe systems per NFPA 13-2025.",
            "status": "assumed",
        },
    ]
}


def _seed(client: TestClient, monkeypatch) -> None:
    fake = FakeClient([tool_turn(["Drafting."], _SEED_EDITS), text_turn(["Done."])])
    monkeypatch.setattr("backend.llm.conversation.get_client", lambda: fake)
    resp = client.post("/api/chat", json={"message": "Start 21 13 13"})
    assert _parse_sse(resp.text)[-1]["type"] == "turn_complete"


def test_set_status_in_tool_schema_enum():
    action = APPLY_SPEC_EDITS_TOOL["input_schema"]["properties"]["edits"][
        "items"
    ]["properties"]["action"]
    assert "set_status" in action["enum"]


def test_manual_replace_edits_transactionally(monkeypatch):
    client = _client()
    _seed(client, monkeypatch)

    resp = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "replace",
                    "target_id": "pt1.a1.p1",
                    "text": "Section includes wet-pipe systems per NFPA 13.",
                    "status": "confirmed",
                }
            ]
        },
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["ok"] is True
    para = data["doc"]["parts"][0]["articles"][0]["paragraphs"][0]
    assert para["text"] == "Section includes wet-pipe systems per NFPA 13."
    assert para["status"] == "confirmed"
    # One new version committed (seed was v1; this is v2).
    assert data["doc"]["version"] == {"index": 2, "count": 3}


def test_manual_set_status_confirms_block(monkeypatch):
    client = _client()
    _seed(client, monkeypatch)

    resp = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_status",
                    "target_id": "pt1.a1.p1",
                    "status": "confirmed",
                }
            ]
        },
    )
    assert resp.status_code == 200
    (applied,) = resp.json()["applied"]
    assert applied == {"action": "set_status", "id": "pt1.a1.p1", "status": "confirmed"}
    para = resp.json()["doc"]["parts"][0]["articles"][0]["paragraphs"][0]
    # Text untouched, status flipped.
    assert para["text"].startswith("Section includes wet-pipe")
    assert para["status"] == "confirmed"


def test_manual_delete_edit(monkeypatch):
    client = _client()
    _seed(client, monkeypatch)
    resp = client.post(
        "/api/doc/edit",
        json={"ops": [{"action": "delete", "target_id": "pt1.a1.p1"}]},
    )
    assert resp.status_code == 200
    assert resp.json()["doc"]["parts"][0]["articles"][0]["paragraphs"] == []


def test_invalid_manual_batch_400_and_unchanged(monkeypatch):
    client = _client()
    _seed(client, monkeypatch)
    before = client.get("/api/doc").json()["doc"]

    resp = client.post(
        "/api/doc/edit",
        json={"ops": [{"action": "delete", "target_id": "does-not-exist"}]},
    )
    assert resp.status_code == 400
    assert "no element" in resp.json()["error"]
    # Document (and version count) unchanged.
    assert client.get("/api/doc").json()["doc"] == before


def test_manual_edit_rejected_while_turn_active(monkeypatch):
    from backend.llm.conversation import stream_user_turn

    fake = FakeClient([tool_turn(["Working. "], _SEED_EDITS), text_turn(["Done."])])
    monkeypatch.setattr("backend.llm.conversation.get_client", lambda: fake)
    session = sessions.get_session()
    client = _client()

    gen = stream_user_turn(session, "Start 21 13 13")
    # Advance to mid-turn (turn_active is set right after begin_turn).
    for event in gen:
        if event["type"] == "doc_patch":
            break
    assert session.turn_active is True

    resp = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {"action": "replace", "target_id": "sec", "text": "HANDS OFF"}
            ]
        },
    )
    assert resp.status_code == 409
    assert "streaming" in resp.json()["error"]

    # Drain the turn; the guard clears and a manual edit now succeeds.
    for _ in gen:
        pass
    assert session.turn_active is False
    ok = client.post(
        "/api/doc/edit",
        json={"ops": [{"action": "set_status", "target_id": "pt1.a1.p1", "status": "confirmed"}]},
    )
    assert ok.status_code == 200


def test_manual_edit_is_undoable(monkeypatch):
    client = _client()
    _seed(client, monkeypatch)
    client.post(
        "/api/doc/edit",
        json={"ops": [{"action": "set_status", "target_id": "pt1.a1.p1", "status": "confirmed"}]},
    )
    undone = client.post("/api/doc/undo")
    assert undone.status_code == 200
    para = undone.json()["doc"]["parts"][0]["articles"][0]["paragraphs"][0]
    assert para["status"] == "assumed"  # back to the seeded value


def test_batch_set_status_is_transactional_and_one_undo_step(monkeypatch):
    """Batch 3, WI2: the review queue's "confirm remaining N in this article"
    affordance sends N set_status ops in one /api/doc/edit call. They apply as
    one transactional batch and land as a single undoable version."""
    client = _client()
    # Seed two assumed paragraphs in one article.
    seed = {
        "edits": [
            {"action": "add_article", "target_id": "pt2", "text": "SPRINKLERS"},
            {
                "action": "add_paragraph",
                "target_id": "pt2.a1",
                "text": "Provide UL-listed sprinklers.",
                "status": "assumed",
            },
            {
                "action": "add_paragraph",
                "target_id": "pt2.a1",
                "text": "Provide guards where subject to damage.",
                "status": "assumed",
            },
        ]
    }
    fake = FakeClient([tool_turn(["Drafting."], seed), text_turn(["Done."])])
    monkeypatch.setattr("backend.llm.conversation.get_client", lambda: fake)
    client.post("/api/chat", json={"message": "draft products"})
    before = client.get("/api/doc").json()["doc"]
    assert before["version"] == {"index": 1, "count": 2}

    resp = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {"action": "set_status", "target_id": "pt2.a1.p1", "status": "confirmed"},
                {"action": "set_status", "target_id": "pt2.a1.p2", "status": "confirmed"},
            ]
        },
    )
    assert resp.status_code == 200
    data = resp.json()
    paras = data["doc"]["parts"][1]["articles"][0]["paragraphs"]
    assert [p["status"] for p in paras] == ["confirmed", "confirmed"]
    # One new version for the whole batch.
    assert data["doc"]["version"] == {"index": 2, "count": 3}

    # A single undo reverts both confirmations at once.
    undone = client.post("/api/doc/undo").json()
    paras = undone["doc"]["parts"][1]["articles"][0]["paragraphs"]
    assert [p["status"] for p in paras] == ["assumed", "assumed"]


def test_batch_set_status_rejects_whole_batch_on_one_bad_op(monkeypatch):
    client = _client()
    _seed(client, monkeypatch)
    before = client.get("/api/doc").json()["doc"]
    resp = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {"action": "set_status", "target_id": "pt1.a1.p1", "status": "confirmed"},
                {"action": "set_status", "target_id": "nope", "status": "confirmed"},
            ]
        },
    )
    assert resp.status_code == 400
    # The good op did not stick — the batch is all-or-nothing.
    assert client.get("/api/doc").json()["doc"] == before


def test_manual_set_project_profile_edit():
    """The panel's project-profile form posts the same op the model's tool
    uses, through the same manual-edit path as any other direct edit — no
    dedicated endpoint, no restricted op vocabulary."""
    client = _client()
    resp = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_project_profile",
                    "target_id": "sec",
                    "city": "Ashburn",
                    "state": "Virginia",
                    "country": "USA",
                    "client": "ExampleCo",
                }
            ]
        },
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["ok"] is True
    assert data["applied"] == [
        {"action": "set_project_profile", "id": "sec", "complete": True}
    ]
    assert data["doc"]["project_profile"] == {
        "city": "Ashburn",
        "state_or_province": "VA",
        "country": "US",
        "client_name": "ExampleCo",
    }
    assert data["profile_complete"] is True


def test_manual_set_project_profile_edit_is_undoable():
    client = _client()
    client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_project_profile",
                    "target_id": "sec",
                    "city": "Ashburn",
                    "state": "VA",
                    "country": "US",
                    "client": "ExampleCo",
                }
            ]
        },
    )
    undone = client.post("/api/doc/undo").json()
    assert undone["doc"]["project_profile"] == {}
    assert undone["profile_complete"] is False


def test_manual_set_project_profile_rejects_unknown_country():
    client = _client()
    resp = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_project_profile",
                    "target_id": "sec",
                    "country": "France",
                }
            ]
        },
    )
    assert resp.status_code == 400
    assert "country" in resp.json()["error"]
    assert sessions.get_session().doc.doc.project_profile == {}


def test_confirming_removes_block_from_assumptions_schedule(monkeypatch):
    client = _client()
    _seed(client, monkeypatch)

    # Before: the assumed block is scheduled.
    doc = Document(io.BytesIO(client.get("/api/export/docx").content))
    assert "ASSUMPTIONS SCHEDULE" in [p.text for p in doc.paragraphs]
    rows_before = [
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    ]
    assert any("wet-pipe systems" in c for c in rows_before)

    # Confirm it, then re-export: it drops off the assumptions schedule.
    client.post(
        "/api/doc/edit",
        json={"ops": [{"action": "set_status", "target_id": "pt1.a1.p1", "status": "confirmed"}]},
    )
    doc2 = Document(io.BytesIO(client.get("/api/export/docx").content))
    rows_after = [
        cell.text for t in doc2.tables for row in t.rows for cell in row.cells
    ]
    assert not any("wet-pipe systems" in c for c in rows_after)


# ---------------------------------------------------------------------------
# Standards manager: add / suppress / restore through /api/doc/edit
# ---------------------------------------------------------------------------


def test_set_standard_suppressed_in_tool_schema_enum():
    action = APPLY_SPEC_EDITS_TOOL["input_schema"]["properties"]["edits"][
        "items"
    ]["properties"]["action"]
    assert "set_standard_suppressed" in action["enum"]


def test_manual_add_standard_edit():
    """The standards manager adds a standard the module does not pin, with a
    title — same op the model's tool uses, through the same manual-edit
    path (no restricted vocabulary)."""
    client = _client()
    resp = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_standard_edition",
                    "target_id": "sec",
                    "standard": "NFPA 30",
                    "edition": "2024",
                    "basis": "on-site flammable storage",
                    "title": "Flammable and Combustible Liquids Code",
                }
            ]
        },
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["ok"] is True
    added = [s for s in data["standards"] if s["name"] == "NFPA 30"]
    assert len(added) == 1
    assert added[0]["is_added"] is True
    assert added[0]["edition"] == "2024"
    assert added[0]["title"] == "Flammable and Combustible Liquids Code"
    assert added[0]["is_suppressed"] is False


def test_manual_suppress_and_restore_standard():
    client = _client()
    # A module-pinned standard is present and not excluded by default.
    before = client.get("/api/doc").json()["standards"]
    assert any(
        s["name"] == "NFPA 2001" and not s["is_suppressed"] for s in before
    )

    suppressed = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_standard_suppressed",
                    "target_id": "sec",
                    "standard": "NFPA 2001",
                    "suppressed": True,
                    "basis": "no clean-agent system in scope",
                }
            ]
        },
    ).json()
    assert suppressed["ok"] is True
    rows = [s for s in suppressed["standards"] if s["name"] == "NFPA 2001"]
    # Exactly one row, now marked excluded (not doubled into the live list).
    assert len(rows) == 1
    assert rows[0]["is_suppressed"] is True
    assert rows[0]["reason"] == "no clean-agent system in scope"

    restored = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_standard_suppressed",
                    "target_id": "sec",
                    "standard": "NFPA 2001",
                    "suppressed": False,
                }
            ]
        },
    ).json()
    rows = [s for s in restored["standards"] if s["name"] == "NFPA 2001"]
    assert len(rows) == 1 and rows[0]["is_suppressed"] is False


def test_manual_suppress_is_undoable():
    client = _client()
    client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_standard_suppressed",
                    "target_id": "sec",
                    "standard": "NFPA 2001",
                    "suppressed": True,
                }
            ]
        },
    )
    undone = client.post("/api/doc/undo").json()
    rows = [s for s in undone["standards"] if s["name"] == "NFPA 2001"]
    assert len(rows) == 1 and rows[0]["is_suppressed"] is False
