"""Import a firm's master, edit it freely, export it looking the same.

This is the contract the appearance-preserving mode makes, stated as tests:

    Everything except the body of ``word/document.xml`` is carried through
    byte-for-byte — headers, footers, styles, theme, fonts, numbering
    definitions, page setup. Inside the body a provision you did not touch is
    a byte-identical clone, a provision you edited keeps its paragraph and
    run properties, preserved blocks (tables, pictures, embedded objects,
    content controls) are emitted verbatim, and a provision you added is
    cloned from its nearest kin.

It replaces the older byte-exact mode as the product's import path because
that promise, while stronger, left almost nothing editable: three of
twenty-seven body operations on a clean master. The byte-exact machinery is
still reachable at ``?mode=source`` for projects that never gave up the
claim, and its own suites still pin it.
"""
from __future__ import annotations

import io
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from fastapi.testclient import TestClient
from lxml import etree

from backend.app import create_app
from backend.spec_doc.importer import parse_master_docx
from backend.spec_doc.model import SpecSection, apply_edits
from backend.spec_doc.source_render import (
    SourceRenderError,
    render_preserving_docx,
)
from tests.docx_fidelity_helpers import DOCX_MEDIA_TYPE

HEADER_TEXT = "ACME ENGINEERING — ISSUED FOR BID"
FOOTER_TEXT = "23 05 48 - 1"


def _master_bytes(*, with_table: bool = True) -> bytes:
    """A master with firm styling, a header, a footer and a schedule."""
    document = Document()
    document.styles["Normal"].font.name = "Cambria"
    document.styles["Normal"].font.size = Pt(11)
    section = document.sections[0]
    section.header.paragraphs[0].text = HEADER_TEXT
    section.footer.paragraphs[0].text = FOOTER_TEXT
    for line in (
        "SECTION 23 05 48",
        "VIBRATION AND SEISMIC CONTROLS",
        "PART 1 - GENERAL",
        "1.1 SUMMARY",
        "A. Section includes vibration isolation for mechanical equipment.",
        "",
        "B. Related requirements are specified elsewhere.",
        "1.2 SCHEDULE",
    ):
        document.add_paragraph(line)
    if with_table:
        table = document.add_table(rows=3, cols=2)
        rows = [
            ("Equipment", "Static Deflection"),
            ("AHU-1", "1 inch"),
            ("Pump P-1", "0.75 inch"),
        ]
        for row, (left, right) in zip(table.rows, rows):
            row.cells[0].text = left
            row.cells[1].text = right
    document.add_paragraph("A. Provide isolators as scheduled.")
    document.add_paragraph("END OF SECTION")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _body_children(payload: bytes) -> list:
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        xml = archive.read("word/document.xml")
    body = etree.fromstring(xml).find(qn("w:body"))
    return [child for child in body.iterchildren() if isinstance(child.tag, str)]


def _texts(payload: bytes) -> list[str]:
    return [
        paragraph.text
        for paragraph in Document(io.BytesIO(payload)).paragraphs
        if paragraph.text.strip()
    ]


def _members(payload: bytes) -> dict[str, bytes]:
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        return {name: archive.read(name) for name in archive.namelist()}


def _render(source: bytes, section: SpecSection, format_map) -> bytes:
    return render_preserving_docx(
        source_bytes=source, format_map=format_map, current=section
    )


def _import(client: TestClient, source: bytes, *, detach: bool = True):
    response = client.post(
        "/api/import/master",
        files={"file": ("office-master.docx", source, DOCX_MEDIA_TYPE)},
        data={"detach": "true" if detach else "false"},
    )
    assert response.status_code == 200, response.text
    return response.json()


# ---------------------------------------------------------------------------
# The package-level promise
# ---------------------------------------------------------------------------


def test_only_the_document_body_is_rewritten(tmp_path):
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)
    section = imported.section
    target = section.parts[0].articles[0].paragraphs[0].uid
    section, _ = apply_edits(
        section,
        [{"action": "replace", "target_id": target, "text": "Rewritten."}],
    )

    exported = _render(source, section, imported.format_map)

    before, after = _members(source), _members(exported)
    assert sorted(before) == sorted(after)
    changed = [name for name in before if before[name] != after[name]]
    assert changed == ["word/document.xml"], changed


def test_the_header_footer_and_fonts_survive(tmp_path):
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)

    exported = _render(source, imported.section, imported.format_map)

    document = Document(io.BytesIO(exported))
    assert document.sections[0].header.paragraphs[0].text == HEADER_TEXT
    assert document.sections[0].footer.paragraphs[0].text == FOOTER_TEXT
    assert document.styles["Normal"].font.name == "Cambria"
    assert document.styles["Normal"].font.size == Pt(11)


def test_an_untouched_document_round_trips_element_for_element(tmp_path):
    """The guarantee the whole design rests on.

    If an unedited provision is not a byte-identical clone then every claim
    about preserved formatting is a guess. This is also the regression guard
    for parsing the body with python-docx's oxml classes: a plain lxml parse
    reports every paragraph as empty, which silently sends untouched
    provisions down the rewrite path and this test straight to red.
    """
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)

    exported = _render(source, imported.section, imported.format_map)

    before = [etree.tostring(el) for el in _body_children(source)]
    after = [etree.tostring(el) for el in _body_children(exported)]
    assert before == after


def test_one_edit_changes_exactly_one_element(tmp_path):
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)
    section = imported.section
    target = section.parts[0].articles[0].paragraphs[1].uid
    section, _ = apply_edits(
        section,
        [
            {
                "action": "replace",
                "target_id": target,
                "text": "Related requirements are specified in Division 22.",
            }
        ],
    )

    exported = _render(source, section, imported.format_map)

    before = [etree.tostring(el) for el in _body_children(source)]
    after = [etree.tostring(el) for el in _body_children(exported)]
    differing = [i for i, (a, b) in enumerate(zip(before, after)) if a != b]
    assert len(differing) == 1
    assert "Division 22" in "".join(
        _body_children(exported)[differing[0]].itertext()
    )


def test_content_the_parse_stopped_at_is_still_exported(tmp_path):
    """``END OF SECTION`` is not a provision, and must not be dropped.

    Under the byte-exact mode nothing could be lost because the upload was
    patched rather than rebuilt. Rebuilding the body makes every body child
    the tree does not model a candidate for silent deletion.
    """
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)

    exported = _render(source, imported.section, imported.format_map)

    assert "END OF SECTION" in _texts(exported)


def test_blank_spacers_travel_with_the_provision_below_them(tmp_path):
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)
    section = imported.section
    article = section.parts[0].articles[0]
    # The spacer sat above "B."; moving B to the top must take it along
    # rather than leaving a gap where B used to be.
    section, _ = apply_edits(
        section,
        [
            {
                "action": "move",
                "target_id": article.paragraphs[1].uid,
                "position": 0,
            }
        ],
    )

    exported = _render(source, section, imported.format_map)

    kinds = [
        "blank" if not "".join(el.itertext()).strip() else "text"
        for el in _body_children(exported)
        if el.tag == qn("w:p")
    ]
    assert "blank" in kinds


# ---------------------------------------------------------------------------
# Preserved blocks
# ---------------------------------------------------------------------------


def test_a_table_is_one_locked_block_not_a_row_per_paragraph(tmp_path):
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)

    schedule = imported.section.parts[0].articles[1]
    locked = [p for p in schedule.paragraphs if p.locked]
    assert len(locked) == 1
    assert locked[0].locked == "table"
    # The grid still reads, one row per line, so the panel and the model can
    # both see what the schedule says.
    assert locked[0].text.splitlines()[0] == "Equipment | Static Deflection"
    assert len(locked[0].text.splitlines()) == 3


def test_a_locked_block_takes_no_label_and_shifts_no_sibling(tmp_path):
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)

    payload = imported.section.to_dict()
    schedule = payload["parts"][0]["articles"][1]["paragraphs"]
    assert [p["label"] for p in schedule] == ["", "A."]

    exported = _render(source, imported.section, imported.format_map)
    assert "A. Provide isolators as scheduled." in _texts(exported)


def test_a_table_can_be_deleted_and_moved_but_not_retyped(tmp_path):
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)
    section = imported.section
    table_uid = section.parts[0].articles[1].paragraphs[0].uid

    from backend.spec_doc.model import SpecEditError

    try:
        apply_edits(
            section,
            [{"action": "replace", "target_id": table_uid, "text": "no"}],
        )
    except SpecEditError as exc:
        assert "preserved Word table" in str(exc)
        assert "move it or delete it" in str(exc)
    else:  # pragma: no cover - the lock is the point of the test
        raise AssertionError("a preserved table accepted a retype")

    moved, _ = apply_edits(
        section, [{"action": "move", "target_id": table_uid, "position": 1}]
    )
    assert moved.parts[0].articles[1].paragraphs[1].locked == "table"

    deleted, _ = apply_edits(
        section, [{"action": "delete", "target_id": table_uid}]
    )
    exported = _render(source, deleted, imported.format_map)
    assert not Document(io.BytesIO(exported)).tables


def test_a_moved_table_is_still_a_table_in_the_export(tmp_path):
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)
    section = imported.section
    table_uid = section.parts[0].articles[1].paragraphs[0].uid
    section, _ = apply_edits(
        section, [{"action": "move", "target_id": table_uid, "position": 1}]
    )

    exported = _render(source, section, imported.format_map)

    document = Document(io.BytesIO(exported))
    assert len(document.tables) == 1
    assert document.tables[0].rows[1].cells[0].text == "AHU-1"
    kinds = [etree.QName(el).localname for el in _body_children(exported)]
    assert kinds.index("tbl") > kinds.index("p")


# ---------------------------------------------------------------------------
# Added content
# ---------------------------------------------------------------------------


def test_a_new_provision_is_cloned_from_its_nearest_kin(tmp_path):
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)
    section = imported.section
    article = section.parts[0].articles[0]
    section, _ = apply_edits(
        section,
        [
            {
                "action": "add_paragraph",
                "target_id": article.uid,
                "text": "Comply with ASHRAE Applications Chapter 49.",
            }
        ],
    )

    exported = _render(source, section, imported.format_map)

    children = _body_children(exported)
    existing = next(
        el for el in children if "Section includes vibration" in "".join(el.itertext())
    )
    added = next(
        el for el in children if "ASHRAE Applications" in "".join(el.itertext())
    )
    # Same paragraph properties as the sibling it was modelled on — that is
    # what "new content looks like the content around it" means.
    def _properties(element):
        found = element.find(qn("w:pPr"))
        return etree.tostring(found) if found is not None else None

    assert _properties(existing) == _properties(added)
    assert "C. Comply with ASHRAE Applications Chapter 49." in _texts(exported)


def test_renumbering_follows_the_document_not_the_upload(tmp_path):
    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)
    section = imported.section
    article = section.parts[0].articles[0]
    section, _ = apply_edits(
        section, [{"action": "delete", "target_id": article.paragraphs[0].uid}]
    )

    exported = _render(source, section, imported.format_map)

    texts = _texts(exported)
    assert "A. Related requirements are specified elsewhere." in texts
    assert not any(text.startswith("B. Related") for text in texts)


# ---------------------------------------------------------------------------
# Binding and failure
# ---------------------------------------------------------------------------


def test_a_format_map_is_refused_beside_bytes_it_does_not_describe(tmp_path):
    source = _master_bytes()
    other = _master_bytes(with_table=False)
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)

    try:
        _render(other, imported.section, imported.format_map)
    except SourceRenderError as exc:
        assert "does not describe" in str(exc)
    else:  # pragma: no cover - the binding is the point of the test
        raise AssertionError("a formatting map was trusted against foreign bytes")


# ---------------------------------------------------------------------------
# End to end, through the API
# ---------------------------------------------------------------------------


def test_import_edit_export_keeps_the_firms_formatting():
    client = TestClient(create_app())
    source = _master_bytes()
    _import(client, source)

    doc = client.get("/api/doc").json()["doc"]
    summary = doc["parts"][0]["articles"][0]
    first = summary["paragraphs"][0]["id"]
    edit = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "replace",
                    "target_id": first,
                    "text": "Section includes seismic restraint and isolation.",
                },
                {
                    "action": "add_paragraph",
                    "target_id": summary["id"],
                    "text": "Submit isolator calculations for review.",
                },
            ]
        },
    )
    assert edit.status_code == 200, edit.text

    exported = client.get("/api/export/docx")
    assert exported.status_code == 200
    payload = exported.content

    before, after = _members(source), _members(payload)
    assert [n for n in before if before[n] != after[n]] == ["word/document.xml"]
    document = Document(io.BytesIO(payload))
    assert document.sections[0].footer.paragraphs[0].text == FOOTER_TEXT
    assert len(document.tables) == 1
    texts = _texts(payload)
    assert "A. Section includes seismic restraint and isolation." in texts
    assert "C. Submit isolator calculations for review." in texts


def test_a_preserved_project_still_preserves_after_save_and_reload():
    client = TestClient(create_app())
    source = _master_bytes()
    _import(client, source)
    saved = client.get("/api/project/save")
    assert saved.status_code == 200

    reopened = TestClient(create_app())
    loaded = reopened.post(
        "/api/project/load-file",
        files={
            "file": (
                "project.baspec",
                saved.content,
                "application/octet-stream",
            )
        },
    )
    assert loaded.status_code == 200, loaded.text

    exported = reopened.get("/api/export/docx")
    assert exported.status_code == 200
    before, after = _members(source), _members(exported.content)
    # Reopened and re-exported without an edit, so at most the body may
    # differ — and in practice nothing does, which is the round trip being
    # exact rather than merely close.
    assert set(n for n in before if before[n] != after[n]) <= {
        "word/document.xml"
    }
    assert (
        Document(io.BytesIO(exported.content))
        .sections[0]
        .header.paragraphs[0]
        .text
        == HEADER_TEXT
    )


def test_normalized_export_stays_available_explicitly():
    client = TestClient(create_app())
    _import(client, _master_bytes())

    normalized = client.get("/api/export/docx", params={"mode": "normalized"})
    assert normalized.status_code == 200
    # A normalized export is Build-a-Spec's own document, so the firm's
    # header does not travel with it — that is the whole difference.
    document = Document(io.BytesIO(normalized.content))
    assert document.sections[0].header.paragraphs[0].text != HEADER_TEXT


# ---------------------------------------------------------------------------
# The header/footer the export never rewrites
# ---------------------------------------------------------------------------


def test_a_stale_footer_section_number_is_reported_not_rewritten():
    """Headers and footers are immutable; a wrong one must still be caught.

    Adapting a master is exactly when the section identifier changes, and a
    spec footer conventionally carries it. Rewriting it would break the
    "we do not touch your header and footer" contract, so the app says so
    instead and names the remedy (Word), rather than letting a stale number
    print on every page of an issued deliverable.
    """
    client = TestClient(create_app())
    _import(client, _master_bytes())

    clean = client.get("/api/doc").json()
    assert not [
        item
        for item in clean["lint"]
        if item["rule"] == "stale_document_identifier"
    ]

    renumbered = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "replace",
                    "target_id": "sec",
                    "text": "SEISMIC CONTROLS",
                    "numbering": "23 05 93",
                }
            ]
        },
    )
    assert renumbered.status_code == 200, renumbered.text

    findings = [
        item
        for item in renumbered.json()["lint"]
        if item["rule"] == "stale_document_identifier"
    ]
    assert len(findings) == 1
    assert "23 05 48" in findings[0]["message"]
    assert "update them in Word" in findings[0]["message"]

    exported = client.get("/api/export/docx")
    assert exported.status_code == 200
    # Reported, never rewritten.
    assert (
        Document(io.BytesIO(exported.content))
        .sections[0]
        .footer.paragraphs[0]
        .text
        == FOOTER_TEXT
    )


def test_the_model_is_told_a_preserved_block_is_not_retypeable(tmp_path):
    from backend.spec_doc.model import outline

    source = _master_bytes()
    path = tmp_path / "master.docx"
    path.write_bytes(source)
    imported = parse_master_docx(path)

    rendered = outline(imported.section, max_text=None)

    assert "[preserved table]" in rendered
    # It still carries an id, because moving and deleting it are allowed.
    table_uid = imported.section.parts[0].articles[1].paragraphs[0].uid
    assert f"[id: {table_uid}]" in rendered
