"""Final QC (Batch 4): lens fan-out, adversarial verification, ops
validation, and the accept/dismiss/readiness/report API — hermetic against
the sequenced fake client."""
from __future__ import annotations

import io
import json
import time

import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend import sessions, settings
from backend.app import create_app
from backend.qc.engine import (
    VERIFICATION_RULE_V4,
    QCFanoutError,
    QCResult,
    qc_version_fingerprint,
    run_final_qc,
)
from backend.qc.runner import QCRunner
from backend.qc.schema import QC_LENSES
from backend.spec_doc.model import DocumentStore
from backend.spec_modules import DEFAULT_MODULE
from tests.docx_fidelity_helpers import (
    DOCX_MEDIA_TYPE,
    TARGET_EDITED_TEXT,
    make_fidelity_master,
)
from tests.fakes import (
    FakeClient,
    SequencedFakeClient,
    qc_findings_response,
    qc_verdict_response,
    text_turn,
    tool_turn,
)

_LENS_KEYS = {lens.lens_id: f"[[QC-LENS:{lens.lens_id}]]" for lens in QC_LENSES}


def _client() -> TestClient:
    return TestClient(create_app())


def _section() -> "DocumentStore":
    """A minimal drafted document (one assumed provision)."""
    store = DocumentStore()
    store.begin_turn()
    store.apply_edits(
        [
            {
                "action": "replace",
                "target_id": "sec",
                "text": "WET-PIPE SPRINKLER SYSTEMS",
                "numbering": "21 13 13",
            },
            {"action": "add_article", "target_id": "pt1", "text": "SUMMARY"},
            {
                "action": "add_paragraph",
                "target_id": "pt1.a1",
                "text": "Comply with NFPA 13-2019 throughout.",
                "status": "assumed",
            },
        ]
    )
    store.commit_turn()
    return store


def _finding(
    title: str,
    issue: str,
    *,
    element_id: str | None = "pt1.a1.p1",
    severity: str = "high",
    ops: list[dict] | None = None,
    source_urls: list[str] | None = None,
) -> dict:
    return {
        "title": title,
        "severity": severity,
        "element_id": element_id,
        "issue": issue,
        "rationale": f"rationale for {title}",
        "source_urls": source_urls or [],
        "proposed_ops": ops,
    }


def _qc_scripts(**per_lens) -> dict[str, list]:
    """Lens scripts; unspecified lenses return zero findings."""
    scripts: dict[str, list] = {}
    for lens_id, key in _LENS_KEYS.items():
        scripts[key] = per_lens.get(
            lens_id, [qc_findings_response(lens_id, findings=[])]
        )
    return scripts


def _run(
    client,
    store,
    *,
    profile=None,
    remembered=None,
    sink=None,
    version_index=0,
    source_guard=None,
):
    return run_final_qc(
        store.doc,
        profile,
        DEFAULT_MODULE,
        client,
        model=settings.QC_MODEL,
        max_tokens=4096,
        version_index=version_index,
        started_at="2026-07-21 10:00",
        finished_at="2026-07-21 10:05",
        source_guard=source_guard,
        remembered_dismissed=remembered,
        event_sink=sink or (lambda _e: None),
    )


# ---------------------------------------------------------------------------
# Phase 1 — lens fan-out
# ---------------------------------------------------------------------------


def test_fanout_runs_all_lenses_grounds_and_survives_defaults():
    store = _section()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Stale NFPA 13 edition",
                        "The draft cites NFPA 13-2019 but the effective edition is 2025.",
                        severity="high",
                        source_urls=["https://nfpa.org/13"],
                    )
                ],
                searched_urls=["https://nfpa.org/13"],
            )
        ],
    )
    # The finding is high → 3-panel; script 3 uphold verdicts.
    scripts["Stale NFPA 13 edition"] = [
        qc_verdict_response(True) for _ in range(3)
    ]
    result = _run(SequencedFakeClient(scripts), store)

    assert [s.lens_id for s in result.lens_statuses] == [
        l.lens_id for l in QC_LENSES
    ]
    assert all(s.status == "completed" for s in result.lens_statuses)
    assert len(result.findings) == 1
    f = result.findings[0]
    assert f.lens_id == "code_compliance"
    assert f.grounded is True
    assert f.accepted_sources == ["https://nfpa.org/13"]
    assert f.finding_id.startswith("qc-")
    assert result.research_profile_present is False


def test_one_failing_lens_does_not_kill_the_run():
    store = _section()
    scripts = _qc_scripts(
        completeness=[RuntimeError("lens boom")],  # non-retryable
    )
    result = _run(SequencedFakeClient(scripts), store)
    statuses = {s.lens_id: s for s in result.lens_statuses}
    assert statuses["completeness"].status == "failed"
    assert "boom" in statuses["completeness"].error
    assert sum(1 for s in result.lens_statuses if s.status == "completed") == 4


def test_all_lenses_failing_raises_fanout_error():
    store = _section()
    scripts = {key: [RuntimeError("dead")] for key in _LENS_KEYS.values()}
    with pytest.raises(QCFanoutError, match="All 5"):
        _run(SequencedFakeClient(scripts), store)


def test_ungrounded_citation_is_kept_but_marked():
    store = _section()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Cites unretrieved page",
                        "Rationale cites a URL never fetched.",
                        severity="medium",
                        source_urls=["https://never-retrieved.example"],
                    )
                ],
                searched_urls=["https://nfpa.org/13"],  # different URL retrieved
            )
        ],
    )
    scripts["Cites unretrieved page"] = [qc_verdict_response(True) for _ in range(2)]
    result = _run(SequencedFakeClient(scripts), store)
    f = result.findings[0]
    assert f.grounded is False
    assert f.accepted_sources == []


# ---------------------------------------------------------------------------
# Phase 2 — adversarial verification
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# The v4 outcome table (Chunk 5.1)
# ---------------------------------------------------------------------------
#
# v3 survived on `upholds >= (size // 2) + 1`, which inverts with panel size:
# a medium needed 2 of 2 while a critical needed only 2 of 3, so the extra
# critical seat bought leniency. v4: unanimous upholds, majority refutation
# refutes, everything between is DISPUTED and escalates to a human.


def _medium_scripts(verdicts: list) -> dict:
    scripts = _qc_scripts(
        enforceability_language=[
            qc_findings_response(
                "enforceability_language",
                findings=[
                    _finding(
                        "Vague language", "Uses 'as required'.", severity="medium"
                    )
                ],
            )
        ],
    )
    scripts["Vague language"] = verdicts
    return scripts


def _high_scripts(verdicts: list) -> dict:
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding("Wrong edition", "Edition mismatch.", severity="high")
                ],
            )
        ],
    )
    scripts["Wrong edition"] = verdicts
    return scripts


def test_two_seat_panel_upholds_only_when_both_seats_agree():
    result = _run(
        SequencedFakeClient(
            _medium_scripts([qc_verdict_response(True), qc_verdict_response(True)])
        ),
        _section(),
    )
    assert len(result.findings) == 1
    assert result.findings[0].verification_outcome == "upheld"
    assert result.disputed == [] and result.refuted == []


def test_two_seat_panel_tie_is_disputed_not_refuted():
    """The v3 behavior this replaces: a 1-1 tie silently went to the
    refuters. A split panel is disagreement, and disagreement is the
    reviewer's information, not something to round away."""
    result = _run(
        SequencedFakeClient(
            _medium_scripts([qc_verdict_response(True), qc_verdict_response(False)])
        ),
        _section(),
    )
    assert result.findings == [] and result.refuted == []
    assert len(result.disputed) == 1
    assert result.disputed[0].title == "Vague language"
    assert result.disputed[0].verification_outcome == "disputed"
    assert result.disputed[0].dispute_reason == "split_panel"


def test_two_seat_panel_refutes_when_both_seats_refute():
    result = _run(
        SequencedFakeClient(
            _medium_scripts([qc_verdict_response(False), qc_verdict_response(False)])
        ),
        _section(),
    )
    assert result.findings == [] and result.disputed == []
    assert len(result.refuted) == 1
    assert result.refuted[0].verification_outcome == "refuted"


def test_three_seat_panel_upholds_unanimously_and_medians_severity():
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(True, severity="critical"),
                    qc_verdict_response(True, severity="critical"),
                    qc_verdict_response(True),
                ]
            )
        ),
        _section(),
    )
    assert len(result.findings) == 1
    # median(["high","critical","critical"]) → critical.
    assert result.findings[0].severity == "critical"
    assert result.findings[0].verification_outcome == "upheld"
    assert len(result.findings[0].verdicts) == 3


def test_three_seat_panel_majority_uphold_is_disputed():
    """2-1 was an outright pass under v3. A life-safety finding two of three
    reviewers believe in is exactly what a human should adjudicate."""
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(True, severity="critical"),
                    qc_verdict_response(True, severity="critical"),
                    qc_verdict_response(False),
                ]
            )
        ),
        _section(),
    )
    assert result.findings == [] and result.refuted == []
    assert len(result.disputed) == 1
    assert result.disputed[0].verification_outcome == "disputed"
    assert result.disputed[0].dispute_reason == "split_panel"


def test_three_seat_panel_majority_refutation_needs_evidence_to_refute():
    """The RF-001 shape: three seats refuted a life-safety-adjacent finding
    having retrieved nothing. Under v4 that escalates instead of vanishing."""
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(False),
                    qc_verdict_response(False),
                    qc_verdict_response(False),
                ]
            )
        ),
        _section(),
    )
    assert result.refuted == []
    assert len(result.disputed) == 1
    assert result.disputed[0].dispute_reason == "insufficient_refutation_evidence"


def test_a_refuting_seat_that_cites_what_it_retrieved_refutes_cleanly():
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(
                        False,
                        evidence=[{"type": "source", "url": "https://a.gov/code"}],
                        searched_urls=["https://a.gov/code"],
                    ),
                    qc_verdict_response(False),
                    qc_verdict_response(False),
                ]
            )
        ),
        _section(),
    )
    assert result.disputed == []
    assert len(result.refuted) == 1
    assert result.refuted[0].verification_outcome == "refuted"
    evidence = result.refuted[0].verdicts[0].refutation_evidence
    assert [e.validated for e in evidence] == [True]


def test_activity_without_a_citation_never_satisfies_the_evidence_gate():
    """One token search must not launder an unevidenced refutation. Queries
    and retrievals are operational records of what a seat DID; only a
    validated citation is evidence."""
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(False, searched_urls=["https://a.gov/code"]),
                    qc_verdict_response(False, searched_urls=["https://b.gov/code"]),
                    qc_verdict_response(False, searched_urls=["https://c.gov/code"]),
                ]
            )
        ),
        _section(),
    )
    assert result.refuted == []
    assert len(result.disputed) == 1
    assert result.disputed[0].dispute_reason == "insufficient_refutation_evidence"
    # The retrievals are still recorded — they are billable activity.
    assert result.disputed[0].verdicts[0].retrieved_sources


def test_a_cited_source_the_seat_never_retrieved_does_not_validate():
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(
                        False,
                        evidence=[{"type": "source", "url": "https://never.example"}],
                    ),
                    qc_verdict_response(False),
                    qc_verdict_response(False),
                ]
            )
        ),
        _section(),
    )
    assert len(result.disputed) == 1
    assert result.disputed[0].dispute_reason == "insufficient_refutation_evidence"
    # Retained and marked, never dropped: that the seat tried to justify
    # itself and cited something unverifiable is part of the audit trail.
    evidence = result.disputed[0].verdicts[0].refutation_evidence
    assert len(evidence) == 1
    assert evidence[0].validated is False
    assert "does not match" in evidence[0].reason


def test_a_document_reference_that_resolves_satisfies_the_gate():
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(
                        False,
                        evidence=[
                            {"type": "document_ref", "reference": "pt1.a1.p1"}
                        ],
                    ),
                    qc_verdict_response(False),
                    qc_verdict_response(False),
                ]
            )
        ),
        _section(),
    )
    assert result.disputed == []
    assert len(result.refuted) == 1
    assert result.refuted[0].verdicts[0].refutation_evidence[0].validated is True


def test_an_unresolvable_document_reference_does_not_satisfy_the_gate():
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(
                        False,
                        evidence=[
                            {"type": "document_ref", "reference": "pt9.a9.p9"}
                        ],
                    ),
                    qc_verdict_response(False),
                    qc_verdict_response(False),
                ]
            )
        ),
        _section(),
    )
    assert len(result.disputed) == 1
    assert result.disputed[0].dispute_reason == "insufficient_refutation_evidence"
    evidence = result.disputed[0].verdicts[0].refutation_evidence
    assert evidence[0].validated is False
    assert "does not resolve" in evidence[0].reason


def test_a_medium_refutation_is_not_evidence_gated():
    """The gate is severity-scoped on purpose — it exists for the findings
    whose false-negative cost is highest, not as a tax on every refutation."""
    result = _run(
        SequencedFakeClient(
            _medium_scripts([qc_verdict_response(False), qc_verdict_response(False)])
        ),
        _section(),
    )
    assert result.disputed == []
    assert len(result.refuted) == 1


def test_an_upholding_seats_evidence_never_opens_the_gate():
    """Only a REFUTING seat's citation can justify a refutation. A 1-2 high
    finding where the upholder cited a retrieved source stays disputed."""
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(
                        True,
                        evidence=[{"type": "source", "url": "https://a.gov/code"}],
                        searched_urls=["https://a.gov/code"],
                    ),
                    qc_verdict_response(False),
                    qc_verdict_response(False),
                ]
            )
        ),
        _section(),
    )
    assert result.refuted == []
    assert len(result.disputed) == 1
    assert result.disputed[0].dispute_reason == "insufficient_refutation_evidence"


def test_a_failed_seat_stays_inconclusive_whatever_the_other_votes_say():
    """Infrastructure failure is never evidence, and never a dispute."""
    scripts = _high_scripts(
        [
            qc_verdict_response(True),
            RuntimeError("verifier died"),
            qc_verdict_response(True),
        ]
    )
    result = _run(SequencedFakeClient(scripts), _section())
    assert result.findings == []
    assert result.refuted == [] and result.disputed == []
    assert len(result.inconclusive) == 1
    assert result.inconclusive[0].verification_outcome == "inconclusive"
    assert result.inconclusive[0].dispute_reason == ""


def test_a_disputed_candidate_blocks_audit_completeness_and_is_not_applicable():
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(True),
                    qc_verdict_response(True),
                    qc_verdict_response(False),
                ]
            )
        ),
        _section(),
    )
    assert len(result.disputed) == 1
    finding = result.disputed[0]
    # The ADJUDICATION is structurally sound — every seat reported and the
    # recorded outcome matches the votes. Disputed is a legitimate outcome,
    # not a broken record.
    assert result.verification_complete() is True
    assert result.is_complete() is True
    # What blocks readiness is the undispositioned dispute, a separate term
    # in exact parallel with an open critical. Folding it into is_complete()
    # deadlocked the dismiss route (PR #103 review).
    assert result.open_disputed_count() == 1
    # Never validated, so nothing about it is auto-applicable.
    assert finding.ops_valid is False
    assert finding.ops_semantic_status in {"not_evaluated", "not_proposed"}
    # And it is not in the actionable queue.
    assert all(f.finding_id != finding.finding_id for f in result.findings)


def test_a_disputed_candidate_can_actually_be_dismissed_with_a_reason():
    """The workflow the drawer and readiness copy tell the user to perform.

    It has to be reachable end to end: the lookup has to find a disputed
    candidate, the dismissal has to stick, and dismissing it has to clear
    the readiness term it was blocking. (Before the PR #103 review it was
    unreachable on all three counts.)
    """
    runner = QCRunner()
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(True),
                    qc_verdict_response(True),
                    qc_verdict_response(False),
                ]
            )
        ),
        _section(),
    )
    runner.restore(result)
    restored = runner.result
    assert restored is not None
    disputed = restored.disputed[0]
    assert restored.open_disputed_count() == 1

    # The lookup reaches it…
    assert restored.finding(disputed.finding_id) is disputed
    # …the dismissal sticks…
    assert (
        runner.dismiss(
            disputed.finding_id,
            "Reviewed the split; the stricter reading is not required here.",
            document_version=3,
            document_fingerprint="f" * 64,
        )
        is True
    )
    assert disputed.status == "dismissed"
    assert disputed.dismiss_reason.startswith("Reviewed the split")
    # …and it stops blocking.
    assert restored.open_disputed_count() == 0
    assert restored.is_complete() is True
    # The audit trail records the disposition.
    assert [e.action for e in disputed.disposition_events] == ["dismissed"]


def test_a_dismissed_dispute_survives_a_save_and_reload():
    """`QCRunner.dismiss` records the id in `dismissed_ids`, so the reload
    reconciliation has to expect it there — computing the expected set from
    survivors alone discarded the entire retained report."""
    runner = QCRunner()
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(True),
                    qc_verdict_response(True),
                    qc_verdict_response(False),
                ]
            )
        ),
        _section(),
    )
    runner.restore(result)
    disputed_id = runner.result.disputed[0].finding_id
    assert runner.dismiss(disputed_id, "Considered and set aside.") is True
    assert disputed_id in runner.result.dismissed_ids

    reloaded = QCResult.from_dict(runner.result.to_dict())
    assert reloaded is not None, "a dismissed dispute must not void the report"
    assert reloaded.disputed[0].status == "dismissed"
    assert reloaded.open_disputed_count() == 0


def test_every_v4_finding_persists_the_rule_that_adjudicated_it():
    result = _run(
        SequencedFakeClient(
            _medium_scripts([qc_verdict_response(True), qc_verdict_response(True)])
        ),
        _section(),
    )
    finding = result.findings[0]
    assert finding.verification_rule == VERIFICATION_RULE_V4
    # v4 upholds only unanimously, so the integer bar equals the panel size.
    assert finding.verification_threshold == finding.verification_panel_size == 2
    assert result.schema_version == 4
    assert result.protocol_version == "final-qc/4"


def test_a_reloaded_v4_report_re_adjudicates_to_the_same_outcome():
    """Serialized, reloaded, and checked against the same helper the live run
    used — so a report either agrees with itself or fails the check."""
    result = _run(
        SequencedFakeClient(
            _high_scripts(
                [
                    qc_verdict_response(True),
                    qc_verdict_response(True),
                    qc_verdict_response(False),
                ]
            )
        ),
        _section(),
    )
    restored = QCResult.from_dict(result.to_dict())
    assert restored is not None
    assert len(restored.disputed) == 1
    assert restored.disputed[0].verification_outcome == "disputed"
    assert restored.disputed[0].dispute_reason == "split_panel"
    # Structurally sound (the reload re-adjudicated to the same outcome);
    # still blocking, via the separate open-dispute term.
    assert restored.verification_complete() is True
    assert restored.open_disputed_count() == 1
    # The seat-level evidence records survive the round trip.
    assert restored.disputed[0].verdicts[2].upholds is False


def test_dead_verifier_counts_as_refuted():
    store = _section()
    scripts = _qc_scripts(
        provenance_hygiene=[
            qc_findings_response(
                "provenance_hygiene",
                findings=[
                    _finding("Risky assumed block", "Assumed default.", severity="low")
                ],
            )
        ],
    )
    # low → 2-panel. One uphold plus one dead verifier is infrastructure-
    # inconclusive, not evidence that substantively refutes the candidate.
    scripts["Risky assumed block"] = [
        qc_verdict_response(True),
        RuntimeError("verifier died"),
    ]
    result = _run(SequencedFakeClient(scripts), store)
    assert result.findings == []
    assert result.refuted == []
    assert len(result.inconclusive) == 1


# ---------------------------------------------------------------------------
# Phase 3 — ops validation
# ---------------------------------------------------------------------------


def test_valid_ops_dry_run_marks_valid_without_mutating_source():
    store = _section()
    before = store.doc.to_dict()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Fixable edition",
                        "Update the reference.",
                        severity="medium",
                        ops=[
                            {
                                "action": "replace",
                                "target_id": "pt1.a1.p1",
                                "text": "Comply with NFPA 13-2025 throughout.",
                                "status": "confirmed",
                            }
                        ],
                    )
                ],
            )
        ],
    )
    scripts["Fixable edition"] = [qc_verdict_response(True) for _ in range(2)]
    result = _run(SequencedFakeClient(scripts), store)
    assert result.findings[0].ops_valid is True
    # The snapshot was NOT mutated by the dry-run.
    assert store.doc.to_dict() == before


def test_invalid_ops_become_advisory_with_reason():
    store = _section()
    scripts = _qc_scripts(
        coordination_consistency=[
            qc_findings_response(
                "coordination_consistency",
                findings=[
                    _finding(
                        "Bad target",
                        "Fix references a missing id.",
                        severity="medium",
                        ops=[
                            {
                                "action": "replace",
                                "target_id": "pt9.a9.p9",  # does not exist
                                "text": "nope",
                            }
                        ],
                    )
                ],
            )
        ],
    )
    scripts["Bad target"] = [qc_verdict_response(True) for _ in range(2)]
    result = _run(SequencedFakeClient(scripts), store)
    assert result.findings[0].ops_valid is False
    assert "pt9.a9.p9" in result.findings[0].ops_invalid_reason


def test_source_backed_ops_require_final_state_preservation_guard(
    tmp_path, monkeypatch
):
    """Semantic validity alone cannot make an imported-DOCX fix applicable."""
    client = _client()
    source = make_fidelity_master(tmp_path)
    imported = client.post(
        "/api/import/master",
        files={
            "file": (
                "qc-source-master.docx",
                source,
                DOCX_MEDIA_TYPE,
            )
        },
    )
    assert imported.status_code == 200, imported.text

    session = sessions.get_session()
    before = session.doc.doc.to_dict()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Unsafe heading rewrite",
                        "Rewrite the imported article heading.",
                        element_id="pt1.a1",
                        severity="medium",
                        ops=[
                            {
                                "action": "replace",
                                "target_id": "pt1.a1",
                                "text": "REVISED SUMMARY",
                            }
                        ],
                    ),
                    _finding(
                        "Safe provision rewrite",
                        "Update one mapped body provision.",
                        severity="medium",
                        ops=[
                            {
                                "action": "replace",
                                "target_id": "pt1.a1.p1",
                                "text": TARGET_EDITED_TEXT,
                                "status": "confirmed",
                            }
                        ],
                    ),
                ],
            )
        ],
    )
    scripts["Unsafe heading rewrite"] = [
        qc_verdict_response(True) for _ in range(2)
    ]
    scripts["Safe provision rewrite"] = [
        qc_verdict_response(True) for _ in range(2)
    ]

    monkeypatch.setattr(
        "backend.app.get_client", lambda: SequencedFakeClient(scripts)
    )
    assert client.post("/api/qc/start").json()["ok"] is True
    snapshot = _wait_qc(client)
    assert snapshot["status"] == "complete"
    findings = {
        finding["title"]: finding
        for finding in snapshot["result"]["findings"]
    }

    unsafe = findings["Unsafe heading rewrite"]
    assert unsafe["ops_valid"] is False
    assert "source-backed edit rejected" in unsafe["ops_invalid_reason"].lower()
    safe = findings["Safe provision rewrite"]
    assert safe["ops_valid"] is True
    assert safe["ops_invalid_reason"] == ""
    # Both validation layers are dry-runs; neither may touch the live session.
    assert session.doc.doc.to_dict() == before


def test_source_backed_qc_fails_closed_when_guard_context_is_incomplete(
    tmp_path, monkeypatch
):
    client = _client()
    source = make_fidelity_master(tmp_path)
    imported = client.post(
        "/api/import/master",
        files={
            "file": (
                "qc-incomplete-source-master.docx",
                source,
                DOCX_MEDIA_TYPE,
            )
        },
    )
    assert imported.status_code == 200, imported.text

    # Keep the active imported baseline but corrupt one required input. The
    # app must still mark this run source-backed; None must not downgrade it
    # to semantic-only proposal validation.
    sessions.get_session().source_docx_map = None
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Apparently safe provision rewrite",
                        "Update one body provision.",
                        severity="medium",
                        ops=[
                            {
                                "action": "replace",
                                "target_id": "pt1.a1.p1",
                                "text": TARGET_EDITED_TEXT,
                                "status": "confirmed",
                            }
                        ],
                    )
                ],
            )
        ],
    )
    scripts["Apparently safe provision rewrite"] = [
        qc_verdict_response(True) for _ in range(2)
    ]
    monkeypatch.setattr(
        "backend.app.get_client", lambda: SequencedFakeClient(scripts)
    )

    assert client.post("/api/qc/start").json()["ok"] is True
    snapshot = _wait_qc(client)
    assert snapshot["status"] == "complete"
    finding = snapshot["result"]["findings"][0]
    assert finding["ops_valid"] is False
    assert "incomplete source-preservation context" in finding[
        "ops_invalid_reason"
    ].lower()


# ---------------------------------------------------------------------------
# Dismiss memory across re-runs (content-addressed ids)
# ---------------------------------------------------------------------------


def test_dismiss_memory_survives_a_rerun():
    store = _section()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[_finding("Persistent finding", "Same each run.", severity="medium")],
            )
        ],
    )
    scripts["Persistent finding"] = [qc_verdict_response(True) for _ in range(2)]
    first = _run(SequencedFakeClient(scripts), store)
    fid = first.findings[0].finding_id
    runner = QCRunner()
    runner.restore(first)
    assert runner.dismiss(
        fid,
        "Reviewed with the engineer of record.",
        document_version=store.index,
        document_fingerprint=qc_version_fingerprint(store.doc),
    )
    remembered = runner.remembered_dismissals()

    scripts2 = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[_finding("Persistent finding", "Same each run.", severity="medium")],
            )
        ],
    )
    scripts2["Persistent finding"] = [qc_verdict_response(True) for _ in range(2)]
    second = _run(
        SequencedFakeClient(scripts2), store, remembered=remembered
    )
    assert second.findings[0].finding_id == fid
    assert second.findings[0].status == "dismissed"
    assert fid in second.dismissed_ids


# ---------------------------------------------------------------------------
# API lifecycle: start gate, apply, dismiss, project round-trip, staleness
# ---------------------------------------------------------------------------


def _seed_doc(client: TestClient, monkeypatch) -> None:
    draft = FakeClient(
        [
            tool_turn(
                ["Drafting."],
                {
                    "edits": [
                        {"action": "add_article", "target_id": "pt1", "text": "SUMMARY"},
                        {
                            "action": "add_paragraph",
                            "target_id": "pt1.a1",
                            "text": "Comply with NFPA 13-2019 throughout.",
                            "status": "assumed",
                        },
                    ]
                },
            ),
            text_turn(["Done."]),
        ]
    )
    monkeypatch.setattr("backend.llm.conversation.get_client", lambda: draft)
    client.post("/api/chat", json={"message": "draft"})


def _wait_qc(client: TestClient, timeout_s: float = 5.0) -> dict:
    deadline = time.monotonic() + timeout_s
    while time.monotonic() < deadline:
        snap = client.get("/api/qc/status").json()
        if snap["status"] in ("complete", "failed"):
            return snap
        time.sleep(0.05)
    raise AssertionError("QC did not settle")


def _apply_scripts_with_fix() -> dict[str, list]:
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Edition fix",
                        "Bump edition.",
                        severity="high",
                        ops=[
                            {
                                "action": "replace",
                                "target_id": "pt1.a1.p1",
                                "text": "Comply with NFPA 13-2025 throughout.",
                                "status": "confirmed",
                            }
                        ],
                    )
                ],
            )
        ],
    )
    scripts["Edition fix"] = [qc_verdict_response(True) for _ in range(3)]
    return scripts


def test_qc_start_gates_and_apply_is_one_undo_step(monkeypatch):
    client = _client()
    # Empty doc → 400.
    assert client.post("/api/qc/start").status_code == 400

    _seed_doc(client, monkeypatch)
    monkeypatch.setattr(
        "backend.app.get_client", lambda: SequencedFakeClient(_apply_scripts_with_fix())
    )
    assert client.post("/api/qc/start").json()["ok"] is True
    snap = _wait_qc(client)
    assert snap["status"] == "complete"
    finding = snap["result"]["findings"][0]
    assert finding["ops_valid"] is True
    fid = finding["finding_id"]

    versions_before = client.get("/api/doc").json()["doc"]["version"]["count"]
    resp = client.post("/api/qc/apply", json={"finding_ids": [fid]}).json()
    assert resp["ok"] is True
    assert resp["outcomes"][fid] == "applied"
    doc = client.get("/api/doc").json()["doc"]
    assert doc["version"]["count"] == versions_before + 1  # exactly one new version
    assert "2025" in doc["parts"][0]["articles"][0]["paragraphs"][0]["text"]
    # The finding is now marked applied.
    assert client.get("/api/qc/status").json()["result"]["findings"][0]["status"] == "applied"

    # One undo step reverts the whole accept-set.
    client.post("/api/doc/undo")
    reverted = client.get("/api/doc").json()["doc"]
    assert "2019" in reverted["parts"][0]["articles"][0]["paragraphs"][0]["text"]


def test_qc_prefers_document_identity_over_legacy_discipline(monkeypatch):
    client = _client()
    client.post(
        "/api/session/reset",
        json={"module_id": "generic", "discipline": "Electrical"},
    )
    _seed_doc(client, monkeypatch)
    identity = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_project_identity",
                    "target_id": "sec",
                    "discipline": "Structural",
                }
            ]
        },
    )
    assert identity.status_code == 200, identity.text
    fake = SequencedFakeClient(_qc_scripts())
    monkeypatch.setattr("backend.app.get_client", lambda: fake)
    assert client.post("/api/qc/start").json()["ok"] is True
    assert _wait_qc(client)["status"] == "complete"
    assert len(fake.requests) == len(QC_LENSES)
    # The discipline rides the cached shared prefix (block 0 of the user turn).
    assert all(
        "<project_discipline>\nStructural\n</project_discipline>"
        in req["messages"][0]["content"][0]["text"]
        and "<project_discipline>\nElectrical\n</project_discipline>"
        not in req["messages"][0]["content"][0]["text"]
        for req in fake.requests
    )


def test_qc_apply_rejects_result_when_doc_moved(monkeypatch):
    client = _client()
    _seed_doc(client, monkeypatch)
    monkeypatch.setattr(
        "backend.app.get_client", lambda: SequencedFakeClient(_apply_scripts_with_fix())
    )
    client.post("/api/qc/start")
    snap = _wait_qc(client)
    fid = snap["result"]["findings"][0]["finding_id"]

    # Delete the target the fix depends on → the fix is now stale.
    client.post(
        "/api/doc/edit",
        json={"ops": [{"action": "delete", "target_id": "pt1.a1.p1"}]},
    )
    response = client.post("/api/qc/apply", json={"finding_ids": [fid]})
    assert response.status_code == 409
    assert "stale" in response.json()["error"].lower()


def test_qc_dismiss_and_project_round_trip_and_staleness(monkeypatch):
    client = _client()
    _seed_doc(client, monkeypatch)
    monkeypatch.setattr(
        "backend.app.get_client", lambda: SequencedFakeClient(_apply_scripts_with_fix())
    )
    client.post("/api/qc/start")
    snap = _wait_qc(client)
    fid = snap["result"]["findings"][0]["finding_id"]

    # Dismiss with a reason.
    dm = client.post(
        "/api/qc/dismiss", json={"finding_id": fid, "reason": "handled offline"}
    ).json()
    assert dm["ok"] is True
    assert dm["qc"]["result"]["dismissed_ids"] == [fid]

    # Project round-trips the QC result.
    project = json.loads(json.dumps(sessions.project_payload(sessions.get_session())))
    assert project["qc_result"]["findings"]
    client.post("/api/session/reset")
    assert client.get("/api/qc/status").json()["status"] == "idle"
    client.post("/api/project/load", json=project)
    restored = client.get("/api/qc/status").json()
    assert restored["status"] == "complete"
    assert restored["result"]["findings"][0]["status"] == "dismissed"
    assert restored["events"][0].get("restored") is True

    # Staleness: change the doc → version_index no longer matches.
    client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {"action": "add_paragraph", "target_id": "pt1.a1", "text": "New line."}
            ]
        },
    )
    ready = client.get("/api/readiness").json()
    qc_check = next(c for c in ready["checks"] if c["id"] == "qc_current")
    assert qc_check["ok"] is False


def test_qc_double_start_conflicts_and_reset_abandons(monkeypatch):
    client = _client()
    _seed_doc(client, monkeypatch)

    import threading

    release = threading.Event()

    class _Blocking:
        def __init__(self):
            self.messages = self

        def stream(self, **_request):
            release.wait(timeout=5)
            raise RuntimeError("aborted")

    monkeypatch.setattr("backend.app.get_client", lambda: _Blocking())
    assert client.post("/api/qc/start").json()["ok"] is True
    assert client.get("/api/qc/status").json()["status"] == "running"
    # Double-start conflicts.
    assert client.post("/api/qc/start").status_code == 409

    old_runner = sessions.get_session().qc
    client.post("/api/session/reset")
    assert client.get("/api/qc/status").json()["status"] == "idle"
    release.set()
    deadline = time.monotonic() + 5
    while time.monotonic() < deadline and not old_runner.is_terminal:
        time.sleep(0.05)
    assert old_runner.status == "failed"
    assert client.get("/api/qc/status").json()["status"] == "idle"


def test_qc_apply_and_start_reject_while_turn_active(monkeypatch):
    client = _client()
    _seed_doc(client, monkeypatch)
    sessions.get_session().turn_active = True
    try:
        assert client.post("/api/qc/start").status_code == 409
        assert client.post("/api/qc/apply", json={"finding_ids": []}).status_code == 409
    finally:
        sessions.get_session().turn_active = False


# ---------------------------------------------------------------------------
# Readiness + report export + usage
# ---------------------------------------------------------------------------


def test_readiness_reflects_state(monkeypatch):
    client = _client()
    _seed_doc(client, monkeypatch)
    ready = client.get("/api/readiness").json()
    checks = {c["id"]: c for c in ready["checks"]}
    # A fresh assumed draft with no research is not ready.
    assert checks["no_assumed_left"]["ok"] is False
    assert checks["research_complete"]["ok"] is False
    assert checks["qc_current"]["ok"] is False
    assert ready["ready"] is False
    # profile_complete is advisory (does not gate on its own).
    assert checks["profile_complete"]["advisory"] is True


def test_qc_report_export_smoke(monkeypatch):
    client = _client()
    _seed_doc(client, monkeypatch)
    # No QC yet → 409.
    assert client.get("/api/qc/export").status_code == 409

    monkeypatch.setattr(
        "backend.app.get_client", lambda: SequencedFakeClient(_apply_scripts_with_fix())
    )
    client.post("/api/qc/start")
    _wait_qc(client)
    resp = client.get("/api/qc/export")
    assert resp.status_code == 200
    texts = [p.text for p in Document(io.BytesIO(resp.content)).paragraphs]
    assert any("FINAL QC AUDIT REPORT" in t for t in texts)
    assert any("Edition fix" in t for t in texts)


def test_qc_usage_rolls_up_under_the_qc_model_pricing(monkeypatch):
    client = _client()
    _seed_doc(client, monkeypatch)
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[_finding("A finding", "issue", severity="medium")],
                tokens={"input": 4000, "output": 800},
            )
        ],
    )
    scripts["A finding"] = [
        qc_verdict_response(True, tokens={"input": 500}),
        qc_verdict_response(True, tokens={"input": 500}),
    ]
    monkeypatch.setattr("backend.app.get_client", lambda: SequencedFakeClient(scripts))
    client.post("/api/qc/start")
    _wait_qc(client)

    usage = client.get("/api/usage").json()
    # Split by the rate each phase was billed at: the streamed lens phase at
    # list price, the batched verification phase at the provider's batch
    # rate. One bucket could only ever be priced at one of the two.
    qc = usage["categories"]["qc"]
    batched = usage["categories"]["qc_batched"]
    assert qc["input_tokens"] == 4000  # the lens phase
    assert qc["output_tokens"] == 800
    assert batched["input_tokens"] == 1000  # 2 x 500 verifier seats
    # The QC model MUST carry its own PRICING entry. A model absent from the
    # table is silently metered at Sonnet 5's rates (``usage_ledger._rates``
    # falls back via ``dict.get``) and the resulting cost_basis still passes
    # every audit-integrity gate — so the wrong number would ship unnoticed.
    assert settings.QC_MODEL in settings.PRICING
    rates = settings.PRICING[settings.QC_MODEL]
    by_category = usage["estimated_cost_usd"]["by_category"]
    assert by_category["qc"] == round(
        4000 * rates["input"] + 800 * rates["output"], 6
    )
    assert by_category["qc_batched"] == round(
        settings.BATCH_COST_MULTIPLIER * 1000 * rates["input"], 6
    )


def test_qc_result_from_dict_degrades_on_malformed_data():
    # A non-numeric version_index / finding_count must degrade to None, never
    # raise — project.load restores QC after the doc/history are swapped in.
    assert (
        QCResult.from_dict(
            {
                "findings": [
                    {"finding_id": "qc-x", "title": "t", "issue": "i", "severity": "high"}
                ],
                "version_index": "not-a-number",
            }
        )
        is None
    )
    assert (
        QCResult.from_dict(
            {"lens_statuses": [{"lens_id": "x", "finding_count": "NaN"}]}
        )
        is None
    )
    assert QCResult.from_dict("garbage") is None


def test_project_load_survives_malformed_qc_result(monkeypatch):
    client = _client()
    _seed_doc(client, monkeypatch)
    project = json.loads(json.dumps(sessions.project_payload(sessions.get_session())))
    # A malformed QC result (non-numeric version_index) would raise in a naive
    # from_dict — the load must still succeed and the doc must still restore.
    project["qc_result"] = {
        "findings": [
            {"finding_id": "qc-x", "title": "t", "issue": "i", "severity": "high"}
        ],
        "version_index": "not-a-number",
    }
    client.post("/api/session/reset")
    resp = client.post("/api/project/load", json=project)
    assert resp.status_code == 200
    assert resp.json()["ok"] is True
    # QC degraded to "not run"; the document still loaded.
    assert client.get("/api/qc/status").json()["status"] == "idle"
    assert client.get("/api/doc").json()["doc"]["parts"][0]["articles"]


def test_qc_result_from_dict_round_trips():
    store = _section()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[_finding("Round trip", "issue", severity="high")],
            )
        ],
    )
    scripts["Round trip"] = [qc_verdict_response(True) for _ in range(3)]
    result = _run(SequencedFakeClient(scripts), store)
    again = QCResult.from_dict(result.to_dict())
    assert again is not None
    assert again.findings[0].title == "Round trip"
    assert again.model == settings.QC_MODEL


def test_qc_proposed_ops_allow_set_standard_suppressed():
    """A QC fix may propose set_standard_suppressed (the standards-manager op):
    it survives normalize_findings/_clean_op with its fields intact, so a
    standards-scope fix QC can now describe is also one Apply QC can enact —
    matching the /api/doc/edit vocabulary the lens reasons from."""
    from backend.qc.schema import QC_OP_ACTIONS, normalize_findings

    assert "set_standard_suppressed" in QC_OP_ACTIONS
    op = {
        "action": "set_standard_suppressed",
        "target_id": "sec",
        "standard": "NFPA 2001",
        "suppressed": True,
        "basis": "no clean-agent system in scope",
    }
    payload = {
        "summary": "s",
        "findings": [
            _finding(
                "Exclude NFPA 2001",
                "No clean-agent system; it should not reach REFERENCES.",
                element_id="sec",
                ops=[op],
            )
        ],
    }
    cleaned = normalize_findings(payload)["findings"][0]["proposed_ops"]
    assert cleaned == [op]
    # An added standard's title also rides through (the other mirrored field).
    add = {
        "action": "set_standard_edition",
        "target_id": "sec",
        "standard": "NFPA 30",
        "edition": "2024",
        "basis": "on-site flammable storage",
        "title": "Flammable and Combustible Liquids Code",
    }
    add_payload = {"summary": "s", "findings": [_finding("Add", "x", ops=[add])]}
    assert normalize_findings(add_payload)["findings"][0]["proposed_ops"] == [add]


def test_qc_proposed_ops_retain_valid_same_parent_move():
    from backend.qc.schema import QC_OP_ACTIONS, normalize_findings

    assert "move" in QC_OP_ACTIONS
    op = {
        "action": "move",
        "target_id": "pt1.a1.p3",
        "position": 0,
    }
    payload = {
        "summary": "s",
        "findings": [
            _finding(
                "Reorder provision",
                "The provision belongs first in its current article.",
                element_id="pt1.a1.p3",
                ops=[op],
            )
        ],
    }
    assert normalize_findings(payload)["findings"][0]["proposed_ops"] == [op]


# ---------------------------------------------------------------------------
# Batch 10: session discipline in the lens user message
# ---------------------------------------------------------------------------


def test_lens_shared_prefix_carries_discipline_only_when_stated():
    from backend.llm.conversation import SessionState
    from backend.qc.engine import _lens_shared_prefix
    from backend.spec_doc.model import SpecSection

    section = SpecSection()
    module = SessionState().module  # the curated default

    without = _lens_shared_prefix(section, module, None)
    assert "<project_discipline>" not in without

    with_discipline = _lens_shared_prefix(section, module, None, "Electrical")
    # The discipline leads the cached prefix, ahead of the standards block.
    assert with_discipline.startswith(
        "<project_discipline>\nElectrical\n</project_discipline>\n\n"
        "<standards_in_effect>"
    )
    # Empty discipline (the curated case) is byte-identical to before.
    assert _lens_shared_prefix(section, module, None, "") == without


def test_lens_shared_prefix_is_identical_across_lenses():
    """The cached block must not vary by lens, or nothing is ever reused."""
    from backend.llm.conversation import SessionState
    from backend.qc.engine import _lens_request_suffix, _lens_shared_prefix
    from backend.spec_doc.model import SpecSection

    section = SpecSection()
    module = SessionState().module
    prefixes = {
        _lens_shared_prefix(section, module, None, "Electrical")
        for _lens in QC_LENSES
    }
    assert len(prefixes) == 1
    # ...and the per-lens bytes really do differ, so the split is load-bearing.
    assert len({_lens_request_suffix(lens) for lens in QC_LENSES}) == len(
        QC_LENSES
    )


def test_qc_requests_cache_the_shared_prefix_across_the_whole_fan_out(
    monkeypatch,
):
    """The document is billed once per lineage, not once per call.

    A QC run is ~40 calls and every one needs the whole section. Before the
    breakpoint moved onto the user turn, each call re-sent the document at
    full input price. The cache is a strict prefix match, so this asserts
    the two things that make the saving real: the shared block carries a
    breakpoint, and it is byte-identical across every call that shares a
    lineage. A reordering regression would show up only as a bill that
    failed to drop, which nothing else would catch.
    """
    client = _client()
    _seed_doc(client, monkeypatch)
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[_finding("A finding", "issue", severity="medium")],
            )
        ],
    )
    scripts["A finding"] = [
        qc_verdict_response(True),
        qc_verdict_response(True),
    ]
    fake = SequencedFakeClient(scripts)
    monkeypatch.setattr("backend.app.get_client", lambda: fake)
    client.post("/api/qc/start")
    assert _wait_qc(client)["status"] == "complete"

    lens_requests, verifier_requests = [], []
    for request in fake.requests:
        blocks = request["messages"][0]["content"]
        # Every QC user turn is [shared prefix (cached), per-call tail].
        assert [block["type"] for block in blocks] == ["text", "text"]
        assert blocks[0]["cache_control"]["type"] == "ephemeral"
        assert "cache_control" not in blocks[1]
        (
            verifier_requests
            if "[[QC-VERIFY:" in blocks[1]["text"]
            else lens_requests
        ).append(request)

    assert len(lens_requests) == len(QC_LENSES)
    assert len(verifier_requests) == 2

    # The four web-toolless lenses share one lineage. code_compliance cannot
    # join it: its tools array carries web search/fetch and tools render
    # ahead of system and messages, so its byte prefix diverges from the
    # start. It still caches across its own pause_turn continuations.
    non_web = [
        request
        for request in lens_requests
        if not any(
            str(tool.get("type", "")).startswith("web_")
            for tool in request["tools"]
        )
    ]
    assert len(non_web) == len(QC_LENSES) - 1
    assert len({request["messages"][0]["content"][0]["text"] for request in non_web}) == 1
    assert len({request["system"][0]["text"] for request in non_web}) == 1

    # Every verifier seat shares one lineage, and its entry has to outlive
    # the 5-minute default: the verification phase runs far longer than that.
    assert (
        len({r["messages"][0]["content"][0]["text"] for r in verifier_requests})
        == 1
    )
    assert all(
        r["messages"][0]["content"][0]["cache_control"]["ttl"] == "1h"
        for r in verifier_requests
    )

    # Four breakpoints is the API ceiling (tools + system + prefix + tail).
    for request in fake.requests:
        blocks = [*request["tools"], *request["system"], *request["messages"][0]["content"]]
        assert sum("cache_control" in block for block in blocks) <= 4

    # Every breakpoint in one request must carry the SAME ttl. The API
    # requires longer-lived cache entries to precede shorter-lived ones in
    # prompt order, and the render order is tools -> system -> messages — so
    # default-ttl tools/system followed by a 1h user block is rejected
    # outright, on every call, with a nonretryable 400 that trips the shared
    # verifier circuit breaker. Uniform-per-request means there is no order
    # to get wrong. A fake client cannot reject a malformed request, so this
    # is the only place that constraint is enforced before it reaches a
    # provider.
    for request in fake.requests:
        ttls = [
            block["cache_control"].get("ttl")
            for block in (
                *request["tools"],
                *request["system"],
                *request["messages"][0]["content"],
            )
            if "cache_control" in block
        ]
        assert ttls, "a QC request with no breakpoint at all caches nothing"
        assert len(set(ttls)) == 1, f"mixed cache TTLs in one request: {ttls}"

    assert all(
        block["cache_control"]["ttl"] == "1h"
        for request in verifier_requests
        for block in (
            *request["tools"],
            *request["system"],
            *request["messages"][0]["content"],
        )
        if "cache_control" in block
    )
    # Lens calls take the 5-minute default throughout: phase 1 is five calls
    # and does not outlive it.
    assert all(
        "ttl" not in block["cache_control"]
        for request in lens_requests
        for block in (
            *request["tools"],
            *request["system"],
            *request["messages"][0]["content"],
        )
        if "cache_control" in block
    )


def test_qc_web_tools_declare_direct_callers_in_both_phases(monkeypatch):
    """Every QC call carrying web tools pins ``allowed_callers: ["direct"]``.

    Both phases matter: the ``code_compliance`` lens does the searching, and
    its findings' verifier seats get their own small allowance. Under the
    provider default (a code-execution caller) a paused call on either would
    need a container id neither phase sends, and the Review Room's live
    query/URL labels would have nothing to render.
    """
    client = _client()
    _seed_doc(client, monkeypatch)
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[_finding("A finding", "issue", severity="medium")],
            )
        ],
    )
    scripts["A finding"] = [qc_verdict_response(True), qc_verdict_response(True)]
    fake = SequencedFakeClient(scripts)
    monkeypatch.setattr("backend.app.get_client", lambda: fake)
    client.post("/api/qc/start")
    assert _wait_qc(client)["status"] == "complete"

    web_requests = [
        request
        for request in fake.requests
        if any(
            str(tool.get("type", "")).startswith("web_") for tool in request["tools"]
        )
    ]
    # One web lens + its two verifier seats. The other four lenses reason
    # from the document alone and carry no web tools at all.
    assert len(web_requests) == 3
    for request in web_requests:
        web_tools = [
            tool
            for tool in request["tools"]
            if str(tool.get("type", "")).startswith("web_")
        ]
        assert {tool["type"] for tool in web_tools} == {
            "web_search_20260209",
            "web_fetch_20260209",
        }
        assert all(tool["allowed_callers"] == ["direct"] for tool in web_tools)
        # QC never sends a locale — only research knows the project's.
        assert all("user_location" not in tool for tool in web_tools)

    assert len(fake.requests) - len(web_requests) == len(QC_LENSES) - 1


def test_qc_tools_are_strict_for_the_configured_model():
    """A QC model missing from the strict allowlist degrades silently.

    ``_STRICT_CAPABLE_MODELS`` gates ``strict: true``; an id absent from it
    just omits the key, so payloads stop being schema-guaranteed with no
    error and no 400 — normalize_* then clamps the damage invisibly.
    """
    from backend.qc.schema import submit_qc_findings_tool, submit_qc_verdict_tool

    for tool in (
        submit_qc_findings_tool(model=settings.QC_MODEL),
        submit_qc_verdict_tool(model=settings.QC_MODEL),
    ):
        assert tool.get("strict") is True, (
            f"{settings.QC_MODEL} is not in _STRICT_CAPABLE_MODELS "
            "(backend/research/schema.py) — QC tools lost their strict shape."
        )

