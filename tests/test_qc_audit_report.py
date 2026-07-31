"""Audit-grade Final QC regression tests.

These tests exercise the persisted, user-visible review record rather than
private model reasoning.  Every model interaction is supplied by the
thread-safe sequenced fake client; no test can reach the network.
"""
from __future__ import annotations

import copy
import io
import threading
import time
from contextlib import contextmanager
from dataclasses import replace
from datetime import datetime, timezone
from types import SimpleNamespace

from docx import Document
from fastapi.testclient import TestClient

from backend import sessions, settings
from backend.app import create_app
from backend.qc.engine import (
    QC_PROTOCOL_VERSION,
    QC_REPORT_SCHEMA_VERSION,
    QCDispositionEvent,
    QCFanoutError,
    QCResult,
    QCSourceGuard,
    qc_input_fingerprint,
    qc_version_fingerprint,
    run_final_qc,
)
from backend.qc.runner import QCRunner
from backend.qc.schema import QC_FINDINGS_SCHEMA, QC_LENSES, normalize_findings
from backend.research.engine import DimensionStatus, RequirementsProfile, ResearchItem
from backend.spec_doc.docx_export import (
    QC_GROUNDING_METHODOLOGY_NOTE,
    qc_pre_remediation_state,
    qc_signoff_state,
    QC_REQUEST_METHODOLOGY_NOTE,
    build_docx,
    build_qc_memo,
    qc_request_population,
    qc_request_population_note,
    qc_research_coverage,
    qc_version_label,
)
from backend.spec_doc.model import DocumentStore
from backend.spec_modules import DEFAULT_MODULE
from tests.fakes import (
    SequencedFakeClient,
    pause_response,
    qc_findings_response,
    qc_verdict_response,
    usage,
)


_LENS_KEYS = {
    lens.lens_id: f"[[QC-LENS:{lens.lens_id}]]" for lens in QC_LENSES
}


def _section() -> DocumentStore:
    store = DocumentStore()
    store.begin_turn()
    store.apply_edits(
        [
            {
                "action": "replace",
                "target_id": "sec",
                "text": "WET-PIPE SPRINKLER SYSTEMS",
                "numbering": "21 13 13",
            },
            {"action": "add_article", "target_id": "pt1", "text": "SUMMARY"},
            {
                "action": "add_paragraph",
                "target_id": "pt1.a1",
                "text": "Comply with NFPA 13-2019 throughout.",
                "status": "assumed",
            },
        ]
    )
    store.commit_turn()
    return store


def _finding(
    title: str,
    issue: str,
    *,
    severity: str = "medium",
    source_urls: list[str] | None = None,
    proposed_ops: list[dict] | None = None,
) -> dict:
    return {
        "title": title,
        "severity": severity,
        "element_id": "pt1.a1.p1",
        "issue": issue,
        "rationale": f"Recorded audit rationale for {title}.",
        "source_urls": list(source_urls or []),
        "proposed_ops": proposed_ops,
    }


def _scripts(**per_lens: list) -> dict[str, list]:
    """Return one independent response queue for every configured lens."""
    return {
        key: per_lens.get(lens_id, [qc_findings_response(lens_id, findings=[])])
        for lens_id, key in _LENS_KEYS.items()
    }


def _run(
    client: SequencedFakeClient,
    store: DocumentStore,
    *,
    profile: RequirementsProfile | None = None,
    module=DEFAULT_MODULE,
    discipline: str = "",
    source_guard: QCSourceGuard | None = None,
) -> QCResult:
    return run_final_qc(
        store.doc,
        profile,
        module,
        client,
        model=settings.QC_MODEL,
        max_tokens=settings.QC_MAX_TOKENS,
        version_index=store.index,
        started_at="2026-07-24T10:00:00+00:00",
        finished_at="2026-07-24T10:05:00+00:00",
        discipline=discipline,
        source_guard=source_guard,
    )


def _failed_verdict_response() -> SimpleNamespace:
    """A completed API response with no verdict payload: one failed seat."""
    return SimpleNamespace(content=[], stop_reason="end_turn", usage=usage())


def _rich_audit_result() -> tuple[DocumentStore, QCResult]:
    """One surviving and one refuted candidate with traceable source evidence."""
    store = _section()
    source_url = "https://example.test/nfpa-13-edition"
    checks = [
        {
            "check": "Verify the cited NFPA 13 edition against retrieved evidence",
            "outcome": "finding",
            "notes": "The draft records an edition that requires reviewer action.",
            "element_ids": ["pt1.a1.p1"],
            "source_urls": [source_url],
        },
        {
            "check": "Check the remaining compliance scope",
            "outcome": "passed",
            "notes": "No additional candidate defect was recorded.",
            "element_ids": ["sec"],
            "source_urls": [],
        },
    ]
    scripts = _scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                summary="Edition evidence reviewed; one candidate survived challenge.",
                reviewed_checks=checks,
                searched_urls=[source_url],
                findings=[
                    _finding(
                        "Upheld source-backed edition issue",
                        "The recorded NFPA 13 edition is not supported by the saved basis.",
                        severity="high",
                        source_urls=[source_url],
                        proposed_ops=[
                            {
                                "action": "replace",
                                "target_id": "pt1.a1.p1",
                                "text": "Comply with the recorded edition of NFPA 13.",
                                "status": "confirmed",
                            }
                        ],
                    ),
                    _finding(
                        "Refuted wording concern",
                        "The candidate alleged that the provision was not enforceable.",
                        severity="medium",
                    ),
                ],
            )
        ]
    )
    scripts["Upheld source-backed edition issue"] = [
        qc_verdict_response(
            True,
            note="The retrieved basis supports the defect.",
            ops_note="The replacement fully resolves the recorded edition conflict.",
        ),
        qc_verdict_response(
            True,
            note="The document and citation conflict.",
            ops_note="The replacement introduces no unresolved choice.",
        ),
        qc_verdict_response(
            True,
            note="The complete correction is safe to apply.",
            ops_note="The complete operation set is adequate.",
        ),
    ]
    scripts["Refuted wording concern"] = [
        qc_verdict_response(False, note="The text is objectively measurable."),
        qc_verdict_response(False, note="No actionable language defect remains."),
    ]
    result = _run(SequencedFakeClient(scripts), store)
    assert result.is_complete(), "fixture must be a complete audit run"
    return store, result


def _profile(*, research_date: str = "2026-07-24") -> RequirementsProfile:
    return RequirementsProfile(
        items=[
            ResearchItem(
                item_id="req-1",
                dimension_id="codes",
                topic="NFPA 13 adoption",
                category="governing_codes",
                requirement="Use the adopted edition recorded by the project team.",
                authority="Example AHJ",
                source_urls=["https://example.test/adoption"],
                accepted_sources=["https://example.test/adoption"],
                grounded=True,
                confidence=0.9,
            )
        ],
        dimension_statuses=[
            DimensionStatus(
                dimension_id="codes",
                title="Codes",
                status="completed",
                item_count=1,
                grounded_count=1,
            )
        ],
        research_date=research_date,
        project={"city": "Example", "state_or_province": "CA", "country": "USA"},
    )


def _document_text(document: Document) -> str:
    """Collect body and table text from the generated report."""
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def test_audit_contract_requires_exactly_five_distinct_lens_records() -> None:
    store = _section()
    result = _run(SequencedFakeClient(_scripts()), store)
    expected = [lens.lens_id for lens in QC_LENSES]

    recorded = [status.lens_id for status in result.lens_statuses]
    assert len(expected) == 5
    assert recorded == expected
    assert len(recorded) == len(set(recorded))
    assert result.coverage_complete() is True
    assert result.is_complete() is True

    missing_payload = copy.deepcopy(result.to_dict())
    missing_payload["lens_statuses"].pop()
    assert QCResult.from_dict(missing_payload) is None

    duplicate_payload = copy.deepcopy(result.to_dict())
    duplicate_payload["lens_statuses"][-1] = copy.deepcopy(
        duplicate_payload["lens_statuses"][0]
    )
    assert QCResult.from_dict(duplicate_payload) is None


def test_verifier_panel_preserves_exact_seats_including_failures() -> None:
    store = _section()
    scripts = _scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Three-seat verification",
                        "A high-severity issue needs a complete adversarial panel.",
                        severity="high",
                    )
                ],
            )
        ]
    )
    # Two upholds meet the numeric threshold, but the failed response prevents
    # a substantive outcome. The candidate remains a first-class
    # infrastructure-inconclusive record and execution coverage is incomplete.
    scripts["Three-seat verification"] = [
        qc_verdict_response(True, note="Uphold one."),
        _failed_verdict_response(),
        qc_verdict_response(True, note="Uphold two."),
    ]

    result = _run(SequencedFakeClient(scripts), store)
    assert result.findings == []
    assert result.refuted == []
    assert len(result.inconclusive) == 1
    finding = result.inconclusive[0]
    assert finding.verification_panel_size == 3
    # v4 upholds only unanimously, so the integer bar equals the panel size.
    assert finding.verification_threshold == 3
    assert finding.verification_outcome == "inconclusive"
    assert len(finding.verdicts) == 3
    assert {verdict.reviewer_index for verdict in finding.verdicts} == {1, 2, 3}
    assert sum(verdict.status == "completed" for verdict in finding.verdicts) == 2
    failed = [verdict for verdict in finding.verdicts if verdict.status == "failed"]
    assert len(failed) == 1
    assert "parseable payload" in failed[0].error
    assert failed[0].upholds is False
    assert result.coverage_complete() is True
    assert result.verification_complete() is False
    assert result.execution_status == "partial"
    assert result.is_complete() is False

    restored = QCResult.from_dict(result.to_dict())
    assert restored is not None
    restored_finding = restored.inconclusive[0]
    assert {verdict.reviewer_index for verdict in restored_finding.verdicts} == {
        1,
        2,
        3,
    }
    assert [v.status for v in restored_finding.verdicts].count("failed") == 1


def test_reviewed_checks_schema_normalization_grounding_and_round_trip() -> None:
    assert "reviewed_checks" in QC_FINDINGS_SCHEMA["required"]
    check_schema = QC_FINDINGS_SCHEMA["properties"]["reviewed_checks"]["items"]
    assert set(check_schema["required"]) == {
        "check",
        "outcome",
        "notes",
        "element_ids",
        "source_urls",
    }

    normalized = normalize_findings(
        {
            "summary": "  saved summary  ",
            "reviewed_checks": [
                {
                    "check": "  Verify the citation  ",
                    "outcome": "FINDING",
                    "notes": "  mismatch recorded  ",
                    "element_ids": [" pt1.a1.p1 ", ""],
                    "source_urls": [" https://example.test/source ", None],
                },
                {"check": "", "outcome": "passed"},
            ],
            "findings": [],
        }
    )
    assert normalized == {
        "summary": "saved summary",
        "reviewed_checks": [
            {
                "check": "Verify the citation",
                "outcome": "finding",
                "notes": "mismatch recorded",
                "element_ids": ["pt1.a1.p1"],
                "source_urls": ["https://example.test/source"],
            }
        ],
        "findings": [],
    }

    store = _section()
    source_url = "https://example.test/source"
    scripts = _scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[],
                searched_urls=[source_url],
                reviewed_checks=[normalized["reviewed_checks"][0]],
            )
        ]
    )
    result = _run(SequencedFakeClient(scripts), store)
    check = result.lens_statuses[0].reviewed_checks[0]
    assert check.check == "Verify the citation"
    assert check.outcome == "finding"
    assert check.element_ids == ["pt1.a1.p1"]
    assert check.source_urls == [source_url]
    assert len(check.source_checks) == 1
    assert check.source_checks[0].url == source_url
    assert check.source_checks[0].accepted is True
    assert "search" in check.source_checks[0].methods

    restored = QCResult.from_dict(result.to_dict())
    assert restored is not None
    assert (
        restored.lens_statuses[0].reviewed_checks[0].to_dict()
        == check.to_dict()
    )


def test_full_input_fingerprint_invalidates_every_material_input_class() -> None:
    store = _section()
    profile = _profile()
    source_guard = QCSourceGuard(
        required=True,
        capability_summary="Mapped body edits are permitted for this snapshot.",
    )
    result = _run(
        SequencedFakeClient(_scripts()),
        store,
        profile=profile,
        discipline="Fire Protection",
        source_guard=source_guard,
    )

    match_args = (
        store.index,
        store.doc,
        profile,
        DEFAULT_MODULE,
        "Fire Protection",
        source_guard,
    )
    assert result.input_fingerprint
    assert result.input_manifest["requirements_research"]["present"] is True
    assert result.matches_inputs(
        *match_args, model=settings.QC_MODEL, max_tokens=settings.QC_MAX_TOKENS
    )

    assert not result.matches_inputs(
        store.index,
        store.doc,
        _profile(research_date="2026-07-25"),
        DEFAULT_MODULE,
        "Fire Protection",
        source_guard,
        model=settings.QC_MODEL,
        max_tokens=settings.QC_MAX_TOKENS,
    )
    assert not result.matches_inputs(
        store.index,
        store.doc,
        profile,
        replace(DEFAULT_MODULE, display_name=f"{DEFAULT_MODULE.display_name} changed"),
        "Fire Protection",
        source_guard,
        model=settings.QC_MODEL,
        max_tokens=settings.QC_MAX_TOKENS,
    )
    assert not result.matches_inputs(
        store.index,
        store.doc,
        profile,
        DEFAULT_MODULE,
        "Specifications",
        source_guard,
        model=settings.QC_MODEL,
        max_tokens=settings.QC_MAX_TOKENS,
    )
    assert not result.matches_inputs(
        store.index,
        store.doc,
        profile,
        DEFAULT_MODULE,
        "Fire Protection",
        QCSourceGuard(
            required=True,
            capability_summary="Body edits are blocked for this snapshot.",
        ),
        model=settings.QC_MODEL,
        max_tokens=settings.QC_MAX_TOKENS,
    )
    assert not result.matches_inputs(
        *match_args, model=f"{settings.QC_MODEL}-changed", max_tokens=settings.QC_MAX_TOKENS
    )
    assert not result.matches_inputs(
        *match_args, model=settings.QC_MODEL, max_tokens=settings.QC_MAX_TOKENS - 1
    )


def test_applied_disposition_survives_round_trip_and_undo_cannot_false_green() -> None:
    store = _section()
    scripts = _scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Apply and undo",
                        "The cited edition requires a recorded correction.",
                        proposed_ops=[
                            {
                                "action": "replace",
                                "target_id": "pt1.a1.p1",
                                "text": "Comply with NFPA 13-2025 throughout.",
                                "status": "confirmed",
                            }
                        ],
                    )
                ],
            )
        ]
    )
    scripts["Apply and undo"] = [
        qc_verdict_response(True),
        qc_verdict_response(True),
    ]
    result = _run(SequencedFakeClient(scripts), store)
    assert result.matches_inputs(
        store.index,
        store.doc,
        None,
        DEFAULT_MODULE,
        model=settings.QC_MODEL,
        max_tokens=settings.QC_MAX_TOKENS,
    )

    runner = QCRunner()
    runner.restore(result)
    store.begin_turn()
    store.apply_edits(result.findings[0].proposed_ops)
    store.commit_turn()
    post_apply_fingerprint = qc_version_fingerprint(store.doc)
    runner.mark_applied(
        [result.findings[0].finding_id],
        document_version=store.index,
        document_fingerprint=post_apply_fingerprint,
    )
    finding = result.findings[0]
    assert finding.status == "applied"
    assert finding.disposition_events[-1].action == "applied"
    assert finding.disposition_events[-1].document_version == store.index
    assert (
        finding.disposition_events[-1].document_fingerprint
        == post_apply_fingerprint
    )

    assert store.undo() is True
    # Undo returns to the exact bytes/version the QC originally reviewed.
    assert result.matches_version(store.index, store.doc) is True
    # The applied event belongs to the later fixed state, so the old report
    # must not become current while still claiming that correction was applied.
    assert not result.matches_inputs(
        store.index,
        store.doc,
        None,
        DEFAULT_MODULE,
        model=settings.QC_MODEL,
        max_tokens=settings.QC_MAX_TOKENS,
    )

    restored = QCResult.from_dict(result.to_dict())
    assert restored is not None
    event = restored.findings[0].disposition_events[-1]
    assert event.action == "applied"
    assert event.document_version == 2
    assert event.document_fingerprint == post_apply_fingerprint


def test_json_export_contains_complete_report_and_authoritative_current_state() -> None:
    store, result = _rich_audit_result()
    session = sessions.get_session()
    session.doc = store
    session.qc.restore(result)
    client = TestClient(create_app())

    response = client.get("/api/qc/export.json")
    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/json")
    assert ".json" in response.headers["content-disposition"].lower()
    payload = response.json()
    assert set(payload) == {"report", "current_state"}

    report = payload["report"]
    assert report["schema_version"] == 4
    assert report["protocol_version"] == "final-qc/4"
    assert report["run_id"].startswith("qc-run-")
    assert report["execution_status"] == "complete"
    assert report["version_fingerprint"] == qc_version_fingerprint(store.doc)
    assert report["input_fingerprint"]
    assert report["input_manifest"]["document"]["section_number"] == "21 13 13"
    assert [item["lens_id"] for item in report["lens_statuses"]] == [
        lens.lens_id for lens in QC_LENSES
    ]
    assert all(item["reviewed_checks"] for item in report["lens_statuses"])
    assert len(report["findings"]) == 1
    assert len(report["refuted"]) == 1
    surviving = report["findings"][0]
    refuted = report["refuted"][0]
    assert surviving["reviewed_ref"]
    assert surviving["reviewed_text"] == "Comply with NFPA 13-2019 throughout."
    assert surviving["verification_panel_size"] == 3
    assert len(surviving["verdicts"]) == 3
    assert surviving["source_checks"][0]["accepted"] is True
    assert surviving["proposed_ops"][0]["action"] == "replace"
    assert refuted["verification_outcome"] == "refuted"
    assert len(refuted["verdicts"]) == 2
    # 5 lenses + 1 grouping call + 5 verifier seats. The run totals are the
    # sum of the component records, and the grouping call is one of them —
    # a reader must be able to reconcile the total without inference.
    consolidation = report["consolidation"]
    assert consolidation["api_request_count"] == 1
    assert report["api_request_count"] == 11
    assert report["model_response_count"] == 11
    assert report["api_request_count"] == (
        sum(lens["api_request_count"] for lens in report["lens_statuses"])
        + consolidation["api_request_count"]
        + sum(
            verdict["api_request_count"]
            for bucket in ("findings", "refuted", "disputed", "inconclusive")
            for finding in report[bucket]
            for verdict in finding["verdicts"]
        )
    )
    assert "usage_totals" in report
    assert "estimated_cost_usd" in report

    current = payload["current_state"]
    assert current["generated_at"].endswith("+00:00")
    assert current["document_version"] == store.index
    assert current["document_fingerprint"] == qc_version_fingerprint(store.doc)
    assert current["stale"] is False
    readiness = current["readiness"]
    assert isinstance(readiness["checks"], list)
    qc_current = next(check for check in readiness["checks"] if check["id"] == "qc_current")
    assert qc_current["ok"] is True
    # Chunk 5.4 split the old collapsed check: coverage is complete, but
    # the surviving findings have not been dispositioned, and THAT is what
    # blocks issue now — the meaning the Word sign-off always had.
    execution = next(
        check
        for check in readiness["checks"]
        if check["id"] == "qc_execution_complete"
    )
    assert execution["ok"] is True
    no_open = next(
        check
        for check in readiness["checks"]
        if check["id"] == "no_open_qc_findings"
    )
    assert no_open["ok"] is False
    assert "still open" in no_open["detail"]
    qc_audit = next(
        check
        for check in readiness["checks"]
        if check["id"] == "qc_audit_complete"
    )
    # The derived alias is exactly the conjunction of the two above.
    assert qc_audit["ok"] is (execution["ok"] and no_open["ok"])


def test_project_restore_keeps_dispositions_synchronized_with_exports() -> None:
    store, result = _rich_audit_result()
    session = sessions.get_session()
    session.doc = store
    session.qc.restore(result)
    project = copy.deepcopy(sessions.project_payload(session))
    assert project["qc_latest_attempt"]["report"]["run_id"] == result.run_id

    client = TestClient(create_app())
    assert client.post("/api/session/reset").status_code == 200
    loaded = client.post("/api/project/load", json=project)
    assert loaded.status_code == 200, loaded.text

    restored_runner = sessions.get_session().qc
    assert restored_runner.result is not None
    assert restored_runner.latest_attempt_result is restored_runner.result
    finding_id = restored_runner.result.findings[0].finding_id

    dismissed = client.post(
        "/api/qc/dismiss",
        json={
            "finding_id": finding_id,
            "reason": "Disposition recorded after project reload.",
        },
    )
    assert dismissed.status_code == 200, dismissed.text

    exported = client.get("/api/qc/export.json")
    assert exported.status_code == 200
    exported_finding = exported.json()["report"]["findings"][0]
    assert exported_finding["status"] == "dismissed"
    assert exported_finding["dismiss_reason"] == (
        "Disposition recorded after project reload."
    )


def test_detailed_docx_preserves_key_sections_findings_evidence_and_votes() -> None:
    store, result = _rich_audit_result()
    runner = QCRunner()
    runner.restore(result)
    survivor = result.findings[0]
    assert runner.dismiss(
        survivor.finding_id,
        "Reviewed with the project engineer; correction tracked in coordination.",
        document_version=store.index,
        document_fingerprint=qc_version_fingerprint(store.doc),
    )

    payload = build_qc_memo(result.to_dict(), store.doc, stale=False)
    document = Document(io.BytesIO(payload))
    text = _document_text(document)

    properties = document.core_properties
    assert properties.title == (
        f"Final QC Audit Report - Section {store.doc.number} | {store.doc.title}"
    )
    assert result.run_id in properties.subject
    assert properties.author == "Build-a-Spec"
    assert properties.last_modified_by == "Build-a-Spec"
    now = datetime.now(timezone.utc)
    for timestamp in (properties.created, properties.modified):
        assert timestamp is not None
        if timestamp.tzinfo is None:
            timestamp = timestamp.replace(tzinfo=timezone.utc)
        assert abs((timestamp.astimezone(timezone.utc) - now).total_seconds()) < 60

    for required_section in (
        "FINAL QC AUDIT REPORT",
        "Executive Status",
        "Run and Input Identity",
        "Input Manifest",
        "Methodology and Interpretation",
        "Lens-by-Lens Audit Trail",
        "Complete Surviving Findings Register",
        "Appendix A: Complete Refuted Candidate Register",
        "Appendix B: Evidence Register",
        "Usage, Requests, and Estimated Cost",
        "Limitations",
        "Reviewer Sign-off",
    ):
        assert required_section in text
    for lens in QC_LENSES:
        assert lens.title in text

    assert result.run_id in text
    assert result.version_fingerprint in text
    assert result.input_fingerprint in text
    assert "Upheld source-backed edition issue" in text
    assert "The recorded NFPA 13 edition is not supported by the saved basis." in text
    assert "Recorded audit rationale for Upheld source-backed edition issue." in text
    assert "Comply with NFPA 13-2019 throughout." in text
    assert "The retrieved basis supports the defect." in text
    assert "https://example.test/nfpa-13-edition" in text
    assert "Refuted wording concern" in text
    assert "The candidate alleged that the provision was not enforceable." in text
    assert "Recorded audit rationale for Refuted wording concern." in text
    assert "The text is objectively measurable." in text
    assert "Reviewed with the project engineer" in text
    assert "replace" in text
    assert "Semantic fix decision: APPROVED" in text
    assert "Proposed fix adequate: APPROVED" in text
    assert "The replacement fully resolves the recorded edition conflict." in text
    assert "3 reviewer record(s)" in text
    assert "2 reviewer record(s)" in text


def test_docx_export_time_qc_controls_cannot_present_a_complete_signoff() -> None:
    store, result = _rich_audit_result()
    base = result.to_dict()
    scenarios = [
        ("latest failed", "failed", None),
        ("latest cancelled", "cancelled", None),
        ("latest partial", "partial", None),
        ("latest running", "running", None),
        ("current QC blocked", "complete", "qc_current"),
        ("audit coverage blocked", "complete", "qc_audit_complete"),
    ]

    for scenario, attempt_status, blocked_check in scenarios:
        report = copy.deepcopy(base)
        checks = [
            {
                "id": check_id,
                "ok": check_id != blocked_check,
                "detail": (
                    "Current export control blocks reliance."
                    if check_id == blocked_check
                    else "Complete."
                ),
            }
            for check_id in ("qc_current", "qc_audit_complete")
        ]
        report["export_current_state"] = {
            "runner": {"status": attempt_status},
            "latest_attempt": {
                "run_id": result.run_id,
                "status": attempt_status,
                "report_available": True,
            },
            "readiness": {
                "ready": blocked_check is None
                and attempt_status not in {"failed", "cancelled", "partial", "running"},
                "checks": checks,
            },
        }
        document = Document(
            io.BytesIO(build_qc_memo(report, store.doc, stale=False))
        )
        paragraph_text = [paragraph.text for paragraph in document.paragraphs]
        all_text = _document_text(document)
        assert "Status: BLOCKED - NOT ISSUE-READY" in paragraph_text, scenario
        assert "Status: COMPLETE" not in paragraph_text, scenario
        assert (
            "QC RECORD COMPLETE - PROFESSIONAL APPROVAL STILL REQUIRED"
            not in all_text
        ), scenario
        assert "HOLD - EXPORT-TIME QC CONTROL BLOCKED" in all_text, scenario


def test_stopped_worker_cannot_resolve_or_emit_into_newer_run(monkeypatch) -> None:
    """Run ownership is a token, not the runner's shared `running` status."""
    store, base = _rich_audit_result()
    first_started = threading.Event()
    release_first = threading.Event()
    second_started = threading.Event()
    release_second = threading.Event()
    calls = 0

    def fake_run(*_args, run_id: str, event_sink, **_kwargs):
        nonlocal calls
        calls += 1
        result = copy.deepcopy(base)
        result.run_id = run_id
        if calls == 1:
            first_started.set()
            release_first.wait(timeout=2)
            event_sink({"type": "old_worker_event"})
            return result
        second_started.set()
        release_second.wait(timeout=2)
        event_sink({"type": "new_worker_event"})
        return result

    monkeypatch.setattr("backend.qc.runner.run_final_qc", fake_run)
    runner = QCRunner()
    kwargs = {
        "section": store.doc,
        "profile": None,
        "module": DEFAULT_MODULE,
        "client": object(),
        "model": settings.QC_MODEL,
        "max_tokens": settings.QC_MAX_TOKENS,
        "version_index": store.index,
    }
    assert runner.start(**kwargs)
    assert first_started.wait(timeout=2)
    first_run_id = runner.latest_attempt_run_id
    assert runner.stop()
    # Stop resolves controls immediately, but replacement is blocked until
    # the paid worker attaches its final report and emits settlement.
    assert not runner.start(**kwargs)
    assert runner.audit_record_snapshot()["runner"]["settling"] is True
    release_first.set()
    deadline = time.monotonic() + 2
    while time.monotonic() < deadline:
        if not runner.audit_record_snapshot()["runner"]["settling"]:
            break
        time.sleep(0.01)
    assert runner.status == "failed"
    assert runner.latest_attempt_run_id == first_run_id
    assert runner.latest_attempt_result is not None
    assert runner.latest_attempt_result.run_id == first_run_id
    assert any(
        event.get("type") == "qc_attempt_settled" for event in runner.events
    )

    assert runner.start(**kwargs)
    assert second_started.wait(timeout=2)
    second_run_id = runner.latest_attempt_run_id
    assert first_run_id != second_run_id

    assert runner.status == "running"
    assert runner.latest_attempt_run_id == second_run_id
    assert not any(
        event.get("type") == "old_worker_event" for event in runner.events
    )

    release_second.set()
    deadline = time.monotonic() + 2
    while time.monotonic() < deadline and runner.status == "running":
        time.sleep(0.01)
    assert runner.status == "complete"
    assert runner.result is not None
    assert runner.result.run_id == second_run_id
    assert any(
        event.get("type") == "new_worker_event" for event in runner.events
    )


def test_failed_fetch_and_abandoned_retry_cannot_ground_final_payload() -> None:
    store = _section()
    failed_url = "https://example.test/failed-fetch"
    failed_response = qc_findings_response(
        "code_compliance",
        findings=[
            _finding(
                "Failed fetch citation",
                "The citation was invoked but never successfully fetched.",
                source_urls=[failed_url],
            )
        ],
        fetches=1,
    )
    failed_response.content[0:0] = [
        SimpleNamespace(
            type="server_tool_use",
            id="fetch-1",
            name="web_fetch",
            input={"url": failed_url},
        ),
        SimpleNamespace(
            type="web_fetch_tool_result",
            tool_use_id="fetch-1",
            content={"type": "web_fetch_tool_result_error", "error_code": "not_found"},
        ),
    ]
    scripts = _scripts(code_compliance=[failed_response])
    scripts["Failed fetch citation"] = [
        qc_verdict_response(True),
        qc_verdict_response(True),
    ]
    failed_result = _run(SequencedFakeClient(scripts), store)
    failed_finding = failed_result.findings[0]
    assert failed_finding.grounded is False
    assert failed_finding.accepted_sources == []
    assert failed_finding.source_checks[0].accepted is False
    lens = failed_result.lens_statuses[0]
    assert all(source.url != failed_url for source in lens.retrieved_sources)
    attempted = next(source for source in lens.attempted_sources if source.url == failed_url)
    assert attempted.accepted is False
    assert "excluded from grounding" in attempted.reason

    old_url = "https://example.test/abandoned-attempt"
    retry_scripts = _scripts(
        code_compliance=[
            pause_response(searched_urls=[old_url]),
            ConnectionResetError("connection reset by peer"),
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Retry-only citation",
                        "The final attempt cited a URL retrieved only by an abandoned retry.",
                        source_urls=[old_url],
                    )
                ],
            ),
        ]
    )
    retry_scripts["Retry-only citation"] = [
        qc_verdict_response(True),
        qc_verdict_response(True),
    ]
    retry_result = _run(SequencedFakeClient(retry_scripts), store)
    retry_finding = retry_result.findings[0]
    assert retry_finding.grounded is False
    retry_lens = retry_result.lens_statuses[0]
    assert all(source.url != old_url for source in retry_lens.retrieved_sources)
    assert any(source.url == old_url for source in retry_lens.attempted_sources)


def test_malformed_json_verdict_is_failed_not_truthy_uphold() -> None:
    store = _section()
    scripts = _scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Malformed verdict candidate",
                        "A tagged JSON fallback supplied a string boolean.",
                    )
                ],
            )
        ]
    )
    malformed = qc_verdict_response(False, tokens={"input": 250})
    malformed.content[0].input["upholds"] = "false"
    scripts["Malformed verdict candidate"] = [
        malformed,
        qc_verdict_response(False),
    ]
    result = _run(SequencedFakeClient(scripts), store)
    assert result.findings == []
    assert result.refuted == []
    finding = result.inconclusive[0]
    assert finding.verification_outcome == "inconclusive"
    failed = [verdict for verdict in finding.verdicts if verdict.status == "failed"]
    assert len(failed) == 1
    assert failed[0].upholds is False
    assert "must be a JSON boolean" in failed[0].error
    assert failed[0].usage_totals["input_tokens"] == 250
    assert result.execution_status == "partial"


def test_duplicate_apply_ids_are_deduplicated_and_dismiss_is_blocked_while_running(
    monkeypatch,
) -> None:
    store = _section()
    scripts = _scripts(
        completeness=[
            qc_findings_response(
                "completeness",
                findings=[
                    _finding(
                        "Add acceptance test",
                        "The execution article lacks the required test record.",
                        proposed_ops=[
                            {
                                "action": "add_paragraph",
                                "target_id": "pt1.a1",
                                "text": "Submit the signed acceptance-test record.",
                                "status": "confirmed",
                            }
                        ],
                    )
                ],
            )
        ]
    )
    scripts["Add acceptance test"] = [
        qc_verdict_response(True),
        qc_verdict_response(True),
    ]
    result = _run(SequencedFakeClient(scripts), store)
    session = sessions.get_session()
    session.doc = store
    session.qc.restore(result)
    client = TestClient(create_app())
    finding_id = result.findings[0].finding_id
    before = len(store.doc.parts[0].articles[0].paragraphs)
    response = client.post(
        "/api/qc/apply",
        json={"finding_ids": [finding_id, finding_id]},
    )
    assert response.status_code == 200
    assert response.json()["outcomes"] == {finding_id: "applied"}
    assert len(store.doc.parts[0].articles[0].paragraphs) == before + 1

    blocking_runner = QCRunner()
    blocking_runner.restore(copy.deepcopy(result))
    blocking_runner.status = "running"
    session.qc = blocking_runner
    dismiss = client.post(
        "/api/qc/dismiss",
        json={"finding_id": finding_id, "reason": "must not race"},
    )
    assert dismiss.status_code == 409


def test_v2_exports_but_apply_and_dismiss_require_current_audit_contract() -> None:
    store, result = _rich_audit_result()
    finding_id = result.findings[0].finding_id
    session = sessions.get_session()
    session.doc = store
    client = TestClient(create_app())

    legacy_payload = copy.deepcopy(result.to_dict())
    legacy_payload["schema_version"] = 2
    legacy_payload["protocol_version"] = "final-qc/2"
    legacy_payload["input_manifest"]["protocol_version"] = "final-qc/2"
    legacy_payload["input_fingerprint"] = qc_input_fingerprint(
        legacy_payload["input_manifest"]
    )
    for raw_finding in [
        *legacy_payload["findings"],
        *legacy_payload["refuted"],
        *legacy_payload["inconclusive"],
    ]:
        raw_finding.pop("ops_semantic_status")
        raw_finding.pop("ops_semantic_reason")
        for raw_verdict in raw_finding["verdicts"]:
            raw_verdict.pop("ops_adequate")
            raw_verdict.pop("ops_note")
    legacy = QCResult.from_dict(legacy_payload)
    assert legacy is not None
    legacy_runner = QCRunner()
    legacy_runner.restore(legacy)
    assert legacy_runner.result is legacy
    session.qc = legacy_runner

    json_export = client.get("/api/qc/export.json")
    word_export = client.get("/api/qc/export")
    assert json_export.status_code == 200
    assert json_export.json()["report"]["schema_version"] == 2
    assert word_export.status_code == 200
    historical_word = _document_text(
        Document(io.BytesIO(word_export.content))
    )
    assert "HISTORICAL REPORT IS NONACTIONABLE" in historical_word

    apply = client.post("/api/qc/apply", json={"finding_ids": [finding_id]})
    dismiss = client.post(
        "/api/qc/dismiss",
        json={"finding_id": finding_id, "reason": "not actionable"},
    )
    assert apply.status_code == 409
    assert dismiss.status_code == 409
    assert "audit-complete" in apply.json()["error"]
    assert "audit-complete" in dismiss.json()["error"]

    settling_runner = QCRunner()
    settling_runner.restore(copy.deepcopy(result))
    with settling_runner._lock:
        settling_runner.status = "failed"
        settling_runner.latest_attempt_status = "cancelled"
        settling_runner._worker_settled = False
    session.qc = settling_runner

    apply = client.post("/api/qc/apply", json={"finding_ids": [finding_id]})
    dismiss = client.post(
        "/api/qc/dismiss",
        json={"finding_id": finding_id, "reason": "must wait"},
    )
    assert apply.status_code == 409
    assert dismiss.status_code == 409
    assert "settling" in apply.json()["error"]
    assert "settling" in dismiss.json()["error"]


def test_readiness_names_partial_cancelled_and_settling_evidence() -> None:
    store, complete = _rich_audit_result()
    session = sessions.get_session()
    session.doc = store
    client = TestClient(create_app())

    partial = copy.deepcopy(complete)
    partial.run_id = "partial-evidence-run"
    partial.execution_status = "partial"
    partial_runner = QCRunner()
    partial_runner.restore(partial)
    session.qc = partial_runner
    readiness = client.get("/api/readiness").json()
    current = next(
        check for check in readiness["checks"] if check["id"] == "qc_current"
    )
    audit = next(
        check
        for check in readiness["checks"]
        if check["id"] == "qc_audit_complete"
    )
    assert current["ok"] is False
    assert "settled partial" in current["detail"].lower()
    assert "paid report is preserved" in current["detail"].lower()
    assert "partial attempt evidence is preserved" in audit["detail"].lower()

    cancelled_runner = QCRunner()
    cancelled_runner.restore(partial)
    cancelled_runner.restore_attempt(
        {
            "run_id": partial.run_id,
            "status": "cancelled",
            "report": partial.to_dict(),
        }
    )
    session.qc = cancelled_runner
    readiness = client.get("/api/readiness").json()
    current = next(
        check for check in readiness["checks"] if check["id"] == "qc_current"
    )
    assert current["ok"] is False
    assert "cancelled" in current["detail"].lower()
    assert "paid report is preserved" in current["detail"].lower()

    contradictory_runner = QCRunner()
    contradictory_runner.restore(copy.deepcopy(complete))
    contradictory = copy.deepcopy(partial)
    contradictory.run_id = complete.run_id
    contradictory_runner.restore_attempt(
        {
            "run_id": complete.run_id,
            "status": "partial",
            "report": contradictory.to_dict(),
        }
    )
    # Even a torn restore that incorrectly reports runner-level completion
    # cannot bypass the explicit latest-attempt status requirement.
    with contradictory_runner._lock:
        contradictory_runner.status = "complete"
        contradictory_runner.error = ""
    session.qc = contradictory_runner
    readiness = client.get("/api/readiness").json()
    current = next(
        check for check in readiness["checks"] if check["id"] == "qc_current"
    )
    assert current["ok"] is False
    assert "settled partial" in current["detail"].lower()

    settling_runner = QCRunner()
    settling_runner.restore(copy.deepcopy(complete))
    with settling_runner._lock:
        settling_runner.status = "failed"
        settling_runner.latest_attempt_status = "cancelled"
        settling_runner._worker_settled = False
    session.qc = settling_runner
    readiness = client.get("/api/readiness").json()
    current = next(
        check for check in readiness["checks"] if check["id"] == "qc_current"
    )
    assert current["ok"] is False
    assert "still settling" in current["detail"].lower()
    assert "already-paid" in current["detail"].lower()


def test_qc_exports_reject_changed_selected_run_identity() -> None:
    store, result = _rich_audit_result()
    session = sessions.get_session()
    session.doc = store
    session.qc.restore(result)
    client = TestClient(create_app())

    for path in ("/api/qc/export", "/api/qc/export.json"):
        matched = client.get(path, params={"run_id": result.run_id})
        assert matched.status_code == 200, (path, matched.text)

        mismatch = client.get(path, params={"run_id": "stale-view-run"})
        assert mismatch.status_code == 409
        assert mismatch.json()["expected_run_id"] == "stale-view-run"
        assert mismatch.json()["selected_run_id"] == result.run_id


def test_apply_commit_and_outcome_audit_share_one_locked_transaction(
    monkeypatch,
) -> None:
    store, result = _rich_audit_result()
    no_ops = copy.deepcopy(result.findings[0])
    no_ops.finding_id = "qc-no-ops-same-transaction"
    no_ops.title = "Advisory without a mechanical operation"
    no_ops.proposed_ops = []
    no_ops.ops_valid = False
    result.findings.append(no_ops)

    session = sessions.get_session()
    session.doc = store
    session.qc.restore(result)
    original_guard = session.session_state_guard
    acquisitions = 0

    @contextmanager
    def guarded_once_per_phase():
        nonlocal acquisitions
        acquisitions += 1
        if acquisitions > 2:
            raise AssertionError(
                "QC apply reacquired session state after committing the document"
            )
        with original_guard():
            yield

    monkeypatch.setattr(session, "session_state_guard", guarded_once_per_phase)
    response = TestClient(create_app()).post(
        "/api/qc/apply",
        json={
            "finding_ids": [result.findings[0].finding_id, no_ops.finding_id]
        },
    )
    assert response.status_code == 200, response.text
    assert acquisitions == 2
    assert response.json()["outcomes"] == {
        result.findings[0].finding_id: "applied",
        no_ops.finding_id: "no_ops",
    }
    assert no_ops.disposition_events[-1].action == "apply_no_ops"


def test_failed_latest_attempt_blocks_readiness_and_is_exported(monkeypatch) -> None:
    store, successful = _rich_audit_result()
    runner = QCRunner()
    runner.restore(successful)
    session = sessions.get_session()
    session.doc = store
    session.qc = runner

    failed = copy.deepcopy(successful)
    failed.execution_status = "failed"
    failed.summary = "All lenses failed; billable activity was preserved."
    failed.findings = []
    failed.refuted = []
    for lens in failed.lens_statuses:
        lens.status = "failed"
        lens.error = "provider unavailable"
        lens.reviewed_checks = []

    def fail_run(*_args, run_id: str, **_kwargs):
        failed.run_id = run_id
        raise QCFanoutError(
            "All 5 QC lenses failed.",
            usage_totals=failed.usage_totals,
            result=failed,
        )

    monkeypatch.setattr("backend.qc.runner.run_final_qc", fail_run)
    assert runner.start(
        section=store.doc,
        profile=None,
        module=DEFAULT_MODULE,
        client=object(),
        model=settings.QC_MODEL,
        max_tokens=settings.QC_MAX_TOKENS,
        version_index=store.index,
    )
    deadline = time.monotonic() + 2
    while time.monotonic() < deadline and runner.status == "running":
        time.sleep(0.01)
    assert runner.status == "failed"
    assert runner.result is successful
    assert runner.latest_attempt_result is failed
    failed_run_id = runner.latest_attempt_run_id

    client = TestClient(create_app())
    snapshot = client.get("/api/qc/status").json()
    assert snapshot["result"]["run_id"] == successful.run_id
    assert snapshot["report"]["run_id"] == failed_run_id
    assert snapshot["report"]["execution_status"] == "failed"
    assert snapshot["report_is_latest_attempt"] is True
    assert snapshot["latest_attempt"]["report_available"] is True
    readiness = client.get("/api/readiness").json()
    qc_check = next(
        check for check in readiness["checks"] if check["id"] == "qc_current"
    )
    assert qc_check["ok"] is False
    assert "latest" in qc_check["detail"].lower()
    execution_check = next(
        check
        for check in readiness["checks"]
        if check["id"] == "qc_execution_complete"
    )
    # The retained success still contains a complete audit record — which is
    # exactly what the split `qc_execution_complete` says — but the separate
    # current/latest-attempt gate prevents it from signing off the document
    # after the failed rerun.
    assert execution_check["ok"] is True
    assert readiness["ready"] is False

    exported = client.get("/api/qc/export.json")
    assert exported.status_code == 200
    payload = exported.json()
    assert payload["report"]["run_id"] == failed_run_id
    assert payload["report"]["execution_status"] == "failed"
    assert payload["last_successful_report"]["run_id"] == successful.run_id
    assert (
        payload["current_state"]["last_successful_report"]["run_id"]
        == successful.run_id
    )
    assert payload["current_state"]["latest_attempt"]["status"] == "failed"
    assert payload["current_state"]["runner"]["status"] == "failed"

    word_export = client.get("/api/qc/export")
    assert word_export.status_code == 200
    word_document = Document(io.BytesIO(word_export.content))
    word_text = _document_text(word_document)
    word_paragraphs = [paragraph.text for paragraph in word_document.paragraphs]
    assert f"Selected report run ID: {failed_run_id}" in word_paragraphs
    assert (
        f"Retained successful run ID: {successful.run_id}" in word_paragraphs
    )
    assert "Retained Prior Successful Report Identity" in word_text
    assert "HISTORICAL SUCCESS - DOES NOT CONTROL CURRENT READINESS" in word_text
    assert "Status: BLOCKED - NOT ISSUE-READY" in word_paragraphs
    assert "Status: COMPLETE" not in word_paragraphs
    assert "QC RECORD COMPLETE - PROFESSIONAL APPROVAL STILL REQUIRED" not in word_text


def test_future_or_tampered_persisted_report_is_not_restored() -> None:
    _store, result = _rich_audit_result()
    future = result.to_dict()
    future["schema_version"] = 999
    assert QCResult.from_dict(future) is None

    tampered = result.to_dict()
    tampered["input_manifest"]["document"]["section_title"] = "TAMPERED"
    assert QCResult.from_dict(tampered) is None

    assert result.cost_basis["rate_model"]
    assert "rates_per_token" in result.cost_basis
    assert all(lens.estimated_cost_usd >= 0 for lens in result.lens_statuses)


# ---------------------------------------------------------------------------
# Audit hardening regressions: disposition identity, persistence validation,
# rationale requirements, failed reruns, and infrastructure-inconclusive work.
# ---------------------------------------------------------------------------


def _run_with_dismissal_memory(
    client: SequencedFakeClient,
    store: DocumentStore,
    remembered: dict[str, dict],
) -> QCResult:
    return run_final_qc(
        store.doc,
        None,
        DEFAULT_MODULE,
        client,
        model=settings.QC_MODEL,
        max_tokens=settings.QC_MAX_TOKENS,
        version_index=store.index,
        started_at="2026-07-24T11:00:00+00:00",
        finished_at="2026-07-24T11:05:00+00:00",
        remembered_dismissed=remembered,
    )


def _dismissal_candidate_scripts(
    *,
    searched: bool,
    verdict_count: int,
    revised_severity: str | None = None,
) -> dict[str, list]:
    title = "Carry-forward identity candidate"
    source_url = "https://example.test/dismissal-evidence"
    scripts = _scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        title,
                        "The same textual issue is being re-evaluated.",
                        severity="medium",
                        source_urls=[source_url],
                    )
                ],
                searched_urls=[source_url] if searched else None,
            )
        ]
    )
    scripts[title] = [
        qc_verdict_response(
            True,
            severity=revised_severity,
            note=f"Completed reviewer seat {index + 1}.",
        )
        for index in range(verdict_count)
    ]
    return scripts


def test_carried_dismissal_requires_same_final_severity_panel_and_grounding(
    monkeypatch,
) -> None:
    store = _section()
    original_panel_size = settings.QC_VERIFIERS_STANDARD
    first = _run(
        SequencedFakeClient(
            _dismissal_candidate_scripts(
                searched=True,
                verdict_count=original_panel_size,
            )
        ),
        store,
    )
    runner = QCRunner()
    runner.restore(first)
    first_finding = first.findings[0]
    assert runner.dismiss(
        first_finding.finding_id,
        "The reviewer accepted this exact residual risk.",
        document_version=store.index,
        document_fingerprint=qc_version_fingerprint(store.doc),
    )
    remembered = runner.remembered_dismissals()
    assert remembered[first_finding.finding_id]["reason"]

    # The raw candidate is text-identical, but adversarial reviewers revise
    # the final severity. A prior risk acceptance cannot silently cover the
    # materially more severe final determination.
    severity_changed = _run_with_dismissal_memory(
        SequencedFakeClient(
            _dismissal_candidate_scripts(
                searched=True,
                verdict_count=original_panel_size,
                revised_severity="high",
            )
        ),
        store,
        remembered,
    )
    assert severity_changed.findings[0].severity == "high"

    # The cited URL is unchanged, but it was not retrieved in the new run.
    # A dismissal of a grounded finding must not carry into an ungrounded one.
    grounding_changed = _run_with_dismissal_memory(
        SequencedFakeClient(
            _dismissal_candidate_scripts(
                searched=False,
                verdict_count=original_panel_size,
            )
        ),
        store,
        remembered,
    )
    assert grounding_changed.findings[0].grounded is False

    # Panel configuration is part of the final reviewed determination even
    # when the raw lens payload is byte-identical.
    expanded_panel_size = original_panel_size + 1
    monkeypatch.setattr(settings, "QC_VERIFIERS_STANDARD", expanded_panel_size)
    panel_changed = _run_with_dismissal_memory(
        SequencedFakeClient(
            _dismissal_candidate_scripts(
                searched=True,
                verdict_count=expanded_panel_size,
            )
        ),
        store,
        remembered,
    )
    assert panel_changed.findings[0].verification_panel_size == expanded_panel_size

    for rerun in (severity_changed, grounding_changed, panel_changed):
        finding = rerun.findings[0]
        assert finding.status == "open"
        assert finding.dismiss_reason == ""
        assert not finding.disposition_events
        assert finding.finding_id not in rerun.dismissed_ids


def _set_nested(payload: dict, path: tuple[object, ...], value: object) -> None:
    cursor: object = payload
    for key in path[:-1]:
        cursor = cursor[key]  # type: ignore[index]
    cursor[path[-1]] = value  # type: ignore[index]


def test_malformed_persisted_enums_and_string_booleans_are_rejected() -> None:
    _store, result = _rich_audit_result()
    mutations = [
        ("execution status", ("execution_status",), "mystery"),
        ("lens status", ("lens_statuses", 0, "status"), "passed"),
        (
            "reviewed-check outcome",
            ("lens_statuses", 0, "reviewed_checks", 0, "outcome"),
            "okay",
        ),
        ("finding severity", ("findings", 0, "severity"), "urgent"),
        ("finding status", ("findings", 0, "status"), "accepted"),
        (
            "verification outcome",
            ("findings", 0, "verification_outcome"),
            "probably",
        ),
        (
            "verdict status",
            ("findings", 0, "verdicts", 0, "status"),
            "successful",
        ),
        (
            "failed verifier in actionable bucket",
            ("findings", 0, "verdicts", 0, "status"),
            "failed",
        ),
        (
            "duplicate verifier seat in actionable bucket",
            ("findings", 0, "verdicts", 0, "reviewer_index"),
            2,
        ),
        (
            "failed verifier in substantively-refuted bucket",
            ("refuted", 0, "verdicts", 0, "status"),
            "failed",
        ),
        (
            "string verdict boolean",
            ("findings", 0, "verdicts", 0, "upholds"),
            "false",
        ),
    ]
    for label, path, value in mutations:
        payload = copy.deepcopy(result.to_dict())
        _set_nested(payload, path, value)
        assert QCResult.from_dict(payload) is None, label

    wrong_majority = copy.deepcopy(result.to_dict())
    for verdict in wrong_majority["refuted"][0]["verdicts"]:
        verdict["upholds"] = True
    assert QCResult.from_dict(wrong_majority) is None


def test_current_schema_pricing_and_aggregate_accounting_are_reconciled() -> None:
    _store, result = _rich_audit_result()
    baseline = result.to_dict()
    assert QCResult.from_dict(copy.deepcopy(baseline)) is not None

    tampered_payloads: list[tuple[str, dict]] = []

    payload = copy.deepcopy(baseline)
    payload["estimated_cost_usd"] += 0.01
    tampered_payloads.append(("aggregate estimate", payload))

    payload = copy.deepcopy(baseline)
    payload["lens_statuses"][0]["estimated_cost_usd"] += 0.01
    tampered_payloads.append(("lens estimate", payload))

    payload = copy.deepcopy(baseline)
    payload["findings"][0]["verdicts"][0]["estimated_cost_usd"] += 0.01
    tampered_payloads.append(("verdict estimate", payload))

    payload = copy.deepcopy(baseline)
    payload["usage_totals"]["tampered_counter"] = 1
    tampered_payloads.append(("aggregate usage", payload))

    payload = copy.deepcopy(baseline)
    payload["api_request_count"] += 1
    tampered_payloads.append(("aggregate API requests", payload))

    payload = copy.deepcopy(baseline)
    payload["model_response_count"] += 1
    tampered_payloads.append(("aggregate model responses", payload))

    priced = copy.deepcopy(baseline)
    priced["lens_statuses"][0]["usage_totals"]["input_tokens"] = (
        priced["lens_statuses"][0]["usage_totals"].get("input_tokens", 0)
        + 100_000
    )
    priced["usage_totals"]["input_tokens"] = (
        priced["usage_totals"].get("input_tokens", 0) + 100_000
    )
    rates = priced["cost_basis"]["rates_per_token"]

    def estimate(usage: dict[str, int]) -> float:
        return round(
            usage.get("input_tokens", 0) * rates["input"]
            + usage.get("output_tokens", 0) * rates["output"]
            + usage.get("cache_read_input_tokens", 0) * rates["cache_read"]
            + usage.get("cache_creation_input_tokens", 0)
            * rates["cache_write"]
            + usage.get("web_search_requests", 0)
            * priced["cost_basis"]["web_search_per_request"],
            6,
        )

    priced["lens_statuses"][0]["estimated_cost_usd"] = estimate(
        priced["lens_statuses"][0]["usage_totals"]
    )
    priced["estimated_cost_usd"] = estimate(priced["usage_totals"])
    assert QCResult.from_dict(copy.deepcopy(priced)) is not None
    priced["cost_basis"]["rates_per_token"]["input"] *= 2
    tampered_payloads.append(("pricing rate", priced))

    payload = copy.deepcopy(baseline)
    payload["cost_basis"]["requested_model"] = "different-model"
    tampered_payloads.append(("requested pricing model", payload))

    payload = copy.deepcopy(baseline)
    payload["cost_basis"]["used_fallback_rate"] = not payload["cost_basis"][
        "used_fallback_rate"
    ]
    tampered_payloads.append(("fallback-rate label", payload))

    payload = copy.deepcopy(baseline)
    payload["cost_basis"] = {}
    tampered_payloads.append(("missing pricing basis", payload))

    for label, payload in tampered_payloads:
        assert QCResult.from_dict(payload) is None, label


# ---------------------------------------------------------------------------
# Per-TTL cache-write accounting (Chunk 4.1)
# ---------------------------------------------------------------------------
#
# Final QC's verifier seats carry a one-hour cache breakpoint (v1.8.0), so
# these records are the first in the app to hold a one-hour subtotal. The
# subtotal lives INSIDE cache_creation_input_tokens; a report that charged
# it at the five-minute rate would under-report a real invoice, and one
# that added it on top would over-report.


def _rates_estimate(basis: dict, usage_totals: dict) -> float:
    """Recompute an estimate the way a persisted report claims it was made."""
    rates = basis["rates_per_token"]
    one_hour = usage_totals.get("cache_creation_1h_input_tokens", 0)
    five_minute = usage_totals.get("cache_creation_input_tokens", 0) - one_hour
    return round(
        usage_totals.get("input_tokens", 0) * rates["input"]
        + usage_totals.get("output_tokens", 0) * rates["output"]
        + usage_totals.get("cache_read_input_tokens", 0) * rates["cache_read"]
        + five_minute * rates["cache_write"]
        + one_hour * rates.get("cache_write_1h", rates["cache_write"])
        + usage_totals.get("web_search_requests", 0)
        * basis["web_search_per_request"]
        + usage_totals.get("web_fetch_requests", 0)
        * basis["web_fetch_per_request"],
        6,
    )


def test_a_verifier_seats_one_hour_cache_subtotal_is_captured_and_priced() -> None:
    store = _section()
    scripts = _scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                summary="One candidate recorded.",
                findings=[
                    _finding(
                        "Priced candidate",
                        "A candidate whose panel wrote a one-hour cache entry.",
                        severity="medium",
                    )
                ],
            )
        ]
    )
    # Both standard seats write a one-hour entry: 900 cache-creation tokens
    # each, 400 of them at the one-hour TTL.
    scripts["Priced candidate"] = [
        qc_verdict_response(
            True,
            note="Upheld.",
            tokens={"cache_write": 900, "cache_write_1h": 400},
        ),
        qc_verdict_response(
            True,
            note="Upheld.",
            tokens={"cache_write": 900, "cache_write_1h": 400},
        ),
    ]
    result = _run(SequencedFakeClient(scripts), store)

    seats = [v for f in result.findings for v in f.verdicts]
    assert seats, "fixture must produce verifier seats"
    for seat in seats:
        assert seat.usage_totals["cache_creation_input_tokens"] == 900
        assert seat.usage_totals["cache_creation_1h_input_tokens"] == 400
        rates = result.cost_basis["rates_per_token"]
        # 500 five-minute + 400 one-hour, each at its own rate — not 900 at
        # either one.
        assert seat.estimated_cost_usd == round(
            500 * rates["cache_write"] + 400 * rates["cache_write_1h"], 6
        )
        assert seat.estimated_cost_usd != round(900 * rates["cache_write"], 6)

    assert result.usage_totals["cache_creation_1h_input_tokens"] == 800
    # Per-record and aggregate accounting both reconcile under the saved
    # five-rate basis, so the record survives a save/load round trip.
    payload = result.to_dict()
    assert payload["cost_basis"]["rates_per_token"]["cache_write_1h"] > 0
    assert "cache_write_treatment" in payload["cost_basis"]
    restored = QCResult.from_dict(copy.deepcopy(payload))
    assert restored is not None
    assert restored.usage_totals["cache_creation_1h_input_tokens"] == 800
    assert restored.cost_basis["rates_per_token"]["cache_write_1h"] == (
        payload["cost_basis"]["rates_per_token"]["cache_write_1h"]
    )


def test_a_legacy_cost_basis_loads_and_reproduces_its_own_estimate() -> None:
    # A report saved before per-TTL pricing carries a four-rate basis and no
    # one-hour subtotal. It must stay readable — and its arithmetic must
    # still check out, which means the whole cache-creation total falls
    # through to the five-minute rate exactly as it did when it was written.
    _store, result = _rich_audit_result()
    payload = copy.deepcopy(result.to_dict())
    payload["cost_basis"]["rates_per_token"].pop("cache_write_1h")
    payload["cost_basis"].pop("cache_write_treatment")

    restored = QCResult.from_dict(copy.deepcopy(payload))
    assert restored is not None
    # Preserved, not rewritten: a cost basis is an immutable claim about how
    # this run was priced, so loading it must not mint a rate it never used.
    assert "cache_write_1h" not in restored.cost_basis["rates_per_token"]
    assert "cache_write_treatment" not in restored.cost_basis
    assert restored.to_dict()["cost_basis"] == payload["cost_basis"]
    assert restored.estimated_cost_usd == _rates_estimate(
        payload["cost_basis"], payload["usage_totals"]
    )


def test_a_legacy_basis_paired_with_a_one_hour_subtotal_prices_it_conservatively(
) -> None:
    # A four-rate basis has no one-hour rate to apply, so the subtotal falls
    # back to cache_write. The record still reconciles — it is simply priced
    # at the rate the run actually recorded.
    _store, result = _rich_audit_result()
    payload = copy.deepcopy(result.to_dict())
    payload["cost_basis"]["rates_per_token"].pop("cache_write_1h")
    payload["cost_basis"].pop("cache_write_treatment")
    lens = payload["lens_statuses"][0]
    for record in (lens["usage_totals"], payload["usage_totals"]):
        record["cache_creation_input_tokens"] = (
            record.get("cache_creation_input_tokens", 0) + 1_000
        )
        record["cache_creation_1h_input_tokens"] = 400
    lens["estimated_cost_usd"] = _rates_estimate(
        payload["cost_basis"], lens["usage_totals"]
    )
    payload["estimated_cost_usd"] = _rates_estimate(
        payload["cost_basis"], payload["usage_totals"]
    )
    assert QCResult.from_dict(copy.deepcopy(payload)) is not None


def test_a_persisted_one_hour_subtotal_above_its_total_is_rejected() -> None:
    # Live provider values are clamped; a persisted report is not. A record
    # claiming more one-hour tokens than it created any cache with cannot
    # describe a real run, so it fails accounting rather than being quietly
    # repaired into something plausible.
    _store, result = _rich_audit_result()
    baseline = result.to_dict()
    assert QCResult.from_dict(copy.deepcopy(baseline)) is not None

    for target in ("lens_statuses", "aggregate"):
        payload = copy.deepcopy(baseline)
        lens = payload["lens_statuses"][0]
        for record in (lens["usage_totals"], payload["usage_totals"]):
            record["cache_creation_input_tokens"] = (
                record.get("cache_creation_input_tokens", 0) + 100
            )
            record["cache_creation_1h_input_tokens"] = 100
        if target == "lens_statuses":
            lens["usage_totals"]["cache_creation_1h_input_tokens"] = 5_000
            payload["usage_totals"]["cache_creation_1h_input_tokens"] = 5_000
        else:
            payload["usage_totals"]["cache_creation_input_tokens"] = 1
        lens["estimated_cost_usd"] = _rates_estimate(
            payload["cost_basis"], lens["usage_totals"]
        )
        payload["estimated_cost_usd"] = _rates_estimate(
            payload["cost_basis"], payload["usage_totals"]
        )
        assert QCResult.from_dict(payload) is None, target


def test_a_half_migrated_cost_basis_is_refused() -> None:
    """The explanation and the rate it explains ship together, or not at all.

    Validating the outer fields and the rate map independently accepts two
    incoherent hybrids. The dangerous one is a basis that keeps
    `cache_write_treatment` — prose promising per-TTL pricing — while
    dropping `cache_write_1h`: one-hour tokens then fall back to the
    five-minute rate and the report's own saved text is a lie about its own
    arithmetic. The mirror case is a current five-rate report that lost its
    required explanatory evidence. Neither can arise from a real run, so
    pairing the shapes costs nothing and refuses both.
    """
    _store, result = _rich_audit_result()
    # deepcopy per case: `to_dict` shallow-copies `cost_basis`, so mutating a
    # payload's `rates_per_token` in place would edit the live result and
    # silently turn the next case into a different shape.
    baseline = result.to_dict()

    keeps_prose_drops_rate = copy.deepcopy(baseline)
    keeps_prose_drops_rate["cost_basis"]["rates_per_token"].pop("cache_write_1h")
    assert QCResult.from_dict(keeps_prose_drops_rate) is None

    keeps_rate_drops_prose = copy.deepcopy(baseline)
    keeps_rate_drops_prose["cost_basis"].pop("cache_write_treatment")
    assert QCResult.from_dict(keeps_rate_drops_prose) is None

    # Both COMPLETE shapes still load — this tightened corruption handling,
    # it did not narrow the supported basis versions to one.
    assert QCResult.from_dict(copy.deepcopy(baseline)) is not None
    legacy = copy.deepcopy(baseline)
    legacy["cost_basis"].pop("cache_write_treatment")
    legacy["cost_basis"]["rates_per_token"].pop("cache_write_1h")
    assert QCResult.from_dict(legacy) is not None


def test_a_cost_basis_with_an_unknown_rate_key_is_still_refused() -> None:
    # Accepting two known shapes must not become accepting any shape: an
    # unrecognized rate would be silently ignored by every consumer while
    # the report claimed it had been applied.
    _store, result = _rich_audit_result()
    # One deepcopied baseline per case: `to_dict` shallow-copies `cost_basis`,
    # so mutating a payload's rate map in place leaks into the next case and
    # lets it pass for the wrong reason.
    baseline = result.to_dict()

    payload = copy.deepcopy(baseline)
    payload["cost_basis"]["rates_per_token"]["cache_write_2h"] = 1e-6
    assert QCResult.from_dict(payload) is None

    payload = copy.deepcopy(baseline)
    payload["cost_basis"]["invented_field"] = "nonsense"
    assert QCResult.from_dict(payload) is None

    payload = copy.deepcopy(baseline)
    payload["cost_basis"]["cache_write_treatment"] = "   "
    assert QCResult.from_dict(payload) is None


def test_current_schema_top_level_identity_agrees_with_hashed_input_manifest() -> None:
    _store, result = _rich_audit_result()
    baseline = result.to_dict()
    assert QCResult.from_dict(copy.deepcopy(baseline)) is not None

    mutations = [
        ("protocol_version", ("protocol_version",), "final-qc/other"),
        ("document version", ("document", "version_index"), 99),
        ("document fingerprint", ("document", "fingerprint"), "a" * 64),
        ("model", ("configuration", "model"), "different-model"),
        ("effort", ("configuration", "effort"), "low"),
        ("maximum tokens", ("configuration", "max_tokens"), 4096),
        ("research presence", ("requirements_research", "present"), True),
    ]
    for label, manifest_path, value in mutations:
        payload = copy.deepcopy(baseline)
        _set_nested(payload["input_manifest"], manifest_path, value)
        # Re-hash the tampered manifest so this specifically exercises the
        # duplicated-claim reconciliation rather than the basic hash check.
        payload["input_fingerprint"] = qc_input_fingerprint(
            payload["input_manifest"]
        )
        assert QCResult.from_dict(payload) is None, label

    malformed_manifest_values = [
        ("document", None),
        ("requirements_research", []),
        ("configuration", "not-an-object"),
    ]
    for key, value in malformed_manifest_values:
        payload = copy.deepcopy(baseline)
        payload["input_manifest"][key] = value
        payload["input_fingerprint"] = qc_input_fingerprint(
            payload["input_manifest"]
        )
        assert QCResult.from_dict(payload) is None, key


def test_dismiss_requires_a_nonempty_trimmed_rationale() -> None:
    store, result = _rich_audit_result()
    runner = QCRunner()
    runner.restore(result)
    session = sessions.get_session()
    session.doc = store
    session.qc = runner
    client = TestClient(create_app())
    finding_id = result.findings[0].finding_id

    invalid_payloads = [
        {"finding_id": finding_id},
        {"finding_id": finding_id, "reason": ""},
        {"finding_id": finding_id, "reason": "   \t  "},
    ]
    for body in invalid_payloads:
        response = client.post("/api/qc/dismiss", json=body)
        assert response.status_code in {400, 422}
        assert result.findings[0].status == "open"
        assert result.findings[0].dismiss_reason == ""
        assert not result.findings[0].disposition_events

    response = client.post(
        "/api/qc/dismiss",
        json={
            "finding_id": finding_id,
            "reason": "  Reviewed with the engineer of record.  ",
        },
    )
    assert response.status_code == 200
    finding = result.findings[0]
    assert finding.status == "dismissed"
    assert finding.dismiss_reason == "Reviewed with the engineer of record."
    assert finding.disposition_events[-1].reason == finding.dismiss_reason


def test_normalized_docx_omits_qc_closing_after_latest_failed_rerun(
    monkeypatch,
) -> None:
    store, successful = _rich_audit_result()
    runner = QCRunner()
    runner.restore(successful)
    session = sessions.get_session()
    session.doc = store
    session.qc = runner

    failed = copy.deepcopy(successful)
    failed.execution_status = "failed"
    failed.summary = "The latest rerun failed before review coverage completed."
    failed.findings = []
    failed.refuted = []
    for lens in failed.lens_statuses:
        lens.status = "failed"
        lens.error = "provider unavailable"
        lens.reviewed_checks = []

    def fail_run(*_args, run_id: str, **_kwargs):
        failed.run_id = run_id
        raise QCFanoutError(
            "All 5 QC lenses failed.",
            usage_totals=failed.usage_totals,
            result=failed,
        )

    monkeypatch.setattr("backend.qc.runner.run_final_qc", fail_run)
    assert runner.start(
        section=store.doc,
        profile=None,
        module=DEFAULT_MODULE,
        client=object(),
        model=settings.QC_MODEL,
        max_tokens=settings.QC_MAX_TOKENS,
        version_index=store.index,
    )
    deadline = time.monotonic() + 2
    while time.monotonic() < deadline and runner.status == "running":
        time.sleep(0.01)
    assert runner.status == "failed"
    assert runner.result is successful
    assert runner.latest_attempt_status == "failed"

    response = TestClient(create_app()).get("/api/export/docx?mode=normalized")
    assert response.status_code == 200
    text = _document_text(Document(io.BytesIO(response.content)))
    assert "FINAL QC SUMMARY" not in text
    assert "Upheld source-backed edition issue" not in text


def test_compact_qc_closing_keeps_candidate_outcomes_distinct() -> None:
    store = _section()
    payload = {
        "model": settings.QC_MODEL,
        "finished_at": "2026-07-24T10:05:00+00:00",
        "version_index": store.index,
        "summary": "All candidate outcomes remain traceable.",
        "findings": [],
        "lens_statuses": [{"lens_id": "code_compliance", "status": "completed"}],
        "refuted": [{"finding_id": "qc-refuted"}],
        "disputed": [{"finding_id": "qc-disputed"}],
        "inconclusive": [{"finding_id": "qc-inconclusive"}],
    }

    text = _document_text(Document(io.BytesIO(build_docx(store.doc, qc_result=payload))))
    assert "No surviving finding remains open." in text
    assert "0 open, 0 applied, 0 dismissed" in text
    # Four distinct outcomes, each named — a disputed candidate is neither a
    # refutation nor an infrastructure failure.
    assert (
        "1 substantively refuted, 1 disputed and awaiting human review, "
        "1 infrastructure-inconclusive"
    ) in text
    assert "every finding was applied or dismissed" not in text


def test_the_word_report_gives_disputed_candidates_their_own_appendix() -> None:
    """A disputed candidate is neither a survivor nor a refutation, so it
    cannot be filed under either without misrepresenting the review."""
    store = _section()
    payload = {
        "schema_version": 4,
        "protocol_version": "final-qc/4",
        "model": settings.QC_MODEL,
        "finished_at": "2026-07-30T10:05:00+00:00",
        "version_index": store.index,
        "summary": "One candidate split its panel.",
        "findings": [],
        "lens_statuses": [{"lens_id": "code_compliance", "status": "completed"}],
        "refuted": [],
        "disputed": [
            {
                "finding_id": "qc-disputed",
                "lens_id": "code_compliance",
                "severity": "high",
                "original_severity": "high",
                "element_id": "pt1.a1.p1",
                "title": "Contested edition reference",
                "issue": "Two reviewers disagreed about the adopted edition.",
                "rationale": "The panel completed but did not agree.",
                "verification_outcome": "disputed",
                "dispute_reason": "insufficient_refutation_evidence",
                "verification_panel_size": 3,
                "proposed_ops": [
                    {
                        "action": "replace",
                        "target_id": "pt1.a1.p1",
                        "text": "Contested replacement.",
                    }
                ],
            }
        ],
        "inconclusive": [],
    }

    text = _document_text(
        Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False)))
    )
    assert "Disputed Candidate Register" in text
    assert "Contested edition reference" in text
    # It is named as awaiting a human, never as a resolved outcome.
    assert "DISPUTED - PANEL COMPLETED WITHOUT AGREEMENT" in text
    # And the under-evidenced reason is spelled out rather than left a token.
    assert "no refuting reviewer cited evidence" in text
    # Its operations are shown but explicitly not actionable.
    assert "NOT EVALUATED - CANDIDATE DISPUTED" in text


def test_a_v3_report_keeps_its_original_majority_outcome_on_load() -> None:
    """A historical 2-of-3 uphold stays upheld. v3 was decided under strict
    majority; re-adjudicating it with v4 rules would rewrite a decision
    nobody re-made, on evidence nobody re-examined."""
    payload = {
        "schema_version": 3,
        "protocol_version": "final-qc/3",
        "findings": [
            {
                "finding_id": "qc-legacy",
                "lens_id": "code_compliance",
                "severity": "high",
                "original_severity": "high",
                "element_id": "pt1.a1.p1",
                "title": "Historic majority uphold",
                "issue": "Recorded under final-qc/3.",
                "rationale": "Two of three reviewers upheld it.",
                "verification_outcome": "upheld",
                "verification_panel_size": 3,
                "verification_threshold": 2,
                "verdicts": [
                    {
                        "upholds": True,
                        "reviewer_index": 1,
                        "status": "completed",
                        "ops_adequate": False,
                        "ops_note": "",
                    },
                    {
                        "upholds": True,
                        "reviewer_index": 2,
                        "status": "completed",
                        "ops_adequate": False,
                        "ops_note": "",
                    },
                    {
                        "upholds": False,
                        "reviewer_index": 3,
                        "status": "completed",
                        "ops_adequate": False,
                        "ops_note": "",
                    },
                ],
            }
        ],
        "lens_statuses": [{"lens_id": "code_compliance", "status": "completed"}],
    }
    restored = QCResult.from_dict(copy.deepcopy(payload))
    assert restored is not None
    assert restored.schema_version == 3
    assert len(restored.findings) == 1
    assert restored.findings[0].verification_outcome == "upheld"
    assert restored.findings[0].verification_threshold == 2
    # Its own rule still validates it — the record is self-consistent.
    assert restored.verification_complete() is True
    # But it is not a v4 record, so it cannot satisfy the current audit gate.
    assert restored.schema_version < 4
    assert restored.findings[0].verification_rule == ""


def test_a_v3_report_is_readable_but_no_longer_current_audit_grade() -> None:
    """The version bump is what makes a v3 record historical rather than
    actionable — readable, exportable, and not a current sign-off."""
    store = _section()
    scripts = _scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding("Live finding", "Real defect.", severity="medium")
                ],
            )
        ],
    )
    scripts["Live finding"] = [
        qc_verdict_response(True),
        qc_verdict_response(True),
    ]
    result = _run(SequencedFakeClient(scripts), store)
    assert result.schema_version == QC_REPORT_SCHEMA_VERSION == 4
    assert result.protocol_version == QC_PROTOCOL_VERSION == "final-qc/4"

    # The same record labelled v3 still loads (history is never dropped)…
    payload = copy.deepcopy(result.to_dict())
    payload["schema_version"] = 3
    payload["protocol_version"] = "final-qc/3"
    legacy = QCResult.from_dict(payload)
    assert legacy is not None
    assert len(legacy.findings) == 1
    assert legacy.findings[0].title == "Live finding"

    # …but it fails the current audit-grade predicate the readiness gate uses.
    assert legacy.schema_version != QC_REPORT_SCHEMA_VERSION
    assert legacy.protocol_version != QC_PROTOCOL_VERSION


def test_infrastructure_failed_verification_is_structurally_inconclusive() -> None:
    store = _section()
    scripts = _scripts(
        enforceability_language=[
            qc_findings_response(
                "enforceability_language",
                findings=[
                    _finding(
                        "Infrastructure-inconclusive candidate",
                        "One required verifier seat did not return a verdict.",
                        severity="medium",
                    )
                ],
            )
        ]
    )
    scripts["Infrastructure-inconclusive candidate"] = [
        qc_verdict_response(True, note="The completed seat upheld the candidate."),
        _failed_verdict_response(),
    ]

    result = _run(SequencedFakeClient(scripts), store)
    assert result.execution_status == "partial"
    assert result.findings == []
    assert result.refuted == []
    assert len(result.inconclusive) == 1
    candidate = result.inconclusive[0]
    assert candidate.verification_outcome == "inconclusive"
    assert {verdict.status for verdict in candidate.verdicts} == {
        "completed",
        "failed",
    }
    assert result.verification_complete() is False
    assert result.is_complete() is False

    payload = result.to_dict()
    assert payload["inconclusive"][0]["verification_outcome"] == "inconclusive"
    restored = QCResult.from_dict(payload)
    assert restored is not None
    assert len(restored.inconclusive) == 1
    assert restored.inconclusive[0].verification_outcome == "inconclusive"

    memo = Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False)))
    memo_text = _document_text(memo)
    assert "Appendix A2: Infrastructure-Inconclusive Candidate Register" in memo_text
    assert "Infrastructure-inconclusive candidate" in memo_text
    assert "NOT EVALUATED - CANDIDATE INFRASTRUCTURE-INCONCLUSIVE" in memo_text
    assert "neither surviving findings nor substantive refutations" in memo_text

    # Historical records used ``default_refuted`` for an incomplete panel.
    # Preserve that raw label for traceability, but never present it as the
    # candidate's substantive QC outcome or execution-coverage bucket.
    legacy_payload = copy.deepcopy(payload)
    legacy_candidate = legacy_payload["inconclusive"].pop()
    legacy_candidate["verification_outcome"] = "default_refuted"
    legacy_payload["refuted"].append(legacy_candidate)
    legacy_memo = Document(
        io.BytesIO(build_qc_memo(legacy_payload, store.doc, stale=False))
    )
    legacy_text = _document_text(legacy_memo)
    assert "Verification outcome: infrastructure-inconclusive" in legacy_text
    assert "Persisted legacy outcome label: default_refuted" in legacy_text
    assert (
        f"Infrastructure-Inconclusive finding {candidate.finding_id} has "
        "incomplete verifier status(es): failed."
    ) in legacy_text
    assert (
        f"Refuted finding {candidate.finding_id} has incomplete verifier "
        "status(es): failed."
    ) not in legacy_text


# ---------------------------------------------------------------------------
# Partial research reads the same in every projection (Chunk 3.3)
# ---------------------------------------------------------------------------


def _research_record(**overrides) -> dict:
    """A captured `input_manifest.requirements_research` record."""
    record = {
        "present": True,
        "dimension_count": 4,
        "declared_dimension_count": 4,
        "completed_dimensions": 4,
        "failed_dimensions": 0,
        "completed_dimension_ids": [
            "governing_codes",
            "ahj_requirements",
            "client_standards",
            "site_environment",
        ],
        "failed_dimension_ids": [],
        "incomplete_required_dimension_ids": [],
        "incomplete_optional_dimension_ids": [],
        "dimension_titles": {
            "governing_codes": "Governing construction codes",
            "ahj_requirements": "Authority-having-jurisdiction requirements",
            "client_standards": "Owner / client and insurer standards",
            "site_environment": "Site and environmental factors",
        },
    }
    record.update(overrides)
    return record


def _result_with_research(record: dict | None, *, present: bool = True) -> dict:
    manifest: dict = {"application_version": settings.VERSION}
    if record is not None:
        manifest["requirements_research"] = record
    return {
        "research_profile_present": present,
        "input_manifest": manifest,
        "lens_statuses": [],
        "findings": [],
        "refuted": [],
        "inconclusive": [],
    }


def test_no_profile_reads_as_no_in_both_the_identity_row_and_limitations() -> None:
    identity, limitation = qc_research_coverage(
        _result_with_research(_research_record(present=False), present=False)
    )
    assert identity == "No"
    assert "No requirements-research profile was available" in limitation


def test_a_complete_profile_says_complete_and_discloses_nothing() -> None:
    identity, limitation = qc_research_coverage(
        _result_with_research(_research_record())
    )
    assert identity == "Yes - complete"
    # Nothing to disclose — a complete profile must not manufacture a
    # limitation, or every clean report would carry noise.
    assert limitation == ""


def test_a_partial_required_profile_names_the_missing_coverage() -> None:
    identity, limitation = qc_research_coverage(
        _result_with_research(
            _research_record(
                completed_dimensions=2,
                failed_dimensions=2,
                completed_dimension_ids=["governing_codes", "client_standards"],
                failed_dimension_ids=["ahj_requirements", "site_environment"],
                incomplete_required_dimension_ids=[
                    "ahj_requirements",
                    "site_environment",
                ],
            )
        )
    )
    # Never a reassuring bare "Yes".
    assert identity == "Yes - partial (2 of 4 areas completed)"
    assert "REQUIRED for issue readiness" in limitation
    assert "Authority-having-jurisdiction requirements" in limitation
    assert "Site and environmental factors" in limitation
    assert "absent from this review, not verified-empty" in limitation


def test_a_partial_optional_profile_quotes_the_declared_rationale() -> None:
    identity, limitation = qc_research_coverage(
        _result_with_research(
            _research_record(
                dimension_count=5,
                declared_dimension_count=5,
                completed_dimensions=4,
                failed_dimensions=1,
                failed_dimension_ids=["seismic_practice"],
                incomplete_optional_dimension_ids=["seismic_practice"],
                dimension_titles={"seismic_practice": "Seismic bracing practice"},
                optional_rationales={
                    "seismic_practice": "rarely governs in this jurisdiction"
                },
            )
        )
    )
    assert identity == "Yes - partial (4 of 5 areas completed)"
    assert "declared optional: rarely governs in this jurisdiction" in limitation
    assert "REQUIRED" not in limitation


def test_a_count_only_legacy_record_says_it_cannot_name_the_gap() -> None:
    identity, limitation = qc_research_coverage(
        _result_with_research(
            {
                "present": True,
                "dimension_count": 4,
                "completed_dimensions": 2,
                "failed_dimensions": 2,
            }
        )
    )
    assert identity == "Yes - partial (2 of 4 areas completed)"
    assert "did not record which areas" in limitation


def test_a_record_less_legacy_report_admits_coverage_is_unknown() -> None:
    identity, limitation = qc_research_coverage(_result_with_research(None))
    assert identity == "Yes - coverage not recorded"
    assert "did not record which research areas completed" in limitation

    absent_identity, absent_limitation = qc_research_coverage(
        _result_with_research(None, present=False)
    )
    assert absent_identity == "No"
    assert "No requirements-research profile was available" in absent_limitation


def test_the_word_report_states_partial_coverage_in_both_places() -> None:
    """End to end through the real writer: identity row AND Limitations."""
    store, result = _rich_audit_result()
    payload = result.to_dict()
    manifest = dict(payload.get("input_manifest") or {})
    manifest["requirements_research"] = _research_record(
        completed_dimensions=3,
        failed_dimensions=1,
        completed_dimension_ids=[
            "governing_codes",
            "client_standards",
            "site_environment",
        ],
        failed_dimension_ids=["ahj_requirements"],
        incomplete_required_dimension_ids=["ahj_requirements"],
    )
    payload["input_manifest"] = manifest
    payload["research_profile_present"] = True

    text = _document_text(Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False))))
    assert "Yes - partial (3 of 4 areas completed)" in text
    assert "Missing coverage REQUIRED for issue readiness" in text
    assert "Authority-having-jurisdiction requirements" in text
    # The bare reassurance is gone. Label and value share one paragraph, so
    # the colon form is what makes this assertion bite.
    assert "Research profile present: Yes\n" not in text


def test_the_exported_report_keeps_its_captured_facts_when_research_moves_on() -> None:
    """A report is an audit of the run's INPUT snapshot.

    Completing the missing dimension in the live session afterwards must not
    retroactively make an already-exported report claim full coverage.
    """
    store, result = _rich_audit_result()
    payload = result.to_dict()
    manifest = dict(payload.get("input_manifest") or {})
    manifest["requirements_research"] = _research_record(
        completed_dimensions=3,
        failed_dimensions=1,
        completed_dimension_ids=[
            "governing_codes",
            "client_standards",
            "site_environment",
        ],
        failed_dimension_ids=["ahj_requirements"],
        incomplete_required_dimension_ids=["ahj_requirements"],
    )
    payload["input_manifest"] = manifest

    # Live research now completes everything…
    session = sessions.get_workspace().session
    from backend.research.engine import DimensionStatus, RequirementsProfile

    session.research.status = "complete"
    session.research.profile_result = RequirementsProfile(
        items=[],
        dimension_statuses=[
            DimensionStatus(dimension_id=d.dimension_id, status="completed", title=d.title)
            for d in session.module.research_dimensions
        ],
        research_date="2026-07-30",
    )

    # …and the captured record still governs the report.
    text = _document_text(Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False))))
    assert "Yes - partial (3 of 4 areas completed)" in text
    assert "Yes - complete" not in text


def test_the_word_readiness_table_carries_the_blocked_research_detail() -> None:
    """Step 7 is a VERIFICATION, not an implementation.

    `docx_export` consumes the readiness checks out of the export state rather
    than re-deriving them, so Chunk 3.2's truthful detail reaches the Word
    report for free — and `ready` is already False when a required area is
    missing, because `research_complete` is a non-advisory check. This pins
    both, so a future refactor that re-derives readiness in the exporter has
    to disagree with a test.
    """
    store, result = _rich_audit_result()
    payload = result.to_dict()
    payload["export_current_state"] = {
        "readiness": {
            "ready": False,
            "checks": [
                {
                    "id": "research_complete",
                    "ok": False,
                    "detail": (
                        "Required research coverage is missing: "
                        "Authority-having-jurisdiction requirements (3 of 4 "
                        "dimensions completed). Press Research again to retry."
                    ),
                    "advisory": False,
                }
            ],
        }
    }

    text = _document_text(
        Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False)))
    )
    # Booleans render as Yes/No in this report, not True/False.
    assert "Issue readiness at export: No" in text
    assert "BLOCK - research_complete" in text
    assert "Required research coverage is missing" in text
    assert "Authority-having-jurisdiction requirements" in text
    # The report cannot claim issue readiness while that check blocks.
    assert "Issue readiness at export: Yes" not in text


def test_a_retired_dimension_does_not_make_declared_coverage_partial() -> None:
    """Raised in review on PR #98, and a real self-contradiction.

    A persisted profile can keep a failed status for a dimension the current
    module no longer declares (`research_coverage` supports exactly that).
    Counting it as a gap rendered "Yes - partial (4 of 4 areas completed)".
    The verdict is about DECLARED coverage — but the stale failure is still
    disclosed, because this report never drops a recorded failure.
    """
    identity, limitation = qc_research_coverage(
        _result_with_research(
            _research_record(
                dimension_count=5,
                declared_dimension_count=4,
                failed_dimensions=1,
                failed_dimension_ids=["retired_axis"],
                dimension_titles={"retired_axis": "Retired seismic axis"},
            )
        )
    )
    assert identity == "Yes - complete"
    assert "partial" not in identity
    assert "Retired seismic axis" in limitation
    assert "no longer declares as coverage" in limitation


def test_a_retired_dimension_rides_along_without_inflating_a_real_gap() -> None:
    identity, limitation = qc_research_coverage(
        _result_with_research(
            _research_record(
                dimension_count=5,
                declared_dimension_count=4,
                completed_dimensions=3,
                failed_dimensions=2,
                completed_dimension_ids=[
                    "governing_codes",
                    "client_standards",
                    "site_environment",
                ],
                failed_dimension_ids=["ahj_requirements", "retired_axis"],
                incomplete_required_dimension_ids=["ahj_requirements"],
                dimension_titles={
                    "ahj_requirements": "Authority-having-jurisdiction requirements",
                    "retired_axis": "Retired seismic axis",
                },
            )
        )
    )
    # The count reflects declared coverage only: 3 of 4, not 3 of 5.
    assert identity == "Yes - partial (3 of 4 areas completed)"
    assert "REQUIRED for issue readiness" in limitation
    assert "Retired seismic axis" in limitation
    assert "no longer declares as coverage" in limitation


def test_a_legacy_record_without_declared_scope_still_lets_failures_lead() -> None:
    """No declared scope to filter against, so the failed ids must still
    drive the verdict — otherwise a pre-3.2 partial report would read clean."""
    identity, limitation = qc_research_coverage(
        _result_with_research(
            {
                "present": True,
                "dimension_count": 4,
                "completed_dimensions": 3,
                "failed_dimensions": 1,
                "completed_dimension_ids": ["a", "b", "c"],
                "failed_dimension_ids": ["ahj_requirements"],
                "dimension_titles": {"ahj_requirements": "AHJ requirements"},
            }
        )
    )
    assert identity == "Yes - partial (3 of 4 areas completed)"
    assert "AHJ requirements" in limitation


# ---------------------------------------------------------------------------
# Chunk 5.3 — version labels and request-count explanations
# ---------------------------------------------------------------------------


def test_the_version_label_states_both_numbers_and_rejects_malformed_input():
    """One convention: the display number AND the stored index.

    Word used to print `v4` in one place and a bare `3` in two others, so a
    reader comparing the reviewed version to the active one was silently
    comparing a 1-based display number with a 0-based stored index.
    """
    assert qc_version_label(0) == "v1 (stored index 0)"
    assert qc_version_label(3) == "v4 (stored index 3)"
    # `True` is an `int` in Python, and "v2" is not what a caller meant.
    assert qc_version_label(True) == "Not recorded"
    assert qc_version_label(False) == "Not recorded"
    for malformed in (None, -1, "3", 3.0, [3]):
        assert qc_version_label(malformed) == "Not recorded"


def test_one_word_report_labels_every_document_index_identically():
    store = _section()
    result = _run(SequencedFakeClient(_scripts()), store)
    payload = result.to_dict()
    payload["export_current_state"] = {
        "generated_at": "2026-07-30T12:00:00+00:00",
        "document_version": 4,
        "document_fingerprint": "f" * 64,
        "stale": True,
        "runner": {"status": "complete", "error": ""},
        "latest_attempt": {"run_id": result.run_id, "status": "complete"},
        "last_successful_report": {
            "run_id": "qc-earlier",
            "version_index": 2,
        },
    }
    text = _document_text(Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=True))))

    # The reviewed version, the active one (4) and the retained one (2) all
    # read in the same convention — derived, so the assertion cannot drift
    # from whichever version the fixture happens to review.
    assert (
        f"Reviewed document version: {qc_version_label(payload['version_index'])}"
        in text
    )
    assert "Active document version: v5 (stored index 4)" in text
    assert "Retained reviewed document version: v3 (stored index 2)" in text
    # And no raw index is left standing in as a version on its own.
    assert "Active document version: 4" not in text
    assert "Retained reviewed document version: 2" not in text


def test_the_run_totals_say_which_records_they_are_the_sum_of():
    store = _section()
    result = _run(SequencedFakeClient(_scripts()), store)
    payload = result.to_dict()

    population = qc_request_population(payload)
    assert population["reconciles"] is True
    assert population["lens_records"] == len(QC_LENSES)
    assert population["api_requests"] == payload["api_request_count"]
    assert population["model_responses"] == payload["model_response_count"]

    note = qc_request_population_note(population)
    assert note.startswith("Sum of ")
    assert f"{len(QC_LENSES)} lens record(s)" in note
    assert "verifier-seat record(s)" in note

    text = _document_text(Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False))))
    assert note in text
    # The relabelled per-record fields, and the note that explains why
    # tokens do not resemble one inference pass.
    assert (
        "Client API requests (streaming calls, including retries and "
        "pause_turn continuations)"
    ) in text
    assert "Final model responses received" in text
    assert QC_REQUEST_METHODOLOGY_NOTE in text


def test_a_v4_consolidated_run_names_the_consolidation_record_in_the_total():
    """The v4 shape: lens + consolidation + verifier records."""
    store = _section()
    result = _run(SequencedFakeClient(_scripts()), store)
    payload = result.to_dict()
    # This fixture's lenses produce candidates at one element, so a real
    # grouping call ran and is part of the population.
    population = qc_request_population(payload)
    assert population["consolidation_records"] == 1
    note = qc_request_population_note(population)
    assert "1 candidate-consolidation record" in note
    assert population["reconciles"] is True


def test_the_historical_five_lens_ninety_five_verifier_shape_reconciles():
    """The plan's worked example: 5 + 95 = 100, explained rather than asserted."""
    payload = {
        "schema_version": 4,
        "lens_statuses": [
            {"api_request_count": 1, "model_response_count": 1}
            for _ in range(5)
        ],
        "findings": [
            {
                "verdicts": [
                    {"api_request_count": 1, "model_response_count": 1}
                    for _ in range(95)
                ]
            }
        ],
        "api_request_count": 100,
        "model_response_count": 100,
    }
    population = qc_request_population(payload)
    assert population["lens_records"] == 5
    assert population["verifier_seats"] == 95
    assert population["consolidation_records"] == 0
    assert population["reconciles"] is True
    note = qc_request_population_note(population)
    assert note == "Sum of 5 lens record(s) + 95 verifier-seat record(s)"


def test_a_disputed_candidates_seats_count_toward_the_population():
    """The plan predates the fourth collection; its seats are still billed."""
    payload = {
        "schema_version": 4,
        "lens_statuses": [{"api_request_count": 1, "model_response_count": 1}],
        "disputed": [
            {
                "verdicts": [
                    {"api_request_count": 1, "model_response_count": 1}
                    for _ in range(3)
                ]
            }
        ],
        "api_request_count": 4,
        "model_response_count": 4,
    }
    population = qc_request_population(payload)
    assert population["verifier_seats"] == 3
    assert population["reconciles"] is True


def test_a_report_whose_parts_do_not_add_up_says_so_instead_of_inventing_one():
    """Better to admit the components are unavailable than to fake a match."""
    payload = {
        # A CURRENT schema, so the mismatch is what fails it — not the
        # legacy gate, which has its own test.
        "schema_version": 4,
        "lens_statuses": [{"api_request_count": 1, "model_response_count": 1}],
        "api_request_count": 100,
        "model_response_count": 100,
    }
    population = qc_request_population(payload)
    assert population["reconciles"] is False
    assert (
        qc_request_population_note(population)
        == "Recorded total; component population unavailable"
    )


def test_a_legacy_report_with_no_records_at_all_does_not_claim_a_sum():
    population = qc_request_population(
        {"schema_version": 4, "api_request_count": 0, "model_response_count": 0}
    )
    assert population["reconciles"] is False
    assert "component population unavailable" in qc_request_population_note(population)


def test_the_word_report_says_grounded_means_retrieval_not_truth():
    store = _section()
    result = _run(SequencedFakeClient(_scripts()), store)
    text = _document_text(
        Document(io.BytesIO(build_qc_memo(result.to_dict(), store.doc, stale=False)))
    )
    assert QC_GROUNDING_METHODOLOGY_NOTE in text
    assert "retrieval confirmation, not truth verification" in text


def test_a_legacy_schema_can_never_claim_a_reconciled_sum():
    """Regression, PR #105 review (Codex P2).

    The persisted-model loaders normalize an absent counter to 0 and a
    serialized record always carries the key afterwards — so on a schema-1
    report every part and the total are 0, the equality holds vacuously,
    and the breakdown was reported as substantiated. Beside a Value column
    that says "Not recorded" for exactly those counters, which made the
    report contradict itself in adjacent cells.
    """
    legacy = {
        "schema_version": 1,
        "lens_statuses": [
            {"api_request_count": 0, "model_response_count": 0}
            for _ in range(5)
        ],
        "findings": [
            {
                "verdicts": [
                    {"api_request_count": 0, "model_response_count": 0}
                    for _ in range(95)
                ]
            }
        ],
        "api_request_count": 0,
        "model_response_count": 0,
    }
    population = qc_request_population(legacy)
    assert population["reconciles"] is False
    assert (
        qc_request_population_note(population)
        == "Recorded total; component population unavailable"
    )
    # The same records under a current schema DO reconcile, so the gate is
    # about recorded provenance and not about the numbers being zero.
    assert qc_request_population({**legacy, "schema_version": 2})["reconciles"] is True


def test_the_legacy_word_report_never_prints_a_sum_beside_not_recorded():
    """The two cells that must not disagree, asserted from the document."""
    store = _section()
    result = _run(SequencedFakeClient(_scripts()), store)
    payload = result.to_dict()
    payload["schema_version"] = 1
    text = _document_text(
        Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False)))
    )
    assert "Recorded total; component population unavailable" in text
    assert "Sum of " not in text


def test_a_malformed_schema_version_fails_closed():
    population = qc_request_population(
        {
            "schema_version": "not a number",
            "lens_statuses": [
                {"api_request_count": 1, "model_response_count": 1}
            ],
            "api_request_count": 1,
            "model_response_count": 1,
        }
    )
    assert population["reconciles"] is False


# ---------------------------------------------------------------------------
# Chunk 5.4 — issue readiness, sign-off consistency, report layering
# ---------------------------------------------------------------------------


def _readiness_by_id(client) -> dict[str, dict]:
    return {
        check["id"]: check
        for check in client.get("/api/readiness").json()["checks"]
    }


def _live_rich_result():
    """A restored rich result on the live session, plus a client."""
    store, result = _rich_audit_result()
    session = sessions.get_session()
    session.doc = store
    return store, result, session, TestClient(create_app())


def test_open_medium_findings_block_issue_readiness_and_name_themselves():
    """The ratified default: the sign-off's meaning wins over the old gate.

    Before this, `qc_audit_complete` gated on open CRITICALS only, so a
    report could say "Issue readiness at export: Yes" beside a sign-off
    reading "OPEN FINDINGS REMAIN" with 25 of them.
    """
    _store, result, session, client = _live_rich_result()
    for finding in result.findings:
        finding.severity = "medium"
        finding.status = "open"
    session.qc.restore(result)

    checks = _readiness_by_id(client)
    assert checks["qc_execution_complete"]["ok"] is True
    assert checks["no_open_qc_findings"]["ok"] is False
    detail = checks["no_open_qc_findings"]["detail"]
    assert "still open" in detail
    assert "medium" in detail
    assert "dismiss it with a reason" in detail
    # The derived alias is exactly the conjunction, never an independent rule.
    assert checks["qc_audit_complete"]["ok"] is False


def test_dispositioning_every_finding_clears_the_findings_gate():
    _store, result, session, client = _live_rich_result()
    for finding in result.findings:
        finding.severity = "medium"
        finding.status = "dismissed"
        finding.dismiss_reason = "Accepted the risk."
    session.qc.restore(result)

    checks = _readiness_by_id(client)
    assert checks["no_open_qc_findings"]["ok"] is True
    assert "applied or dismissed" in checks["no_open_qc_findings"]["detail"]
    assert checks["qc_audit_complete"]["ok"] is True


def test_an_undispositioned_dispute_blocks_the_findings_gate():
    _store, result, session, client = _live_rich_result()
    for finding in result.findings:
        finding.status = "dismissed"
        finding.dismiss_reason = "Accepted the risk."
    disputed = replace(result.findings[0])
    disputed.finding_id = "qc-disputed-fixture"
    disputed.status = "open"
    disputed.verification_outcome = "disputed"
    disputed.dispute_reason = "split_panel"
    disputed.original_severity = "medium"
    disputed.severity = "medium"
    disputed.verification_panel_size = 2
    disputed.verification_threshold = 2
    disputed.proposed_ops = []
    disputed.ops_semantic_status = "not_proposed"
    disputed.ops_semantic_reason = "No proposed operations were supplied."
    disputed.ops_valid = False
    # One uphold, one refute: the split IS what makes it disputed, and the
    # runner refuses a result whose recorded outcome its seats do not imply.
    disputed.verdicts = [
        replace(disputed.verdicts[0], reviewer_index=1, upholds=True,
                ops_adequate=False, revised_severity=""),
        replace(disputed.verdicts[0], reviewer_index=2, upholds=False,
                ops_adequate=False, revised_severity=""),
    ]
    result.disputed = [disputed]
    session.qc.restore(result)

    checks = _readiness_by_id(client)
    assert checks["no_open_qc_findings"]["ok"] is False
    assert "disputed" in checks["no_open_qc_findings"]["detail"].lower()


def test_the_masthead_and_the_signoff_can_never_contradict_each_other():
    """Asserted from the RENDERED document, both verdicts extracted."""
    store, result = _rich_audit_result()
    for finding in result.findings:
        finding.severity = "medium"
        finding.status = "open"
    payload = result.to_dict()
    payload["export_current_state"] = {
        "generated_at": "2026-07-31T12:00:00+00:00",
        "document_version": payload["version_index"],
        "document_fingerprint": payload["version_fingerprint"],
        "stale": False,
        "runner": {"status": "complete", "error": ""},
        "latest_attempt": {"run_id": payload["run_id"], "status": "complete"},
        # A pre-5.4 payload: it really does claim readiness beside open
        # findings, which is the contradiction this must make unrenderable.
        "readiness": {"ready": True, "checks": []},
    }
    text = _document_text(
        Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False)))
    )

    assert "REVIEW REQUIRED - OPEN FINDINGS REMAIN" in text
    # The masthead cannot say Yes while the sign-off says findings remain.
    assert "Issue readiness at export: Yes" not in text
    assert "Issue readiness at export: No" in text
    assert "READINESS RECORDED BEFORE THE CURRENT GATE" in text


def test_a_clean_report_says_yes_in_both_places():
    store, result = _rich_audit_result()
    for finding in result.findings:
        finding.status = "dismissed"
        finding.dismiss_reason = "Accepted the risk."
    payload = result.to_dict()
    payload["export_current_state"] = {
        "generated_at": "2026-07-31T12:00:00+00:00",
        "document_version": payload["version_index"],
        "document_fingerprint": payload["version_fingerprint"],
        "stale": False,
        "runner": {"status": "complete", "error": ""},
        "latest_attempt": {"run_id": payload["run_id"], "status": "complete"},
        "readiness": {"ready": True, "checks": []},
    }
    text = _document_text(
        Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False)))
    )
    assert "OPEN FINDINGS REMAIN" not in text
    assert "Issue readiness at export: Yes" in text


def test_the_executive_layer_names_the_blocking_checks_and_the_open_queue():
    store, result = _rich_audit_result()
    for finding in result.findings:
        finding.severity = "medium"
        finding.status = "open"
    payload = result.to_dict()
    payload["export_current_state"] = {
        "generated_at": "2026-07-31T12:00:00+00:00",
        "document_version": payload["version_index"],
        "document_fingerprint": payload["version_fingerprint"],
        "stale": False,
        "runner": {"status": "complete", "error": ""},
        "latest_attempt": {"run_id": payload["run_id"], "status": "complete"},
        "readiness": {
            "ready": False,
            "checks": [
                {
                    "id": "no_open_qc_findings",
                    "ok": False,
                    "detail": "2 surviving finding(s) still open (2 medium).",
                    "advisory": False,
                },
                {
                    "id": "profile_complete",
                    "ok": False,
                    "detail": "Project profile is incomplete.",
                    "advisory": True,
                },
            ],
        },
    }
    text = _document_text(
        Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False)))
    )
    # The verdict, the named blocking check, and the queue itself.
    assert "Sign-off verdict:" in text
    assert "Issue readiness: No" in text
    assert (
        "Blocking: no_open_qc_findings — 2 surviving finding(s) still open"
        in text
    )
    # An ADVISORY failing check is not presented as blocking. (The annex's
    # full checklist still lists it, which is why this asserts the
    # executive layer's own bullet form rather than the bare id.)
    assert "Blocking: profile_complete" not in text
    assert "profile_complete" in text
    assert "Open Queue" in text
    assert "Run identity" in text
    assert "Estimated run cost" in text
    for finding in result.findings:
        assert finding.title in text


def test_an_applied_fix_makes_later_exports_disclose_pre_remediation():
    store, result = _rich_audit_result()
    target = result.findings[0]
    target.status = "applied"
    target.disposition_events = [
        QCDispositionEvent(
            action="applied",
            at="2026-07-31T12:00:00+00:00",
            # The apply committed a NEW version; the report reviewed an
            # earlier one, and that difference IS the evidence.
            document_version=result.version_index + 1,
            document_fingerprint="f" * 64,
        )
    ]
    payload = result.to_dict()
    state = qc_pre_remediation_state(payload)
    assert state["pre_remediation"] is True
    assert state["latest_version"] == result.version_index + 1

    text = _document_text(
        Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False)))
    )
    assert "PRE-REMEDIATION - DOCUMENT HAS BEEN MODIFIED SINCE THIS REVIEW" in text
    assert "1 fix(es) from this report were applied" in text


def test_a_report_with_no_applied_fixes_discloses_nothing():
    store, result = _rich_audit_result()
    payload = result.to_dict()
    assert qc_pre_remediation_state(payload)["pre_remediation"] is False
    text = _document_text(
        Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False)))
    )
    assert "PRE-REMEDIATION" not in text


def test_a_legacy_v3_report_keeps_its_historical_rendering():
    store, result = _rich_audit_result()
    payload = result.to_dict()
    payload["schema_version"] = 3
    payload["protocol_version"] = "final-qc/3"
    # It still renders, still carries its own facts, and is not upgraded.
    text = _document_text(
        Document(io.BytesIO(build_qc_memo(payload, store.doc, stale=False)))
    )
    assert "FINAL QC AUDIT REPORT" in text
    assert "final-qc/3" in text
