"""Cross-lens candidate consolidation (Chunk 5.2).

Five lenses reviewing one document routinely raise the SAME defect in
different words, and before this each variant bought its own verifier panel.
These tests pin the two halves of the design: eligible duplicates share one
panel, and NOTHING can lose a candidate — not an invalid partition, not a
failed call, not the feature being switched off.
"""
from __future__ import annotations

import io
import json

import pytest
from docx import Document

from backend import settings
from backend.qc.engine import (
    CONSOLIDATION_OPS_IDENTICAL,
    CONSOLIDATION_OPS_NONE,
    CONSOLIDATION_OPS_ORIGINAL,
    CONSOLIDATION_OPS_RECONCILED,
    CONSOLIDATION_OPS_UNRECONCILED,
    CONSOLIDATION_STATUS_COMPLETE,
    CONSOLIDATION_STATUS_FAILED,
    CONSOLIDATION_STATUS_SKIPPED,
    QCResult,
    _mint_origin_id,
)
from backend.spec_doc.docx_export import build_qc_memo
from tests.fakes import (
    SequencedFakeClient,
    qc_consolidation_response,
    qc_findings_response,
    qc_verdict_response,
)
from tests.test_qc import _finding, _qc_scripts, _run, _section

CONSOLIDATE = "[[QC-CONSOLIDATE:"


def _group(
    members: list[int],
    *,
    title: str | None = None,
    issue: str | None = None,
    rationale: str | None = None,
    why: str | None = None,
    ops: list[dict] | None = None,
) -> dict:
    return {
        "member_indexes": members,
        "canonical_title": title,
        "canonical_issue": issue,
        "canonical_rationale": rationale,
        "grouping_rationale": why,
        "reconciled_ops": ops,
    }


def _scripts_with_grouping(groups: list[dict], **per_lens) -> dict[str, list]:
    """Lens scripts plus one scripted answer for the single grouping call."""
    scripts = _qc_scripts(**per_lens)
    scripts[CONSOLIDATE] = [qc_consolidation_response(groups)]
    return scripts


def _upheld(count: int) -> list:
    return [qc_verdict_response(True, note="Real defect.") for _ in range(count)]


def _replace_op(text: str, target: str = "pt1.a1.p1") -> dict:
    return {"action": "replace", "target_id": target, "text": text}


# ---------------------------------------------------------------------------
# Eligible duplicates share one panel
# ---------------------------------------------------------------------------


def test_three_variants_at_one_element_buy_exactly_one_panel():
    """The headline case: one defect, one panel, three claims all retained."""
    store = _section()
    scripts = _scripts_with_grouping(
        [
            _group(
                [0, 1, 2],
                title="Stale sprinkler edition",
                issue="The provision cites a superseded NFPA 13 edition.",
                rationale="All three lenses reached the same defect.",
                why="One edition correction disposes of every variant.",
            )
        ],
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Stale edition cited",
                        "NFPA 13-2019 is superseded.",
                        severity="high",
                        source_urls=["https://nfpa.org/13"],
                    )
                ],
                searched_urls=["https://nfpa.org/13"],
            )
        ],
        completeness=[
            qc_findings_response(
                "completeness",
                findings=[
                    _finding(
                        "Edition not current",
                        "The cited edition is not the one in effect.",
                        severity="medium",
                    )
                ],
            )
        ],
        provenance_hygiene=[
            qc_findings_response(
                "provenance_hygiene",
                findings=[
                    _finding(
                        "Assumed edition is stale",
                        "An assumed block carries an out-of-date edition.",
                        severity="medium",
                    )
                ],
            )
        ],
    )
    # ONE panel of three seats (max original severity is high), not three
    # panels: seven verdicts scripted would over-supply, three exactly fits.
    scripts["[[QC-VERIFY:"] = _upheld(3)
    result = _run(SequencedFakeClient(scripts), store)

    assert len(result.findings) == 1
    finding = result.findings[0]
    assert finding.title == "Stale sprinkler edition"
    assert len(finding.verdicts) == 3
    assert len(finding.candidate_origins) == 3

    consolidation = result.consolidation
    assert consolidation is not None
    assert consolidation.status == CONSOLIDATION_STATUS_COMPLETE
    assert consolidation.raw_candidate_count() == 3
    assert consolidation.grouped_candidate_count() == 1
    assert consolidation.panels_avoided() == 2

    # Every original claim survives verbatim, with its own lens and wording.
    origins = result.origins_for(finding)
    assert [origin.lens_id for origin in origins] == [
        "code_compliance",
        "completeness",
        "provenance_hygiene",
    ]
    assert [origin.title for origin in origins] == [
        "Stale edition cited",
        "Edition not current",
        "Assumed edition is stale",
    ]
    # And the union of their evidence, without inventing any.
    assert finding.source_urls == ["https://nfpa.org/13"]


def test_every_origin_claim_reaches_the_serialized_and_rendered_report():
    store = _section()
    scripts = _scripts_with_grouping(
        [
            _group(
                [0, 1],
                title="Missing service access",
                issue="No access requirement is stated for the valve assembly.",
                why="One inserted access requirement resolves both.",
            )
        ],
        coordination_consistency=[
            qc_findings_response(
                "coordination_consistency",
                findings=[
                    _finding(
                        "Access omitted from execution",
                        "PART 3 never requires servicing clearance.",
                        source_urls=["https://example.org/access"],
                    )
                ],
                searched_urls=["https://example.org/access"],
            )
        ],
        completeness=[
            qc_findings_response(
                "completeness",
                findings=[
                    _finding(
                        "No clearance provision",
                        "The section omits a maintenance clearance requirement.",
                    )
                ],
            )
        ],
    )
    scripts["[[QC-VERIFY:"] = _upheld(3)
    result = _run(SequencedFakeClient(scripts), store)

    payload = result.to_dict()
    serialized = json.dumps(payload)
    for claim in ("Access omitted from execution", "No clearance provision"):
        assert claim in serialized

    memo = build_qc_memo(payload, store.doc, stale=False)
    document = Document(io.BytesIO(memo))
    paragraphs = "\n".join(p.text for p in document.paragraphs)
    rendered = paragraphs + "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "Original Lens Claims" in paragraphs
    for claim in ("Access omitted from execution", "No clearance provision"):
        assert claim in paragraphs
    # The origin's evidence is cited by register number inline and resolved
    # to the full URL exactly once, in the Appendix B evidence table.
    assert "https://example.org/access" in rendered
    assert "https://example.org/access" not in paragraphs
    assert "evidence E-" in paragraphs
    # The consolidation section states what it bought.
    assert "Cross-Lens Candidate Consolidation" in paragraphs
    assert "Verifier panels avoided" in paragraphs


def test_maximum_original_severity_sizes_the_shared_panel():
    """A medium and a critical claim merged face the CRITICAL panel.

    Consolidation must never let the more severe claim inherit the smaller
    panel of the one it was merged with.
    """
    store = _section()
    scripts = _scripts_with_grouping(
        [_group([0, 1], title="Life-safety defect", issue="Merged.", why="One fix.")],
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[_finding("Low-key wording", "Minor.", severity="medium")],
            )
        ],
        completeness=[
            qc_findings_response(
                "completeness",
                findings=[
                    _finding("Missing life-safety rule", "Severe.", severity="critical")
                ],
            )
        ],
    )
    scripts["[[QC-VERIFY:"] = _upheld(settings.QC_VERIFIERS_CRITICAL)
    result = _run(SequencedFakeClient(scripts), store)

    finding = result.findings[0]
    assert finding.original_severity == "critical"
    assert finding.verification_panel_size == settings.QC_VERIFIERS_CRITICAL
    assert finding.verification_threshold == settings.QC_VERIFIERS_CRITICAL


# ---------------------------------------------------------------------------
# Hard compatibility gates grouping before any model sees it
# ---------------------------------------------------------------------------


def test_the_same_title_at_different_elements_is_never_grouped():
    """Different write scopes are never one defect, however alike they read.

    No grouping call is even made — the bucketing is deterministic and runs
    before the model, so the fake would raise if one were attempted with no
    matching bucket.
    """
    store = _section()
    store.begin_turn()
    store.apply_edits(
        [{"action": "add_paragraph", "target_id": "pt1.a1", "text": "Second block."}]
    )
    store.commit_turn()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding("Identical title", "A.", element_id="pt1.a1.p1"),
                    _finding("Identical title", "B.", element_id="pt1.a1.p2"),
                ],
            )
        ],
    )
    scripts["[[QC-VERIFY:"] = _upheld(6)
    result = _run(SequencedFakeClient(scripts), store)

    assert len(result.findings) == 2
    assert result.consolidation.status == CONSOLIDATION_STATUS_SKIPPED
    assert result.consolidation.panels_avoided() == 0
    assert "hard-compatible write scope" in result.consolidation.fallback_reason
    assert result.consolidation.api_request_count == 0


def test_section_level_candidates_need_shared_evidence_to_be_eligible():
    """Two section-level claims with no overlapping source stay separate.

    A section-level candidate has no anchor to overlap, so similarity alone
    would be the only thing left — which is exactly what must not be enough.
    """
    store = _section()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Scope gap",
                        "The section scope is incomplete.",
                        element_id=None,
                        source_urls=["https://example.org/a"],
                    )
                ],
                searched_urls=["https://example.org/a"],
            )
        ],
        completeness=[
            qc_findings_response(
                "completeness",
                findings=[
                    _finding(
                        "Scope gap",
                        "The section scope is incomplete.",
                        element_id=None,
                        source_urls=["https://example.org/b"],
                    )
                ],
                searched_urls=["https://example.org/b"],
            )
        ],
    )
    scripts["[[QC-VERIFY:"] = _upheld(6)
    result = _run(SequencedFakeClient(scripts), store)

    assert len(result.findings) == 2
    assert result.consolidation.api_request_count == 0
    assert result.consolidation.panels_avoided() == 0


def test_section_level_candidates_sharing_a_source_become_eligible():
    store = _section()
    scripts = _scripts_with_grouping(
        [_group([0, 1], title="Scope gap", issue="Merged scope gap.", why="One fix.")],
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding(
                        "Scope gap A",
                        "Incomplete scope.",
                        element_id=None,
                        source_urls=["https://example.org/shared"],
                    )
                ],
                searched_urls=["https://example.org/shared"],
            )
        ],
        completeness=[
            qc_findings_response(
                "completeness",
                findings=[
                    _finding(
                        "Scope gap B",
                        "Incomplete scope.",
                        element_id=None,
                        source_urls=["https://example.org/shared"],
                    )
                ],
                searched_urls=["https://example.org/shared"],
            )
        ],
    )
    scripts["[[QC-VERIFY:"] = _upheld(3)
    result = _run(SequencedFakeClient(scripts), store)

    assert len(result.findings) == 1
    assert result.consolidation.panels_avoided() == 1


# ---------------------------------------------------------------------------
# Operation reconciliation — never apply two members' fixes for one defect
# ---------------------------------------------------------------------------


def _two_ops_at_one_element(first: str, second: str) -> dict:
    return _scripts_with_grouping(
        [
            _group(
                [0, 1],
                title="Missing access requirement",
                issue="No servicing access is required.",
                why="One inserted requirement resolves both.",
            )
        ],
        coordination_consistency=[
            qc_findings_response(
                "coordination_consistency",
                findings=[
                    _finding(
                        "Access omitted (execution)",
                        "PART 3 omits it.",
                        ops=[_replace_op(first)],
                    )
                ],
            )
        ],
        completeness=[
            qc_findings_response(
                "completeness",
                findings=[
                    _finding(
                        "Access omitted (scope)",
                        "Scope omits it.",
                        ops=[_replace_op(second)],
                    )
                ],
            )
        ],
    )


def test_identical_member_operations_pass_through_unchanged():
    store = _section()
    text = "Provide clear service access to each valve assembly."
    scripts = _two_ops_at_one_element(text, text)
    scripts["[[QC-VERIFY:"] = [
        qc_verdict_response(True, note="Real.", ops_adequate=True) for _ in range(3)
    ]
    result = _run(SequencedFakeClient(scripts), store)

    finding = result.findings[0]
    assert finding.ops_source == CONSOLIDATION_OPS_IDENTICAL
    assert finding.proposed_ops == [_replace_op(text)]
    assert finding.ops_valid is True


def test_the_sf009_sf023_shape_yields_one_reconciled_remediation():
    """The reviewed run's clearest duplicate pair.

    Same element, same missing requirement, DIFFERENT inserted wording. An
    identical-ops eligibility gate would have kept them apart and re-created
    the duplicate provision after apply; grouping by defect and reconciling
    afterwards produces exactly one write.
    """
    store = _section()
    reconciled = "Provide clear service access to each valve assembly."
    scripts = _two_ops_at_one_element(
        "Provide access for servicing.", "Provide service clearance."
    )
    scripts[CONSOLIDATE] = [
        qc_consolidation_response(
            [
                _group(
                    [0, 1],
                    title="Missing access requirement",
                    issue="No servicing access is required.",
                    why="One inserted requirement resolves both.",
                    ops=[_replace_op(reconciled)],
                )
            ]
        )
    ]
    scripts["[[QC-VERIFY:"] = [
        qc_verdict_response(True, note="Real.", ops_adequate=True) for _ in range(3)
    ]
    result = _run(SequencedFakeClient(scripts), store)

    assert len(result.findings) == 1
    finding = result.findings[0]
    assert finding.ops_source == CONSOLIDATION_OPS_RECONCILED
    # Exactly ONE write, not both members' operations in sequence.
    assert finding.proposed_ops == [_replace_op(reconciled)]
    assert len(finding.verdicts) == 3
    # Both alternatives remain readable on the immutable originals.
    alternatives = [
        origin.proposed_ops for origin in result.origins_for(finding)
    ]
    assert alternatives == [
        [_replace_op("Provide access for servicing.")],
        [_replace_op("Provide service clearance.")],
    ]


def test_an_unreconciled_group_survives_advisory_only_with_alternatives():
    store = _section()
    scripts = _two_ops_at_one_element(
        "Provide access for servicing.", "Provide service clearance."
    )
    scripts["[[QC-VERIFY:"] = [
        qc_verdict_response(True, note="Real.", ops_adequate=True) for _ in range(3)
    ]
    result = _run(SequencedFakeClient(scripts), store)

    finding = result.findings[0]
    assert finding.ops_source == CONSOLIDATION_OPS_UNRECONCILED
    # Advisory: Apply can never enact one member's fix for a defect the
    # others described differently.
    assert finding.proposed_ops == []
    assert finding.ops_valid is False
    assert [origin.proposed_ops for origin in result.origins_for(finding)] == [
        [_replace_op("Provide access for servicing.")],
        [_replace_op("Provide service clearance.")],
    ]


def test_a_reconciliation_reaching_outside_the_members_scope_is_refused():
    """The grouping call groups; it does not become an editing pass."""
    store = _section()
    scripts = _two_ops_at_one_element(
        "Provide access for servicing.", "Provide service clearance."
    )
    scripts[CONSOLIDATE] = [
        qc_consolidation_response(
            [
                _group(
                    [0, 1],
                    title="Missing access requirement",
                    issue="No servicing access is required.",
                    why="One fix.",
                    ops=[_replace_op("Rewrite something else.", target="pt3")],
                )
            ]
        )
    ]
    scripts["[[QC-VERIFY:"] = _upheld(3)
    result = _run(SequencedFakeClient(scripts), store)

    finding = result.findings[0]
    assert finding.ops_source == CONSOLIDATION_OPS_UNRECONCILED
    assert finding.proposed_ops == []


def test_a_group_whose_members_proposed_nothing_records_that_distinctly():
    store = _section()
    scripts = _scripts_with_grouping(
        [_group([0, 1], title="Merged", issue="Merged issue.", why="One fix.")],
        code_compliance=[
            qc_findings_response(
                "code_compliance", findings=[_finding("A", "A issue.")]
            )
        ],
        completeness=[
            qc_findings_response("completeness", findings=[_finding("B", "B issue.")])
        ],
    )
    scripts["[[QC-VERIFY:"] = _upheld(3)
    result = _run(SequencedFakeClient(scripts), store)

    assert result.findings[0].ops_source == CONSOLIDATION_OPS_NONE


# ---------------------------------------------------------------------------
# Every failure path falls back to singletons, losing nothing
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    "groups, why",
    [
        ([{"member_indexes": [0]}], "an index left unaccounted for"),
        ([{"member_indexes": [0]}, {"member_indexes": [5]}], "an unknown index"),
        (
            [{"member_indexes": [0, 1]}, {"member_indexes": [1]}],
            "an index in two groups",
        ),
        ([], "no groups at all"),
        (
            [{"member_indexes": [0, 1], "canonical_title": "", "canonical_issue": ""}],
            "a merge with no canonical wording",
        ),
    ],
)
def test_an_invalid_partition_falls_back_to_singletons_without_losing_a_candidate(
    groups, why
):
    store = _section()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding("First", "One.", severity="medium"),
                    _finding("Second", "Two.", severity="medium"),
                ],
            )
        ],
    )
    scripts[CONSOLIDATE] = [qc_consolidation_response(groups)]
    scripts["[[QC-VERIFY:"] = _upheld(4)
    result = _run(SequencedFakeClient(scripts), store)

    assert len(result.findings) == 2, why
    assert {f.title for f in result.findings} == {"First", "Second"}
    assert result.consolidation.status == CONSOLIDATION_STATUS_FAILED
    assert result.consolidation.error
    assert result.consolidation.raw_candidate_count() == 2
    assert result.consolidation.panels_avoided() == 0
    # The failed call was still billed, and still reconciles.
    assert result.consolidation.api_request_count == 1
    assert all(
        f.ops_source == CONSOLIDATION_OPS_ORIGINAL for f in result.findings
    )


def test_a_grouping_call_that_produces_no_payload_falls_back_and_keeps_its_cost():
    store = _section()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding("First", "One.", severity="medium"),
                    _finding("Second", "Two.", severity="medium"),
                ],
            )
        ],
    )
    # No tool call at all — the parse-failure shape.
    scripts[CONSOLIDATE] = [qc_consolidation_response(None)]
    scripts["[[QC-VERIFY:"] = _upheld(4)
    result = _run(SequencedFakeClient(scripts), store)

    assert len(result.findings) == 2
    assert result.consolidation.status == CONSOLIDATION_STATUS_FAILED
    assert result.consolidation.api_request_count == 1
    assert result.consolidation.model_response_count == 1
    # And the run totals still reconcile to their component records.
    assert result.api_request_count == (
        sum(status.api_request_count for status in result.lens_statuses)
        + result.consolidation.api_request_count
        + sum(
            verdict.api_request_count
            for finding in result.findings
            for verdict in finding.verdicts
        )
    )


def test_a_raised_grouping_call_never_reaches_the_run(monkeypatch):
    store = _section()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding("First", "One.", severity="medium"),
                    _finding("Second", "Two.", severity="medium"),
                ],
            )
        ],
    )
    scripts[CONSOLIDATE] = [RuntimeError("provider exploded")]
    scripts["[[QC-VERIFY:"] = _upheld(4)
    result = _run(SequencedFakeClient(scripts), store)

    assert len(result.findings) == 2
    assert result.consolidation.status == CONSOLIDATION_STATUS_FAILED
    assert "provider exploded" in result.consolidation.error


def test_consolidation_disabled_reviews_every_candidate_on_its_own_panel(
    monkeypatch,
):
    monkeypatch.setattr(settings, "QC_CONSOLIDATION", False)
    store = _section()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding("First", "One.", severity="medium"),
                    _finding("Second", "Two.", severity="medium"),
                ],
            )
        ],
    )
    scripts["[[QC-VERIFY:"] = _upheld(4)
    result = _run(SequencedFakeClient(scripts), store)

    assert len(result.findings) == 2
    assert result.consolidation.status == CONSOLIDATION_STATUS_SKIPPED
    assert result.consolidation.api_request_count == 0
    assert (
        result.input_manifest["configuration"]["consolidation_enabled"] is False
    )


def test_an_oversized_bucket_says_so_instead_of_silently_truncating(monkeypatch):
    monkeypatch.setattr(settings, "QC_CONSOLIDATION_MAX_BUCKET", 2)
    store = _section()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding("A", "One.", severity="medium"),
                    _finding("B", "Two.", severity="medium"),
                    _finding("C", "Three.", severity="medium"),
                ],
            )
        ],
    )
    scripts["[[QC-VERIFY:"] = _upheld(6)
    result = _run(SequencedFakeClient(scripts), store)

    assert len(result.findings) == 3
    assert result.consolidation.api_request_count == 0
    assert "limit 2" in result.consolidation.fallback_reason


# ---------------------------------------------------------------------------
# Determinism, identity, and dismiss memory
# ---------------------------------------------------------------------------


def test_group_membership_and_ids_do_not_depend_on_lens_completion_order():
    """Two runs whose lenses finish in opposite orders agree exactly."""

    def run_with(order: list[str]) -> QCResult:
        store = _section()
        per_lens = {
            lens_id: [
                qc_findings_response(
                    lens_id,
                    findings=[
                        _finding(
                            f"{lens_id} claim",
                            f"{lens_id} issue.",
                            severity="medium",
                        )
                    ],
                )
            ]
            for lens_id in order
        }
        scripts = _scripts_with_grouping(
            [
                _group(
                    [0, 1],
                    title="Merged",
                    issue="Merged issue.",
                    why="One fix.",
                )
            ],
            **per_lens,
        )
        scripts["[[QC-VERIFY:"] = _upheld(2)
        return _run(SequencedFakeClient(scripts), store)

    first = run_with(["code_compliance", "completeness"])
    second = run_with(["completeness", "code_compliance"])

    assert first.findings[0].finding_id == second.findings[0].finding_id
    assert (
        first.findings[0].candidate_origins
        == second.findings[0].candidate_origins
    )
    assert [origin.origin_id for origin in first.consolidation.origins] == [
        origin.origin_id for origin in second.consolidation.origins
    ]


def test_an_origin_id_ignores_position_and_tracks_content():
    """Content-addressed, never ordinal — the property dismiss memory rests on."""
    claim = {
        "element_id": "pt1.a1.p1",
        "title": "Stale edition",
        "issue": "Superseded.",
        "rationale": "Because.",
        "severity": "high",
        "source_urls": ["https://nfpa.org/13"],
        "proposed_ops": [],
    }
    assert _mint_origin_id("code_compliance", claim) == _mint_origin_id(
        "code_compliance", dict(claim)
    )
    moved = {**claim, "issue": "Superseded, materially."}
    assert _mint_origin_id("code_compliance", claim) != _mint_origin_id(
        "code_compliance", moved
    )
    assert _mint_origin_id("code_compliance", claim) != _mint_origin_id(
        "completeness", claim
    )


def test_a_rerun_that_regroups_a_defect_does_not_inherit_its_dismissal():
    """Changed membership is a different thing to have set aside."""
    store = _section()

    def run(groups: list[dict], remembered=None) -> QCResult:
        scripts = _scripts_with_grouping(
            groups,
            code_compliance=[
                qc_findings_response(
                    "code_compliance",
                    findings=[_finding("A", "A issue.", severity="medium")],
                )
            ],
            completeness=[
                qc_findings_response(
                    "completeness",
                    findings=[_finding("B", "B issue.", severity="medium")],
                )
            ],
        )
        scripts["[[QC-VERIFY:"] = _upheld(4)
        return _run(SequencedFakeClient(scripts), store, remembered=remembered)

    merged = run(
        [_group([0, 1], title="Merged", issue="Merged issue.", why="One fix.")]
    )
    assert len(merged.findings) == 1
    dismissed_id = merged.findings[0].finding_id
    remembered = {
        dismissed_id: {
            "reason": "Accepted the risk.",
            "events": [
                {
                    "action": "dismissed",
                    "reason": "Accepted the risk.",
                    "at": "2026-07-30T10:00:00+00:00",
                    "document_version": 1,
                    "document_fingerprint": "abc",
                }
            ],
        }
    }

    # Same defects, but the rerun keeps them apart. Neither singleton may
    # inherit the merged candidate's dismissal.
    split = run([_group([0]), _group([1])], remembered=remembered)
    assert len(split.findings) == 2
    assert all(finding.status == "open" for finding in split.findings)
    assert split.dismissed_ids == []

    # The same grouping DOES carry it, which is what makes the above a
    # membership rule rather than dismissal being broken outright.
    again = run(
        [_group([0, 1], title="Merged", issue="Merged issue.", why="One fix.")],
        remembered=remembered,
    )
    assert again.findings[0].finding_id == dismissed_id
    assert again.findings[0].status == "dismissed"


def test_membership_alone_changes_the_finding_id_when_the_claim_does_not():
    """The precise reason the hash carries origin identities.

    A consolidated claim's top-level wording is CANONICAL rather than any
    one lens's, so a group whose canonical text happens to reproduce one
    member's own words verbatim would hash identically to that member
    standing alone — same lens, same element, same severity, same panel
    shape. Without the membership in the hash the two are one id, and a
    dismissal of "this one narrow claim" would silently carry to a
    candidate that now also stands for a second lens's finding.
    """
    store = _section()

    def run(groups: list[dict]) -> QCResult:
        scripts = _scripts_with_grouping(
            groups,
            code_compliance=[
                qc_findings_response(
                    "code_compliance",
                    findings=[_finding("A", "A issue.", severity="medium")],
                )
            ],
            completeness=[
                qc_findings_response(
                    "completeness",
                    findings=[_finding("B", "B issue.", severity="medium")],
                )
            ],
        )
        scripts["[[QC-VERIFY:"] = _upheld(4)
        return _run(SequencedFakeClient(scripts), store)

    # The group's canonical claim reproduces candidate 0's wording exactly:
    # same title, issue, rationale, severity, element, sources and ops. The
    # only thing that differs from the singleton is who it speaks for.
    merged = run(
        [
            _group(
                [0, 1],
                title="A",
                issue="A issue.",
                rationale="rationale for A",
                why="One fix.",
            )
        ]
    )
    split = run([_group([0]), _group([1])])

    merged_finding = merged.findings[0]
    singleton = next(f for f in split.findings if f.title == "A")
    # Same claim, byte for byte.
    assert (merged_finding.title, merged_finding.issue, merged_finding.rationale) == (
        singleton.title,
        singleton.issue,
        singleton.rationale,
    )
    assert merged_finding.severity == singleton.severity
    assert merged_finding.element_id == singleton.element_id
    assert merged_finding.lens_id == singleton.lens_id
    assert len(merged_finding.verdicts) == len(singleton.verdicts)
    # Different membership, therefore a different thing to dismiss.
    assert merged_finding.candidate_origins != singleton.candidate_origins
    assert merged_finding.finding_id != singleton.finding_id


def _twice(claim: dict) -> dict:
    """One lens emitting the SAME normalized finding twice.

    ``normalize_findings`` deduplicates nothing — the payload is untrusted
    model output — so this is a shape the engine really can be handed.
    """
    return _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance", findings=[dict(claim), dict(claim)]
            )
        ],
    )


def test_a_lens_emitting_one_claim_twice_still_reloads(monkeypatch):
    """Regression, PR #104 review (Codex P2).

    Both records content-addressed to ONE origin id, and a duplicate origin
    id is exactly what the reload partition check refuses — so the run
    finished, serialized, and then had the whole paid report discarded the
    next time the project was opened. Silent loss of a paid report is the
    worst failure this record has.
    """
    store = _section()
    scripts = _twice(_finding("Same claim", "Same issue.", severity="medium"))
    scripts["[[QC-VERIFY:"] = _upheld(4)
    result = _run(SequencedFakeClient(scripts), store)

    origin_ids = [origin.origin_id for origin in result.consolidation.origins]
    assert len(origin_ids) == 2, "neither claim may be dropped"
    assert len(set(origin_ids)) == 2
    # The suffix is a disambiguator on the SAME content hash, not a new one.
    assert origin_ids[1].startswith(origin_ids[0])

    # The whole point: the report survives a save and reload.
    restored = QCResult.from_dict(result.to_dict())
    assert restored is not None
    assert len(restored.consolidation.origins) == 2


def test_the_same_repeat_also_used_to_collide_on_the_finding_id():
    """The pre-existing half of the same bug, closed by the same fix.

    Two byte-identical claims from one lens minted one `finding_id` on
    master too — and `from_dict`'s duplicate-id check discarded the report
    for that alone, before consolidation existed. Unique origins feed the
    finding hash, so those ids now diverge as well.
    """
    store = _section()
    scripts = _twice(_finding("Same claim", "Same issue.", severity="medium"))
    scripts["[[QC-VERIFY:"] = _upheld(4)
    result = _run(SequencedFakeClient(scripts), store)

    finding_ids = [f.finding_id for f in result.findings]
    assert len(finding_ids) == 2
    assert len(set(finding_ids)) == 2


def test_a_repeated_claims_suffix_does_not_shift_with_unrelated_candidates():
    """The ordinal-independence `_mint_origin_id` exists for survives.

    The suffix counts byte-identical EARLIER claims only. An unrelated
    candidate appearing in a rerun must not renumber it, or dismissals
    would churn on defects nothing about which changed.
    """
    store = _section()
    dup = _finding("Same claim", "Same issue.", severity="medium")

    def run(extra: list[dict]) -> QCResult:
        scripts = _qc_scripts(
            code_compliance=[
                qc_findings_response(
                    "code_compliance", findings=[dict(dup), dict(dup)]
                )
            ],
            completeness=[
                qc_findings_response("completeness", findings=extra)
            ],
        )
        scripts["[[QC-VERIFY:"] = _upheld(6)
        return _run(SequencedFakeClient(scripts), store)

    without = run([])
    with_extra = run(
        [
            _finding(
                "Unrelated",
                "Something else entirely.",
                element_id="pt1.a1",
                severity="low",
            )
        ]
    )
    repeated = [
        origin.origin_id
        for origin in without.consolidation.origins
        if origin.lens_id == "code_compliance"
    ]
    still = [
        origin.origin_id
        for origin in with_extra.consolidation.origins
        if origin.lens_id == "code_compliance"
    ]
    assert repeated == still


def test_three_identical_claims_disambiguate_without_running_out_of_room():
    store = _section()
    claim = _finding("Same claim", "Same issue.", severity="medium")
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[dict(claim), dict(claim), dict(claim)],
            )
        ],
    )
    scripts["[[QC-VERIFY:"] = _upheld(6)
    result = _run(SequencedFakeClient(scripts), store)

    origin_ids = [origin.origin_id for origin in result.consolidation.origins]
    assert len(set(origin_ids)) == 3
    assert QCResult.from_dict(result.to_dict()) is not None


def test_a_single_member_group_keeps_its_original_claim_verbatim():
    """The grouping call is not licensed to rewrite one lens's finding."""
    store = _section()
    scripts = _qc_scripts(
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[
                    _finding("Real title", "Real issue.", severity="medium"),
                    _finding("Other", "Other issue.", severity="medium"),
                ],
            )
        ],
    )
    scripts[CONSOLIDATE] = [
        qc_consolidation_response(
            [
                _group([0], title="REWRITTEN", issue="REWRITTEN ISSUE."),
                _group([1], title="ALSO REWRITTEN", issue="ALSO REWRITTEN."),
            ]
        )
    ]
    scripts["[[QC-VERIFY:"] = _upheld(4)
    result = _run(SequencedFakeClient(scripts), store)

    assert {f.title for f in result.findings} == {"Real title", "Other"}
    assert {f.issue for f in result.findings} == {"Real issue.", "Other issue."}


# ---------------------------------------------------------------------------
# Live events reconcile with the persisted record
# ---------------------------------------------------------------------------


def test_live_counts_reconcile_with_the_final_report():
    store = _section()
    scripts = _scripts_with_grouping(
        [_group([0, 1], title="Merged", issue="Merged issue.", why="One fix.")],
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[_finding("A", "A issue.", severity="medium")],
            )
        ],
        completeness=[
            qc_findings_response(
                "completeness",
                findings=[_finding("B", "B issue.", severity="medium")],
            )
        ],
    )
    scripts["[[QC-VERIFY:"] = _upheld(2)
    events: list[dict] = []
    result = _run(SequencedFakeClient(scripts), store, sink=events.append)

    started = [e for e in events if e["type"] == "consolidation_started"]
    complete = [e for e in events if e["type"] == "consolidation_complete"]
    assert len(started) == 1
    assert len(complete) == 1
    assert started[0]["raw_candidate_count"] == 2
    assert started[0]["eligible_candidate_count"] == 2
    assert complete[0]["status"] == CONSOLIDATION_STATUS_COMPLETE
    assert complete[0]["raw_candidate_count"] == result.consolidation.raw_candidate_count()
    assert (
        complete[0]["grouped_candidate_count"]
        == result.consolidation.grouped_candidate_count()
    )
    assert complete[0]["panels_avoided"] == result.consolidation.panels_avoided()

    # Consolidation runs BETWEEN the lenses and the roster, and the roster
    # is built from what it emitted.
    order = [e["type"] for e in events]
    assert order.index("consolidation_complete") < order.index(
        "verification_started"
    )
    roster = next(e for e in events if e["type"] == "verification_started")
    assert roster["total_candidates"] == 1
    assert roster["candidates"][0]["origin_count"] == 2


# ---------------------------------------------------------------------------
# Persistence
# ---------------------------------------------------------------------------


def test_the_consolidation_record_round_trips_and_is_partition_checked():
    store = _section()
    scripts = _scripts_with_grouping(
        [_group([0, 1], title="Merged", issue="Merged issue.", why="One fix.")],
        code_compliance=[
            qc_findings_response(
                "code_compliance",
                findings=[_finding("A", "A issue.", severity="medium")],
            )
        ],
        completeness=[
            qc_findings_response(
                "completeness",
                findings=[_finding("B", "B issue.", severity="medium")],
            )
        ],
    )
    scripts["[[QC-VERIFY:"] = _upheld(2)
    result = _run(SequencedFakeClient(scripts), store)
    payload = result.to_dict()

    restored = QCResult.from_dict(payload)
    assert restored is not None
    assert restored.consolidation is not None
    assert len(restored.consolidation.origins) == 2
    assert restored.origins_for(restored.findings[0])[0].title == "A"

    # An origin dropped from the record is a lost audit claim, not a
    # cosmetic omission: the whole report must refuse to load.
    broken = json.loads(json.dumps(payload))
    broken["consolidation"]["origins"] = broken["consolidation"]["origins"][:1]
    assert QCResult.from_dict(broken) is None

    # So is a candidate whose references no longer match any group.
    broken = json.loads(json.dumps(payload))
    broken["findings"][0]["candidate_origins"] = ["qco-nonexistent"]
    assert QCResult.from_dict(broken) is None

    # And a record claiming the run grouped when the manifest says it did not.
    broken = json.loads(json.dumps(payload))
    broken["consolidation"] = None
    assert QCResult.from_dict(broken) is None
