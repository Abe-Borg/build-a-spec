"""Focused contracts for QC preflight and multi-finding apply safety."""
from __future__ import annotations

import io

from docx import Document
from fastapi.testclient import TestClient

from backend import sessions
from backend.app import create_app
from backend.qc.engine import QCFinding, QCResult
from backend.qc.op_conflicts import plan_qc_operation_batch
from backend.qc.preflight import module_section_compatibility
from backend.spec_doc.model import DocumentStore, SpecSection
from backend.spec_modules.generic import GENERIC
from backend.spec_modules.hyperscale_fire import HYPERSCALE_FIRE


_DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _client() -> TestClient:
    return TestClient(create_app())


def _store() -> DocumentStore:
    store = DocumentStore()
    store.begin_turn()
    store.apply_edits(
        [
            {
                "action": "replace",
                "target_id": "sec",
                "text": "WET-PIPE SPRINKLER SYSTEMS",
                "numbering": "21 13 13",
            },
            {"action": "add_article", "target_id": "pt1", "text": "SUMMARY"},
            {
                "action": "add_paragraph",
                "target_id": "pt1.a1",
                "text": "Provide a complete system.",
                "status": "confirmed",
            },
            {
                "action": "add_paragraph",
                "target_id": "pt1.a1",
                "text": "Coordinate all interfaces.",
                "status": "confirmed",
            },
        ]
    )
    store.commit_turn()
    return store


def _finding(finding_id: str, operations: list[dict]) -> QCFinding:
    return QCFinding(
        finding_id=finding_id,
        lens_id="coordination_consistency",
        severity="high",
        element_id=str(operations[0].get("target_id") or ""),
        title=f"Finding {finding_id}",
        issue="A focused endpoint-test issue.",
        rationale="A focused endpoint-test rationale.",
        proposed_ops=operations,
        ops_semantic_status="approved",
        ops_semantic_reason="Every verifier approved the complete operation set.",
        ops_valid=True,
        verification_outcome="upheld",
    )


def _install_result(findings: list[QCFinding]) -> None:
    runner = sessions.get_session().qc
    runner.result = QCResult(findings=findings)
    runner.status = "complete"


def _bypass_result_contract(monkeypatch) -> None:
    """Keep these endpoint tests focused on post-audit apply behavior."""
    monkeypatch.setattr("backend.app._qc_result_is_audit_complete", lambda _r: True)
    monkeypatch.setattr(
        "backend.app._qc_matches_current_inputs",
        lambda _session, _result, **_kwargs: True,
    )


def _mismatch_master_bytes() -> bytes:
    document = Document()
    for text in (
        "SECTION 23 74 13",
        "PACKAGED, OUTDOOR, CENTRAL-STATION AIR-HANDLING UNITS",
        "PART 1 - GENERAL",
        "1.1 SUMMARY",
        "A. Provide packaged rooftop air-handling units.",
        "PART 2 - PRODUCTS",
        "2.1 UNITS",
        "A. Factory assemble units.",
        "PART 3 - EXECUTION",
        "3.1 INSTALLATION",
        "A. Install units in accordance with manufacturer instructions.",
        "END OF SECTION 23 74 13",
    ):
        document.add_paragraph(text)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def test_compatibility_statuses_use_exact_normalized_section_numbers() -> None:
    section = SpecSection.empty()
    section.number = "21-13-13"
    section.title = "WET-PIPE SPRINKLER SYSTEMS"
    matched = module_section_compatibility(section, HYPERSCALE_FIRE)
    assert matched["status"] == "match"
    assert matched["section_number"] == "21 13 13"
    assert matched["allowed_sections"]

    section.number = "23 74 13"
    assert module_section_compatibility(section, HYPERSCALE_FIRE)["status"] == (
        "mismatch"
    )

    section.number = ""
    assert module_section_compatibility(section, HYPERSCALE_FIRE)["status"] == (
        "unknown"
    )
    assert module_section_compatibility(section, GENERIC)["status"] == (
        "not_applicable"
    )


def test_import_warning_status_and_acknowledged_qc_start(monkeypatch) -> None:
    client = _client()
    reset = client.post(
        "/api/session/reset", json={"module_id": "hyperscale_fire"}
    )
    assert reset.status_code == 200

    imported = client.post(
        "/api/import/master",
        files={
            "file": (
                "mechanical-master.docx",
                _mismatch_master_bytes(),
                _DOCX_MEDIA_TYPE,
            )
        },
    )
    assert imported.status_code == 200, imported.text
    payload = imported.json()
    mismatch_warnings = [
        warning
        for warning in payload["warnings"]
        if "outside the" in warning and "catalog" in warning
    ]
    assert len(mismatch_warnings) == 1
    assert sessions.get_session().import_report["warnings"] == payload["warnings"]

    status = client.get("/api/qc/status").json()
    compatibility = status["module_section_compatibility"]
    assert compatibility["status"] == "mismatch"
    assert compatibility["section_number"] == "23 74 13"
    assert compatibility["message"] == mismatch_warnings[0]

    runner_calls: list[dict] = []
    monkeypatch.setattr("backend.app.get_client", lambda: object())
    monkeypatch.setattr(
        sessions.get_session().qc,
        "start",
        lambda **kwargs: runner_calls.append(kwargs) or True,
    )

    refused = client.post("/api/qc/start")
    assert refused.status_code == 409
    refused_payload = refused.json()
    assert refused_payload["code"] == "module_section_mismatch"
    assert refused_payload["module_section_compatibility"] == compatibility
    assert runner_calls == []

    accepted = client.post(
        "/api/qc/start", json={"acknowledge_scope_mismatch": True}
    )
    assert accepted.status_code == 200, accepted.text
    assert accepted.json() == {"ok": True}
    assert len(runner_calls) == 1


def test_nonidentical_same_target_ops_conflict_without_mutation(monkeypatch) -> None:
    session = sessions.get_session()
    session.doc = _store()
    first = _finding(
        "f-1",
        [
            {
                "action": "replace",
                "target_id": "pt1.a1.p1",
                "text": "Use the first revision.",
                "status": "confirmed",
            }
        ],
    )
    second = _finding(
        "f-2",
        [
            {
                "action": "replace",
                "target_id": "pt1.a1.p1",
                "text": "Use the competing revision.",
                "status": "confirmed",
            }
        ],
    )
    _install_result([first, second])
    _bypass_result_contract(monkeypatch)
    before = session.doc.doc.to_dict()
    versions_before = len(session.doc.versions)

    response = _client().post(
        "/api/qc/apply", json={"finding_ids": ["f-1", "f-2"]}
    )
    assert response.status_code == 409
    body = response.json()
    assert body["code"] == "qc_operation_conflict"
    assert body["finding_ids"] == ["f-1", "f-2"]
    assert body["write_keys"] == ["element:pt1.a1.p1:*"]
    assert body["conflicts"]
    assert session.doc.doc.to_dict() == before
    assert len(session.doc.versions) == versions_before
    assert first.status == second.status == "open"
    assert first.disposition_events == second.disposition_events == []


def test_structural_collection_conflict_is_atomic(monkeypatch) -> None:
    session = sessions.get_session()
    session.doc = _store()
    deletion = _finding(
        "delete",
        [{"action": "delete", "target_id": "pt1.a1"}],
    )
    addition = _finding(
        "add",
        [
            {
                "action": "add_paragraph",
                "target_id": "pt1.a1",
                "text": "Add a competing sibling.",
                "status": "confirmed",
            }
        ],
    )
    _install_result([deletion, addition])
    _bypass_result_contract(monkeypatch)
    before = session.doc.doc.to_dict()

    response = _client().post(
        "/api/qc/apply", json={"finding_ids": ["delete", "add"]}
    )
    assert response.status_code == 409
    assert response.json()["write_keys"] == ["element:pt1.a1:*"]
    assert session.doc.doc.to_dict() == before


def test_positioned_same_parent_additions_remain_conflicts() -> None:
    section = _store().doc
    positioned = plan_qc_operation_batch(
        section,
        [
            (
                "positioned-a",
                [
                    {
                        "action": "add_paragraph",
                        "target_id": "pt1.a1",
                        "text": "Insert at the first position.",
                        "position": 0,
                    }
                ],
            ),
            (
                "positioned-b",
                [
                    {
                        "action": "add_paragraph",
                        "target_id": "pt1.a1",
                        "text": "Compete for the first position.",
                        "position": 0,
                    }
                ],
            ),
        ],
    )
    assert len(positioned.conflicts) == 1
    assert positioned.conflicts[0]["finding_ids"] == [
        "positioned-a",
        "positioned-b",
    ]

    positioned_and_append = plan_qc_operation_batch(
        section,
        [
            (
                "positioned",
                [
                    {
                        "action": "add_article",
                        "target_id": "pt1",
                        "text": "POSITIONED ARTICLE",
                        "position": 0,
                    }
                ],
            ),
            (
                "append",
                [
                    {
                        "action": "add_article",
                        "target_id": "pt1",
                        "text": "APPENDED ARTICLE",
                    }
                ],
            ),
        ],
    )
    assert len(positioned_and_append.conflicts) == 1
    assert positioned_and_append.conflicts[0]["finding_ids"] == [
        "positioned",
        "append",
    ]


def test_generated_id_dependent_additions_remain_atomic(monkeypatch) -> None:
    cases = [
        [
            _finding(
                "dependent-a",
                [
                    {
                        "action": "add_article",
                        "target_id": "pt1",
                        "text": "DEPENDENT ARTICLE A",
                    },
                    {
                        "action": "add_paragraph",
                        "target_id": "pt1.a2",
                        "text": "Child intended for article A.",
                    },
                ],
            ),
            _finding(
                "dependent-b",
                [
                    {
                        "action": "add_article",
                        "target_id": "pt1",
                        "text": "DEPENDENT ARTICLE B",
                    },
                    {
                        "action": "add_paragraph",
                        "target_id": "pt1.a2",
                        "text": "Child intended for article B.",
                    },
                ],
            ),
        ],
        [
            _finding(
                "singleton-first",
                [
                    {
                        "action": "add_article",
                        "target_id": "pt1",
                        "text": "STANDALONE ARTICLE",
                    }
                ],
            ),
            _finding(
                "dependent-second",
                [
                    {
                        "action": "add_article",
                        "target_id": "pt1",
                        "text": "DEPENDENT ARTICLE",
                    },
                    {
                        "action": "add_paragraph",
                        "target_id": "pt1.a2",
                        "text": "Child intended for the dependent article.",
                    },
                ],
            ),
        ],
    ]

    for findings in cases:
        session = sessions.get_session()
        session.doc = _store()
        _install_result(findings)
        _bypass_result_contract(monkeypatch)
        before = session.doc.doc.to_dict()
        versions_before = len(session.doc.versions)
        finding_ids = [finding.finding_id for finding in findings]

        response = _client().post(
            "/api/qc/apply", json={"finding_ids": finding_ids}
        )

        assert response.status_code == 409, response.text
        body = response.json()
        assert body["code"] == "qc_operation_conflict"
        assert body["finding_ids"] == finding_ids
        assert session.doc.doc.to_dict() == before
        assert len(session.doc.versions) == versions_before
        assert all(finding.status == "open" for finding in findings)
        assert all(not finding.disposition_events for finding in findings)


def test_identical_ops_execute_once_and_mark_every_finding(monkeypatch) -> None:
    session = sessions.get_session()
    session.doc = _store()
    operation = {
        "action": "add_paragraph",
        "target_id": "pt1.a1",
        "text": "Submit a single acceptance record.",
        "status": "confirmed",
    }
    first = _finding("same-1", [operation])
    # A different key order and a schema-style null remain the same operation.
    second = _finding(
        "same-2",
        [
            {
                "status": "confirmed",
                "text": "Submit a single acceptance record.",
                "position": None,
                "target_id": "pt1.a1",
                "action": "add_paragraph",
            }
        ],
    )
    _install_result([first, second])
    _bypass_result_contract(monkeypatch)
    before_count = len(session.doc.doc.parts[0].articles[0].paragraphs)
    versions_before = len(session.doc.versions)

    response = _client().post(
        "/api/qc/apply", json={"finding_ids": ["same-1", "same-2"]}
    )
    assert response.status_code == 200, response.text
    assert response.json()["outcomes"] == {
        "same-1": "applied",
        "same-2": "applied",
    }
    assert len(session.doc.doc.parts[0].articles[0].paragraphs) == before_count + 1
    assert len(session.doc.versions) == versions_before + 1
    assert first.status == second.status == "applied"


def test_different_standard_settings_remain_compatible(monkeypatch) -> None:
    session = sessions.get_session()
    session.doc = _store()
    nfpa_13 = _finding(
        "nfpa-13",
        [
            {
                "action": "set_standard_edition",
                "target_id": "sec",
                "standard": "nfpa  13",
                "edition": "2022",
                "basis": "Adopted project basis.",
            }
        ],
    )
    nfpa_24 = _finding(
        "nfpa-24",
        [
            {
                "action": "set_standard_edition",
                "target_id": "sec",
                "standard": "NFPA 24",
                "edition": "2022",
                "basis": "Adopted project basis.",
            }
        ],
    )
    _install_result([nfpa_13, nfpa_24])
    _bypass_result_contract(monkeypatch)

    response = _client().post(
        "/api/qc/apply", json={"finding_ids": ["nfpa-13", "nfpa-24"]}
    )
    assert response.status_code == 200, response.text
    assert session.doc.doc.edition_overrides.keys() == {"NFPA 13", "NFPA 24"}


def test_stale_finding_does_not_block_a_compatible_finding(monkeypatch) -> None:
    session = sessions.get_session()
    session.doc = _store()
    stale = _finding(
        "stale",
        [
            {
                "action": "replace",
                "target_id": "pt9.a9.p9",
                "text": "This target is gone.",
            }
        ],
    )
    compatible = _finding(
        "compatible",
        [
            {
                "action": "set_standard_edition",
                "target_id": "sec",
                "standard": "NFPA 13",
                "edition": "2022",
                "basis": "Adopted project basis.",
            }
        ],
    )
    _install_result([stale, compatible])
    _bypass_result_contract(monkeypatch)

    response = _client().post(
        "/api/qc/apply", json={"finding_ids": ["stale", "compatible"]}
    )
    assert response.status_code == 200, response.text
    assert response.json()["outcomes"] == {
        "stale": "stale",
        "compatible": "applied",
    }
    assert stale.status == "open"
    assert stale.disposition_events[-1].action == "apply_stale"
    assert compatible.status == "applied"
    assert session.doc.doc.edition_overrides["NFPA 13"]["edition"] == "2022"


def test_semantic_rejection_cannot_reach_the_mutation_path(monkeypatch) -> None:
    session = sessions.get_session()
    session.doc = _store()
    finding = _finding(
        "rejected",
        [
            {
                "action": "replace",
                "target_id": "pt1.a1.p1",
                "text": "This must not be written.",
            }
        ],
    )
    finding.ops_semantic_status = "rejected"
    _install_result([finding])
    _bypass_result_contract(monkeypatch)
    before = session.doc.doc.to_dict()

    response = _client().post(
        "/api/qc/apply", json={"finding_ids": ["rejected"]}
    )
    assert response.status_code == 200
    assert response.json()["outcomes"] == {"rejected": "no_ops"}
    assert session.doc.doc.to_dict() == before
    assert finding.status == "open"
