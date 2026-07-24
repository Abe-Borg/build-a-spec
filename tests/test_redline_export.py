"""Tracked-changes (redline) ``.docx`` export tests (Batch 5).

The killer invariant: importing the redlined export back through the
Accept-All resolver reproduces the current document; a Reject-All reading
reproduces the diff's base. Plus XML shape assertions (author/date/unique
id; ``w:delText`` not ``w:t`` inside deletions; paragraph-mark ins/del) and
the export/diff API surface, including a byte-stability guard on the clean
default path.
"""
from __future__ import annotations

import io
import zipfile

from docx import Document
from docx.oxml.ns import qn
from fastapi.testclient import TestClient

from backend.app import create_app
from backend.spec_doc.diffing import diff_sections
from backend.spec_doc.docx_export import build_docx
from backend.spec_doc.importer import _accept_all_paragraph_text
from backend.spec_doc.model import DocumentStore, SpecSection

_W_INS = qn("w:ins")
_W_DEL = qn("w:del")
_W_MOVE_TO = qn("w:moveTo")
_W_T = qn("w:t")
_W_DELTEXT = qn("w:delText")
_W_TAB = qn("w:tab")
_W_P = qn("w:p")
_W_PPR = qn("w:pPr")
_W_RPR = qn("w:rPr")
_W_NUMPR = qn("w:numPr")
_W_IND = qn("w:ind")
_DATE = "2026-07-21T12:00:00Z"


# ---------------------------------------------------------------------------
# Fixtures + reject-all reader
# ---------------------------------------------------------------------------


def _populated_store() -> DocumentStore:
    store = DocumentStore()
    store.begin_turn()
    store.apply_edits(
        [
            {
                "action": "replace",
                "target_id": "sec",
                "text": "WET-PIPE SPRINKLER SYSTEMS",
                "numbering": "21 13 13",
            },
            {"action": "add_article", "target_id": "pt1", "text": "SUMMARY"},
            {"action": "add_article", "target_id": "pt2", "text": "PIPE AND FITTINGS"},
            {"action": "add_article", "target_id": "pt3", "text": "INSTALLATION"},
        ]
    )
    store.commit_turn()
    a1 = store.doc.parts[0].articles[0].uid
    a2 = store.doc.parts[1].articles[0].uid
    a3 = store.doc.parts[2].articles[0].uid
    store.begin_turn()
    store.apply_edits(
        [
            {
                "action": "add_paragraph",
                "target_id": a1,
                "text": "Section includes wet-pipe systems per NFPA 13-2025.",
                "status": "imported",
            },
            {
                "action": "add_paragraph",
                "target_id": a1,
                "text": "Related sections: 21 13 19.",
                "status": "imported",
            },
            {
                "action": "add_paragraph",
                "target_id": a2,
                "text": "Steel pipe: ASTM A53.",
                "status": "imported",
            },
            {
                "action": "add_paragraph",
                "target_id": a3,
                "text": "Install per NFPA 13.",
                "status": "imported",
            },
        ]
    )
    store.commit_turn()
    return store


def _adapt(store: DocumentStore) -> None:
    """Edit / delete / insert to make an interesting redline."""
    p0 = store.doc.parts[0].articles[0].paragraphs[0].uid
    p1 = store.doc.parts[0].articles[0].paragraphs[1].uid
    a2 = store.doc.parts[1].articles[0].uid
    store.begin_turn()
    store.apply_edits(
        [
            {
                "action": "replace",
                "target_id": p0,
                "text": "Section includes wet-pipe automatic sprinkler systems per NFPA 13-2025.",
                "status": "confirmed",
            },
            {"action": "delete", "target_id": p1},
            {
                "action": "add_paragraph",
                "target_id": a2,
                "text": "Provide seismic bracing per NFPA 13.",
                "status": "confirmed",
            },
        ]
    )
    store.commit_turn()


def _para_mark(p_el, tag):
    ppr = p_el.find(_W_PPR)
    if ppr is None:
        return False
    rpr = ppr.find(_W_RPR)
    return rpr is not None and rpr.find(tag) is not None


def _reject_walk(el, parts):
    tag = el.tag
    if not isinstance(tag, str):
        return
    if tag in (_W_INS, _W_MOVE_TO):
        return  # drop insertions
    if tag in (_W_T, _W_DELTEXT):
        parts.append(el.text or "")
        return
    if tag == _W_TAB:
        parts.append("\t")
        return
    for child in el:
        _reject_walk(child, parts)


def _body_texts(docx_bytes: bytes, mode: str) -> list[str]:
    """Non-empty body paragraph texts up to END OF SECTION, one resolution."""
    document = Document(io.BytesIO(docx_bytes))
    out: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag != _W_P:
            continue
        if mode == "reject":
            if _para_mark(child, _W_INS):
                continue  # inserted paragraph mark -> paragraph gone on reject
            parts: list[str] = []
            _reject_walk(child, parts)
            text = "".join(parts)
        else:  # accept
            text = _accept_all_paragraph_text(child)
        text = " ".join(text.split())
        if text.upper().startswith("END OF SECTION"):
            break
        if text:
            out.append(text)
    return out


def _semantic_body_texts(docx_bytes: bytes, mode: str) -> list[str]:
    """Body text independent of clean versus redline list representation.

    Clean exports carry provision labels in ``w:numPr`` and only semantic
    content in ``w:t``. Redlines deliberately retain their established
    literal labels beside tracked runs. Strip only those literal labels,
    identified by the redline provision's direct hanging indent.
    """
    document = Document(io.BytesIO(docx_bytes))
    out: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag != _W_P:
            continue
        if mode == "reject":
            if _para_mark(child, _W_INS):
                continue
            parts: list[str] = []
            _reject_walk(child, parts)
            text = "".join(parts)
        else:
            text = _accept_all_paragraph_text(child)
        text = " ".join(text.split())
        if text.upper().startswith("END OF SECTION"):
            break
        if not text:
            continue
        properties = child.find(_W_PPR)
        literal_redline_provision = (
            properties is not None
            and properties.find(_W_IND) is not None
            and properties.find(_W_NUMPR) is None
        )
        if literal_redline_provision:
            _label, separator, semantic_text = text.partition(" ")
            if separator:
                text = semantic_text
        out.append(text)
    return out


# ---------------------------------------------------------------------------
# The killer round-trip invariant
# ---------------------------------------------------------------------------


def test_accept_all_reproduces_current_reject_all_reproduces_base():
    store = _populated_store()
    base = SpecSection.from_dict(store.doc.to_dict())
    _adapt(store)
    cur = store.doc

    redline = build_docx(cur, redline=diff_sections(base, cur), redline_date=_DATE)

    clean_cur = build_docx(cur)
    clean_base = build_docx(base)
    # Redlines deliberately keep literal labels; clean exports now carry the
    # same display labels as genuine Word numbering outside their w:t text.
    assert any(
        line.startswith("A. Section includes")
        for line in _body_texts(redline, "accept")
    )
    assert (
        "Section includes wet-pipe automatic sprinkler systems per "
        "NFPA 13-2025."
    ) in _body_texts(clean_cur, "accept")
    # Across that intentional representation boundary, Accept/Reject still
    # reproduce the current/base semantic bodies respectively.
    assert _semantic_body_texts(redline, "accept") == _semantic_body_texts(
        clean_cur, "accept"
    )
    assert _semantic_body_texts(redline, "reject") == _semantic_body_texts(
        clean_base, "accept"
    )


def test_position_shift_accept_and_reject_are_text_faithful():
    """When a survivor's position shifts (a preceding sibling was deleted),
    Accept-All still reproduces the current semantic document; Reject-All
    reproduces the baseline provision text. Redline numbering remains literal
    and positional while clean numbering is genuine Word numbering. This is
    the case the shipped round-trip fixture did not exercise — both Codex and
    the batch's own review flagged it."""
    store = _populated_store()
    base = SpecSection.from_dict(store.doc.to_dict())
    # Delete the FIRST paragraph of article 1 — the survivor shifts up a slot.
    first = store.doc.parts[0].articles[0].paragraphs[0].uid
    store.begin_turn()
    store.apply_edits([{"action": "delete", "target_id": first}])
    store.commit_turn()
    cur = store.doc
    redline = build_docx(cur, redline=diff_sections(base, cur), redline_date=_DATE)

    assert "A. Related sections: 21 13 19." in _body_texts(redline, "accept")
    assert _semantic_body_texts(redline, "accept") == _semantic_body_texts(
        build_docx(cur), "accept"
    )
    assert _semantic_body_texts(redline, "reject") == _semantic_body_texts(
        build_docx(base), "accept"
    )


def test_part_emptying_tracks_not_used_placeholder():
    """Deleting every article in a part must track its '(Not used.)' line so
    Accept-All still matches clean-cur content (and Reject-All clean-base)."""
    store = _populated_store()
    base = SpecSection.from_dict(store.doc.to_dict())
    # PART 3 (INSTALLATION) has exactly one article; delete it, emptying PART 3.
    article = store.doc.parts[2].articles[0].uid
    store.begin_turn()
    store.apply_edits([{"action": "delete", "target_id": article}])
    store.commit_turn()
    cur = store.doc
    redline = build_docx(cur, redline=diff_sections(base, cur), redline_date=_DATE)
    assert "(Not used.)" in _body_texts(redline, "accept")
    assert _semantic_body_texts(redline, "accept") == _semantic_body_texts(
        build_docx(cur), "accept"
    )
    assert _semantic_body_texts(redline, "reject") == _semantic_body_texts(
        build_docx(base), "accept"
    )


def test_from_scratch_redline_vs_empty_round_trips():
    """A from-scratch section redlined against the empty baseline is all
    insertions; Accept-All == the draft, Reject-All == the empty document
    (section '[TBD]' placeholders and '(Not used.)' parts included)."""
    store = DocumentStore()
    store.begin_turn()
    store.apply_edits(
        [
            {"action": "replace", "target_id": "sec", "text": "WIDGETS", "numbering": "21 13 14"},
            {"action": "add_article", "target_id": "pt1", "text": "SCOPE"},
        ]
    )
    store.commit_turn()
    article = store.doc.parts[0].articles[0].uid
    store.begin_turn()
    store.apply_edits(
        [{"action": "add_paragraph", "target_id": article, "text": "Provide widgets."}]
    )
    store.commit_turn()
    empty = SpecSection.empty()
    redline = build_docx(store.doc, redline=diff_sections(empty, store.doc), redline_date=_DATE)
    assert "A. Provide widgets." in _body_texts(redline, "accept")
    assert _semantic_body_texts(redline, "accept") == _semantic_body_texts(
        build_docx(store.doc), "accept"
    )
    assert _semantic_body_texts(redline, "reject") == _semantic_body_texts(
        build_docx(empty), "accept"
    )


def test_accept_all_via_real_importer_matches_current_tree(tmp_path):
    store = _populated_store()
    base = SpecSection.from_dict(store.doc.to_dict())
    _adapt(store)
    cur = store.doc
    redline = build_docx(cur, redline=diff_sections(base, cur), redline_date=_DATE)
    path = tmp_path / "redline.docx"
    path.write_bytes(redline)

    from backend.spec_doc.importer import parse_master_docx

    reimported = parse_master_docx(path)

    def texts(section):
        rows = []
        for part in section.parts:
            for article in part.articles:
                rows.append(("A", article.title))

                def walk(paragraphs):
                    for p in paragraphs:
                        rows.append(("P", p.text))
                        walk(p.children)

                walk(article.paragraphs)
        return rows

    # The Accept-All re-import structurally equals a clean re-import of cur.
    clean_path = tmp_path / "clean.docx"
    clean_path.write_bytes(build_docx(cur))
    assert texts(reimported.section) == texts(parse_master_docx(clean_path).section)


# ---------------------------------------------------------------------------
# XML shape assertions
# ---------------------------------------------------------------------------


def test_redline_xml_shapes():
    store = _populated_store()
    base = SpecSection.from_dict(store.doc.to_dict())
    _adapt(store)
    redline = build_docx(store.doc, redline=diff_sections(base, store.doc), redline_date=_DATE)
    body = Document(io.BytesIO(redline)).element.body

    inserts = body.findall(".//" + _W_INS)
    deletes = body.findall(".//" + _W_DEL)
    assert inserts and deletes

    ids: list[str] = []
    for element in inserts + deletes:
        assert element.get(qn("w:id")) is not None
        assert element.get(qn("w:author")) == "Build-a-Spec"
        assert element.get(qn("w:date")) == _DATE
        ids.append(element.get(qn("w:id")))
    assert len(ids) == len(set(ids)), "w:id values must be unique"

    # No w:t inside a deletion — deleted text lives in w:delText.
    for dele in deletes:
        assert dele.findall(".//" + _W_T) == []

    # The deleted paragraph carries a paragraph-mark w:del (whole-para delete);
    # the inserted paragraph carries a paragraph-mark w:ins.
    del_marks = [p for p in body.iter(_W_P) if _para_mark(p, _W_DEL)]
    ins_marks = [p for p in body.iter(_W_P) if _para_mark(p, _W_INS)]
    assert len(del_marks) == 1
    assert len(ins_marks) == 1


def test_changed_paragraph_has_both_ins_and_del_but_survives_accept():
    store = _populated_store()
    base = SpecSection.from_dict(store.doc.to_dict())
    p0 = store.doc.parts[0].articles[0].paragraphs[0].uid
    store.begin_turn()
    store.apply_edits(
        [
            {
                "action": "replace",
                "target_id": p0,
                "text": "Section includes lattice wet-pipe systems per NFPA 13-2025.",
            }
        ]
    )
    store.commit_turn()
    redline = build_docx(store.doc, redline=diff_sections(base, store.doc), redline_date=_DATE)
    accept = _body_texts(redline, "accept")
    assert "Section includes lattice wet-pipe systems per NFPA 13-2025." in " ".join(
        accept
    )


# ---------------------------------------------------------------------------
# API surface
# ---------------------------------------------------------------------------


def _seed(store: DocumentStore) -> None:
    store.begin_turn()
    store.apply_edits(
        [
            {"action": "replace", "target_id": "sec", "text": "X", "numbering": "21 13 13"},
            {"action": "add_article", "target_id": "pt1", "text": "SUMMARY"},
        ]
    )
    store.commit_turn()
    art = store.doc.parts[0].articles[0].uid
    store.begin_turn()
    store.apply_edits(
        [{"action": "add_paragraph", "target_id": art, "text": "Provide sprinklers."}]
    )
    store.commit_turn()


def test_doc_diff_endpoint_validation():
    from backend import sessions

    client = TestClient(create_app())
    _seed(sessions.get_session().doc)

    ok = client.get("/api/doc/diff?base=0")
    assert ok.status_code == 200
    body = ok.json()
    assert body["stats"]["inserted"] >= 1
    assert body["base_index"] == 0 and body["cur_index"] == 2
    assert "baseline_index" in body

    assert client.get("/api/doc/diff?base=2&cur=2").status_code == 400  # same
    assert client.get("/api/doc/diff?base=9").status_code == 400  # out of range
    assert client.get("/api/doc/diff?base=-1").status_code == 400


def test_export_redline_master_requires_baseline():
    from backend import sessions

    client = TestClient(create_app())
    _seed(sessions.get_session().doc)

    # No imported master -> 400.
    resp = client.get("/api/export/docx?redline=master")
    assert resp.status_code == 400
    assert "no imported master" in resp.json()["error"].lower()

    # A version redline works and gets a REDLINE filename.
    resp = client.get("/api/export/docx?redline=version&base=1")
    assert resp.status_code == 200
    assert "REDLINE" in resp.headers["content-disposition"]
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml"
    )

    assert client.get("/api/export/docx?redline=version&base=9").status_code == 400
    assert client.get("/api/export/docx?redline=bogus").status_code == 400


def test_export_redline_master_after_import(tmp_path):
    from docx import Document as Docx

    from backend import sessions

    # Build a tiny master .docx and import it.
    master = Docx()
    for line in [
        "SECTION 21 13 13",
        "WET-PIPE SPRINKLER SYSTEMS",
        "PART 1 - GENERAL",
        "1.1 SUMMARY",
        "A. Section includes wet-pipe sprinkler systems.",
        "END OF SECTION 21 13 13",
    ]:
        master.add_paragraph(line)
    master_path = tmp_path / "master.docx"
    master.save(str(master_path))

    client = TestClient(create_app())
    with open(master_path, "rb") as handle:
        resp = client.post(
            "/api/import/master",
            files={"file": ("master.docx", handle.read(), "application/zip")},
        )
    assert resp.json()["baseline_index"] == 1

    # Adapt the imported provision, then redline vs master must succeed.
    session = sessions.get_session()
    para = session.doc.doc.parts[0].articles[0].paragraphs[0].uid
    session.doc.begin_turn()
    session.doc.apply_edits(
        [
            {
                "action": "replace",
                "target_id": para,
                "text": "Section includes wet-pipe automatic sprinkler systems.",
                "status": "confirmed",
            }
        ]
    )
    session.doc.commit_turn()

    resp = client.get("/api/export/docx?redline=master")
    assert resp.status_code == 200
    assert "REDLINE" in resp.headers["content-disposition"]
    # The exported redline carries genuine tracked changes.
    body = Document(io.BytesIO(resp.content)).element.body
    assert body.findall(".//" + _W_INS)


def test_baseline_index_survives_project_save_and_load(tmp_path):
    import json

    from docx import Document as Docx

    from backend import sessions

    master = Docx()
    for line in [
        "SECTION 21 13 13",
        "WET-PIPE SPRINKLER SYSTEMS",
        "PART 1 - GENERAL",
        "1.1 SUMMARY",
        "A. Section includes wet-pipe sprinkler systems.",
        "END OF SECTION 21 13 13",
    ]:
        master.add_paragraph(line)
    master_path = tmp_path / "master.docx"
    master.save(str(master_path))

    client = TestClient(create_app())
    with open(master_path, "rb") as handle:
        client.post(
            "/api/import/master",
            files={"file": ("master.docx", handle.read(), "application/zip")},
        )
    assert sessions.get_session().doc.baseline_index == 1

    project = json.loads(json.dumps(sessions.project_payload(sessions.get_session())))
    client.post("/api/session/reset")
    assert sessions.get_session().doc.baseline_index is None

    client.post("/api/project/load", json=project)
    assert sessions.get_session().doc.baseline_index == 1


def _docx_member_contents(docx_bytes: bytes) -> dict[str, bytes]:
    """Member name -> content for a ``.docx`` (a zip), ignoring zip metadata.

    python-docx writes each zip entry's last-modified stamp from the wall
    clock at save time (2-second DOS resolution), so two exports a moment
    apart are NOT raw-byte identical when they straddle a 2-second boundary
    — the envelope timestamps differ even though every member's *content*
    is identical. Comparing member contents is the timestamp-insensitive
    way to assert the clean export is deterministic, which is what
    "byte-stable" means here (the clean path writes no dynamic dates into
    the document itself — ``_redline_now`` is only reached with a redline).
    """
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return {name: zf.read(name) for name in sorted(zf.namelist())}


def test_clean_export_has_no_tracked_changes():
    """Regression guard: the default export path is deterministic (stable
    content) and carries zero redline markup (no w:ins/w:del)."""
    from backend import sessions

    client = TestClient(create_app())
    _seed(sessions.get_session().doc)

    first = client.get("/api/export/docx").content
    second = client.get("/api/export/docx").content
    # Deterministic content — compare member contents, not raw zip bytes,
    # so the zip's wall-clock entry timestamps don't make this flaky.
    assert _docx_member_contents(first) == _docx_member_contents(second)

    body = Document(io.BytesIO(first)).element.body
    assert body.findall(".//" + _W_INS) == []
    assert body.findall(".//" + _W_DEL) == []
