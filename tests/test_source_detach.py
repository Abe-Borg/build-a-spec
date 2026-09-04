"""Importing an existing spec and then actually editing it.

Source-preserving mode promises the exported ``.docx`` is a byte-exact clone
of the upload with only approved text slices changed. Every editing
restriction on an imported master serves that promise, and a package carrying
tracked changes, macros, an embedded OLE object, or enforced protection
cannot be patched at all, so it is frozen outright.

Two things are covered here. First, a frozen document now says WHY it is
frozen and what the user can do about it. Second, "Edit freely" lets the user
give up the preservation promise for a document and edit it like any other —
without giving up the exact original, the saved project, or redline vs
master, all of which keep working because detaching drops the claim and none
of the evidence.
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend import sessions
from backend.app import create_app
from backend.spec_doc.model import DocumentStore, SpecSection
from backend.spec_doc.source_mapping import (
    source_blocker_message,
    source_blocker_remedy,
)
from tests.conftest import settle_capability_sweep
from tests.docx_fidelity_helpers import (
    DOCX_MEDIA_TYPE,
    add_document_protection,
    add_tracked_change,
    make_fidelity_master,
)


@pytest.fixture
def client():
    sessions.reset_session()
    with TestClient(create_app()) as api:
        yield api
    sessions.reset_session()


def _plain_master() -> bytes:
    document = Document()
    for text in (
        "SECTION 23 05 48",
        "VIBRATION AND SEISMIC CONTROLS FOR HVAC",
        "PART 1 - GENERAL",
        "1.1 SECTION INCLUDES",
        "A. Vibration isolation devices for HVAC equipment.",
        "B. Seismic restraints for HVAC equipment.",
        "END OF SECTION 23 05 48",
    ):
        document.add_paragraph(text)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _import(client: TestClient, source: bytes) -> dict:
    response = client.post(
        "/api/import/master",
        files={"file": ("23 05 48.docx", source, DOCX_MEDIA_TYPE)},
    )
    assert response.status_code == 200, response.text
    settle_capability_sweep()
    doc = client.get("/api/doc")
    assert doc.status_code == 200, doc.text
    return doc.json()


def _first_paragraph_id(payload: dict) -> str:
    for part in payload["doc"]["parts"]:
        for article in part["articles"]:
            if article["paragraphs"]:
                return article["paragraphs"][0]["id"]
    raise AssertionError("the imported master has no paragraph")


def _body_operations(payload: dict) -> list[tuple[str, str, dict]]:
    capabilities = payload["source_capabilities"]
    assert capabilities is not None
    return [
        (uid, name, capability)
        for uid, element in capabilities["elements"].items()
        for name, capability in element.items()
        if name in {"replace_text", "delete", "move", "add_paragraph", "add_article"}
    ]


def _load(client: TestClient, package: bytes):
    return client.post(
        "/api/project/load-file",
        files={"file": ("project.baspec", package, "application/octet-stream")},
    )


# ---------------------------------------------------------------------------
# 1. A frozen document says why, and what to do about it
# ---------------------------------------------------------------------------


def test_a_frozen_package_names_its_cause_and_a_user_actionable_remedy(
    client, tmp_path
):
    """The symptom was a read-only document with no explanation anywhere.

    Tracked changes are the worst of these: the importer shows the clean
    Accept-All text, so nothing on screen hints that the package is still
    revision-bearing and therefore permanently unpatchable.
    """
    payload = _import(client, add_tracked_change(make_fidelity_master(tmp_path)))
    capabilities = payload["source_capabilities"]

    assert capabilities["status"] == "pass_through_only"
    assert [cause["blocker"] for cause in capabilities["causes"]] == [
        "tracked_changes"
    ]
    cause = capabilities["causes"][0]
    assert cause["message"] == source_blocker_message("tracked_changes")
    # The remedy is the half that was missing: this one is user-fixable.
    assert "Accept or reject the tracked changes" in cause["remedy"]
    assert "import again" in cause["remedy"]

    # And it really is frozen, so the remedy is not cosmetic.
    assert not any(capability["allowed"] for _uid, _name, capability in
                   _body_operations(payload))


def test_every_package_wide_cause_carries_a_remedy_naming_the_way_out(client):
    """Whatever the cause, the user is never left with a dead end.

    A cause we have no specific Word advice for must still say that Edit
    freely exists, or the honest "we cannot patch this" degrades back into
    the silent read-only document this work exists to remove.
    """
    for blocker in (
        "tracked_changes",
        "document_protection",
        "active_content",
        "signed_package",
        "unsafe_relationship_scan",
        "unsupported_raw_zip_layout",
        "some_future_blocker",
    ):
        assert "Edit freely" in source_blocker_remedy(blocker)


def test_document_protection_is_reported_as_its_own_cause(client, tmp_path):
    payload = _import(client, add_document_protection(_plain_master()))
    causes = payload["source_capabilities"]["causes"]
    assert [cause["blocker"] for cause in causes] == ["document_protection"]
    assert "Restrict Editing" in causes[0]["remedy"]


def test_a_clean_master_names_no_cause(client):
    """A ready report must not manufacture a fault to explain."""
    payload = _import(client, _plain_master())
    assert payload["source_capabilities"]["status"] == "ready"
    assert payload["source_capabilities"]["causes"] == []


def test_a_pending_sweep_is_not_reported_as_a_package_fault(client):
    """`pending` denies everything too, but the package is not at fault.

    The panel says "still working" there. Naming a cause would tell the user
    their file is broken while the app is merely still thinking.
    """
    response = client.post(
        "/api/import/master",
        files={"file": ("m.docx", _plain_master(), DOCX_MEDIA_TYPE)},
    )
    assert response.status_code == 200, response.text
    # Deliberately read BEFORE settling: this is the pending window.
    capabilities = client.get("/api/doc/capabilities").json()["source_capabilities"]
    if capabilities["status"] == "pending":
        assert capabilities["causes"] == []
    settle_capability_sweep()


# ---------------------------------------------------------------------------
# 2. Edit freely
# ---------------------------------------------------------------------------


def test_detaching_unlocks_every_edit_on_a_permanently_frozen_master(
    client, tmp_path
):
    """The headline case: the worst possible package becomes fully editable."""
    payload = _import(client, add_tracked_change(make_fidelity_master(tmp_path)))
    uid = _first_paragraph_id(payload)

    blocked = client.post(
        "/api/doc/edit",
        json={"ops": [{"action": "replace", "target_id": uid, "text": "Edited."}]},
    )
    assert blocked.status_code == 400
    assert "tracked_changes" in blocked.json()["error"]

    detached = client.post("/api/doc/detach-source")
    assert detached.status_code == 200, detached.text
    assert detached.json()["detached"] is True
    # No source scope at all any more -- ordinary semantic editing.
    assert detached.json()["source_capabilities"] is None

    for ops in (
        [{"action": "replace", "target_id": uid, "text": "Edited provision."}],
        [{"action": "set_status", "target_id": uid, "status": "confirmed"}],
        [{"action": "replace", "target_id": "sec", "text": "NEW TITLE",
          "numbering": "23 05 48"}],
        [{"action": "add_article", "target_id": "pt2", "text": "NEW ARTICLE"}],
    ):
        result = client.post("/api/doc/edit", json={"ops": ops})
        assert result.status_code == 200, (ops, result.text)


def test_detaching_keeps_the_exact_original_downloadable(client, tmp_path):
    """Detaching drops the promise, not the evidence."""
    source = add_tracked_change(make_fidelity_master(tmp_path))
    _import(client, source)
    assert client.post("/api/doc/detach-source").status_code == 200

    original = client.get("/api/import/original")
    assert original.status_code == 200
    assert original.content == source


def test_a_detached_document_exports_with_its_formatting_by_default(
    client, tmp_path
):
    """An imported project defaults to source mode; a detached one must not.

    Without this the export keeps selecting a mode whose preconditions were
    just given up, and every download 409s. A detached document with its
    formatting map still present takes the appearance-preserving render —
    the product's import path — and the payload says so, because the panel
    cannot infer it: detaching keeps every artifact the flag depends on.
    """
    source = add_tracked_change(make_fidelity_master(tmp_path))
    _import(client, source)
    assert client.post("/api/doc/detach-source").status_code == 200
    payload = client.get("/api/doc").json()
    assert payload["preserved_export_available"] is True
    assert payload["source_preservation"]["status"] == "detached"
    assert payload["source_preservation"]["exact_original_available"] is True

    export = client.get("/api/export/docx")
    assert export.status_code == 200, export.text
    assert len(export.content) > 0
    # The firm's own styles part travelled through untouched: that is the
    # appearance-preserving mode, not the normalized re-render.
    import io
    import zipfile

    with zipfile.ZipFile(io.BytesIO(source)) as archive:
        original_styles = archive.read("word/styles.xml")
    with zipfile.ZipFile(io.BytesIO(export.content)) as archive:
        assert archive.read("word/styles.xml") == original_styles


def test_a_project_without_a_formatting_map_says_so_in_the_payload(client):
    """A from-scratch document has nothing to rebuild; the flag is false."""
    payload = client.get("/api/doc").json()
    assert payload["preserved_export_available"] is False
    assert payload["source_preservation"] is None


def test_asking_for_source_export_after_detaching_says_why_it_is_gone(
    client, tmp_path
):
    """Every artifact is still present, so the generic message would be false."""
    _import(client, make_fidelity_master(tmp_path))
    assert client.post("/api/doc/detach-source").status_code == 200

    response = client.get("/api/export/docx", params={"mode": "source"})
    assert response.status_code == 409
    error = response.json()["error"]
    assert "detached from its Word original" in error
    assert "download the exact original upload" in error


def test_redline_against_the_master_survives_detaching(client, tmp_path):
    """The reason detach retains the baseline instead of clearing it.

    A user who imports, edits freely, and then wants to show their client
    what changed is the whole point of importing an existing spec.
    """
    _import(client, make_fidelity_master(tmp_path))
    assert client.post("/api/doc/detach-source").status_code == 200
    payload = client.get("/api/doc").json()
    assert payload["baseline_index"] is not None

    uid = _first_paragraph_id(payload)
    edit = client.post(
        "/api/doc/edit",
        json={"ops": [{"action": "replace", "target_id": uid,
                       "text": "Materially different provision text."}]},
    )
    assert edit.status_code == 200, edit.text

    redline = client.get("/api/export/docx", params={"redline": "master"})
    assert redline.status_code == 200, redline.text
    assert "REDLINE" in redline.headers["content-disposition"]


def test_detaching_survives_save_and_reload(client, tmp_path):
    """Detached must stay detached, stay editable, and keep the original.

    The package validates source/map/baseline against each other on load, so
    a detach that stripped artifacts would produce a file that rejects
    itself. Keeping them and persisting one flag is what makes this work.
    """
    source = add_tracked_change(make_fidelity_master(tmp_path))
    _import(client, source)
    assert client.post("/api/doc/detach-source").status_code == 200

    package = client.get("/api/project/save")
    assert package.status_code == 200, package.text
    assert client.post("/api/session/reset").status_code == 200

    loaded = _load(client, package.content)
    assert loaded.status_code == 200, loaded.text
    settle_capability_sweep()

    payload = client.get("/api/doc").json()
    assert payload["source_capabilities"] is None, "reload must stay detached"
    assert payload["source_available"] is True, "the original must come back"
    assert client.get("/api/import/original").content == source

    uid = _first_paragraph_id(payload)
    edit = client.post(
        "/api/doc/edit",
        json={"ops": [{"action": "replace", "target_id": uid, "text": "Edited."}]},
    )
    assert edit.status_code == 200, edit.text


def test_a_detached_project_reopens_after_edits_the_source_gate_would_refuse(
    client, tmp_path
):
    """The point of detaching is to make edits source mode would reject.

    The loader re-validates every retained version from the baseline forward
    against the preservation boundary, which is right for an attached project
    — a forged redo version must not enter the session and become active
    later. For a detached one it is the wrong question entirely: exceeding
    that boundary is exactly what the user asked for, so re-imposing it at
    load turns "edit freely, save, reopen" into a project that cannot be
    opened at all.

    Deliberately edits a HEADING, which `heading_change` refuses on every
    imported document however clean the package is.
    """
    source = add_tracked_change(make_fidelity_master(tmp_path))
    payload = _import(client, source)
    assert client.post("/api/doc/detach-source").status_code == 200

    uid = _first_paragraph_id(payload)
    for ops in (
        [{"action": "replace", "target_id": "sec", "text": "RETITLED SECTION",
          "numbering": "23 05 48"}],
        [{"action": "replace", "target_id": uid,
          "text": "Wholly rewritten provision text."}],
        [{"action": "add_article", "target_id": "pt2", "text": "NEW ARTICLE"}],
    ):
        assert client.post("/api/doc/edit", json={"ops": ops}).status_code == 200

    package = client.get("/api/project/save")
    assert package.status_code == 200, package.text
    assert client.post("/api/session/reset").status_code == 200

    loaded = _load(client, package.content)
    assert loaded.status_code == 200, loaded.text
    settle_capability_sweep()

    reopened = client.get("/api/doc").json()
    assert reopened["source_capabilities"] is None
    assert reopened["doc"]["section"]["title"] == "RETITLED SECTION"
    assert client.get("/api/import/original").content == source


def test_an_attached_project_still_validates_every_retained_version(
    client, tmp_path
):
    """The detached skip must not weaken the attached path it sits beside.

    A forged version past the baseline that exceeds the preservation boundary
    is still refused, because an attached project promises every retained
    state can be exported through source mode.
    """
    import json
    import zipfile
    import io as _io

    _import(client, make_fidelity_master(tmp_path))
    package = client.get("/api/project/save").content

    # Forge a retained version with a rewritten heading -- something the
    # source gate refuses -- while leaving the project ATTACHED.
    with zipfile.ZipFile(_io.BytesIO(package), "r") as archive:
        members = {name: archive.read(name) for name in archive.namelist()}
    project = json.loads(members["project.json"])
    assert project["doc"].get("source_detached") is False
    forged = json.loads(json.dumps(project["doc"]["versions"][-1]))
    forged["title"] = "FORGED HEADING"
    project["doc"]["versions"].append(forged)
    project["doc"]["index"] = len(project["doc"]["versions"]) - 1
    members["project.json"] = json.dumps(project).encode()

    output = _io.BytesIO()
    with zipfile.ZipFile(output, "w") as archive:
        for name, data in members.items():
            archive.writestr(name, data)

    assert client.post("/api/session/reset").status_code == 200
    rejected = _load(client, output.getvalue())
    assert rejected.status_code >= 400, rejected.text


def test_the_payload_states_detachment_because_it_cannot_be_inferred(
    client, tmp_path
):
    """A client must not have to guess, and cannot.

    Detaching keeps `source_available` true and `baseline_index` set (that is
    what leaves the exact original downloadable and redline vs master
    working) while `source_capabilities` goes null. That triple is
    byte-identical to a source-backed document whose report has not been
    delivered — so a client inferring scope from the retained artifacts reads
    a freely-editable document as locked and disables the very controls the
    user just enabled.
    """
    payload = _import(client, make_fidelity_master(tmp_path))
    assert payload["source_detached"] is False

    detached = client.post("/api/doc/detach-source").json()
    assert detached["source_detached"] is True
    # The artifacts that make the flag necessary are all still present.
    assert detached["source_available"] is True
    assert detached["baseline_index"] is not None
    assert detached["source_capabilities"] is None

    assert client.get("/api/doc").json()["source_detached"] is True


def test_detaching_twice_is_a_harmless_no_op(client, tmp_path):
    _import(client, make_fidelity_master(tmp_path))
    assert client.post("/api/doc/detach-source").status_code == 200
    again = client.post("/api/doc/detach-source")
    assert again.status_code == 200, again.text
    assert again.json()["detached"] is True


def test_detaching_without_an_imported_original_is_refused(client):
    """A from-scratch document has nothing to detach from."""
    response = client.post("/api/doc/detach-source")
    assert response.status_code == 409
    assert "no imported Word original" in response.json()["error"]


def test_detaching_is_refused_while_a_model_turn_streams(client, tmp_path):
    """A streaming turn's edits were validated under source scope.

    Letting the rules change underneath a turn that is about to commit is
    the same hazard that makes a mid-turn manual edit a 409.
    """
    _import(client, make_fidelity_master(tmp_path))
    session = sessions.get_session()
    session.turn_active = True
    try:
        response = client.post("/api/doc/detach-source")
        assert response.status_code == 409
        assert "model turn is streaming" in response.json()["error"]
    finally:
        session.turn_active = False


# ---------------------------------------------------------------------------
# 3. The persisted flag itself
# ---------------------------------------------------------------------------


def test_a_project_saved_before_edit_freely_existed_loads_attached():
    """Absence means attached, which is what those files were."""
    store = DocumentStore()
    store.load({"versions": [SpecSection.empty().to_dict()], "index": 0})
    assert store.source_detached is False


@pytest.mark.parametrize("value", ["true", 1, {}, [], None, "yes"])
def test_only_an_explicit_true_detaches_on_load(value):
    """Fail closed: a malformed flag must never hand out edit permissions.

    Every other guard in the source gate fails closed, and this one decides
    whether the gate runs at all.
    """
    store = DocumentStore()
    store.load(
        {
            "versions": [SpecSection.empty().to_dict()],
            "index": 0,
            "source_detached": value,
        }
    )
    assert store.source_detached is False


def test_the_flag_round_trips_through_the_store():
    store = DocumentStore()
    store.source_detached = True
    restored = DocumentStore()
    restored.load(store.to_dict())
    assert restored.source_detached is True


def test_undo_and_redo_never_flip_the_detach_decision(client, tmp_path):
    """It describes the document's relationship to its source, not content.

    Versions are section snapshots; this rides the store beside
    ``baseline_index`` precisely so stepping through history cannot silently
    re-lock a document the user chose to edit freely.
    """
    payload = _import(client, make_fidelity_master(tmp_path))
    assert client.post("/api/doc/detach-source").status_code == 200
    uid = _first_paragraph_id(payload)
    assert client.post(
        "/api/doc/edit",
        json={"ops": [{"action": "replace", "target_id": uid, "text": "Edited."}]},
    ).status_code == 200

    undone = client.post("/api/doc/undo")
    assert undone.status_code == 200, undone.text
    assert undone.json()["source_capabilities"] is None
    redone = client.post("/api/doc/redo")
    assert redone.status_code == 200, redone.text
    assert redone.json()["source_capabilities"] is None


def test_a_fresh_import_re_arms_source_preservation(client, tmp_path):
    """Detaching is one-way for a document, not for the session.

    Importing again is the supported way back, so the new package must not
    inherit the previous document's decision.
    """
    _import(client, make_fidelity_master(tmp_path))
    assert client.post("/api/doc/detach-source").status_code == 200
    assert client.post("/api/session/reset").status_code == 200

    payload = _import(client, _plain_master())
    assert payload["source_capabilities"] is not None
    assert payload["source_capabilities"]["status"] == "ready"


def test_importing_again_after_undoing_to_an_empty_page_re_arms_preservation(
    client, tmp_path
):
    """The re-arm that a session reset would otherwise hide.

    Import is gated on the CURRENT document being empty, and undo walks back
    to the empty version 0 while the detach flag rides the store untouched.
    So a second import is reachable without a reset — and it must come back
    source-preserving, not silently inherit the previous document's decision
    and hand out edit permissions the new package never granted.
    """
    _import(client, make_fidelity_master(tmp_path))
    assert client.post("/api/doc/detach-source").status_code == 200

    # Walk back to the empty page the import landed on top of.
    for _ in range(10):
        if not client.get("/api/doc").json()["doc"]["parts"][0]["articles"]:
            break
        assert client.post("/api/doc/undo").status_code == 200

    payload = _import(client, _plain_master())
    assert payload["source_capabilities"] is not None, (
        "a re-imported master must be source-backed again"
    )
    assert payload["source_capabilities"]["status"] == "ready"


def test_a_detached_document_stops_offering_source_preserving_export(
    client, tmp_path
):
    """`preservation_ready` is what the panel's export menu keys on.

    Leaving it true would advertise a source-preserving export that the
    endpoint now refuses -- a menu entry that always errors.
    """
    payload = _import(client, make_fidelity_master(tmp_path))
    assert payload["preservation_ready"] is True

    detached = client.post("/api/doc/detach-source")
    assert detached.status_code == 200, detached.text
    assert detached.json()["preservation_ready"] is False
    assert client.get("/api/doc").json()["preservation_ready"] is False
