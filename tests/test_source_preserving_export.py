"""P1a acceptance tests for source-backed DOCX export.

P1a is intentionally narrow. An imported package is an immutable base; a
source export either returns it exactly or applies one proven-safe body-text
patch. Unsupported semantic operations are rejected before they can enter the
version history. The normalized renderer remains an explicit separate mode.
"""
from __future__ import annotations

import io

import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend import sessions
from backend.app import create_app
from backend.qc.engine import (
    QCFinding,
)
from tests.fakes import (
    audit_grade_qc_result,
    FakeClient,
    text_turn,
    tool_turn,
)
from tests.docx_render_harness import (
    DocxRenderHarness,
    RENDERER_SKIP_REASON,
    renderer_is_configured,
)
from tests.docx_fidelity_helpers import (
    DOCX_MEDIA_TYPE,
    TARGET_EDITED_SOURCE_TEXT,
    TARGET_EDITED_TEXT,
    TARGET_MODEL_TEXT,
    add_active_content_marker,
    add_document_protection,
    add_signature_origin_marker,
    add_tracked_change,
    assert_only_target_text_changed,
    assert_untouched_parts_identical,
    assert_valid_docx_package,
    make_fidelity_master,
    make_table_projection_master,
    paragraph_text,
)


def _client() -> TestClient:
    return TestClient(create_app())


def _import_master(client: TestClient, source: bytes):
    response = client.post(
        "/api/import/master",
        files={
            "file": (
                "client-fidelity-master.docx",
                source,
                DOCX_MEDIA_TYPE,
            )
        },
    )
    assert response.status_code == 200, response.text
    return response.json()


def _source_export(client: TestClient):
    return client.get("/api/export/docx", params={"mode": "source"})


def _replace_target(client: TestClient):
    return client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "replace",
                    "target_id": "pt1.a1.p1",
                    "text": TARGET_EDITED_TEXT,
                    "status": "confirmed",
                }
            ]
        },
    )


def test_noop_source_export_is_byte_identical_and_normalized_is_explicit(
    tmp_path,
):
    client = _client()
    source = make_fidelity_master(tmp_path)
    imported = _import_master(client, source)
    assert imported["source_available"] is True
    assert imported["preservation_ready"] is True

    preserving = _source_export(client)
    assert preserving.status_code == 200
    assert preserving.content == source
    assert_valid_docx_package(preserving.content)

    # Imported projects default to the source-preserving path. A caller must
    # explicitly request the old semantic reconstruction.
    default_export = client.get("/api/export/docx")
    assert default_export.status_code == 200
    assert default_export.content == source

    normalized = client.get("/api/export/docx", params={"mode": "normalized"})
    assert normalized.status_code == 200
    assert normalized.content != source
    assert_valid_docx_package(normalized.content)
    assert "ASSUMPTIONS SCHEDULE" in [
        paragraph.text for paragraph in Document(io.BytesIO(normalized.content)).paragraphs
    ]


def test_app_only_edits_do_not_dirty_the_source_docx(tmp_path):
    client = _client()
    source = make_fidelity_master(tmp_path)
    _import_master(client, source)

    response = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "set_status",
                    "target_id": "pt1.a1.p1",
                    "status": "confirmed",
                },
                {
                    "action": "set_project_profile",
                    "target_id": "sec",
                    "city": "Ashburn",
                    "state": "VA",
                    "country": "US",
                    "client": "Fixture Client",
                },
                {
                    "action": "set_standard_edition",
                    "target_id": "sec",
                    "standard": "NFPA 13",
                    "edition": "2022",
                    "basis": "Acceptance-test project basis",
                },
            ]
        },
    )
    assert response.status_code == 200, response.text
    assert response.json()["doc"]["version"] == {"index": 2, "count": 3}
    exported = _source_export(client)
    assert exported.status_code == 200
    assert exported.content == source


def test_one_text_node_replace_is_surgical_and_original_stays_immutable(
    tmp_path,
):
    client = _client()
    source = make_fidelity_master(tmp_path)
    _import_master(client, source)

    edited_response = _replace_target(client)
    assert edited_response.status_code == 200, edited_response.text
    assert edited_response.json()["preservation_ready"] is True
    paragraph = edited_response.json()["doc"]["parts"][0]["articles"][0][
        "paragraphs"
    ][0]
    assert paragraph["text"] == TARGET_EDITED_TEXT

    first_export = _source_export(client)
    assert first_export.status_code == 200, first_export.text
    edited = first_export.content
    assert edited != source
    assert paragraph_text(source).endswith("NFPA 13-2019.")
    assert paragraph_text(edited) == TARGET_EDITED_SOURCE_TEXT

    assert_untouched_parts_identical(source, edited)
    assert_only_target_text_changed(source, edited)
    assert_valid_docx_package(edited)

    # Export is deterministic and the recovery endpoint never becomes the
    # edited copy.
    second_export = _source_export(client)
    assert second_export.status_code == 200
    assert second_export.content == edited
    original = client.get("/api/import/original")
    assert original.status_code == 200
    assert original.content == source

    reopened = Document(io.BytesIO(edited))
    header_text = "\n".join(
        paragraph.text
        for section in reopened.sections
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        )
        for paragraph in header.paragraphs
    )
    footer_text = "\n".join(
        paragraph.text
        for section in reopened.sections
        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )
        for paragraph in footer.paragraphs
    )
    assert "CLIENT MASTER HEADER | KEEP EXACT" in header_text
    assert "CLIENT FIRST-PAGE HEADER | KEEP EXACT" in header_text
    assert "CLIENT CONFIDENTIAL | PAGE" in footer_text


def test_undo_to_import_baseline_is_exact_source_and_redo_replays_patch(tmp_path):
    client = _client()
    source = make_fidelity_master(tmp_path)
    _import_master(client, source)
    assert _replace_target(client).status_code == 200
    edited = _source_export(client).content

    undone = client.post("/api/doc/undo")
    assert undone.status_code == 200
    assert undone.json()["doc"]["parts"][0]["articles"][0]["paragraphs"][0][
        "text"
    ] == TARGET_MODEL_TEXT
    assert _source_export(client).content == source

    redone = client.post("/api/doc/redo")
    assert redone.status_code == 200
    assert _source_export(client).content == edited


@pytest.mark.parametrize(
    "operation",
    [
        pytest.param(
            {
                "action": "add_paragraph",
                "target_id": "pt1.a1",
                "text": "A new provision is not source-safe in P1a.",
                "status": "confirmed",
            },
            id="add-paragraph",
        ),
        pytest.param(
            {"action": "delete", "target_id": "pt1.a1.p1"},
            id="delete-paragraph",
        ),
        pytest.param(
            {
                "action": "replace",
                "target_id": "pt1.a1",
                "text": "RENAMED SUMMARY",
            },
            id="replace-article-heading",
        ),
        pytest.param(
            {
                "action": "replace",
                "target_id": "sec",
                "text": "RENAMED SECTION TITLE",
            },
            id="replace-section-heading",
        ),
    ],
)
def test_unsupported_source_operations_fail_before_commit(
    tmp_path, operation
):
    client = _client()
    source = make_fidelity_master(tmp_path)
    _import_master(client, source)
    before = client.get("/api/doc").json()["doc"]

    rejected = client.post("/api/doc/edit", json={"ops": [operation]})
    assert rejected.status_code == 400
    error = rejected.json()["error"].lower()
    assert "source-backed edit rejected" in error
    assert client.get("/api/doc").json()["doc"] == before
    assert _source_export(client).content == source


def test_multi_run_target_fails_closed_instead_of_flattening_formatting(tmp_path):
    client = _client()
    source = make_fidelity_master(tmp_path, split_target_runs=True)
    _import_master(client, source)
    before = client.get("/api/doc").json()["doc"]

    rejected = _replace_target(client)
    assert rejected.status_code == 400
    assert client.get("/api/doc").json()["doc"] == before
    preserving = _source_export(client)
    assert preserving.status_code == 200
    assert preserving.content == source


def test_model_tool_path_uses_the_same_source_guard(tmp_path, monkeypatch):
    client = _client()
    source = make_fidelity_master(tmp_path)
    _import_master(client, source)
    before = client.get("/api/doc").json()["doc"]
    fake = FakeClient(
        [
            tool_turn(
                ["Trying the requested source edit."],
                {
                    "edits": [
                        {"action": "delete", "target_id": "pt1.a1.p1"}
                    ]
                },
            ),
            text_turn(["That structural edit is outside the safe P1a boundary."]),
        ]
    )
    monkeypatch.setattr("backend.llm.conversation.get_client", lambda: fake)

    response = client.post(
        "/api/chat", json={"message": "Delete the first imported provision."}
    )
    assert response.status_code == 200
    assert client.get("/api/doc").json()["doc"] == before
    tool_results = [
        block
        for message in sessions.get_session().history
        for block in message.get("content", [])
        if block.get("type") == "tool_result"
    ]
    assert len(tool_results) == 1
    assert tool_results[0]["is_error"] is True
    assert "source-backed edit rejected" in tool_results[0]["content"].lower()
    assert _source_export(client).content == source


def test_qc_apply_path_uses_the_same_source_guard(tmp_path):
    client = _client()
    source = make_fidelity_master(tmp_path)
    _import_master(client, source)
    before = client.get("/api/doc").json()["doc"]
    finding_id = "qc-source-guard"
    session = sessions.get_session()
    result = audit_grade_qc_result(
        session,
        [
            QCFinding(
                finding_id=finding_id,
                lens_id="constructability",
                severity="high",
                element_id="pt1.a1.p1",
                title="Remove the provision",
                issue="Fixture finding proposes a structural source edit.",
                rationale="Exercise the central guard.",
                proposed_ops=[
                    {"action": "delete", "target_id": "pt1.a1.p1"}
                ],
                ops_valid=True,
            )
        ],
    )
    session.qc.restore(result)

    rejected = client.post(
        "/api/qc/apply", json={"finding_ids": [finding_id]}
    )
    assert rejected.status_code == 400
    assert "source-backed edit rejected" in rejected.json()["error"].lower()
    assert client.get("/api/doc").json()["doc"] == before
    assert session.qc.result.finding(finding_id).status == "open"
    assert _source_export(client).content == source


def test_hyperlink_target_fails_closed_instead_of_rewriting_relationship_markup(
    tmp_path,
):
    client = _client()
    source = make_fidelity_master(tmp_path)
    _import_master(client, source)
    before = client.get("/api/doc").json()["doc"]

    rejected = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "replace",
                    "target_id": "pt1.a1.p3",
                    "text": "Replace a paragraph that contains a hyperlink.",
                }
            ]
        },
    )
    assert rejected.status_code == 400
    assert "[complex_paragraph_markup]" in rejected.json()["error"]
    assert client.get("/api/doc").json()["doc"] == before
    assert _source_export(client).content == source


@pytest.mark.parametrize(
    ("source_mutator", "expected_blocker"),
    [
        pytest.param(
            add_document_protection,
            "document_protection",
            id="document-protection",
        ),
        pytest.param(add_tracked_change, "tracked_changes", id="tracked-changes"),
        pytest.param(
            add_signature_origin_marker,
            "signed_package",
            id="signed-package",
        ),
        pytest.param(
            add_active_content_marker,
            "active_content",
            id="active-content",
        ),
    ],
)
def test_global_mutation_blockers_allow_exact_noop_but_reject_text_edit(
    tmp_path, source_mutator, expected_blocker
):
    client = _client()
    source = source_mutator(make_fidelity_master(tmp_path))
    _import_master(client, source)

    # Passing the source through does not invalidate signatures, protection,
    # or revision semantics, so it remains exact.
    assert _source_export(client).content == source
    before = client.get("/api/doc").json()["doc"]
    rejected = _replace_target(client)
    assert rejected.status_code == 400
    assert f"[{expected_blocker}]" in rejected.json()["error"]
    assert client.get("/api/doc").json()["doc"] == before
    assert _source_export(client).content == source


def test_table_projection_is_read_only_and_preserved_as_a_table(tmp_path):
    client = _client()
    source = make_table_projection_master(tmp_path)
    imported = _import_master(client, source)
    assert imported["imported_block_count"] == 1
    before = client.get("/api/doc").json()["doc"]

    rejected = client.post(
        "/api/doc/edit",
        json={
            "ops": [
                {
                    "action": "replace",
                    "target_id": "pt2.a1.p1",
                    "text": "Flattened table text must never overwrite the table.",
                }
            ]
        },
    )
    assert rejected.status_code == 400
    assert "[table_projection]" in rejected.json()["error"]
    assert client.get("/api/doc").json()["doc"] == before
    exported = _source_export(client)
    assert exported.status_code == 200
    assert exported.content == source
    assert len(Document(io.BytesIO(exported.content)).tables) == 1


def test_source_mode_without_an_import_never_silently_normalizes():
    client = _client()
    preserving = _source_export(client)
    assert preserving.status_code == 409
    assert preserving.json()["ok"] is False

    normalized = client.get("/api/export/docx", params={"mode": "normalized"})
    assert normalized.status_code == 200
    assert_valid_docx_package(normalized.content)


@pytest.mark.skipif(
    not renderer_is_configured(),
    reason=RENDERER_SKIP_REASON,
)
def test_rendered_edit_keeps_page_furniture_and_limits_visual_change(tmp_path):
    """Optional same-renderer visual regression for the surgical fixture.

    This is deliberately paired with the package-level assertions above.
    Raster comparison can catch layout drift, but cannot prove that Word
    relationships, numbering, or custom parts survived.
    """

    pytest.importorskip("PIL.Image")
    renderer = DocxRenderHarness.from_environment()

    client = _client()
    source = make_fidelity_master(tmp_path)
    _import_master(client, source)
    assert _replace_target(client).status_code == 200
    edited = _source_export(client).content

    source_rendered = renderer.render_bytes(
        source,
        work_dir=tmp_path,
        stem="render-source",
    )
    edited_rendered = renderer.render_bytes(
        edited,
        work_dir=tmp_path,
        stem="render-edited",
    )
    comparison = renderer.compare(source_rendered, edited_rendered)
    comparison.assert_changed_pages({1})
    comparison.assert_page_furniture_unchanged()
    comparison.assert_changes_within_body()
    # A same-length edition change should affect a tiny body area, never
    # trigger page-wide reflow.
    comparison.assert_diff_budget(
        max_changed_pixel_fraction=0.01,
        max_bbox_area_fraction=0.02,
    )
