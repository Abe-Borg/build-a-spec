"""P1b acceptance tests for controlled source-backed body structure edits.

Structural editing is intentionally much narrower than text editing.  The
only writable island in these tests is a flat, contiguous run of simple,
direct-body paragraphs that all use one genuine Word ``w:numPr`` definition.
Everything around that island remains opaque and byte-preserved.
"""
from __future__ import annotations

import io
import posixpath
import zipfile
from copy import deepcopy

import pytest
from docx import Document
from docx.oxml.ns import qn
from fastapi.testclient import TestClient
from lxml import etree

import backend.spec_doc.source_patch as source_patch_module
from backend.app import create_app
from backend.spec_doc.importer import parse_master_docx
from backend.spec_doc.model import Paragraph
from backend.spec_doc.source_patch import SourcePatchError, validate_source_transition
from backend.spec_doc.xml_lexical import (
    XmlPatch,
    apply_xml_patches,
    build_source_xml_index,
    decoded_slice_byte_span,
    encode_word_text,
)
from tests.docx_fidelity_helpers import (
    DOCX_MEDIA_TYPE,
    NUMBERED_ISLAND_PARA_IDS,
    NUMBERED_ISLAND_TEXTS,
    add_active_content_marker,
    add_document_protection,
    add_paragraph_property_change,
    add_signature_origin_marker,
    add_tracked_change,
    assert_untouched_parts_identical,
    assert_valid_docx_package,
    document_xml,
    make_fidelity_master,
    make_numbered_island_master,
    make_table_projection_master,
    remove_numbering_relationship,
    replace_numbering_content_type,
    rewrite_zip_members,
)


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_NS = {"w": _W_NS, "w14": _W14_NS}
_UTF8_BOM = b"\xef\xbb\xbf"
_NEW_TEXT = "Provide tamper switches at all supervised valves."
_RANGE_OR_REVISION_TAGS = {
    "bookmarkStart",
    "bookmarkEnd",
    "commentRangeStart",
    "commentRangeEnd",
    "customXmlInsRangeStart",
    "customXmlInsRangeEnd",
    "customXmlDelRangeStart",
    "customXmlDelRangeEnd",
    "moveFromRangeStart",
    "moveFromRangeEnd",
    "moveToRangeStart",
    "moveToRangeEnd",
    "permStart",
    "permEnd",
    "ins",
    "del",
    "moveFrom",
    "moveTo",
}


def _client() -> TestClient:
    return TestClient(create_app())


def _import_master(client: TestClient, source: bytes):
    response = client.post(
        "/api/import/master",
        files={
            "file": (
                "client-numbered-island-master.docx",
                source,
                DOCX_MEDIA_TYPE,
            )
        },
    )
    assert response.status_code == 200, response.text
    return response.json()


def _edit(client: TestClient, *ops: dict):
    return client.post("/api/doc/edit", json={"ops": list(ops)})


def _source_export(client: TestClient):
    return client.get("/api/export/docx", params={"mode": "source"})


def _body(payload: bytes):
    root = etree.fromstring(document_xml(payload))
    body = root.find(qn("w:body"))
    assert body is not None
    return body


def _canonical(element) -> bytes:
    return etree.tostring(
        element,
        method="c14n",
        exclusive=True,
        with_comments=False,
    )


def _direct_paragraph_by_id(payload: bytes, para_id: str):
    matches = _body(payload).xpath(
        f'./w:p[@w14:paraId="{para_id}"]', namespaces=_NS
    )
    assert len(matches) == 1, f"Expected one direct body paragraph {para_id}"
    return matches[0]


def _direct_paragraph_by_text(payload: bytes, text: str):
    matches = [
        paragraph
        for paragraph in _body(payload).xpath("./w:p", namespaces=_NS)
        if "".join(paragraph.xpath(".//w:t/text()", namespaces=_NS)) == text
    ]
    assert len(matches) == 1, f"Expected one direct body paragraph with {text!r}"
    return matches[0]


def _direct_numbered_paragraphs(payload: bytes) -> list:
    return _body(payload).xpath("./w:p[w:pPr/w:numPr]", namespaces=_NS)


def _paragraph_text(paragraph) -> str:
    return "".join(paragraph.xpath(".//w:t/text()", namespaces=_NS))


def _paragraph_id(paragraph) -> str:
    return paragraph.get(qn("w14:paraId"), "")


def _serialize(element) -> bytes:
    return etree.tostring(
        element,
        encoding="UTF-8",
        xml_declaration=True,
        standalone=True,
    )


def _body_elements_with_indices(payload: bytes) -> tuple[bytes, object, list]:
    xml = document_xml(payload)
    index = build_source_xml_index(xml)
    root = etree.fromstring(xml)
    body = root.find(qn("w:body"))
    assert body is not None
    children = [child for child in body.iterchildren() if isinstance(child.tag, str)]
    assert len(children) == len(index.body_children)
    return xml, index, children


def _body_child_index_by_id(payload: bytes, para_id: str) -> int:
    _xml, _index, children = _body_elements_with_indices(payload)
    matches = [
        child_index
        for child_index, child in enumerate(children)
        if child.tag == qn("w:p") and child.get(qn("w14:paraId"), "") == para_id
    ]
    assert len(matches) == 1, f"Expected one body paragraph {para_id}"
    return matches[0]


def _raw_direct_paragraph_by_text(payload: bytes, text: str) -> bytes:
    xml, index, children = _body_elements_with_indices(payload)
    matches = [
        child_index
        for child_index, child in enumerate(children)
        if child.tag == qn("w:p")
        and "".join(child.xpath(".//w:t/text()", namespaces=_NS)) == text
    ]
    assert len(matches) == 1, f"Expected one raw paragraph with {text!r}"
    child = index.body_child(matches[0])
    return xml[child.element_span.start : child.element_span.end]


def _source_island_bytes(payload: bytes) -> tuple[
    bytes,
    object,
    tuple[int, ...],
    tuple[bytes, ...],
    tuple[bytes, ...],
]:
    xml, index, _children = _body_elements_with_indices(payload)
    indices = tuple(
        _body_child_index_by_id(payload, para_id)
        for para_id in NUMBERED_ISLAND_PARA_IDS
    )
    assert indices == tuple(range(indices[0], indices[0] + len(indices)))
    elements = tuple(
        xml[
            index.body_child(child_index).element_span.start :
            index.body_child(child_index).element_span.end
        ]
        for child_index in indices
    )
    gaps = tuple(
        xml[
            index.body_gaps[child_index + 1].start :
            index.body_gaps[child_index + 1].end
        ]
        for child_index in indices[:-1]
    )
    return xml, index, indices, elements, gaps


def _expected_island_rewrite_xml(
    source: bytes,
    desired_elements: tuple[bytes, ...],
) -> bytes:
    """Independent oracle for the documented fixed-gap slot strategy."""
    xml, index, indices, _source_elements, gaps = _source_island_bytes(source)
    slots: list[list[bytes]] = [[] for _item in indices]
    if len(desired_elements) <= len(slots):
        for position, element in enumerate(desired_elements):
            slots[position].append(element)
    else:
        for position in range(len(slots) - 1):
            slots[position].append(desired_elements[position])
        slots[-1].extend(desired_elements[len(slots) - 1 :])
    replacement_parts: list[bytes] = []
    for position, slot in enumerate(slots):
        replacement_parts.extend(slot)
        if position < len(gaps):
            replacement_parts.append(gaps[position])
    first = index.body_child(indices[0]).element_span
    last = index.body_child(indices[-1]).element_span
    return (
        xml[: first.start]
        + b"".join(replacement_parts)
        + xml[last.end :]
    )


def _assert_exact_island_rewrite(
    source: bytes,
    output: bytes,
    desired_elements: tuple[bytes, ...],
) -> None:
    source_xml, index, indices, _elements, gaps = _source_island_bytes(source)
    expected = _expected_island_rewrite_xml(source, desired_elements)
    output_xml = document_xml(output)
    assert output_xml == expected

    # These narrower checks make a failure explain whether the lexical
    # envelope or the island reconstruction drifted.
    first = index.body_child(indices[0]).element_span
    last = index.body_child(indices[-1]).element_span
    assert output_xml[: first.start] == source_xml[: first.start]
    assert output_xml.endswith(source_xml[last.end :])
    expected_region = expected[
        first.start : len(expected) - len(source_xml[last.end :])
    ]
    for gap in gaps:
        assert expected_region.count(gap) >= 1


def _with_internal_island_gaps(
    payload: bytes,
    first_gap: bytes,
    second_gap: bytes,
) -> bytes:
    xml, index, indices, _elements, _gaps = _source_island_bytes(payload)
    replacements = []
    for child_index, replacement in zip(indices[:-1], (first_gap, second_gap)):
        span = index.body_gaps[child_index + 1]
        replacements.append(
            XmlPatch(
                span.start,
                span.end,
                replacement,
                f"gap-{child_index}",
                "test_fixture_gap",
            )
        )
    rewritten = apply_xml_patches(xml, replacements)
    return rewrite_zip_members(
        payload,
        replacements={"word/document.xml": rewritten},
    )


def _with_distinct_whitespace_gaps(payload: bytes) -> bytes:
    return _with_internal_island_gaps(
        payload,
        b"\r\n \t",
        b"\n    ",
    )


def _with_unknown_clone_property(payload: bytes) -> bytes:
    """Add schema-unproven WML uniformly so island signatures still agree."""
    root = etree.fromstring(document_xml(payload))
    for para_id in NUMBERED_ISLAND_PARA_IDS:
        matches = root.xpath(
            f'.//w:p[@w14:paraId="{para_id}"]/w:pPr',
            namespaces=_NS,
        )
        assert len(matches) == 1
        unknown = etree.SubElement(matches[0], qn("w:futureBehavior"))
        unknown.set(qn("w:val"), "1")
    return rewrite_zip_members(
        payload,
        replacements={"word/document.xml": _serialize(root)},
    )


def _raw_paragraph_with_text_patch(
    payload: bytes,
    para_id: str,
    replacement_text: str,
) -> bytes:
    xml, index, _children = _body_elements_with_indices(payload)
    child_index = _body_child_index_by_id(payload, para_id)
    child = index.body_child(child_index)
    text_node = index.word_text(child_index, 0)
    span = decoded_slice_byte_span(xml, text_node, 0, len(text_node.decoded_text))
    replacement = encode_word_text(
        replacement_text,
        raw_prefix=xml[text_node.content_span.start : span.start],
        raw_suffix=xml[span.end : text_node.content_span.end],
    )
    raw = xml[child.element_span.start : child.element_span.end]
    return apply_xml_patches(
        raw,
        (
            XmlPatch(
                span.start - child.element_span.start,
                span.end - child.element_span.start,
                replacement,
                para_id,
                "replace_text",
            ),
        ),
    )


def _make_paragraph_local_word_prefix_master(tmp_path) -> tuple[bytes, bytes, bytes]:
    """Numbered source whose paragraph prefix is not bound at ``w:body``."""
    base = make_numbered_island_master(tmp_path)
    base_root = etree.fromstring(document_xml(base))
    num_ids = base_root.xpath(
        "./w:body/w:p/w:pPr/w:numPr/w:numId/@w:val",
        namespaces=_NS,
    )
    assert num_ids
    num_id = str(num_ids[0]).encode("ascii")
    word_uri = _W_NS.encode("ascii")
    w14_uri = _W14_NS.encode("ascii")
    p_pr = b"".join(
        (
            b"<quill:pPr><quill:numPr>",
            b"<quill:ilvl quill:val='0'/>",
            b"<quill:numId quill:val='",
            num_id,
            b"'/></quill:numPr></quill:pPr>",
        )
    )
    r_pr = (
        b"<quill:rPr><quill:b/><quill:color quill:val='1F4E79'/>"
        b"</quill:rPr>"
    )

    def heading(text: bytes) -> bytes:
        return (
            b"<rootw:p><rootw:r><rootw:t>"
            + text
            + b"</rootw:t></rootw:r></rootw:p>"
        )

    def numbered(para_id: str, text: str, ordinal: int) -> bytes:
        return b"".join(
            (
                b"<quill:p idmeta:paraId='",
                para_id.encode("ascii"),
                b"' quill:rsidR='00ABC",
                str(ordinal).encode("ascii"),
                b"' xmlns:quill='",
                word_uri,
                b"'>",
                p_pr,
                b"<quill:r quill:rsidR='00123",
                str(ordinal).encode("ascii"),
                b"'>",
                r_pr,
                b"<quill:t>",
                text.encode("utf-8"),
                b"</quill:t></quill:r></quill:p>",
            )
        )

    source_xml = b"".join(
        (
            _UTF8_BOM,
            b"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\r\n",
            b"<?before-document lexical='exact'?>\r\n",
            b"<rootw:document data-z='last' xmlns:unused=\"urn:unused\" ",
            b"xmlns:idmeta='",
            w14_uri,
            b"' xmlns:rootw=\"",
            word_uri,
            b"\" data-a=\"first\">\r\n",
            b"<rootw:body data-two='2' data-one=\"1\">\r\n",
            heading(b"SECTION 21 13 15"),
            b"\r\n",
            heading(b"LOCAL PREFIX NUMBERING"),
            b"\r\n",
            heading(b"PART 1 - GENERAL"),
            b"\r\n",
            heading(b"1.1 PREFIX SAFETY"),
            b"\r\n<!-- before island --><?island-boundary before?>\r\n",
            numbered(NUMBERED_ISLAND_PARA_IDS[0], NUMBERED_ISLAND_TEXTS[0], 0),
            b"\r\n \t",
            numbered(NUMBERED_ISLAND_PARA_IDS[1], NUMBERED_ISLAND_TEXTS[1], 1),
            b"\n    ",
            numbered(NUMBERED_ISLAND_PARA_IDS[2], NUMBERED_ISLAND_TEXTS[2], 2),
            b"\r\n<?island-boundary after?><!-- after island -->\r\n",
            heading(b"END OF SECTION 21 13 15"),
            b"\r\n<rootw:sectPr/>\r\n",
            b"</rootw:body>\r\n</rootw:document>",
        )
    )
    source = rewrite_zip_members(
        base,
        replacements={"word/document.xml": source_xml},
    )
    return source, p_pr, r_pr


def _expected_local_prefix_new_paragraph(
    p_pr: bytes,
    r_pr: bytes,
    text: str,
) -> bytes:
    return b"".join(
        (
            b'<quill:p xmlns:quill="',
            _W_NS.encode("ascii"),
            b'">',
            p_pr,
            b"<quill:r>",
            r_pr,
            b"<quill:t>",
            encode_word_text(text),
            b"</quill:t></quill:r></quill:p>",
        )
    )


def _with_relocated_header_numbering_use(payload: bytes) -> bytes:
    """Reference the island numId from a correctly typed non-.xml header."""
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        document = etree.fromstring(archive.read("word/document.xml"))
        relationships = etree.fromstring(
            archive.read("word/_rels/document.xml.rels")
        )
        content_types = etree.fromstring(archive.read("[Content_Types].xml"))

        numbered = document.find(f".//{{{_W_NS}}}numPr")
        assert numbered is not None
        header_relationships = [
            relationship
            for relationship in relationships.findall(
                f"{{{_REL_NS}}}Relationship"
            )
            if relationship.get("Type", "").endswith("/header")
        ]
        assert header_relationships
        relationship = header_relationships[0]
        old_target = relationship.get("Target")
        assert old_target
        old_part = posixpath.normpath(posixpath.join("word", old_target))
        header = etree.fromstring(archive.read(old_part))

    paragraph = etree.SubElement(header, f"{{{_W_NS}}}p")
    p_pr = etree.SubElement(paragraph, f"{{{_W_NS}}}pPr")
    p_pr.append(deepcopy(numbered))
    run = etree.SubElement(paragraph, f"{{{_W_NS}}}r")
    text = etree.SubElement(run, f"{{{_W_NS}}}t")
    text.text = "Header list counter reference"

    new_part = "client-data/list-counter-header.dat"
    relationship.set("Target", f"../{new_part}")
    header_overrides = [
        override
        for override in content_types.findall(f"{{{_CT_NS}}}Override")
        if override.get("PartName") == f"/{old_part}"
    ]
    assert len(header_overrides) == 1
    header_overrides[0].set("PartName", f"/{new_part}")
    return rewrite_zip_members(
        payload,
        replacements={
            "word/_rels/document.xml.rels": _serialize(relationships),
            "[Content_Types].xml": _serialize(content_types),
        },
        omit={old_part},
        additions=[(new_part, _serialize(header))],
    )


def _with_non_whitespace_body_tail(payload: bytes) -> bytes:
    root = etree.fromstring(document_xml(payload))
    target = root.xpath(
        f'.//w:p[@w14:paraId="{NUMBERED_ISLAND_PARA_IDS[0]}"]',
        namespaces=_NS,
    )
    assert len(target) == 1
    target[0].tail = "UNMAPPED-BODY-SENTINEL"
    return rewrite_zip_members(
        payload,
        replacements={"word/document.xml": _serialize(root)},
    )


def _outside_island(payload: bytes, *, added_text: str | None = None) -> list[bytes]:
    retained: list[bytes] = []
    for child in _body(payload):
        if child.tag == qn("w:p"):
            if _paragraph_id(child) in NUMBERED_ISLAND_PARA_IDS:
                continue
            if added_text is not None and _paragraph_text(child) == added_text:
                continue
        retained.append(_canonical(child))
    return retained


def _assert_unrelated_body_and_package_unchanged(
    before: bytes,
    after: bytes,
    *,
    added_text: str | None = None,
) -> None:
    assert_untouched_parts_identical(before, after)
    assert _outside_island(before) == _outside_island(after, added_text=added_text)
    before_body = _body(before)
    after_body = _body(after)
    before_sect = before_body.find(qn("w:sectPr"))
    after_sect = after_body.find(qn("w:sectPr"))
    assert before_sect is not None and after_sect is not None
    assert _canonical(after_sect) == _canonical(before_sect)
    assert after_body[-1] is after_sect, "w:sectPr must remain the final body child"


def _assert_original_paragraphs_exact(
    before: bytes,
    after: bytes,
    *,
    retained_ids: tuple[str, ...] = NUMBERED_ISLAND_PARA_IDS,
) -> None:
    for para_id in retained_ids:
        assert _canonical(_direct_paragraph_by_id(after, para_id)) == _canonical(
            _direct_paragraph_by_id(before, para_id)
        )


def _assert_rejected_atomically(client: TestClient, operation: dict) -> str:
    before = client.get("/api/doc").json()["doc"]
    source_before = _source_export(client)
    assert source_before.status_code == 200
    rejected = _edit(client, operation)
    assert rejected.status_code == 400, rejected.text
    assert client.get("/api/doc").json()["doc"] == before
    assert _source_export(client).content == source_before.content
    return rejected.json()["error"]


def test_numbered_island_noop_is_exact_source_bytes(tmp_path):
    client = _client()
    source = make_numbered_island_master(tmp_path)
    imported = _import_master(client, source)

    assert imported["preservation_ready"] is True
    assert _source_export(client).content == source
    assert client.get("/api/export/docx").content == source


@pytest.mark.parametrize(
    "gap_markup",
    [
        pytest.param(b"\r\n<!-- internal island comment -->\r\n", id="comment"),
        pytest.param(b"\n<?internal-island keep='opaque'?>\n", id="pi"),
    ],
)
def test_comment_or_pi_inside_structural_gap_fails_closed(tmp_path, gap_markup):
    source = _with_internal_island_gaps(
        make_numbered_island_master(tmp_path),
        gap_markup,
        b"\r\n  ",
    )
    client = _client()
    _import_master(client, source)
    assert _source_export(client).content == source

    error = _assert_rejected_atomically(
        client,
        {"action": "move", "target_id": "pt1.a1.p3", "position": 0},
    )
    assert "[noncontiguous_structural_island]" in error


def test_local_word_prefix_is_redeclared_only_on_synthesized_paragraph(tmp_path):
    source, p_pr, r_pr = _make_paragraph_local_word_prefix_master(tmp_path)
    client = _client()
    imported = _import_master(client, source)
    assert imported["source_preservation"]["status"] == "ready"

    added = _edit(
        client,
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "position": 1,
            "text": _NEW_TEXT,
            "status": "confirmed",
        },
    )
    assert added.status_code == 200, added.text
    output = _source_export(client).content
    assert_valid_docx_package(output)

    synthesized = _raw_direct_paragraph_by_text(output, _NEW_TEXT)
    expected_new = _expected_local_prefix_new_paragraph(p_pr, r_pr, _NEW_TEXT)
    assert synthesized == expected_new
    source_elements = _source_island_bytes(source)[3]
    _assert_exact_island_rewrite(
        source,
        output,
        (
            source_elements[0],
            expected_new,
            source_elements[1],
            source_elements[2],
        ),
    )
    assert _source_export(client).content == output
    assert Document(io.BytesIO(output)).paragraphs


def test_add_clones_only_safe_format_shape_and_preserves_real_numbering(tmp_path):
    client = _client()
    source = _with_distinct_whitespace_gaps(
        make_numbered_island_master(tmp_path)
    )
    _import_master(client, source)

    added = _edit(
        client,
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "position": 1,
            "text": _NEW_TEXT,
            "status": "confirmed",
        },
    )
    assert added.status_code == 200, added.text
    assert added.json()["applied"][0]["id"] == "pt1.a1.p4"

    exported = _source_export(client)
    assert exported.status_code == 200, exported.text
    output = exported.content
    assert output != source
    assert_valid_docx_package(output)
    _assert_unrelated_body_and_package_unchanged(
        source, output, added_text=_NEW_TEXT
    )
    _assert_original_paragraphs_exact(source, output)
    source_elements = _source_island_bytes(source)[3]
    synthesized_raw = _raw_direct_paragraph_by_text(output, _NEW_TEXT)
    _assert_exact_island_rewrite(
        source,
        output,
        (
            source_elements[0],
            synthesized_raw,
            source_elements[1],
            source_elements[2],
        ),
    )

    numbered = _direct_numbered_paragraphs(output)
    assert [_paragraph_text(paragraph) for paragraph in numbered] == [
        NUMBERED_ISLAND_TEXTS[0],
        _NEW_TEXT,
        NUMBERED_ISLAND_TEXTS[1],
        NUMBERED_ISLAND_TEXTS[2],
    ]
    # The marker is still produced by Word numbering, never literal text.
    assert all(
        not _paragraph_text(paragraph).startswith(("A. ", "B. ", "C. ", "D. "))
        for paragraph in numbered
    )

    template = _direct_paragraph_by_id(source, NUMBERED_ISLAND_PARA_IDS[0])
    synthesized = _direct_paragraph_by_text(output, _NEW_TEXT)
    assert _canonical(synthesized.find(qn("w:pPr"))) == _canonical(
        template.find(qn("w:pPr"))
    )
    template_run = template.find(qn("w:r"))
    new_run = synthesized.find(qn("w:r"))
    assert template_run is not None and new_run is not None
    assert _canonical(new_run.find(qn("w:rPr"))) == _canonical(
        template_run.find(qn("w:rPr"))
    )
    assert synthesized.attrib == {}
    assert new_run.attrib == {}
    assert [etree.QName(child).localname for child in synthesized] == ["pPr", "r"]
    assert [etree.QName(child).localname for child in new_run] == ["rPr", "t"]
    assert len(synthesized.xpath(".//w:t", namespaces=_NS)) == 1
    assert not any(
        etree.QName(element).localname in _RANGE_OR_REVISION_TAGS
        for element in synthesized.iter()
    )
    assert not any(
        etree.QName(attribute).localname.startswith("rsid")
        for element in synthesized.iter()
        for attribute in element.attrib
    )

    # Export is deterministic and neither numbering definitions nor any
    # source relationship/style/media/header/footer part was regenerated.
    assert _source_export(client).content == output


def test_multiple_additions_preserve_every_original_element_and_gap(tmp_path):
    first_text = "Provide a first inserted supervised device."
    middle_text = "Provide a middle inserted supervised device."
    last_text = "Provide a last inserted supervised device."
    source = _with_distinct_whitespace_gaps(
        make_numbered_island_master(tmp_path)
    )
    client = _client()
    _import_master(client, source)

    added = _edit(
        client,
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "position": 0,
            "text": first_text,
        },
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "position": 2,
            "text": middle_text,
        },
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "text": last_text,
        },
    )
    assert added.status_code == 200, added.text
    output = _source_export(client).content
    source_elements = _source_island_bytes(source)[3]
    desired = (
        _raw_direct_paragraph_by_text(output, first_text),
        source_elements[0],
        _raw_direct_paragraph_by_text(output, middle_text),
        source_elements[1],
        source_elements[2],
        _raw_direct_paragraph_by_text(output, last_text),
    )
    _assert_exact_island_rewrite(source, output, desired)
    assert _source_export(client).content == output


def test_delete_keeps_surviving_numbered_paragraph_xml_exact(tmp_path):
    client = _client()
    source = _with_distinct_whitespace_gaps(
        make_numbered_island_master(tmp_path)
    )
    _import_master(client, source)

    deleted = _edit(client, {"action": "delete", "target_id": "pt1.a1.p2"})
    assert deleted.status_code == 200, deleted.text
    output = _source_export(client).content
    assert_valid_docx_package(output)
    _assert_unrelated_body_and_package_unchanged(source, output)
    _assert_original_paragraphs_exact(
        source,
        output,
        retained_ids=(NUMBERED_ISLAND_PARA_IDS[0], NUMBERED_ISLAND_PARA_IDS[2]),
    )
    source_elements = _source_island_bytes(source)[3]
    _assert_exact_island_rewrite(
        source,
        output,
        (source_elements[0], source_elements[2]),
    )
    assert _body(output).xpath(
        f'count(./w:p[@w14:paraId="{NUMBERED_ISLAND_PARA_IDS[1]}"])',
        namespaces=_NS,
    ) == 0.0
    assert [_paragraph_text(paragraph) for paragraph in _direct_numbered_paragraphs(output)] == [
        NUMBERED_ISLAND_TEXTS[0],
        NUMBERED_ISLAND_TEXTS[2],
    ]


def test_delete_entire_numbered_island_preserves_every_boundary_block(tmp_path):
    client = _client()
    source = _with_distinct_whitespace_gaps(
        make_numbered_island_master(tmp_path)
    )
    _import_master(client, source)

    deleted = _edit(
        client,
        *(
            {"action": "delete", "target_id": f"pt1.a1.p{index}"}
            for index in (1, 2, 3)
        ),
    )
    assert deleted.status_code == 200, deleted.text
    output = _source_export(client).content
    assert_valid_docx_package(output)
    _assert_unrelated_body_and_package_unchanged(source, output)
    assert _direct_numbered_paragraphs(output) == []
    _assert_exact_island_rewrite(source, output, ())

    # Once every source anchor is gone, a later addition has no surviving
    # neighbor that can prove which island/template owns it. Undo is the safe
    # way back; the exporter does not guess from an operation history.
    rejected = _edit(
        client,
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "text": _NEW_TEXT,
        },
    )
    assert rejected.status_code == 400
    assert "[ambiguous_structural_insert]" in rejected.json()["error"]
    assert _source_export(client).content == output

    assert client.post("/api/doc/undo").status_code == 200
    assert _source_export(client).content == source


@pytest.mark.parametrize(
    ("target_id", "position", "previous_position", "desired_order"),
    [
        pytest.param(
            "pt1.a1.p3",
            0,
            2,
            (2, 0, 1),
            id="last-to-first",
        ),
        pytest.param(
            "pt1.a1.p1",
            2,
            0,
            (1, 2, 0),
            id="first-to-last",
        ),
    ],
)
def test_move_reorders_whole_original_elements_without_rewriting_them(
    tmp_path,
    monkeypatch,
    target_id,
    position,
    previous_position,
    desired_order,
):
    client = _client()
    source = _with_distinct_whitespace_gaps(
        make_numbered_island_master(tmp_path)
    )
    _import_master(client, source)

    def serializer_must_not_run(*_args, **_kwargs):
        raise AssertionError("structural source export used the tree serializer")

    monkeypatch.setattr(
        source_patch_module,
        "_serialize_tree",
        serializer_must_not_run,
        raising=False,
    )

    moved = _edit(
        client,
        {"action": "move", "target_id": target_id, "position": position},
    )
    assert moved.status_code == 200, moved.text
    assert moved.json()["applied"] == [
        {
            "action": "move",
            "id": target_id,
            "position": position,
            "previous_position": previous_position,
        }
    ]

    output = _source_export(client).content
    assert_valid_docx_package(output)
    _assert_unrelated_body_and_package_unchanged(source, output)
    _assert_original_paragraphs_exact(source, output)
    source_elements = _source_island_bytes(source)[3]
    desired_elements = tuple(source_elements[index] for index in desired_order)
    _assert_exact_island_rewrite(source, output, desired_elements)
    assert [_paragraph_id(paragraph) for paragraph in _direct_numbered_paragraphs(output)] == [
        NUMBERED_ISLAND_PARA_IDS[index] for index in desired_order
    ]


def test_structural_undo_returns_exact_source_and_redo_is_deterministic(tmp_path):
    client = _client()
    source = _with_distinct_whitespace_gaps(
        make_numbered_island_master(tmp_path)
    )
    _import_master(client, source)
    assert _edit(
        client,
        {"action": "move", "target_id": "pt1.a1.p3", "position": 0},
    ).status_code == 200
    moved = _source_export(client).content
    source_elements = _source_island_bytes(source)[3]
    _assert_exact_island_rewrite(
        source,
        moved,
        (source_elements[2], source_elements[0], source_elements[1]),
    )
    assert _source_export(client).content == moved

    assert client.post("/api/doc/undo").status_code == 200
    assert _source_export(client).content == source
    assert client.post("/api/doc/redo").status_code == 200
    assert _source_export(client).content == moved


def test_composed_final_state_add_move_delete_and_text_patch_is_surgical(tmp_path):
    """The exporter derives one safe final plan across mixed edit kinds."""

    client = _client()
    source = _with_distinct_whitespace_gaps(
        make_numbered_island_master(tmp_path)
    )
    _import_master(client, source)
    edited_text = "Provide monitored control valves and tamper switches."

    composed = _edit(
        client,
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "position": 1,
            "text": _NEW_TEXT,
            "status": "confirmed",
        },
        {"action": "move", "target_id": "pt1.a1.p3", "position": 0},
        {"action": "delete", "target_id": "pt1.a1.p2"},
        {
            "action": "replace",
            "target_id": "pt1.a1.p1",
            "text": edited_text,
            "status": "confirmed",
        },
    )
    assert composed.status_code == 200, composed.text

    output = _source_export(client).content
    assert_valid_docx_package(output)
    _assert_unrelated_body_and_package_unchanged(
        source, output, added_text=_NEW_TEXT
    )
    assert [_paragraph_text(paragraph) for paragraph in _direct_numbered_paragraphs(output)] == [
        NUMBERED_ISLAND_TEXTS[2],
        edited_text,
        _NEW_TEXT,
    ]
    source_elements = _source_island_bytes(source)[3]
    patched_first = _raw_paragraph_with_text_patch(
        source,
        NUMBERED_ISLAND_PARA_IDS[0],
        edited_text,
    )
    synthesized_raw = _raw_direct_paragraph_by_text(output, _NEW_TEXT)
    _assert_exact_island_rewrite(
        source,
        output,
        (source_elements[2], patched_first, synthesized_raw),
    )

    # The moved surviving original stays completely byte/XML-equivalent.
    assert _canonical(
        _direct_paragraph_by_id(output, NUMBERED_ISLAND_PARA_IDS[2])
    ) == _canonical(_direct_paragraph_by_id(source, NUMBERED_ISLAND_PARA_IDS[2]))
    assert not _body(output).xpath(
        f'./w:p[@w14:paraId="{NUMBERED_ISLAND_PARA_IDS[1]}"]', namespaces=_NS
    )

    # The other surviving original changes in exactly one existing w:t.
    source_edited = _direct_paragraph_by_id(source, NUMBERED_ISLAND_PARA_IDS[0])
    output_edited = _direct_paragraph_by_id(output, NUMBERED_ISLAND_PARA_IDS[0])
    source_text_nodes = source_edited.xpath(".//w:t", namespaces=_NS)
    output_text_nodes = output_edited.xpath(".//w:t", namespaces=_NS)
    assert len(source_text_nodes) == len(output_text_nodes) == 1
    assert source_text_nodes[0].text == NUMBERED_ISLAND_TEXTS[0]
    assert output_text_nodes[0].text == edited_text
    expected_edited = deepcopy(source_edited)
    expected_edited.xpath(".//w:t", namespaces=_NS)[0].text = edited_text
    assert _canonical(output_edited) == _canonical(expected_edited)

    # The new paragraph still consists only of the proven formatting shape
    # plus one fresh text run; no source identity or revision metadata leaks.
    synthesized = _direct_paragraph_by_text(output, _NEW_TEXT)
    template = _direct_paragraph_by_id(source, NUMBERED_ISLAND_PARA_IDS[0])
    assert _canonical(synthesized.find(qn("w:pPr"))) == _canonical(
        template.find(qn("w:pPr"))
    )
    assert synthesized.attrib == {}
    new_run = synthesized.find(qn("w:r"))
    assert new_run is not None and new_run.attrib == {}
    template_run = template.find(qn("w:r"))
    assert template_run is not None
    assert _canonical(new_run.find(qn("w:rPr"))) == _canonical(
        template_run.find(qn("w:rPr"))
    )
    assert [etree.QName(child).localname for child in synthesized] == ["pPr", "r"]
    assert [etree.QName(child).localname for child in new_run] == ["rPr", "t"]
    assert not any(
        etree.QName(element).localname in _RANGE_OR_REVISION_TAGS
        for element in synthesized.iter()
    )
    assert _source_export(client).content == output


@pytest.mark.parametrize(
    "operation",
    [
        pytest.param(
            {"action": "delete", "target_id": "pt1.a1.p1"},
            id="delete-manual-label",
        ),
        pytest.param(
            {"action": "move", "target_id": "pt1.a1.p1", "position": 1},
            id="move-manual-label",
        ),
        pytest.param(
            {
                "action": "add_paragraph",
                "target_id": "pt1.a1",
                "position": 1,
                "text": _NEW_TEXT,
            },
            id="add-to-manual-label-island",
        ),
    ],
)
def test_manual_label_structural_edits_fail_closed(tmp_path, operation):
    client = _client()
    source = make_fidelity_master(tmp_path)
    _import_master(client, source)
    error = _assert_rejected_atomically(client, operation)
    assert any(
        blocker in error
        for blocker in (
            "[manual_label_structural_change]",
            "[ambiguous_structural_insert]",
            "[unsafe_structural_island]",
        )
    )


@pytest.mark.parametrize("separator", ["empty", "sdt"])
def test_noncontiguous_numbered_island_is_rejected(tmp_path, separator):
    client = _client()
    source = make_numbered_island_master(tmp_path, separator=separator)
    _import_master(client, source)
    error = _assert_rejected_atomically(
        client,
        {"action": "move", "target_id": "pt1.a1.p3", "position": 0},
    )
    assert "[noncontiguous_structural_island]" in error


def test_mixed_num_id_island_is_rejected(tmp_path):
    client = _client()
    source = make_numbered_island_master(tmp_path, mixed_num_id=True)
    _import_master(client, source)
    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt1.a1.p2"},
    )
    assert "[mixed_numbering_island]" in error


@pytest.mark.parametrize("invalid_num_id", ["zero", "dangling"])
def test_unproven_word_numbering_is_pass_through_only(tmp_path, invalid_num_id):
    client = _client()
    source = make_numbered_island_master(
        tmp_path, invalid_num_id=invalid_num_id
    )
    _import_master(client, source)
    assert _source_export(client).content == source

    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt1.a1.p2"},
    )
    assert "[automatic_numbering_required]" in error


def test_orphan_numbering_part_is_not_treated_as_a_word_list(tmp_path):
    client = _client()
    source = remove_numbering_relationship(make_numbered_island_master(tmp_path))
    _import_master(client, source)
    assert _source_export(client).content == source

    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt1.a1.p2"},
    )
    assert "[automatic_numbering_required]" in error


def test_mistyped_numbering_part_is_not_treated_as_a_word_list(tmp_path):
    client = _client()
    source = replace_numbering_content_type(
        make_numbered_island_master(tmp_path),
        "application/xml",
    )
    _import_master(client, source)
    assert _source_export(client).content == source

    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt1.a1.p2"},
    )
    assert "[unsafe_revision_scan]" in error


def test_constant_level_marker_is_not_claimed_as_auto_incrementing(tmp_path):
    client = _client()
    source = make_numbered_island_master(
        tmp_path,
        constant_level_text="CONSTANT-LABEL",
    )
    _import_master(client, source)
    assert _source_export(client).content == source

    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt1.a1.p2"},
    )
    assert "[automatic_numbering_required]" in error


def test_real_word_bullet_level_remains_a_safe_static_marker_exception(tmp_path):
    client = _client()
    source = make_numbered_island_master(
        tmp_path,
        constant_level_text="•",
        level_number_format="bullet",
    )
    _import_master(client, source)

    added = _edit(
        client,
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "position": 1,
            "text": _NEW_TEXT,
        },
    )
    assert added.status_code == 200, added.text
    output = _source_export(client).content
    assert_valid_docx_package(output)
    assert _paragraph_text(_direct_paragraph_by_text(output, _NEW_TEXT)) == _NEW_TEXT


def test_numbering_instance_cannot_leak_beyond_the_structural_island(tmp_path):
    client = _client()
    source = make_numbered_island_master(
        tmp_path,
        separator="empty_after_second",
    )
    _import_master(client, source)

    # The first two provisions are locally contiguous, but the third still
    # references their document-wide numId after an opaque body boundary.
    # Deleting either local item would visibly renumber that untouched item.
    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt1.a1.p1"},
    )
    assert "[numbering_instance_not_isolated]" in error


def test_numbering_use_in_correctly_typed_non_xml_header_blocks_island(tmp_path):
    client = _client()
    source = _with_relocated_header_numbering_use(
        make_numbered_island_master(tmp_path)
    )
    _import_master(client, source)

    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt1.a1.p1"},
    )
    assert "[numbering_instance_not_isolated]" in error


def test_unmapped_body_character_data_makes_source_pass_through_only(tmp_path):
    client = _client()
    source = _with_non_whitespace_body_tail(
        make_numbered_island_master(tmp_path)
    )
    _import_master(client, source)
    assert _source_export(client).content == source

    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt1.a1.p1"},
    )
    assert "[unsafe_document_xml]" in error


def test_final_state_gate_rejects_text_that_needs_xml_space_metadata(tmp_path):
    source = make_numbered_island_master(tmp_path)
    parsed = parse_master_docx(tmp_path / "client-numbered-island-master.docx")
    assert parsed.source_map is not None
    current = deepcopy(parsed.section)
    article = current.parts[0].articles[0]
    article.paragraphs.insert(
        1,
        Paragraph(uid="pt1.a1.p4", text="  padded project text  "),
    )

    with pytest.raises(SourcePatchError) as exc_info:
        validate_source_transition(
            source_bytes=source,
            source_map=parsed.source_map,
            baseline=parsed.section,
            current=current,
        )
    assert exc_info.value.blocker == "unsupported_edge_whitespace"


def test_add_requires_one_unambiguous_format_template_but_move_does_not(tmp_path):
    client = _client()
    source = _with_distinct_whitespace_gaps(
        make_numbered_island_master(tmp_path, inconsistent_format=True)
    )
    _import_master(client, source)

    error = _assert_rejected_atomically(
        client,
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "position": 1,
            "text": _NEW_TEXT,
        },
    )
    assert "[ambiguous_structural_template]" in error

    # Reordering does not synthesize formatting: each original paragraph is
    # moved whole, so heterogeneous local formatting remains byte/XML exact.
    moved = _edit(
        client,
        {"action": "move", "target_id": "pt1.a1.p3", "position": 0},
    )
    assert moved.status_code == 200, moved.text
    output = _source_export(client).content
    _assert_original_paragraphs_exact(source, output)
    _assert_unrelated_body_and_package_unchanged(source, output)
    source_elements = _source_island_bytes(source)[3]
    _assert_exact_island_rewrite(
        source,
        output,
        (source_elements[2], source_elements[0], source_elements[1]),
    )


def test_unknown_word_property_blocks_addition_but_not_exact_move(tmp_path):
    source = _with_distinct_whitespace_gaps(
        _with_unknown_clone_property(make_numbered_island_master(tmp_path))
    )
    client = _client()
    _import_master(client, source)

    error = _assert_rejected_atomically(
        client,
        {
            "action": "add_paragraph",
            "target_id": "pt1.a1",
            "position": 1,
            "text": _NEW_TEXT,
        },
    )
    assert "[ambiguous_structural_template]" in error

    moved = _edit(
        client,
        {"action": "move", "target_id": "pt1.a1.p3", "position": 0},
    )
    assert moved.status_code == 200, moved.text
    output = _source_export(client).content
    source_elements = _source_island_bytes(source)[3]
    _assert_exact_island_rewrite(
        source,
        output,
        (source_elements[2], source_elements[0], source_elements[1]),
    )


def test_nested_numbered_subtree_is_rejected_instead_of_orphaned(tmp_path):
    client = _client()
    source = make_numbered_island_master(tmp_path, ilvls=(0, 1, 0))
    imported = _import_master(client, source)
    first = imported["doc"]["parts"][0]["articles"][0]["paragraphs"][0]
    assert [child["id"] for child in first["children"]] == ["pt1.a1.p1.p1"]

    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt1.a1.p1"},
    )
    assert "[nested_structural_change]" in error


@pytest.mark.parametrize("complex_middle", ["field", "hyperlink"])
def test_complex_numbered_paragraph_cannot_be_structurally_mutated(
    tmp_path, complex_middle
):
    client = _client()
    source = make_numbered_island_master(
        tmp_path, complex_middle=complex_middle
    )
    _import_master(client, source)
    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt1.a1.p2"},
    )
    assert (
        "[unsafe_structural_island]" in error
        or "[complex_paragraph_markup]" in error
        or "[complex_run_markup]" in error
    )


def test_table_projection_cannot_be_deleted_as_if_it_were_a_paragraph(tmp_path):
    client = _client()
    source = make_table_projection_master(tmp_path)
    _import_master(client, source)
    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt2.a1.p1"},
    )
    assert "[table_projection]" in error or "[unsafe_structural_island]" in error


def test_move_schema_cannot_reparent_across_semantic_parents(tmp_path):
    client = _client()
    source = make_numbered_island_master(tmp_path)
    _import_master(client, source)
    error = _assert_rejected_atomically(
        client,
        {
            "action": "move",
            "target_id": "pt1.a1.p3",
            "position": 0,
            "parent_id": "pt2",
        },
    )
    assert "move" in error.lower()
    assert "parent_id" in error or "unsupported" in error.lower()


@pytest.mark.parametrize(
    ("source_mutator", "expected_blocker"),
    [
        pytest.param(
            add_document_protection,
            "document_protection",
            id="document-protection",
        ),
        pytest.param(add_tracked_change, "tracked_changes", id="tracked-changes"),
        pytest.param(
            add_paragraph_property_change,
            "tracked_changes",
            id="paragraph-property-change",
        ),
        pytest.param(
            add_signature_origin_marker,
            "signed_package",
            id="signed-package",
        ),
        pytest.param(
            add_active_content_marker,
            "active_content",
            id="active-content",
        ),
    ],
)
def test_global_blockers_allow_structural_noop_but_reject_mutation(
    tmp_path, source_mutator, expected_blocker
):
    client = _client()
    source = source_mutator(make_numbered_island_master(tmp_path))
    _import_master(client, source)
    assert _source_export(client).content == source

    error = _assert_rejected_atomically(
        client,
        {"action": "delete", "target_id": "pt1.a1.p2"},
    )
    assert f"[{expected_blocker}]" in error
    # The recovery copy remains the immutable original package as well.
    assert client.get("/api/import/original").content == source
    assert Document(io.BytesIO(source)).sections
